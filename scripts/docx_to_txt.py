#!/usr/bin/env python3
"""
Convert a .docx file to plain text (.txt), preserving tables as TSV.

Usage:
    python docx_to_txt.py input.docx
    python docx_to_txt.py input.docx output.txt
    python docx_to_txt.py docs/*.docx          # batch mode
"""

import sys
import os
import docx


def docx_to_txt(input_path):
    """Read a .docx and return its full text content."""
    doc = docx.Document(input_path)
    lines = []

    # Track element order (paragraphs + tables interleaved)
    body = doc.element.body
    para_idx = 0
    table_idx = 0

    for child in body:
        tag = child.tag.split('}')[-1]

        if tag == 'p':
            if para_idx < len(doc.paragraphs):
                p = doc.paragraphs[para_idx]
                text = p.text.strip()
                style = p.style.name if p.style else ''

                # Add heading markers
                if 'Heading 1' in style:
                    lines.append(f"\n{'='*60}")
                    lines.append(text)
                    lines.append('='*60)
                elif 'Heading 2' in style:
                    lines.append(f"\n{'-'*40}")
                    lines.append(text)
                    lines.append('-'*40)
                elif text:
                    lines.append(text)
                else:
                    lines.append('')

            para_idx += 1

        elif tag == 'tbl':
            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                lines.append('')
                for i, row in enumerate(table.rows):
                    cells = [cell.text.replace('\n', ' ').strip()
                             for cell in row.cells]
                    lines.append('\t'.join(cells))
                    if i == 0:
                        lines.append('\t'.join(['---'] * len(cells)))
                lines.append('')

            table_idx += 1

    return '\n'.join(lines)


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    files = sys.argv[1:]

    for input_path in files:
        if not os.path.exists(input_path):
            print(f"  ! File not found: {input_path}")
            continue

        if not input_path.endswith('.docx'):
            print(f"  ! Not a .docx file: {input_path}")
            continue

        # Output path: same name with .txt, or second arg if single file
        if len(files) == 1 and len(sys.argv) == 3 and not sys.argv[2].endswith('.docx'):
            output_path = sys.argv[2]
        else:
            output_path = input_path.replace('.docx', '.txt')

        text = docx_to_txt(input_path)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)

        word_count = len(text.split())
        print(f"  > {os.path.basename(input_path)} -> {os.path.basename(output_path)} ({word_count} words)")


if __name__ == '__main__':
    main()
