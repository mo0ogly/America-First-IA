from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_section_header(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_subsection_header(doc, text):
    doc.add_heading(text, level=2)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def generate_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('Refonte Profil LinkedIn - Fabrice Pizzi', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Pivot vers l\'AI Security Senior\nCible : Head of AI Security, Principal / Staff Research Engineer LLM Safety, Distinguished Red Team Lead.')
    
    # 1. DIAGNOSTIC STRATEGIQUE
    add_section_header(doc, '1. DIAGNOSTIC STRATEGIQUE')
    
    add_subsection_header(doc, '1.1 Etat actuel du profil')
    doc.add_paragraph('Headline actuelle reperee via indexation Google : "Cyberthreats Information Management". Cette formulation est un sous-emploi massif. Elle ne capture ni la seniorite (28 ans d\'XP, 3 mandats CISO, double laureat RSSI de l\'Annee), ni le pivot AI Security qui est la valeur differenciante actuelle, ni les credentials academiques (M2 Sorbonne en cours, recherche sur agents IA autonomes, indice CACI, couverture ZATAZ).')
    
    add_subsection_header(doc, '1.2 Probleme central')
    doc.add_paragraph('Le profil envoie deux signaux contradictoires : un opacite institutionnelle ANSSI (normal, contrainte OPSEC) face a une production de recherche publique tres riche (140+ pages OpenClaw, 103 pages CACI, lab Da Vinci, plateformes LIA-Scan et Wattson). Resultat : un recruteur AI security qui te decouvre via tes publications ne trouve pas, sur ton LinkedIn, le seniority signal qui justifierait un Head of / Principal / Distinguished. A l\'inverse, un recruteur cyberdefense classique ne percoit pas le pivot IA.')
    
    add_subsection_header(doc, '1.3 Hypothese de positionnement')
    doc.add_paragraph('Construire une narration unique : "L\'analyste etatique qui applique la doctrine de cyberdefense nationale a la securite des agents IA autonomes". C\'est un angle non-occupe sur le marche AI safety, ou la majorite des chercheurs vient soit du ML academique, soit de la securite produit, rarement de la cyberdefense souveraine avec experience operationnelle G7 / JO / COP.')
    
    # 2. HEADLINES
    add_section_header(doc, '2. HEADLINES | 3 VERSIONS A A/B TESTER')
    
    add_subsection_header(doc, 'Version A | Hybride senior cyber + recherche AEGIS (Recommandee)')
    add_paragraph(doc, 'AI Security & Agentic Systems Threat Modeling | Author AEGIS Lab (98 attack templates, SVC 6D + Sep(M)) | Senior Cyber Defense ANSSI 10+ yrs | Sorbonne M2 IA-Cyber | 3x CISO | RSSI of the Year x2', bold=True)
    
    add_subsection_header(doc, 'Version B | Offensif chercheur (mo0ogly forward)')
    add_paragraph(doc, 'LLM Red Teamer & Agentic AI Threat Researcher | Author AEGIS Lab Da Vinci v4.2 (Zenodo DOI) | 98 attack templates, SVC 6D + Sep(M) | Senior Cyber ANSSI | mo0ogly | CISSP', bold=True)
    
    add_subsection_header(doc, 'Version C | Strategique souverainete')
    add_paragraph(doc, 'AI Security Strategy | Sovereign Compute & Agentic Threats | Author CACI Index (US-EU compute asymmetry) + AEGIS Red Team Lab | ANSSI senior | 3x CISO | EU AI Act, NIS2, DORA', bold=True)
    
    # 3. SECTION "INFOS" (ABOUT) - FR
    add_section_header(doc, '3. SECTION "INFOS" (ABOUT) | VERSION FR')
    doc.add_paragraph('Vingt-huit ans en cybersecurite operationnelle. Dix a l\'ANSSI sur la cyberdefense nationale (G7, Jeux Olympiques, COP, elections nationales et europeennes). Trois mandats CISO anterieurs (Eiffage 71 000 collaborateurs, ministere du Travail, BIPOP). Double laureat du trophee RSSI de l\'Annee.')
    doc.add_paragraph('Depuis 2024, je pivote vers la securite des systemes d\'IA autonomes via un programme de recherche personnel (handle : mo0ogly). Pas par opportunisme : la doctrine de cyberdefense souveraine que j\'applique depuis dix ans (kill chain, RETEX, gestion de crise multi-acteurs) est exactement ce qui manque dans la litterature AI safety, encore largement dominee par des angles ML academiques.')
    
    add_paragraph(doc, 'Programme de recherche AEGIS', bold=True)
    doc.add_paragraph('  131 papiers analyses (corpus RAG ChromaDB, pipeline PDCA auto-declenchant), 8 conjectures C1-C8 (4 saturees a 10-10), 28 decouvertes D-001 a D-028. Cadre formel : SVC 6 dimensions, score de separation Sep(M) (Zverev et al. ICLR 2025), Integrity(S) sur DY-AGENT, Delta-0 baseline.')
    doc.add_paragraph('  D-001 Triple Convergence (article publie) : preuve que delta-0 (RLHF) + delta-1 (system prompt) + delta-2 (juges LLM) sont simultanement vulnerables. delta-3 (verification formelle externe) est la seule couche survivante.')
    doc.add_paragraph('  D-024 HyDE Self-Amplification : contribution originale, 96.7% ASR (29/30) sur llama-3.1-8b. Aucun papier du corpus (P001-P121) n\'avait identifie HyDE comme vecteur d\'attaque endogene pre-retrieval.')
    doc.add_paragraph('  AEGIS Lab | Da Vinci v4.2 (DOI 10.5281/zenodo.19854891) : POC live d\'un robot chirurgical Da Vinci sous attaque LLM, 98 templates d\'attaque API-served, 4 scenarios MITRE ATT&CK (T1565.001 Slow Poison, T1486 Ransomware, T1059.009 Defense bypass), 66 techniques de defense (Prevention / Detection / Response / Measurement), multi-LLM PromptForge (Claude Opus / Sonnet / Haiku, GPT-4o, Gemini, Grok, Llama via Groq / Ollama).')
    
    add_paragraph(doc, 'Productions complementaires', bold=True)
    doc.add_paragraph('  Operation OpenClaw (2025, 140+ pages, couvert par ZATAZ) : kill chain complete d\'un agent IA autonome malveillant, modele de defense a 5 couches. Support M2 Sorbonne.')
    doc.add_paragraph('  Indice CACI | AI for Americans First (103 pages, 157 notes, trilingue FR/EN/PT-BR) : mesure de l\'asymetrie compute US/EU, scenarios prospectifs 2026-2030.')
    doc.add_paragraph('  LIA-Scan : plateforme d\'audit 300+ technologies, 230+ frameworks (DORA, NIS2, ANSSI), pipeline agentique de detection CVE.')
    
    add_paragraph(doc, 'Stack : Python (FastAPI / Flask), Go, TypeScript, React + Three.js, ChromaDB, Docker, Redis, PostgreSQL, n8n, Anthropic / OpenAI / Groq / Ollama.', bold=True)

    # 4. ABOUT SECTION - EN
    add_section_header(doc, '4. ABOUT SECTION | ENGLISH VERSION')
    doc.add_paragraph('Twenty-eight years in operational cybersecurity. Ten at ANSSI running national cyber defense operations. Three CISO mandates before that. Two-time winner of the French CISO of the Year award.')
    doc.add_paragraph('Since 2024 I have been pivoting toward AI agent security through a personal research program (handle: mo0ogly).')
    
    # 5. FEATURED
    add_section_header(doc, '5. SECTION FEATURED | 5 ITEMS A EPINGLER')
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Description'
    
    items = [
        ('Triple Convergence Paper (D-001)', 'Article de recherche AEGIS.'),
        ('AEGIS Lab | Da Vinci v4.2', 'POC live d\'un robot Da Vinci sous attaque LLM.'),
        ('Operation OpenClaw', 'Kill chain agent IA malveillant, 140+ pages.'),
        ('Indice CACI', 'AI for Americans First, 103 pages.'),
        ('LIA-Scan', 'Plateforme d\'audit 300+ technos.')
    ]
    
    for item, desc in items:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = desc

    # 6. EXPERIENCES
    add_section_header(doc, '6. REFORMULATION DES EXPERIENCES')
    add_subsection_header(doc, '6.1 ANSSI | Senior Cyber Defense Expert')
    doc.add_paragraph('Lead of national cyber defense operations and contributor to AI security threat intelligence.')
    
    add_subsection_header(doc, '6.3 Independent Researcher (mo0ogly)')
    doc.add_paragraph('Self-funded research on AI Agent Security & LLM Red Team (handle: mo0ogly). AEGIS Lab, OpenClaw, CACI.')

    # 7. COMPETENCES
    add_section_header(doc, '7. COMPETENCES | TOP 50')
    doc.add_paragraph('• AI Agent Security (Pinned)\n• LLM Red Teaming (Pinned)\n• Threat Intelligence (Pinned)\n• Adversarial Machine Learning\n• Prompt Injection\n• EU AI Act\n• NIS2\n• CISSP')

    # 10. PLAYBOOK OUTREACH
    add_section_header(doc, '10. PLAYBOOK OUTREACH | 4 SEMAINES')
    doc.add_paragraph('Semaine 1 : Social Proof\nSemaine 2 : Posts engageants\nSemaine 3 : Outreach DM\nSemaine 4 : Conversion')

    # Save
    doc.save('Refonte_LinkedIn_Fabrice_Pizzi.docx')
    print("Document généré : Refonte_LinkedIn_Fabrice_Pizzi.docx")

if __name__ == "__main__":
    generate_document()
