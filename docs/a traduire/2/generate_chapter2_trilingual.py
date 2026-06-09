"""

import sys
from pathlib import Path
sys.path.append(str(Path(__file__).resolve().parents[1]))
import caci_data_helper
m = caci_data_helper.get_metrics()
us_share = m['us_compute_share']
us_eu_caci = m['us_eu_caci_power_ratio']
us_eu_raw = m['us_eu_raw_ratio']
eu_sov = m['eu_sovereignty_ratio']
caci_scores = {k: v['caci_power_phys'] for k, v in m['country_results'].items()}

def fmt_fr(val, decimals=1):
    return f"{val:.{decimals}f}".replace(".", ",")

def fmt_en(val, decimals=1):
    return f"{val:.{decimals}f}"

Chapter II - Methodology - trilingual generator (EN / FR / PT-BR).

Generates the Chapter II .docx for the doctoral study
"AI for Americans First - AI Protectionism, Energy and Semiconductors:
US/Europe Divergence Trajectories 2024-2030".

Key revisions (April 2026 dashboard alignment):
    1. Cover banner: {fmt_en(us_share, 1)}% / 1.59x / {fmt_en(us_eu_caci, 2)}:1 (replaces 77% / 1.6x / 3.4:1,
       all aligned with Chapter I and the live dashboard).
    2. Section 2.4.3 - Energy costs realigned to dashboard CSV
       (USA $85, France $115, Germany $140 per MWh) with explicit note
       that figures reflect hyperscaler-PPA-adjusted prices, not raw
       Eurostat industrial tariffs. The phrasing "EU costs typically
       2 to 3 times US" is replaced with "1.4 to 1.7 times US"
       to match the dashboard reference values.
    3. Section 2.4.3 - National compute shares updated to operational
       Epoch AI snapshot: USA {fmt_en(us_share, 1)}%, China 12.8%, EU 4.4%, Norway 1.1%,
       Japan 1.0% (Apr 2026, was May 2025 figures).
    4. Section 2.4.3 - Normalised CACI scores rebuilt on Power Mode
       F_total: France ~25, Germany ~5, EU ~29 (US = 100), with a
       primary US/EU ratio of {fmt_en(us_eu_caci, 2)}:1.
    5. Section 2.4.6 NEW - "Reproducible numbers - April 2026 snapshot"
       block, explicitly listing the four formula inputs per region and
       the resulting CACI scores so the reader can replay the calculation
       from the public dashboard CSVs.

Author: Fabrice Pizzi (Universite Paris-Sorbonne, M2 Economic Intelligence).
Build : python3 generate_chapter2_trilingual.py
"""

from __future__ import annotations

import sys
from pathlib import Path
sys.path.append(str(Path(__file__).resolve().parents[1]))
import caci_data_helper
m = caci_data_helper.get_metrics()
us_share = m['us_compute_share']
us_eu_caci = m['us_eu_caci_power_ratio']
us_eu_raw = m['us_eu_raw_ratio']
eu_sov = m['eu_sovereignty_ratio']
caci_scores = {k: v['caci_power_phys'] for k, v in m['country_results'].items()}

def fmt_fr(val, decimals=1):
    return f"{val:.{decimals}f}".replace(".", ",")

def fmt_en(val, decimals=1):
    return f"{val:.{decimals}f}"


import logging
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
log = logging.getLogger("chapter2_gen")


# ---------------------------------------------------------------------------
# Visual identity (shared with Chapter I)
# ---------------------------------------------------------------------------

NAVY = RGBColor(0x1A, 0x27, 0x44)
GOLD = RGBColor(0xB8, 0x92, 0x2F)
GREY = RGBColor(0x55, 0x55, 0x55)


@dataclass
class LangPack:
    """Container for one language version of Chapter II."""

    code: str
    filename: str
    cover_title: str
    cover_subtitle: str
    cover_blurb: str
    cover_chip_lines: list[str]
    cover_meta: str
    cover_keywords_label: str
    cover_keywords: str
    chapter_label: str
    chapter_title: str
    sections: list[tuple[str, list[str]]] = field(default_factory=list)
    notes_label: str = "Notes"
    notes: list[str] = field(default_factory=list)
    license_block: list[str] = field(default_factory=list)
    page_footer: str = ""


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------


def set_run(run, *, font="Calibri", size=11, bold=False, italic=False, color=None):
    """Apply consistent typography to a docx run."""
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text, *, style=None, align=None, space_after=6, **run_kwargs):
    """Add a paragraph and return it (text may be empty for spacers)."""
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run(run, **run_kwargs)
    return p


def add_heading(doc, text, level):
    """Custom heading rendering (avoids built-in colored heading styles)."""
    sizes = {1: 22, 2: 16, 3: 13, 4: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, font="Calibri", size=sizes.get(level, 11),
            bold=True, color=NAVY)
    return p


def add_cover(doc, lp: LangPack):
    """Render the cover block (title, KPIs, metadata)."""
    add_paragraph(doc, "", space_after=0)
    banner = {
        "EN": "RESEARCH STUDY - FEBRUARY 2026",
        "FR": "ETUDE DE RECHERCHE - FEVRIER 2026",
        "PT-BR": "ESTUDO DE PESQUISA - FEVEREIRO DE 2026",
    }.get(lp.code, "RESEARCH STUDY - FEBRUARY 2026")
    add_paragraph(doc, banner,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, bold=True, color=GREY, space_after=4)
    add_paragraph(doc, lp.cover_title,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=26, bold=True, color=NAVY, space_after=4)
    add_paragraph(doc, lp.cover_subtitle,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, italic=True, color=GREY, space_after=4)
    add_paragraph(doc, lp.cover_blurb,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=11, italic=True, color=GREY, space_after=18)

    table = doc.add_table(rows=1, cols=3)
    table.autofit = True
    for i, line in enumerate(lp.cover_chip_lines):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        set_run(run, size=11, bold=True, color=NAVY)

    add_paragraph(doc, "", space_after=8)
    add_paragraph(doc, "Fabrice Pizzi",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, bold=True, color=NAVY, space_after=2)
    add_paragraph(doc, "Universite Paris-Sorbonne",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, color=GREY, space_after=2)
    master_line = {
        "EN": "Master 2 Economic Intelligence - Intelligence Warfare",
        "FR": "Master 2 Intelligence Economique - Intelligence Warfare",
        "PT-BR": "Mestrado 2 Inteligencia Economica - Intelligence Warfare",
    }.get(lp.code, "Master 2 Economic Intelligence - Intelligence Warfare")
    add_paragraph(doc, master_line,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, italic=True, color=GREY, space_after=14)
    add_paragraph(doc, lp.cover_meta,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, italic=True, color=GREY, space_after=10)
    add_paragraph(doc, f"{lp.cover_keywords_label}: {lp.cover_keywords}",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=9, italic=True, color=GREY, space_after=24)


def add_chapter_header(doc, lp: LangPack):
    """Render the chapter header strip below the cover."""
    add_paragraph(doc, lp.chapter_label,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=11, bold=True, color=GOLD, space_after=2)
    add_paragraph(doc, lp.chapter_title,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=18, bold=True, color=NAVY, space_after=18)


def render_sections(doc, lp: LangPack):
    """Render the body sections; heading level is inferred from numbering."""
    for heading, paragraphs in lp.sections:
        # detect level from the prefix
        h = heading.strip()
        if h.startswith(("2.1 ", "2.2 ", "2.3 ", "2.4 ", "2.5 ", "2.6 ")):
            add_heading(doc, heading, 2)
        elif h.startswith(("2.2.", "2.4.")):
            add_heading(doc, heading, 3)
        else:
            add_heading(doc, heading, 4)
        for para in paragraphs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.25
            run = p.add_run(para)
            set_run(run, size=11, color=RGBColor(0x20, 0x20, 0x20))


def render_notes(doc, lp: LangPack):
    """Render the footnotes block at the end of the chapter."""
    add_heading(doc, lp.notes_label, 2)
    for i, note in enumerate(lp.notes, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run_id = p.add_run(f"{i}. ")
        set_run(run_id, size=9, bold=True, color=GOLD)
        run_txt = p.add_run(note)
        set_run(run_txt, size=9, color=GREY)


def render_license(doc, lp: LangPack):
    """Render the license disclaimer block."""
    doc.add_paragraph()
    for line in lp.license_block:
        add_paragraph(doc, line,
                      align=WD_ALIGN_PARAGRAPH.LEFT,
                      size=8, italic=True, color=GREY, space_after=2)
    add_paragraph(doc, lp.page_footer,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=8, italic=True, color=GREY, space_after=0)


def build(lp: LangPack, out_dir: Path) -> Path:
    """Build the .docx file for one language pack."""
    log.info("Building Chapter II [%s] -> %s", lp.code, lp.filename)
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    add_cover(doc, lp)
    add_chapter_header(doc, lp)
    render_sections(doc, lp)
    render_notes(doc, lp)
    render_license(doc, lp)

    out = out_dir / lp.filename
    doc.save(out)
    log.info("Saved %s", out)
    return out


# ===========================================================================
# Content - English
# ===========================================================================

EN = LangPack(
    code="EN",
    filename="Chapter_II_Methodologie_EN.docx",
    cover_title="AI FOR AMERICANS FIRST",
    cover_subtitle="AI Protectionism, Energy and Semiconductors: US/Europe Divergence Trajectories 2024-2030",
    cover_blurb="Integrated Geostrategic and Economic Analysis - Chapter II",
    cover_chip_lines=[
        f"{fmt_en(us_share, 1)}% global operational AI compute = USA",
        "1.59x energy cost EU/US",
        f"{fmt_en(us_eu_caci, 2)}:1 CACI ratio US/EU (Power Mode)",
    ],
    cover_meta="Paris - February 2026  |  7 chapters  |  4 prospective scenarios  |  3 geographic zones",
    cover_keywords_label="Keywords",
    cover_keywords=("artificial intelligence, technology protectionism, semiconductors, "
                    "export controls, sovereign compute, AI geopolitics, France, "
                    "United States, China"),
    chapter_label="CHAPTER II",
    chapter_title="Methodology",
    notes_label="Notes",
    license_block=[
        "License and Disclaimer. This work, 'AI for Americans First,' is made available under the terms of the America-First-IA Creative Commons Attribution - NonCommercial - ShareAlike 4.0 International License (CC BY-NC-SA 4.0).",
        "You are free to share and adapt the material for non-commercial purposes, provided you give appropriate credit to Fabrice Pizzi (Universite Paris-Sorbonne) and distribute your contributions under the same license. This document is provided for educational and research purposes only.",
        "Public dashboard: https://mo0ogly.github.io/America-First-IA/dashboard/",
        "Repository: https://github.com/mo0ogly/America-First-IA",
    ],
    page_footer="AI for Americans First - Fabrice Pizzi - Chapter II",
)

EN.sections = [
    ("2.1 General approach: mixed-method multi-scenario prospective analysis", [
        "This study combines a retrospective empirical analysis (2020-2026 diagnosis) with a prospective scenario-based projection (2026-2030). This two-pronged architecture responds to the nature of the phenomenon under investigation: AI technology protectionism is simultaneously an observable fact (export controls, Section 232 tariffs) and an ongoing process whose future trajectory depends on discretionary political variables, European strategic responses and partially unpredictable technological developments.",
        "The retrospective component employs a descriptive quantitative method, based on the aggregation and cross-referencing of data from institutional sources (IEA, SIA/WSTS, Eurostat, EIA), industry reports (McKinsey, Deloitte, Epoch AI) and regulatory documents (Federal Register, BIS, White House). The objective is to establish a rigorous, sourced factual foundation covering three dimensions: data center energy consumption, the semiconductor market, and installed AI compute capacity by region.",
        "The prospective component draws on the scenario planning tradition as formalized by Schwartz (1991) and practiced at Royal Dutch/Shell since the 1970s.[1] This method, belonging to the Intuitive Logics school (Bradfield et al., 2005), consists of constructing plausible and internally consistent scenarios - not to predict the future, but to explore the space of possibilities and evaluate the robustness of different strategies against divergent environmental developments.[2] It is particularly suited to situations characterized by high political and technological uncertainty, where classical econometric models reach their limits, which is precisely the case with AI technology protectionism.",
        "Justification of the methodological choice. Three reasons underpin the choice of scenario methods over pure econometric modelling. First, the key analytical variables are largely political and discretionary: the decision of a US president to impose or not impose GPU quotas on Europe cannot be modelled through a regression function. Second, the interactions between the energy, technological and geopolitical dimensions are non-linear and systemic: a restriction on GPUs can, through cascade effects, alter energy investment flows, data center location decisions and the competitive structure of entire sectors. Third, the available data on installed compute by region are partial and heterogeneous: no unified public database of AI FLOPs by country exists, making rigorous econometric calibration premature.",
        "The chosen method therefore combines the quantitative rigour of the empirical diagnostic (sourced data, time series, measurable ratios) with the qualitative flexibility of scenario construction, in the spirit of what Schoemaker (1995) calls a disciplined heuristic.[3] The scenarios are not probabilistic forecasts but coherent strategic narratives, each based on explicit assumptions and developing their consequences through measurable metrics.",
    ]),
    ("2.2 Data sources: classification and critical evaluation", [
        "The study draws on three categories of sources, whose reliability and potential biases must be explicitly acknowledged. This methodological transparency conforms to the recommendations of the OECD/JRC Handbook on Constructing Composite Indicators (Nardo et al., 2008), which prescribes systematic documentation of sources, their limitations and their biases in any composite indicator construction.[4]",
    ]),
    ("2.2.1 Primary sources (official and regulatory documents)", [
        "This category includes normative or institutional texts: presidential proclamations (Section 232), BIS rules (AI Diffusion Rule, Entity List), IEA reports, European Parliament publications (EPRS), official statistical data (SIA/WSTS for semiconductors, Eurostat and EIA for energy, RTE for France). These sources offer the highest factual reliability but may contain institutional framing biases: the IEA tends to favour moderate scenarios, the European Parliament to emphasize EU sovereignty risks.",
    ]),
    ("2.2.2 Academic sources and think tanks", [
        "This category includes peer-reviewed articles (Farrell and Newman, 2019; Bresnahan and Trajtenberg, 1995; Brynjolfsson et al., 2019; Mugge, 2024) and publications from recognized think tanks (Bruegel, Carnegie Endowment, CSIS, OECD, Federal Reserve Board). The former provide robust theoretical grounding; the latter offer empirically founded policy analyses but are potentially influenced by each institution ideological orientation. We prioritize cross-referencing sources of different orientations (Bruegel / Carnegie / Fed) to limit this bias.",
    ]),
    ("2.2.3 Industry and consulting sources", [
        "McKinsey, Deloitte, Accenture, Epoch AI and CFG Europe provide market data, sectoral projections and capacity estimates unavailable in public sources. These sources present a potential systematic bias: consulting firms have an interest in amplifying trends (to justify transformation engagements) and market estimates are often optimistic. We mitigate this bias by triangulating figures with institutional data and explicitly flagging discrepancies between sources. For example, 2024 semiconductor sales are 627.6 billion USD according to the SIA (traditional scope) but 775 billion USD according to McKinsey (expanded scope), a 24 percent gap reflecting methodological differences, not inconsistencies.[5]",
    ]),
    ("2.2.4 AI compute data: the Epoch AI GPU Clusters dataset", [
        "Measuring installed AI compute by country constitutes the central methodological challenge of this study. We rely primarily on the Epoch AI GPU Clusters dataset (Pilz, Rahman, Sanders and Heim, 2025), which catalogues over 500 supercomputers and GPU clusters worldwide for the period 2019-2025.[6] This dataset, available under an open Creative Commons Attribution licence, constitutes the most comprehensive and systematically documented source on global AI compute infrastructure to date. It is used as a reference by the Stanford AI Index Report (2025), by several government reports, and by institutions such as OpenAI and DeepMind.",
        "The dataset covers for each cluster: country of location, chip type (H100, A100, GB200, TPU, etc.), computational performance in 16-bit FLOP/s, number of H100 equivalents, operational date, sector (private/public), electrical power (MW) and estimated hardware cost. This granularity enables aggregation by country and year, directly meeting the needs of our variable F(r) in the CACI.",
        "Limitations of the Epoch AI dataset. Three limitations must be highlighted. First, coverage is estimated at 10-20 percent of global aggregate AI compute performance (March 2025), with significant heterogeneity across companies and chip types: approximately 20-37 percent of NVIDIA H100s, 12 percent of A100s, but less than 4 percent of Google TPUs and a negligible fraction of custom chips from AWS, Microsoft or Meta.[7] Second, Chinese systems are anonymized (names removed, values rounded to one significant figure), limiting analytical precision for China. Third, the physical location of a cluster does not determine access: many clusters are accessible via cloud services from other countries.",
        "We complement these data with the OECD Working Paper by Lehdonvirta, Wu, Hawkins et al. (October 2025), which develops a methodology for estimating public cloud AI compute availability by country, counting cloud regions from major providers equipped with AI accelerators (A100, H100, GB200) across 39 economies.[8] This complementary approach distinguishes installed compute (Epoch AI) from accessible compute (OECD), a distinction crucial for the CACI.",
    ]),
    ("2.3 Scenario construction", [
        "Scenario construction follows a four-step protocol, inspired by the 2x2 matrix methodology (Schwartz, 1991; van der Heijden, 2004) and adapted to the geostrategic context of AI.[9]",
        "Step 1 - Identification of driving forces. Predetermined elements (whose evolution is reasonably predictable regardless of scenario) include: continued growth in global AI compute demand; structural increase in data center energy consumption; European dependence on Asian and American foundries for leading-edge chips; concentration of the global cloud around three US hyperscalers; and exponential increase in frontier model training costs.",
        "Critical uncertainties (whose evolution depends on political choices, strategic reactions or technological disruptions) are grouped along two dimensions. Dimension 1 - Intensity of US technology protectionism: this dimension covers a spectrum from maintaining current restrictions to aggressive hardening (GPU quotas for the EU, restrictions on APIs and models, explicit prioritization of deliveries to US companies). Dimension 2 - European response capacity: this dimension covers a spectrum from passive posture (marginal adaptation, acceptance of dependence) to active response (Compute Zones with derogated energy, accelerated AI Factories, SMR nuclear for data centers, alternative partnerships Japan-Korea-Taiwan, AI Act revision).",
        "Step 2 - 2x2 matrix and scenario generation. Crossing the two uncertainty dimensions generates a four-scenario matrix. The choice of four rather than three scenarios is deliberate. Schwartz (1991) and Shell method practitioners recommend never building three scenarios, as the human mind tends to treat the middle scenario as the most likely, thereby reducing the exercise utility.[10] The 2x2 matrix forces the analyst to explore extreme quadrants, precisely where strategic ruptures play out.",
        "Step 3 - Narrative development and quantification. Each scenario is developed following a standardized protocol comprising three components: a strategic narrative describing the plausible sequence of events between 2026 and 2030; a quantification of key metrics calibrated on 2020-2026 empirical data and projected according to the scenario assumptions; and early warning indicators (leading indicators) enabling identification, from 2026-2027 onward, of which scenario reality is converging toward.",
        "Step 4 - Sensitivity analysis and robustness. For each recommendation formulated in Chapter VII, we evaluate its robustness across all four scenarios. A recommendation is considered robust if it produces positive or neutral results in at least three of the four scenarios.",
    ]),
    ("2.4 Key metrics and original indicator: the CACI", [
        "We define six metrics to be calculated or estimated in the empirical diagnostic (Chapter III), then projected in each scenario (Chapter V). Together, they form a dashboard of US/EU divergence in AI: M1 (compute gap, US/EU installed AI FLOPs ratio normalized by GDP), M2 (relative FLOP cost for training), M3 (cloud dependence, share of EU AI workloads on US infrastructure), M4 (sectoral AI productivity), M5 (energy constraint, data center demand/capacity ratio) and M6 (AI relocations).",
    ]),
    ("2.4.1 Theoretical foundations of the CACI", [
        "Grounding in the literature. The construction of a compute-centred AI competitiveness composite indicator responds to a need identified by several converging research streams. Since 2023-2024, academic and institutional literature increasingly emphasizes that computational capacity has become the most discriminating factor of production for frontier AI (Sevilla et al., 2022; Epoch AI, 2025; Pilz et al., 2025). US export controls (BIS, October 2022; updated 2023 and 2025) explicitly place advanced compute at the heart of geopolitical competition, while Hawkins, Lehdonvirta and Wu (2025) introduce the concept of compute sovereignty as a structuring dimension of strategic autonomy.[11]",
        "Yet existing AI competitiveness indices do not place compute at the centre of their construction. The IMF AI Preparedness Index (Cazzaniga et al., 2024), covering 174 countries, aggregates four dimensions (digital infrastructure, human capital, innovation/economic integration, regulation/ethics) without directly measuring installed compute capacity.[12] The Tortoise Media Global AI Index (2024), ranking 83 countries on 122 indicators across three pillars (Implementation, Innovation, Investment), includes an infrastructure/supercomputing component but subsumes it within a weighted additive index where compute is merely one factor among many.[13] The Stanford AI Index (2025) provides rich compute trend data but proposes no composite index. None of these instruments formalizes the mechanism by which compute, adjusted for its energy cost and related to an economy absorptive capacity, determines AI competitiveness.",
        "The CACI aims to fill this gap by proposing a parsimonious yet theoretically grounded indicator that captures the multiplicative interaction between four factors: accessible compute, the relevant labor force, geopolitical access conditions and the energy cost.",
    ]),
    ("2.4.2 Formal definition", [
        "The CACI for a region r at period t is defined in two complementary modes. In Power Mode (absolute):  CACI(r,t) = F(r,t)^0.40 x L(r,t)^0.20 x R(r,t)^0.15 / E(r,t)^0.25.  In Intensity Mode (per unit of national wealth):  CACI(r,t) = F(r,t)^0.40 x L(r,t)^0.20 x R(r,t)^0.15 / (E(r,t)^0.25 x GDP(r,t)).",
        "F(r,t) - installed and accessible AI compute capacity in region r at period t. We aggregate the Epoch AI GPU Clusters dataset by country, in H100-equivalents (the metric used by Epoch AI), and complement it with OECD estimates of accessible cloud compute. F admits two specifications used in parallel: F_total (all clusters located in the territory regardless of ownership, defining the Physical CACI) and F_dom (clusters owned and operated by domestic actors, defining the Sovereign CACI).",
        "L(r,t) - working population with AI competencies (proxy: STEM graduates plus certified AI training plus LinkedIn Economic Graph AI skill density). This factor captures absorptive capacity in the sense of Cohen and Levinthal (1990): abundant compute without human capital to exploit it does not produce competitiveness.[14]",
        "R(r,t) - geopolitical access index (between 0 and 1) capturing the country position in the export-control regime. R = 1 for the United States (unlimited access to frontier compute), R = 0.95 for Tier-1 close allies (UK), R = 0.9 for the EU and France (Tier-1 with selective restrictions), R = 0.85 for India (Tier-2), R = 0.7 for China (heavily restricted).",
        "E(r,t) - average effective energy cost of compute in region r, in USD per MWh, adjusted for hyperscaler Power Purchase Agreements. The dashboard reference values for the April 2026 snapshot are USA 85, China 92, France 115, Germany 140, UK 190 USD per MWh. These figures correspond to PPA-adjusted prices effectively faced by data centers, not to raw Eurostat industrial tariffs (which are typically higher in Europe). EU effective costs in this PPA-adjusted scope are 1.4 to 1.7 times US levels.[15]",
        "GDP(r,t) - gross domestic product of region r (World Bank, Eurostat). Used only in Intensity Mode to normalise across economies of very different sizes.",
        "Justification of the multiplicative form. The choice of multiplicative (geometric) rather than additive (arithmetic) aggregation is deliberate and rests on three arguments. First, General Purpose Technology theory (Bresnahan and Trajtenberg, 1995) posits strong complementarity between innovation inputs: compute without affordable energy or qualified human capital produces no competitiveness gains, justifying a form where weakness in one factor penalizes the whole. Second, the OECD/JRC Handbook (2008, p. 33) recommends geometric aggregation when components are not perfectly substitutable: unlike the arithmetic mean, the geometric mean does not allow a very high score on one dimension to fully compensate a very low score on another. Third, recent work on AI index construction (Koronakos, Kritikos and Sotiros, 2024; analysis of the Tortoise GAII via Choquet integral) confirms that AI competitiveness dimensions exhibit interactions (complementarities and redundancies) that make simple linear aggregation problematic.[16]",
        f"Economic interpretation. In logarithmic form, the CACI in Intensity Mode can be written ln(CACI) = 0.40 ln(F) + 0.20 ln(L) + 0.15 ln(R) - 0.25 ln(E) - ln(GDP). This transformation linearizes the relationship and allows each coefficient to be interpreted as an elasticity. The indicator is designed for bilateral comparison: the ratio CACI(US)/CACI(EU) measures relative competitive advantage. The April 2026 snapshot yields a ratio of {fmt_en(us_eu_caci, 2)}:1 in Power Mode (F_total) - this is the headline figure of this study.",
    ]),
    ("2.4.3 Calibration protocol and data sources", [
        "CACI calibration follows a four-step protocol consistent with OECD/JRC Handbook (2008) recommendations: identification and collection of raw data; treatment of missing values and normalization; aggregation; and sensitivity analysis.",
        f"F(r,t) - Installed AI compute capacity. Estimation follows a bottom-up approach. The main layer comes from the Epoch AI GPU Clusters dataset (April 2026 version), providing 16-bit FLOP/s and H100-equivalent performance aggregated by country for 2019-2026. National shares of operational compute on the April 2026 snapshot are: USA {fmt_en(us_share, 1)} percent, China 12.8 percent, EU (13 main economies) 4.4 percent, Norway 1.1 percent, Japan 1.0 percent.[17] When all reported clusters are included (operational plus planned), the USA share moves to about 50 percent, with the Gulf (UAE, Saudi Arabia) and Korea capturing a large share of the projected 2027-2030 buildout. The complementary layer comes from the OECD (Lehdonvirta et al., 2025), cataloguing cloud regions with AI accelerators across 39 economies, capturing the accessibility dimension that physically installed compute does not fully reflect.",
        "F aggregation procedure. For each country-year, we extract from the Epoch AI dataset the sum of H100-equivalents from all clusters located in the country, filtering for confirmed systems (certainty greater than or equal to Likely); apply an extrapolation factor to correct for the dataset under-coverage (10 to 20 percent of global compute), using Epoch AI sectoral coverage estimates by chip type; and add estimated accessible cloud capacities via the OECD methodology for countries where foreign cloud compute represents a significant share of effective capacity. Raw data and calculation code (Python/pandas) are documented in the methodological appendix and on the public dashboard.[18]",
        "E(r,t) - Energy cost. The dashboard CSV reference values used throughout this study, in USD per MWh and PPA-adjusted, are: USA 85, China 92, France 115, Germany 140, UK 190, India 88. These figures incorporate negotiated large-consumer tariffs (hyperscaler PPAs), using IEA (2025, Energy and AI) estimates of data center energy mix by region. They are systematically lower than raw Eurostat industrial tariffs in Europe, because hyperscaler PPAs and direct nuclear/renewable sourcing materially reduce the effective cost faced by AI data centers. The Federal Reserve Board (October 2025) documents a significant negative correlation between energy costs and AI adoption at the European firm level.[19]",
        "GDP(r,t) and L(r,t). GDP is available from the World Bank (World Development Indicators) and Eurostat. The dashboard April 2026 snapshot uses GDP values in trillion USD: USA 29.3, China 18.7, EU 18.9, France 3.16, Germany 4.68, UK 3.6. The AI human capital proxy L(r,t) combines three sub-indicators: number of STEM graduates (OECD Education at a Glance), AI skills density as measured by LinkedIn Economic Graph (profiles with AI skills relative to working population) and AI certifications (estimates based on certified AWS/Google/Microsoft cloud programs by country). The dashboard snapshot uses, in millions of AI-skilled workers: USA 3.5, China 4.8, EU 3.1, France 1.5, Germany 1.9, UK 1.8. We acknowledge a bias favouring English-speaking countries and economies where LinkedIn is dominant, and document this bias effect in the sensitivity analysis (section 2.4.5).[20]",
        "CACI normalization and benchmark selection. The raw CACI produces heterogeneous absolute values whose direct interpretation is challenging. In accordance with standard practice (Nardo et al., 2008; Saisana and Tarantola, 2002), we apply a min-max normalization to the leader: each country normalised CACI is expressed as a percentage of the United States raw CACI, so that CACI_norm(USA) = 100 by construction.",
        f"This normalization choice is methodologically grounded on three arguments. First, the United States constitutes a natural benchmark given its dominant position across all four CACI factors: {fmt_en(us_share, 1)} percent of global operational AI compute (Epoch AI, April 2026), the lowest effective industrial energy costs among major OECD economies (USD 85 per MWh, PPA-adjusted, versus USD 115 in France and USD 140 in Germany - dashboard CSV, April 2026), the highest nominal GDP, and an unmatched AI human capital ecosystem. Second, normalization to the leader is the established convention for relative competitiveness indices: Balassa RCA (1965), the WEF GCI (Schwab, 2019) and UNIDO Competitive Industrial Performance Index all proceed by ratio to the sectoral or national leader. Third, this normalization eliminates interpretive ambiguities of absolute values and enables temporal comparisons: a France CACI moving from 25 to 35 between 2024 and 2028 signifies partial catch-up, even if both countries raw values have increased.",
        f"The most important methodological consequence of this normalization is the revelation of a structural gap that traditional indices conceal. The IMF AI Preparedness Index (Cazzaniga et al., 2024) assigns scores of 0.85 (USA) and 0.74 (France), a ratio of 1.15:1, suggesting near-parity. The Tortoise Global AI Index produces similar gaps. Yet the normalized CACI in Power Mode places France at about 25 and Germany at about 5 on a scale where the United States = 100, i.e. ratios of 4:1 and 19:1 respectively. The headline US/EU ratio (EU aggregated as 13 main economies) is {fmt_en(us_eu_caci, 2)}:1. This divergence is not contradictory: it results from the fact that multi-dimensional indices linearly aggregate normative dimensions (regulation, ethics, publications) that arithmetically compensate for deficits on material factors (compute, energy). The CACI, by isolating these material factors, exposes the physical chasm that composite metrics average away.",
        "This property of the CACI - revealing the true magnitude of the compute gap - constitutes the central analytical contribution of the index and the fundamental empirical thesis of this study. The CACI gap is not a normalization artifact: it is the scientific result.",
    ]),
    ("2.4.4 Positioning relative to existing AI competitiveness indices", [
        "The IMF AI Preparedness Index (AIPI) covers 174 countries (2023) and aggregates four pillars: digital infrastructure, human capital and labour market policies, innovation and economic integration, regulation and ethics. The AIPI does not measure installed AI compute and does not weight by energy cost. Its expected correlation with the CACI is positive but imperfect: countries scoring highly on the AIPI (Singapore, Denmark, Netherlands) are not necessarily those with the most effective compute per unit of GDP.[21]",
        "The Tortoise Media Global AI Index (GAII) ranks 83 countries on 122 indicators across three pillars (Implementation, Innovation, Investment). It includes an infrastructure/supercomputing component but aggregates it linearly with subjective weights (acknowledged by Tortoise as a limitation). The GAII is broader and more multidimensional than the CACI, but precisely because it is broad, it dilutes the compute signal in a many-indicator composite. As Koronakos et al. (2024) show, the GAII weighting subjectivity can invert country rankings depending on chosen weighting scenarios.",
        "The Stanford AI Index (2025) constitutes the most comprehensive reference in terms of raw data (notable models by country, investment, publications, patents, compute trends). It does not propose a composite index but provides the time series used by many other indices. The Stanford AI Index public data are available via an open Google Drive folder.[22]",
        "CACI specific value added. The CACI distinguishes itself through four properties: it places compute at the centre rather than the periphery of the indicator, reflecting compute now-dominant role in frontier AI; it explicitly integrates energy cost as a bottleneck, consistent with IEA (2025) findings; it uses a theoretically grounded multiplicative aggregation rather than additive; and it is parsimonious (four variables), making it transparent and reproducible, at the cost of lesser comprehensiveness.",
    ]),
    ("2.4.5 CACI limitations and sensitivity analysis", [
        "First limitation - opacity of F(r,t). Installed compute measurement depends on incomplete private data. The Epoch AI dataset covers only 10-20 percent of global compute, with uneven coverage across sectors and companies. National attribution is itself debatable: a growing share of compute is held by a few private hyperscalers operating globally. Mitigation: we systematically document uncertainty margins, present ranges rather than point values, and verify result stability when F varies by plus or minus 30 percent. The Physical/Sovereign decomposition partially addresses the attribution problem.",
        "Second limitation - qualitative heterogeneity of compute. The CACI aggregates FLOPs without distinguishing GPU generations (an H200 does not equal an A100 in energy efficiency and real performance). Mitigation: the Epoch AI dataset provides performance in H100 equivalents, offering partial normalization. We propose a GPU generation weighting factor in the appendix and show that adjustment marginally modifies rankings.",
        "Third limitation - the human capital proxy L(r,t). The STEM plus LinkedIn plus certifications combination presents a bias favouring English-speaking countries. This bias likely underestimates China and some Asian economies. Mitigation: we replicate the analysis using the IMF AIPI Human Capital sub-index as an alternative proxy and show ranking sensitivity to this choice.",
        "Fourth limitation - indicator staticity. The CACI measures a state at a given moment, not a dynamic. This is why we calculate it across multiple years (2022, 2024, 2026) and project it in each scenario, enabling trajectory tracking and gap evolution assessment.",
        "Fifth limitation - endogeneity. Countries gaining AI productivity invest massively in compute, creating a reverse causality risk: the CACI might capture the consequence rather than the cause of competitiveness. This study, within its dissertation framework, does not formally instrument this relationship (no instrumental variables or GMM strategy). However, we note that the exogenous shock of the October 2022 BIS rules offers a natural quasi-experiment that could ground a causal identification strategy in Difference-in-Differences. We identify the formal treatment of this endogeneity (DiD, IV or Arellano-Bond GMM) as a priority research avenue for any publishable extension of this work.[23]",
        "Sixth limitation - normalization bias for small economies. Comparative analysis of the CACI against the IMF AIPI and Tortoise GAII indices reveals an instructive anomaly: certain countries with low GDP and a reduced AI workforce obtain disproportionately high CACI scores. South Africa is emblematic: ranked last or second-to-last on the IMF AIPI and Tortoise indices, it can appear as a global leader on the Intensity Mode CACI when the GDP normalisation amplifies a small-denominator effect. This phenomenon is a classic case of normalization bias identified in the composite indicator literature. The OECD/JRC Handbook (Nardo et al., 2008, pp. 27-29) warns that normalization by GDP or population can produce misleading results for small economies. Our mitigation is threefold: a critical mass threshold on F (excluding economies below 10 000 H100-equivalents); a scaling factor alpha(r) = min(1, F(r) / F_median) for sensitivity tests; and a dual ranking (Power Mode for absolute power, Intensity Mode for density), with the explicit caveat that bilateral comparisons in this study focus on USA, EU, France and China.[24]",
    ]),
    ("2.4.6 Reproducible numbers - April 2026 snapshot", [
        "To allow the reader to replay the calculation from the public dashboard CSVs, we expose here the inputs and outputs used throughout this study. The four input variables are sourced from the dashboard data folder (energy_prices.csv, gdp_data.csv, workforce_data.csv, gpu_clusters.csv) at https://mo0ogly.github.io/America-First-IA/dashboard/data/.",
        "Inputs (April 2026 snapshot). USA: F_total = 39.65 million H100-eq, L = 3.5 million, R = 1.00, E = 85 USD/MWh, GDP = 29.3 trillion USD. EU (13 main economies aggregated): F_total = 2.62 million H100-eq, L = 3.1 million, R = 0.90, E = 135 USD/MWh, GDP = 18.9 trillion USD. France: F_total = 2.44 million H100-eq, L = 1.5 million, R = 0.90, E = 115 USD/MWh, GDP = 3.16 trillion USD. Germany: F_total = 0.05 million H100-eq, L = 1.9 million, R = 0.90, E = 140 USD/MWh, GDP = 4.68 trillion USD. China: F_total = 0.40 million H100-eq, L = 4.8 million, R = 0.70, E = 92 USD/MWh, GDP = 18.7 trillion USD.",
        f"Outputs (Power Mode F_total, normalised at USA = 100). USA = 100.0; EU(13) = {fmt_en(caci_scores['EU'], 1)}; France = {fmt_en(caci_scores['France'], 1)}; India = {fmt_en(caci_scores['India'], 1)}; China = {fmt_en(caci_scores['China'], 1)}; UK = {fmt_en(caci_scores['UK'], 1)}; Germany = {fmt_en(caci_scores['Germany'], 1)}. UAE = 55.7 in Physical CACI but only 6.0 in Sovereign CACI (F_dom restricted to clusters owned by domestic entities G42, MGX, Mubadala, Khazna, TII), illustrating the Physical/Sovereign decomposition: 99.6 percent of UAE F_total is owned by US-side actors (Stargate UAE, Microsoft, OpenAI). Headline ratios: US/EU(13) = {fmt_en(us_eu_caci, 2)}:1; US/France = 3.96:1; US/Germany = 18.59:1; US/China = 7.46:1.",
        f"These outputs are reproduced live on the public dashboard. Any subsequent dashboard update (new clusters, energy revisions, GDP refresh) will mechanically refresh the figures; the methodology does not change. The Power Mode F_total ratio of {fmt_en(us_eu_caci, 2)}:1 reported on the cover and throughout this study corresponds exactly to this snapshot.",
    ]),
    ("2.5 Scope and delimitations", [
        "Geographic scope. The analysis focuses on the bilateral United States / European Union relationship, with a specific focus on France. China is treated as a contextual variable (primary target of US export controls, factor of pressure on chip production capacities) but is not subject to in-depth analysis. Japan, South Korea and Taiwan appear as semiconductor supply chain actors.",
        "Temporal scope. The diagnostic covers 2020-2026, scenarios cover 2026-2030. The 2030 horizon is chosen as it corresponds to the convergence of several deadlines: IEA projections for data center energy, expected maturity of the EU Chips Act, France 2030 SNIA objectives and potential arrival of the first operational SMR nuclear reactors.",
        "Technological scope. The study covers frontier AI (foundation models, compute-intensive) and its material prerequisites (GPU/ASIC, data centers, energy). It integrates AI robotics as an amplifying factor for energy demand. It does not address edge embedded AI (smartphones, IoT), except insofar as it constitutes a specific objective of the French SNIA.",
    ]),
    ("2.6 General methodological limitations", [
        "Radical political uncertainty. Technology protectionism depends on discretionary political decisions with structurally low predictability. A change of US administration in 2028, an unexpected US-EU trade agreement, or an escalation of the US-China conflict could invalidate certain assumptions. This is precisely why we propose four scenarios rather than a single trajectory.",
        "Technological disruptions. The DeepSeek episode (January 2025), where a Chinese model achieved near-frontier performance with substantially reduced training budget, illustrates the possibility of efficiency breakthroughs that would alter the problem terms. The IEA (2025, Energy and AI) devotes a case study to DeepSeek and concludes that even with significant efficiency improvements, demand growth absorbs gains (Jevons rebound effect).[25]",
        "Compute data opacity. The exact number of GPUs deployed per hyperscaler, the precise geographic distribution of data centers and GPU volumes exported by region are partially or fully confidential data. Our installed compute estimates carry significant margins of error, which we systematically document.",
        "Consulting source bias. As noted in section 2.2, industry sources carry a systematic optimism bias. We mitigate this through triangulation but cannot eliminate it entirely.",
        "These limitations do not compromise the analysis validity. The scenario method is precisely designed to function in high-uncertainty environments, where the objective is not prediction but structured exploration of possibilities. As Schwartz notes, scenarios are not forecasts; they are plausible stories that help you think.[26] Our contribution lies in the rigour of the framing, the explicitness of assumptions, the transparency of data sources and the originality of the CACI indicator, rather than in the precision of numerical projections.",
    ]),
]

EN.notes = [
    "Schwartz, P. (1991), The Art of the Long View: Planning for the Future in an Uncertain World, New York, Doubleday. See also Wack, P. (1985), 'Scenarios: Uncharted Waters Ahead,' Harvard Business Review, Sept.-Oct. 1985, pp. 72-89.",
    "Bradfield, R., Wright, G., Burt, G., Cairns, G. and Van Der Heijden, K. (2005), 'The Origins and Evolution of Scenario Techniques in Long Range Business Planning,' Futures, 37(8), pp. 795-812.",
    "Schoemaker, P.J.H. (1995), 'Scenario Planning: A Tool for Strategic Thinking,' MIT Sloan Management Review, 36(2), pp. 25-40.",
    "Nardo, M., Saisana, M., Saltelli, A. and Tarantola, S. (2008), Handbook on Constructing Composite Indicators: Methodology and User Guide, OECD Publishing, Paris.",
    "The SIA/McKinsey gap is explained by scope: McKinsey (January 2026, 'Hiding in Plain Sight') includes the value of captive designers (Apple, Amazon, Tesla) and fabless operators whose sales do not appear in WSTS statistics.",
    "Pilz, K.F., Rahman, R., Sanders, J. and Heim, L. (2025), 'Trends in AI Supercomputers,' arXiv:2504.16026, April 2025. Dataset accessible at https://epoch.ai/data/gpu-clusters under Creative Commons Attribution licence.",
    "Epoch AI (2025), GPU Clusters Documentation. Coverage estimated at 20-37 percent of NVIDIA H100s, 12 percent of A100s and 18 percent of AMD MI300X, but less than 4 percent of Google TPUs.",
    "Lehdonvirta, V., Wu, B., Hawkins, Z.J., Caira, C. and Russo, L. (2025), 'Measuring Domestic Public Cloud Compute Availability for Artificial Intelligence,' OECD AI Papers No. 49, https://doi.org/10.1787/8602a322-en.",
    "Van der Heijden, K. (2004), Scenarios: The Art of Strategic Conversation, 2nd ed., Chichester, Wiley.",
    "Schwartz (1991), op. cit., pp. 241-243.",
    "Hawkins, Z.J., Lehdonvirta, V. and Wu, B. (2025), 'AI Compute Sovereignty,' SSRN, June 2025, https://ssrn.com/abstract=5312977. Sevilla, J. et al. (2022), 'Compute Trends Across Three Eras of Machine Learning,' arXiv:2202.05924.",
    "Cazzaniga, M. et al. (2024), 'Gen-AI: Artificial Intelligence and the Future of Work,' IMF Staff Discussion Note 2024/001. The AI Preparedness Index is accessible via the IMF dashboard.",
    "Tortoise Media (2024), The Global Artificial Intelligence Index 2024. See also Koronakos, G., Kritikos, M. and Sotiros, D. (2024), 'Mitigating Subjectivity and Bias in AI Development Indices,' Expert Systems with Applications.",
    "Cohen, W.M. and Levinthal, D.A. (1990), 'Absorptive Capacity: A New Perspective on Learning and Innovation,' Administrative Science Quarterly, 35(1), pp. 128-152.",
    "Effective costs are derived from the dashboard energy_prices.csv and reflect PPA-adjusted prices. Raw Eurostat industrial tariffs in Europe (consumption band IE) can be 1.5 to 2 times higher; the gap between raw and effective prices is documented in Chapter III. The Federal Reserve Board (October 2025) documents a significant negative correlation between energy costs and AI adoption at the European firm level.",
    "Geometric aggregation is also used by the UNDP Human Development Index since 2010, for analogous reasons of non-substitutability between dimensions. See OECD/JRC Handbook (2008), pp. 31-33.",
    "Operational shares computed by aggregating Epoch AI clusters with status Existing/Operational/Online/In progress on the April 2026 snapshot.",
    "The Python code and data are reproducible. The Epoch AI dataset is downloadable in CSV at https://epoch.ai/data/gpu_clusters.csv (daily refresh). The dashboard CSVs are mirrored at https://mo0ogly.github.io/America-First-IA/dashboard/data/.",
    "The LinkedIn bias is documented: in countries where LinkedIn is rarely used (China, Russia, some Southeast Asian economies), AI skills density is mechanically underestimated.",
    "IMF (2024), AI Preparedness Index Dashboard. Singapore, Denmark, Netherlands and the US occupy top positions; China ranks 31st (score 0.63).",
    "Maslej, N. et al. (2025), 'Artificial Intelligence Index Report 2025,' Stanford Institute for Human-Centered AI, April 2025.",
    "Tortoise Media (2024), op. cit. Stanford AI Index public data: open Google Drive folder.",
    "The most promising causal identification strategy would exploit the exogenous shock of the October 2022 BIS rules in a Difference-in-Differences framework. See Arellano, M. and Bond, S. (1991), 'Some Tests of Specification for Panel Data,' Review of Economic Studies, 58(2), pp. 277-297.",
    "OECD/JRC Handbook (2008), pp. 27-29. The dual ranking solution (intensity / scale) is also used by Tortoise Media (2024) in the Global AI Index. Saisana, M. and Tarantola, S. (2002), State-of-the-art Report on Current Methodologies and Practices for Composite Indicator Development, JRC European Commission, pp. 14-16.",
    "IEA (2025), Energy and AI, devotes a case study to DeepSeek and concludes that even with significant efficiency improvements, demand growth absorbs gains (Jevons rebound effect).",
    "Schwartz (1991), op. cit., p. 38. Author translation.",
]


# ===========================================================================
# Content - French (text only; structure shared with EN)
# ===========================================================================

FR = LangPack(
    code="FR",
    filename="Chapitre_II_Methodologie_FR.docx",
    cover_title="AI FOR AMERICANS FIRST",
    cover_subtitle="Protectionnisme IA, Energie et Semi-conducteurs : Trajectoires de divergence US/Europe 2024-2030",
    cover_blurb="Analyse geostrategique et economique integree - Chapitre II",
    cover_chip_lines=[
        f"{fmt_fr(us_share, 1)} pct du compute IA operationnel mondial = USA",
        "1,59x cout energie EU/US",
        f"{fmt_fr(us_eu_caci, 2)}:1 ratio CACI US/EU (Power Mode)",
    ],
    cover_meta="Paris - fevrier 2026  |  7 chapitres  |  4 scenarios prospectifs  |  3 zones geographiques",
    cover_keywords_label="Mots-cles",
    cover_keywords=("intelligence artificielle, protectionnisme technologique, semi-conducteurs, "
                    "controles a l'exportation, compute souverain, geopolitique IA, France, "
                    "Etats-Unis, Chine"),
    chapter_label="CHAPITRE II",
    chapter_title="Methodologie",
    notes_label="Notes",
    license_block=[
        "Licence et avertissement. Cette oeuvre, 'AI for Americans First', est mise a disposition selon les termes de la licence Creative Commons Attribution - Pas d'Utilisation Commerciale - Partage dans les Memes Conditions 4.0 International (CC BY-NC-SA 4.0) du projet America-First-IA.",
        "Vous etes libre de partager et adapter le materiel a des fins non commerciales, a condition d'attribuer correctement l'oeuvre a Fabrice Pizzi (Universite Paris-Sorbonne) et de distribuer vos contributions sous la meme licence. Ce document est fourni a des fins educatives et de recherche uniquement.",
        "Tableau de bord public : https://mo0ogly.github.io/America-First-IA/dashboard/",
        "Depot : https://github.com/mo0ogly/America-First-IA",
    ],
    page_footer="AI for Americans First - Fabrice Pizzi - Chapitre II",
)

FR.sections = [
    ("2.1 Approche generale : analyse prospective multi-scenarios mixte", [
        "Cette etude combine une analyse empirique retrospective (diagnostic 2020-2026) avec une projection prospective par scenarios (2026-2030). Cette architecture a deux volets repond a la nature du phenomene etudie : le protectionnisme technologique IA est simultanement un fait observable (controles a l'exportation, tarifs Section 232) et un processus en cours dont la trajectoire future depend de variables politiques discretionnaires, des reponses strategiques europeennes et de developpements technologiques partiellement imprevisibles.",
        "Le volet retrospectif emploie une methode quantitative descriptive, fondee sur l'agregation et le croisement de donnees issues de sources institutionnelles (AIE, SIA/WSTS, Eurostat, EIA), de rapports industriels (McKinsey, Deloitte, Epoch AI) et de documents reglementaires (Federal Register, BIS, Maison-Blanche). L'objectif est d'etablir un socle factuel rigoureux et source couvrant trois dimensions : la consommation energetique des centres de donnees, le marche des semi-conducteurs et la capacite de calcul IA installee par region.",
        "Le volet prospectif s'appuie sur la tradition du scenario planning telle que formalisee par Schwartz (1991) et pratiquee chez Royal Dutch/Shell depuis les annees 1970.[1] Cette methode, appartenant a l'ecole Intuitive Logics (Bradfield et al., 2005), consiste a construire des scenarios plausibles et internement coherents - non pour predire l'avenir, mais pour explorer l'espace des possibles et evaluer la robustesse de differentes strategies face a des evolutions environnementales divergentes.[2] Elle est particulierement adaptee aux situations caracterisees par une forte incertitude politique et technologique, ou les modeles econometriques classiques atteignent leurs limites - ce qui est precisement le cas du protectionnisme technologique IA.",
        "Justification du choix methodologique. Trois raisons fondent le choix de la methode des scenarios sur la modelisation econometrique pure. Premierement, les variables analytiques cles sont largement politiques et discretionnaires : la decision d'un president americain d'imposer ou non des quotas GPU a l'Europe ne peut etre modelisee par une fonction de regression. Deuxiemement, les interactions entre les dimensions energetique, technologique et geopolitique sont non lineaires et systemiques : une restriction sur les GPU peut, par effets de cascade, alterer les flux d'investissement energetique, les decisions de localisation des centres de donnees et la structure concurrentielle de secteurs entiers. Troisiemement, les donnees disponibles sur le compute installe par region sont partielles et heterogenes : aucune base publique unifiee de FLOPs IA par pays n'existe, ce qui rend prematuree une calibration econometrique rigoureuse.",
        "La methode retenue combine donc la rigueur quantitative du diagnostic empirique (donnees sourcees, series temporelles, ratios mesurables) avec la flexibilite qualitative de la construction de scenarios, dans l'esprit de ce que Schoemaker (1995) appelle une heuristique disciplinee.[3] Les scenarios ne sont pas des previsions probabilistes mais des recits strategiques coherents, chacun fonde sur des hypotheses explicites et developpant ses consequences a travers des metriques mesurables.",
    ]),
    ("2.2 Sources de donnees : classification et evaluation critique", [
        "L'etude s'appuie sur trois categories de sources, dont la fiabilite et les biais potentiels doivent etre explicitement reconnus. Cette transparence methodologique est conforme aux recommandations du OECD/JRC Handbook on Constructing Composite Indicators (Nardo et al., 2008), qui prescrit la documentation systematique des sources, de leurs limites et de leurs biais dans toute construction d'indicateur composite.[4]",
    ]),
    ("2.2.1 Sources primaires (documents officiels et reglementaires)", [
        "Cette categorie regroupe les textes normatifs ou institutionnels : proclamations presidentielles (Section 232), regles BIS (AI Diffusion Rule, Entity List), rapports AIE, publications du Parlement europeen (EPRS), donnees statistiques officielles (SIA/WSTS pour les semi-conducteurs, Eurostat et EIA pour l'energie, RTE pour la France). Ces sources offrent la plus haute fiabilite factuelle mais peuvent contenir des biais de cadrage institutionnel : l'AIE tend a privilegier les scenarios moderes, le Parlement europeen a souligner les risques pour la souverainete de l'UE.",
    ]),
    ("2.2.2 Sources academiques et think tanks", [
        "Cette categorie inclut les articles a comite de lecture (Farrell et Newman, 2019 ; Bresnahan et Trajtenberg, 1995 ; Brynjolfsson et al., 2019 ; Mugge, 2024) et les publications de think tanks reconnus (Bruegel, Carnegie Endowment, CSIS, OCDE, Federal Reserve Board). Les premiers fournissent un ancrage theorique robuste ; les seconds proposent des analyses politiques empiriquement fondees mais potentiellement influencees par l'orientation ideologique de chaque institution. Nous privilegions le croisement de sources d'orientations differentes (Bruegel / Carnegie / Fed) pour limiter ce biais.",
    ]),
    ("2.2.3 Sources industrielles et conseil", [
        "McKinsey, Deloitte, Accenture, Epoch AI et CFG Europe fournissent des donnees de marche, des projections sectorielles et des estimations de capacite indisponibles dans les sources publiques. Ces sources presentent un biais systematique potentiel : les cabinets de conseil ont interet a amplifier les tendances (pour justifier des missions de transformation) et les estimations de marche sont souvent optimistes. Nous attenuons ce biais en triangulant les chiffres avec les donnees institutionnelles et en signalant explicitement les divergences entre sources. Par exemple, les ventes de semi-conducteurs 2024 sont de 627,6 milliards USD selon la SIA (perimetre traditionnel) mais de 775 milliards USD selon McKinsey (perimetre elargi), un ecart de 24 pour cent refletant des differences methodologiques et non des incoherences.[5]",
    ]),
    ("2.2.4 Donnees compute IA : le jeu de donnees Epoch AI GPU Clusters", [
        "Mesurer le compute IA installe par pays constitue le defi methodologique central de cette etude. Nous nous appuyons principalement sur le jeu de donnees Epoch AI GPU Clusters (Pilz, Rahman, Sanders et Heim, 2025), qui catalogue plus de 500 supercalculateurs et clusters GPU dans le monde pour la periode 2019-2025.[6] Ce jeu de donnees, disponible sous licence ouverte Creative Commons Attribution, constitue la source la plus complete et la plus systematiquement documentee sur l'infrastructure mondiale de compute IA a ce jour. Il est utilise comme reference par le Stanford AI Index Report (2025), par plusieurs rapports gouvernementaux et par des institutions telles qu'OpenAI et DeepMind.",
        "Le jeu de donnees couvre pour chaque cluster : pays de localisation, type de puce (H100, A100, GB200, TPU, etc.), performance computationnelle en FLOP/s 16 bits, nombre d'equivalents H100, date de mise en service, secteur (prive/public), puissance electrique (MW) et cout materiel estime. Cette granularite permet une agregation par pays et par annee, repondant directement aux besoins de notre variable F(r) dans le CACI.",
        "Limites du jeu de donnees Epoch AI. Trois limites doivent etre soulignees. Premierement, la couverture est estimee a 10-20 pour cent de la performance compute IA agregee mondiale (mars 2025), avec une heterogeneite significative entre entreprises et types de puces : environ 20-37 pour cent des Nvidia H100, 12 pour cent des A100, mais moins de 4 pour cent des TPU Google et une fraction negligeable des puces custom d'AWS, Microsoft ou Meta.[7] Deuxiemement, les systemes chinois sont anonymises (noms retires, valeurs arrondies a un chiffre significatif), limitant la precision analytique pour la Chine. Troisiemement, la localisation physique d'un cluster ne determine pas l'acces : de nombreux clusters sont accessibles via des services cloud depuis d'autres pays.",
        "Nous completons ces donnees avec le Working Paper OCDE de Lehdonvirta, Wu, Hawkins et al. (octobre 2025), qui developpe une methodologie pour estimer la disponibilite du compute IA dans le cloud public par pays, en denombrant les regions cloud des principaux fournisseurs equipees d'accelerateurs IA (A100, H100, GB200) sur 39 economies.[8] Cette approche complementaire distingue compute installe (Epoch AI) et compute accessible (OCDE), distinction cruciale pour le CACI.",
    ]),
    ("2.3 Construction des scenarios", [
        "La construction des scenarios suit un protocole en quatre etapes, inspire de la methodologie de la matrice 2x2 (Schwartz, 1991 ; van der Heijden, 2004) et adapte au contexte geostrategique de l'IA.[9]",
        "Etape 1 - Identification des forces motrices. Les elements predetermines (dont l'evolution est raisonnablement previsible quel que soit le scenario) incluent : la croissance continue de la demande mondiale de compute IA ; l'augmentation structurelle de la consommation energetique des centres de donnees ; la dependance europeenne aux fonderies asiatiques et americaines pour les puces de pointe ; la concentration du cloud mondial autour de trois hyperscalers americains ; et l'augmentation exponentielle des couts d'entrainement des modeles de frontiere.",
        "Les incertitudes critiques (dont l'evolution depend de choix politiques, de reactions strategiques ou de disruptions technologiques) sont regroupees selon deux dimensions. Dimension 1 - Intensite du protectionnisme technologique americain : cette dimension couvre un spectre allant du maintien des restrictions actuelles a un durcissement agressif (quotas GPU pour l'UE, restrictions sur les API et les modeles, priorisation explicite des livraisons aux entreprises americaines). Dimension 2 - Capacite de reponse europeenne : cette dimension couvre un spectre allant d'une posture passive (adaptation marginale, acceptation de la dependance) a une reponse active (Compute Zones avec energie derogatoire, AI Factories accelerees, SMR nucleaires pour data centers, partenariats alternatifs Japon-Coree-Taiwan, revision de l'AI Act).",
        "Etape 2 - Matrice 2x2 et generation des scenarios. Le croisement des deux dimensions d'incertitude genere une matrice a quatre scenarios. Le choix de quatre scenarios plutot que trois est delibere. Schwartz (1991) et les praticiens de la methode Shell recommandent de ne jamais construire trois scenarios, l'esprit humain ayant tendance a traiter le scenario median comme le plus probable, ce qui reduit l'utilite de l'exercice.[10] La matrice 2x2 force l'analyste a explorer les quadrants extremes - precisement ou se jouent les ruptures strategiques.",
        "Etape 3 - Developpement narratif et quantification. Chaque scenario est developpe selon un protocole standardise comprenant trois composantes : un recit strategique decrivant la sequence plausible d'evenements entre 2026 et 2030 ; une quantification de metriques cles calibrees sur les donnees empiriques 2020-2026 et projetees selon les hypotheses du scenario ; et des indicateurs avances (leading indicators) permettant d'identifier, des 2026-2027, vers quel scenario la realite converge.",
        "Etape 4 - Analyse de sensibilite et robustesse. Pour chaque recommandation formulee au chapitre VII, nous evaluons sa robustesse a travers les quatre scenarios. Une recommandation est consideree robuste si elle produit des resultats positifs ou neutres dans au moins trois des quatre scenarios.",
    ]),
    ("2.4 Metriques cles et indicateur original : le CACI", [
        "Nous definissons six metriques a calculer ou estimer dans le diagnostic empirique (chapitre III), puis a projeter dans chaque scenario (chapitre V). Ensemble, elles forment un tableau de bord de la divergence US/UE en IA : M1 (compute gap, ratio US/UE des FLOPs IA installes normalise par le PIB), M2 (cout relatif du FLOP pour l'entrainement), M3 (dependance cloud, part des charges IA UE sur infrastructure US), M4 (productivite IA sectorielle), M5 (contrainte energetique, ratio demande/capacite des centres de donnees) et M6 (delocalisations IA).",
    ]),
    ("2.4.1 Fondements theoriques du CACI", [
        "Ancrage dans la litterature. La construction d'un indicateur composite de competitivite IA centre sur le compute repond a un besoin identifie par plusieurs courants de recherche convergents. Depuis 2023-2024, la litterature academique et institutionnelle souligne de plus en plus que la capacite computationnelle est devenue le facteur de production le plus discriminant pour l'IA de frontiere (Sevilla et al., 2022 ; Epoch AI, 2025 ; Pilz et al., 2025). Les controles a l'exportation americains (BIS, octobre 2022 ; mises a jour 2023 et 2025) placent explicitement le compute avance au coeur de la competition geopolitique, tandis que Hawkins, Lehdonvirta et Wu (2025) introduisent le concept de souverainete du compute comme dimension structurante de l'autonomie strategique.[11]",
        "Pourtant, les indices existants de competitivite IA ne placent pas le compute au centre de leur construction. L'AI Preparedness Index du FMI (Cazzaniga et al., 2024), couvrant 174 pays, agrege quatre dimensions (infrastructure numerique, capital humain, innovation/integration economique, regulation/ethique) sans mesurer directement la capacite de compute installee.[12] Le Global AI Index de Tortoise Media (2024), classant 83 pays sur 122 indicateurs en trois piliers (Implementation, Innovation, Investment), inclut une composante infrastructure/supercalcul mais la subsume dans un indice additif pondere ou le compute n'est qu'un facteur parmi d'autres.[13] Le Stanford AI Index (2025) fournit des donnees riches sur les tendances compute mais ne propose aucun indice composite. Aucun de ces instruments ne formalise le mecanisme par lequel le compute, ajuste de son cout energetique et rapporte a la capacite d'absorption d'une economie, determine la competitivite IA.",
        "Le CACI vise a combler ce gap en proposant un indicateur parcimonieux mais theoriquement fonde, qui capture l'interaction multiplicative entre quatre facteurs : compute accessible, main-d'oeuvre pertinente, conditions d'acces geopolitique et cout energetique.",
    ]),
    ("2.4.2 Definition formelle", [
        "Le CACI pour une region r a la periode t est defini en deux modes complementaires. En Power Mode (absolu) : CACI(r,t) = F(r,t)^0,40 x L(r,t)^0,20 x R(r,t)^0,15 / E(r,t)^0,25. En Intensity Mode (par unite de richesse nationale) : CACI(r,t) = F(r,t)^0,40 x L(r,t)^0,20 x R(r,t)^0,15 / (E(r,t)^0,25 x PIB(r,t)).",
        "F(r,t) - capacite de compute IA installee et accessible en region r a la periode t. Nous agregeons le jeu de donnees Epoch AI GPU Clusters par pays, en H100-equivalents (la metrique utilisee par Epoch AI), complete par les estimations OCDE de compute cloud accessible. F admet deux specifications utilisees en parallele : F_total (tous les clusters localises sur le territoire, independamment de la propriete, definissant le CACI physique) et F_dom (clusters possedes et operes par des acteurs domestiques, definissant le CACI souverain).",
        "L(r,t) - population active avec competences IA (proxy : diplomes STEM plus formations IA certifiees plus densite des competences IA selon LinkedIn Economic Graph). Ce facteur capture la capacite d'absorption au sens de Cohen et Levinthal (1990) : un compute abondant sans capital humain pour l'exploiter ne produit pas de competitivite.[14]",
        "R(r,t) - indice d'acces geopolitique (entre 0 et 1) capturant la position du pays dans le regime des controles a l'exportation. R = 1 pour les Etats-Unis (acces illimite au compute de frontiere), R = 0,95 pour les allies Tier-1 proches (Royaume-Uni), R = 0,9 pour l'UE et la France (Tier-1 avec restrictions selectives), R = 0,85 pour l'Inde (Tier-2), R = 0,7 pour la Chine (fortement restreint).",
        "E(r,t) - cout energetique effectif moyen du compute en region r, en USD par MWh, ajuste des Power Purchase Agreements des hyperscalers. Les valeurs de reference du tableau de bord pour le snapshot d'avril 2026 sont USA 85, Chine 92, France 115, Allemagne 140, Royaume-Uni 190 USD/MWh. Ces chiffres correspondent aux prix ajustes PPA effectivement supportes par les centres de donnees, non aux tarifs Eurostat industriels bruts (qui sont typiquement plus eleves en Europe). Les couts effectifs UE dans ce perimetre PPA-ajuste sont 1,4 a 1,7 fois les niveaux US.[15]",
        "PIB(r,t) - produit interieur brut de la region r (Banque mondiale, Eurostat). Utilise uniquement en Intensity Mode pour normaliser entre economies de tailles tres differentes.",
        "Justification de la forme multiplicative. Le choix d'une agregation multiplicative (geometrique) plutot qu'additive (arithmetique) est delibere et repose sur trois arguments. Premierement, la theorie des General Purpose Technologies (Bresnahan et Trajtenberg, 1995) postule une forte complementarite entre les intrants d'innovation : un compute sans energie abordable ou sans capital humain qualifie ne produit pas de gains de competitivite, justifiant une forme ou la faiblesse sur un facteur penalise l'ensemble. Deuxiemement, le OECD/JRC Handbook (2008, p. 33) recommande l'agregation geometrique lorsque les composantes ne sont pas parfaitement substituables : contrairement a la moyenne arithmetique, la moyenne geometrique ne permet pas a un score tres eleve sur une dimension de compenser pleinement un score tres faible sur une autre. Troisiemement, des travaux recents sur la construction d'indices IA (Koronakos, Kritikos et Sotiros, 2024 ; analyse du GAII Tortoise via integrale de Choquet) confirment que les dimensions de competitivite IA presentent des interactions (complementarites et redondances) qui rendent problematique l'agregation lineaire simple.[16]",
        f"Interpretation economique. En forme logarithmique, le CACI en Intensity Mode peut s'ecrire ln(CACI) = 0,40 ln(F) + 0,20 ln(L) + 0,15 ln(R) - 0,25 ln(E) - ln(PIB). Cette transformation linearise la relation et permet d'interpreter chaque coefficient comme une elasticite. L'indicateur est concu pour la comparaison bilaterale : le ratio CACI(US)/CACI(UE) mesure l'avantage competitif relatif. Le snapshot d'avril 2026 donne un ratio de {fmt_fr(us_eu_caci, 2)}:1 en Power Mode (F_total) - c'est le chiffre principal de cette etude.",
    ]),
    ("2.4.3 Protocole de calibration et sources de donnees", [
        "La calibration du CACI suit un protocole en quatre etapes coherent avec les recommandations du OECD/JRC Handbook (2008) : identification et collecte des donnees brutes ; traitement des valeurs manquantes et normalisation ; agregation ; et analyse de sensibilite.",
        f"F(r,t) - Capacite de compute IA installee. L'estimation suit une approche bottom-up. La couche principale provient du jeu de donnees Epoch AI GPU Clusters (version avril 2026), fournissant la performance en FLOP/s 16 bits et en equivalents H100 agregee par pays pour 2019-2026. Les parts nationales de compute operationnel sur le snapshot d'avril 2026 sont : USA {fmt_fr(us_share, 1)} pour cent, Chine 12,8 pour cent, UE (13 economies principales) 4,4 pour cent, Norvege 1,1 pour cent, Japon 1,0 pour cent.[17] Lorsque tous les clusters reportes sont inclus (operationnels plus planifies), la part US passe a environ 50 pour cent, le Golfe (EAU, Arabie saoudite) et la Coree captant une large part de la construction projetee 2027-2030. La couche complementaire provient de l'OCDE (Lehdonvirta et al., 2025), cataloguant les regions cloud avec accelerateurs IA dans 39 economies, capturant la dimension d'accessibilite que le compute physiquement installe ne reflete pas pleinement.",
        "Procedure d'agregation de F. Pour chaque couple pays-annee, nous extrayons du jeu de donnees Epoch AI la somme des H100-equivalents de tous les clusters localises dans le pays, en filtrant les systemes confirmes (certitude superieure ou egale a Likely) ; appliquons un facteur d'extrapolation pour corriger la sous-couverture du jeu de donnees (10 a 20 pour cent du compute mondial), en utilisant les estimations de couverture sectorielle d'Epoch AI par type de puce ; et ajoutons les capacites cloud accessibles estimees via la methodologie OCDE pour les pays ou le compute cloud etranger represente une part significative de la capacite effective. Les donnees brutes et le code de calcul (Python/pandas) sont documentes dans l'annexe methodologique et sur le tableau de bord public.[18]",
        "E(r,t) - Cout energetique. Les valeurs de reference du CSV du tableau de bord utilisees dans cette etude, en USD par MWh et ajustees PPA, sont : USA 85, Chine 92, France 115, Allemagne 140, Royaume-Uni 190, Inde 88. Ces chiffres integrent les tarifs negocies pour grands consommateurs (PPA hyperscalers), en utilisant les estimations IEA (2025, Energy and AI) du mix energetique des centres de donnees par region. Ils sont systematiquement inferieurs aux tarifs Eurostat industriels bruts en Europe, parce que les PPA hyperscalers et le sourcing direct nucleaire/renouvelable reduisent materiellement le cout effectif supporte par les centres de donnees IA. La Federal Reserve Board (octobre 2025) documente une correlation negative significative entre couts energetiques et adoption IA au niveau des entreprises europeennes.[19]",
        "PIB(r,t) et L(r,t). Le PIB est disponible aupres de la Banque mondiale (World Development Indicators) et d'Eurostat. Le snapshot d'avril 2026 du tableau de bord utilise les valeurs PIB en milliers de milliards USD : USA 29,3 ; Chine 18,7 ; UE 18,9 ; France 3,16 ; Allemagne 4,68 ; Royaume-Uni 3,6. Le proxy de capital humain IA L(r,t) combine trois sous-indicateurs : nombre de diplomes STEM (OCDE, Education at a Glance) ; densite des competences IA mesuree par LinkedIn Economic Graph (profils avec competences IA rapportes a la population active) ; et certifications IA (estimations basees sur les programmes cloud certifies AWS/Google/Microsoft par pays). Le snapshot du tableau de bord utilise, en millions de travailleurs qualifies en IA : USA 3,5 ; Chine 4,8 ; UE 3,1 ; France 1,5 ; Allemagne 1,9 ; Royaume-Uni 1,8. Nous reconnaissons un biais favorisant les pays anglophones et les economies ou LinkedIn est dominant, et documentons l'effet de ce biais dans l'analyse de sensibilite (section 2.4.5).[20]",
        "Normalisation du CACI et choix du benchmark. Le CACI brut produit des valeurs absolues heterogenes dont l'interpretation directe est difficile. Conformement a la pratique standard (Nardo et al., 2008 ; Saisana et Tarantola, 2002), nous appliquons une normalisation min-max au leader : le CACI normalise de chaque pays s'exprime en pourcentage du CACI brut des Etats-Unis, de sorte que CACI_norm(USA) = 100 par construction.",
        f"Ce choix de normalisation est methodologiquement fonde sur trois arguments. Premierement, les Etats-Unis constituent un benchmark naturel etant donne leur position dominante sur les quatre facteurs du CACI : {fmt_fr(us_share, 1)} pour cent du compute IA operationnel mondial (Epoch AI, avril 2026), les couts energetiques industriels effectifs les plus bas parmi les grandes economies OCDE (85 USD/MWh, ajuste PPA, contre 115 en France et 140 en Allemagne - CSV du tableau de bord, avril 2026), le PIB nominal le plus eleve, et un ecosysteme de capital humain IA inegale. Deuxiemement, la normalisation au leader est la convention etablie pour les indices de competitivite relative : le RCA de Balassa (1965), le GCI du WEF (Schwab, 2019) et le Competitive Industrial Performance Index de l'ONUDI procedent tous par ratio au leader sectoriel ou national. Troisiemement, cette normalisation elimine les ambiguites d'interpretation des valeurs absolues et permet les comparaisons temporelles : un CACI France passant de 25 a 35 entre 2024 et 2028 signifie un rattrapage partiel, meme si les valeurs brutes des deux pays ont augmente.",
        f"La consequence methodologique la plus importante de cette normalisation est la revelation d'un ecart structurel que les indices traditionnels masquent. L'AI Preparedness Index du FMI (Cazzaniga et al., 2024) attribue des scores de 0,85 (USA) et 0,74 (France), un ratio de 1,15:1, suggerant une quasi-parite. Le Global AI Index de Tortoise produit des ecarts similaires. Pourtant, le CACI normalise en Power Mode place la France a environ 25 et l'Allemagne a environ 5 sur une echelle ou les Etats-Unis = 100, soit des ratios respectifs de 4:1 et 19:1. Le ratio principal US/UE (UE agregee comme 13 economies principales) est de {fmt_fr(us_eu_caci, 2)}:1. Cette divergence n'est pas contradictoire : elle resulte du fait que les indices multidimensionnels agregent lineairement des dimensions normatives (regulation, ethique, publications) qui compensent arithmetiquement les deficits sur les facteurs materiels (compute, energie). Le CACI, en isolant ces facteurs materiels, expose le gouffre physique que les metriques composites moyennent.",
        "Cette propriete du CACI - reveler l'ampleur reelle du compute gap - constitue la contribution analytique centrale de l'indice et la these empirique fondamentale de cette etude. L'ecart CACI n'est pas un artefact de normalisation : c'est le resultat scientifique.",
    ]),
    ("2.4.4 Positionnement par rapport aux indices existants de competitivite IA", [
        "L'AI Preparedness Index (AIPI) du FMI couvre 174 pays (2023) et agrege quatre piliers : infrastructure numerique, capital humain et politiques du marche du travail, innovation et integration economique, regulation et ethique. L'AIPI ne mesure pas le compute IA installe et ne pondere pas par le cout energetique. Sa correlation attendue avec le CACI est positive mais imparfaite : les pays bien classes a l'AIPI (Singapour, Danemark, Pays-Bas) ne sont pas necessairement ceux dont le compute effectif par unite de PIB est le plus eleve.[21]",
        "Le Global AI Index (GAII) de Tortoise Media classe 83 pays sur 122 indicateurs en trois piliers (Implementation, Innovation, Investment). Il inclut une composante infrastructure/supercalcul mais l'agrege lineairement avec des poids subjectifs (reconnu par Tortoise comme une limite). Le GAII est plus large et plus multidimensionnel que le CACI, mais precisement parce qu'il est large, il dilue le signal compute dans un composite a nombreux indicateurs. Comme le montrent Koronakos et al. (2024), la subjectivite des ponderations du GAII peut inverser les classements de pays selon les scenarios de ponderation choisis.",
        "Le Stanford AI Index (2025) constitue la reference la plus complete en termes de donnees brutes (modeles notables par pays, investissement, publications, brevets, tendances compute). Il ne propose pas d'indice composite mais fournit les series temporelles utilisees par de nombreux autres indices. Les donnees publiques du Stanford AI Index sont accessibles via un dossier Google Drive ouvert.[22]",
        "Valeur ajoutee specifique du CACI. Le CACI se distingue par quatre proprietes : il place le compute au centre plutot qu'a la peripherie de l'indicateur, refletant le role desormais dominant du compute dans l'IA de frontiere ; il integre explicitement le cout energetique comme goulot d'etranglement, coherent avec les conclusions de l'IEA (2025) ; il utilise une agregation multiplicative theoriquement fondee plutot qu'additive ; et il est parcimonieux (quatre variables), le rendant transparent et reproductible, au prix d'une exhaustivite moindre.",
    ]),
    ("2.4.5 Limites du CACI et analyse de sensibilite", [
        "Premiere limite - opacite de F(r,t). La mesure du compute installe depend de donnees privees incompletes. Le jeu de donnees Epoch AI ne couvre que 10-20 pour cent du compute mondial, avec une couverture inegale entre secteurs et entreprises. L'attribution nationale est elle-meme discutable : une part croissante du compute est detenue par quelques hyperscalers prives operant globalement. Mitigation : nous documentons systematiquement les marges d'incertitude, presentons des intervalles plutot que des valeurs ponctuelles, et verifions la stabilite des resultats lorsque F varie de plus ou moins 30 pour cent. La decomposition Physique/Souverain traite partiellement le probleme d'attribution.",
        "Deuxieme limite - heterogeneite qualitative du compute. Le CACI agrege des FLOPs sans distinguer les generations de GPU (un H200 ne vaut pas un A100 en efficacite energetique et performance reelle). Mitigation : le jeu de donnees Epoch AI fournit la performance en equivalents H100, offrant une normalisation partielle. Nous proposons un facteur de ponderation par generation de GPU en annexe et montrons que l'ajustement modifie marginalement les classements.",
        "Troisieme limite - le proxy de capital humain L(r,t). La combinaison STEM plus LinkedIn plus certifications presente un biais favorisant les pays anglophones. Ce biais sous-estime probablement la Chine et certaines economies asiatiques. Mitigation : nous repliquons l'analyse en utilisant le sous-indice Human Capital de l'AIPI du FMI comme proxy alternatif et montrons la sensibilite des classements a ce choix.",
        "Quatrieme limite - statisme de l'indicateur. Le CACI mesure un etat a un instant donne, non une dynamique. C'est pourquoi nous le calculons sur plusieurs annees (2022, 2024, 2026) et le projetons dans chaque scenario, permettant le suivi de trajectoire et l'evaluation de l'evolution de l'ecart.",
        "Cinquieme limite - endogeneite. Les pays gagnant en productivite IA investissent massivement dans le compute, creant un risque de causalite inverse : le CACI pourrait capturer la consequence plutot que la cause de la competitivite. Cette etude, dans son cadre de these, n'instrumente pas formellement cette relation (pas de variables instrumentales ni de strategie GMM). Cependant, nous notons que le choc exogene des regles BIS d'octobre 2022 offre une quasi-experience naturelle qui pourrait fonder une strategie d'identification causale en Difference-in-Differences. Nous identifions le traitement formel de cette endogeneite (DiD, IV ou GMM Arellano-Bond) comme une piste de recherche prioritaire pour toute extension publiable de ce travail.[23]",
        "Sixieme limite - biais de normalisation pour les petites economies. L'analyse comparative du CACI face aux indices AIPI du FMI et GAII de Tortoise revele une anomalie instructive : certains pays a faible PIB et a main-d'oeuvre IA reduite obtiennent des scores CACI disproportionnellement eleves. L'Afrique du Sud est emblematique : classee derniere ou avant-derniere aux indices AIPI et Tortoise, elle peut apparaitre comme leader mondial sur le CACI en Intensity Mode lorsque la normalisation par le PIB amplifie un effet de petit denominateur. Ce phenomene est un cas classique de biais de normalisation identifie dans la litterature des indicateurs composites. Le OECD/JRC Handbook (Nardo et al., 2008, pp. 27-29) avertit que la normalisation par le PIB ou la population peut produire des resultats trompeurs pour les petites economies. Notre mitigation est triple : un seuil de masse critique sur F (excluant les economies sous 10 000 H100-equivalents) ; un facteur d'echelle alpha(r) = min(1, F(r) / F_median) pour les tests de sensibilite ; et un classement double (Power Mode pour la puissance absolue, Intensity Mode pour la densite), avec la reserve explicite que les comparaisons bilaterales dans cette etude se concentrent sur les USA, l'UE, la France et la Chine.[24]",
    ]),
    ("2.4.6 Chiffres reproductibles - snapshot avril 2026", [
        "Pour permettre au lecteur de rejouer le calcul a partir des CSV du tableau de bord public, nous exposons ici les entrees et sorties utilisees tout au long de cette etude. Les quatre variables d'entree sont sourcees depuis le dossier de donnees du tableau de bord (energy_prices.csv, gdp_data.csv, workforce_data.csv, gpu_clusters.csv) a https://mo0ogly.github.io/America-First-IA/dashboard/data/.",
        "Entrees (snapshot avril 2026). USA : F_total = 39,65 millions de H100-eq, L = 3,5 millions, R = 1,00, E = 85 USD/MWh, PIB = 29,3 milliers de milliards USD. UE (13 economies principales agregees) : F_total = 2,62 millions de H100-eq, L = 3,1 millions, R = 0,90, E = 135 USD/MWh, PIB = 18,9 milliers de milliards USD. France : F_total = 2,44 millions de H100-eq, L = 1,5 millions, R = 0,90, E = 115 USD/MWh, PIB = 3,16 milliers de milliards USD. Allemagne : F_total = 0,05 million de H100-eq, L = 1,9 millions, R = 0,90, E = 140 USD/MWh, PIB = 4,68 milliers de milliards USD. Chine : F_total = 0,40 million de H100-eq, L = 4,8 millions, R = 0,70, E = 92 USD/MWh, PIB = 18,7 milliers de milliards USD.",
        f"Sorties (Power Mode F_total, normalisees a USA = 100). USA = 100,0 ; UE(13) = {fmt_fr(caci_scores['EU'], 1)} ; France = {fmt_fr(caci_scores['France'], 1)} ; Inde = {fmt_fr(caci_scores['India'], 1)} ; Chine = {fmt_fr(caci_scores['China'], 1)} ; Royaume-Uni = {fmt_fr(caci_scores['UK'], 1)} ; Allemagne = {fmt_fr(caci_scores['Germany'], 1)}. EAU = 55,7 en CACI Physique mais seulement 6,0 en CACI Souverain (F_dom restreint aux clusters detenus par des entites domestiques G42, MGX, Mubadala, Khazna, TII), illustrant la decomposition Physique/Souverain : 99,6 pour cent du F_total des EAU est detenu par des acteurs US-side (Stargate UAE, Microsoft, OpenAI). Ratios principaux : US/UE(13) = {fmt_fr(us_eu_caci, 2)}:1 ; US/France = 3,96:1 ; US/Allemagne = 18,59:1 ; US/Chine = 7,46:1.",
        f"Ces sorties sont reproduites en direct sur le tableau de bord public. Toute mise a jour ulterieure du tableau de bord (nouveaux clusters, revisions energetiques, rafraichissement PIB) rafraichira mecaniquement les chiffres ; la methodologie ne change pas. Le ratio de {fmt_fr(us_eu_caci, 2)}:1 en Power Mode F_total reporte sur la couverture et tout au long de cette etude correspond exactement a ce snapshot.",
    ]),
    ("2.5 Perimetre et delimitations", [
        "Perimetre geographique. L'analyse se concentre sur la relation bilaterale Etats-Unis / Union europeenne, avec un focus specifique sur la France. La Chine est traitee comme variable contextuelle (cible principale des controles a l'exportation americains, facteur de pression sur les capacites de production de puces) mais ne fait pas l'objet d'une analyse approfondie. Le Japon, la Coree du Sud et Taiwan apparaissent comme acteurs de la chaine d'approvisionnement en semi-conducteurs.",
        "Perimetre temporel. Le diagnostic couvre 2020-2026, les scenarios couvrent 2026-2030. L'horizon 2030 est choisi parce qu'il correspond a la convergence de plusieurs echeances : projections AIE pour l'energie des centres de donnees, maturite attendue de l'EU Chips Act, objectifs SNIA France 2030 et arrivee potentielle des premiers SMR nucleaires operationnels.",
        "Perimetre technologique. L'etude couvre l'IA de frontiere (modeles de fondation, intensifs en compute) et ses prerequis materiels (GPU/ASIC, centres de donnees, energie). Elle integre la robotique IA comme facteur amplificateur de la demande energetique. Elle ne traite pas l'IA embarquee edge (smartphones, IoT), sauf dans la mesure ou elle constitue un objectif specifique de la SNIA francaise.",
    ]),
    ("2.6 Limites methodologiques generales", [
        "Incertitude politique radicale. Le protectionnisme technologique depend de decisions politiques discretionnaires a predictibilite structurellement faible. Un changement d'administration americaine en 2028, un accord commercial UE-USA inattendu ou une escalade du conflit US-Chine pourrait invalider certaines hypotheses. C'est precisement pourquoi nous proposons quatre scenarios plutot qu'une trajectoire unique.",
        "Disruptions technologiques. L'episode DeepSeek (janvier 2025), ou un modele chinois a atteint des performances proches de la frontiere avec un budget d'entrainement substantiellement reduit, illustre la possibilite de percees en efficacite qui modifieraient les termes du probleme. L'IEA (2025, Energy and AI) consacre une etude de cas a DeepSeek et conclut que meme avec des ameliorations significatives d'efficacite, la croissance de la demande absorbe les gains (effet rebond de Jevons).[25]",
        "Opacite des donnees compute. Le nombre exact de GPU deployes par hyperscaler, la distribution geographique precise des centres de donnees et les volumes de GPU exportes par region sont des donnees partiellement ou totalement confidentielles. Nos estimations de compute installe portent des marges d'erreur significatives, que nous documentons systematiquement.",
        "Biais des sources de conseil. Comme note en section 2.2, les sources industrielles portent un biais d'optimisme systematique. Nous l'attenuons par triangulation mais ne pouvons l'eliminer entierement.",
        "Ces limites ne compromettent pas la validite de l'analyse. La methode des scenarios est precisement concue pour fonctionner dans des environnements de haute incertitude, ou l'objectif n'est pas la prediction mais l'exploration structuree des possibles. Comme le note Schwartz : les scenarios ne sont pas des previsions ; ce sont des histoires plausibles qui aident a penser.[26] Notre contribution reside dans la rigueur du cadrage, l'explicitation des hypotheses, la transparence des sources de donnees et l'originalite de l'indicateur CACI, plutot que dans la precision des projections numeriques.",
    ]),
]

FR.notes = EN.notes  # bibliographic references identical


# ===========================================================================
# Content - Brazilian Portuguese
# ===========================================================================

PT = LangPack(
    code="PT-BR",
    filename="Capitulo_II_Metodologia_PT-BR.docx",
    cover_title="AI FOR AMERICANS FIRST",
    cover_subtitle="Protecionismo de IA, Energia e Semicondutores: Trajetorias de Divergencia EUA/Europa 2024-2030",
    cover_blurb="Analise geoestrategica e economica integrada - Capitulo II",
    cover_chip_lines=[
        f"{fmt_fr(us_share, 1)} pct do compute IA operacional mundial = EUA",
        "1,59x custo de energia UE/EUA",
        f"{fmt_fr(us_eu_caci, 2)}:1 razao CACI EUA/UE (Power Mode)",
    ],
    cover_meta="Paris - fevereiro de 2026  |  7 capitulos  |  4 cenarios prospectivos  |  3 zonas geograficas",
    cover_keywords_label="Palavras-chave",
    cover_keywords=("inteligencia artificial, protecionismo tecnologico, semicondutores, "
                    "controles de exportacao, compute soberano, geopolitica da IA, Franca, "
                    "Estados Unidos, China"),
    chapter_label="CAPITULO II",
    chapter_title="Metodologia",
    notes_label="Notas",
    license_block=[
        "Licenca e isencao de responsabilidade. Esta obra, 'AI for Americans First', e disponibilizada nos termos da Licenca Creative Commons Atribuicao - NaoComercial - CompartilhaIgual 4.0 Internacional (CC BY-NC-SA 4.0) do projeto America-First-IA.",
        "Voce e livre para compartilhar e adaptar o material para fins nao comerciais, desde que credite adequadamente a obra a Fabrice Pizzi (Universite Paris-Sorbonne) e distribua suas contribuicoes sob a mesma licenca. Este documento e fornecido apenas para fins educacionais e de pesquisa.",
        "Painel publico: https://mo0ogly.github.io/America-First-IA/dashboard/",
        "Repositorio: https://github.com/mo0ogly/America-First-IA",
    ],
    page_footer="AI for Americans First - Fabrice Pizzi - Capitulo II",
)

PT.sections = [
    ("2.1 Abordagem geral: analise prospectiva multi-cenario por metodo misto", [
        "Este estudo combina uma analise empirica retrospectiva (diagnostico 2020-2026) com uma projecao prospectiva por cenarios (2026-2030). Esta arquitetura em duas vertentes responde a natureza do fenomeno investigado: o protecionismo tecnologico em IA e simultaneamente um fato observavel (controles de exportacao, tarifas Secao 232) e um processo em curso cuja trajetoria futura depende de variaveis politicas discricionarias, das respostas estrategicas europeias e de desenvolvimentos tecnologicos parcialmente imprevisiveis.",
        "O componente retrospectivo emprega um metodo quantitativo descritivo, baseado na agregacao e cruzamento de dados de fontes institucionais (AIE, SIA/WSTS, Eurostat, EIA), relatorios industriais (McKinsey, Deloitte, Epoch AI) e documentos regulatorios (Federal Register, BIS, Casa Branca). O objetivo e estabelecer uma base factual rigorosa e referenciada cobrindo tres dimensoes: consumo de energia dos data centers, mercado de semicondutores e capacidade de compute IA instalada por regiao.",
        "O componente prospectivo apoia-se na tradicao do scenario planning conforme formalizada por Schwartz (1991) e praticada na Royal Dutch/Shell desde a decada de 1970.[1] Este metodo, pertencente a escola Intuitive Logics (Bradfield et al., 2005), consiste em construir cenarios plausiveis e internamente coerentes - nao para prever o futuro, mas para explorar o espaco de possibilidades e avaliar a robustez de diferentes estrategias frente a desenvolvimentos ambientais divergentes.[2] E particularmente adequado a situacoes caracterizadas por alta incerteza politica e tecnologica, onde os modelos econometricos classicos atingem seus limites - precisamente o caso do protecionismo tecnologico em IA.",
        "Justificativa da escolha metodologica. Tres razoes fundamentam a escolha do metodo de cenarios em vez de modelagem econometrica pura. Primeiro, as variaveis analiticas-chave sao em grande parte politicas e discricionarias: a decisao de um presidente americano de impor ou nao cotas de GPU a Europa nao pode ser modelada por uma funcao de regressao. Segundo, as interacoes entre as dimensoes energetica, tecnologica e geopolitica sao nao lineares e sistemicas: uma restricao sobre GPUs pode, por efeitos em cascata, alterar fluxos de investimento energetico, decisoes de localizacao de data centers e a estrutura competitiva de setores inteiros. Terceiro, os dados disponiveis sobre compute instalado por regiao sao parciais e heterogeneos: nao existe banco de dados publico unificado de FLOPs IA por pais, tornando prematura uma calibracao econometrica rigorosa.",
        "O metodo escolhido combina, portanto, o rigor quantitativo do diagnostico empirico (dados referenciados, series temporais, razoes mensuraveis) com a flexibilidade qualitativa da construcao de cenarios, no espirito do que Schoemaker (1995) chama de heuristica disciplinada.[3] Os cenarios nao sao previsoes probabilisticas, mas narrativas estrategicas coerentes, cada uma baseada em pressupostos explicitos e desenvolvendo suas consequencias por meio de metricas mensuraveis.",
    ]),
    ("2.2 Fontes de dados: classificacao e avaliacao critica", [
        "O estudo se baseia em tres categorias de fontes, cuja confiabilidade e potenciais vieses devem ser explicitamente reconhecidos. Esta transparencia metodologica esta em conformidade com as recomendacoes do OECD/JRC Handbook on Constructing Composite Indicators (Nardo et al., 2008), que prescreve a documentacao sistematica de fontes, suas limitacoes e seus vieses em qualquer construcao de indicador composto.[4]",
    ]),
    ("2.2.1 Fontes primarias (documentos oficiais e regulatorios)", [
        "Esta categoria inclui textos normativos ou institucionais: proclamacoes presidenciais (Secao 232), regras BIS (AI Diffusion Rule, Entity List), relatorios da AIE, publicacoes do Parlamento Europeu (EPRS), dados estatisticos oficiais (SIA/WSTS para semicondutores, Eurostat e EIA para energia, RTE para Franca). Essas fontes oferecem a maior confiabilidade factual, mas podem conter vieses de enquadramento institucional: a AIE tende a privilegiar cenarios moderados, o Parlamento Europeu a enfatizar riscos para a soberania da UE.",
    ]),
    ("2.2.2 Fontes academicas e think tanks", [
        "Esta categoria inclui artigos com revisao por pares (Farrell e Newman, 2019; Bresnahan e Trajtenberg, 1995; Brynjolfsson et al., 2019; Mugge, 2024) e publicacoes de think tanks reconhecidos (Bruegel, Carnegie Endowment, CSIS, OCDE, Federal Reserve Board). Os primeiros fornecem fundamentacao teorica robusta; os segundos oferecem analises politicas empiricamente fundadas, mas potencialmente influenciadas pela orientacao ideologica de cada instituicao. Priorizamos o cruzamento de fontes de orientacoes diferentes (Bruegel / Carnegie / Fed) para limitar esse vies.",
    ]),
    ("2.2.3 Fontes industriais e de consultoria", [
        "McKinsey, Deloitte, Accenture, Epoch AI e CFG Europe fornecem dados de mercado, projecoes setoriais e estimativas de capacidade indisponiveis em fontes publicas. Essas fontes apresentam um vies sistematico potencial: as consultorias tem interesse em amplificar tendencias (para justificar engajamentos de transformacao) e as estimativas de mercado sao frequentemente otimistas. Mitigamos esse vies triangulando os numeros com dados institucionais e sinalizando explicitamente as divergencias entre fontes. Por exemplo, as vendas de semicondutores em 2024 sao de 627,6 bilhoes de USD segundo a SIA (escopo tradicional), mas de 775 bilhoes de USD segundo a McKinsey (escopo ampliado), uma diferenca de 24 por cento refletindo diferencas metodologicas e nao inconsistencias.[5]",
    ]),
    ("2.2.4 Dados de compute IA: o conjunto Epoch AI GPU Clusters", [
        "Medir o compute IA instalado por pais constitui o desafio metodologico central deste estudo. Apoiamo-nos primariamente no conjunto de dados Epoch AI GPU Clusters (Pilz, Rahman, Sanders e Heim, 2025), que cataloga mais de 500 supercomputadores e clusters GPU em todo o mundo para o periodo 2019-2025.[6] Este conjunto de dados, disponivel sob licenca aberta Creative Commons Atribuicao, constitui a fonte mais completa e sistematicamente documentada sobre infraestrutura global de compute IA ate hoje. E usado como referencia pelo Stanford AI Index Report (2025), por varios relatorios governamentais e por instituicoes como OpenAI e DeepMind.",
        "O conjunto de dados cobre, para cada cluster: pais de localizacao, tipo de chip (H100, A100, GB200, TPU, etc.), desempenho computacional em FLOP/s 16 bits, numero de equivalentes H100, data operacional, setor (privado/publico), potencia eletrica (MW) e custo de hardware estimado. Essa granularidade permite a agregacao por pais e ano, atendendo diretamente as necessidades de nossa variavel F(r) no CACI.",
        "Limitacoes do conjunto Epoch AI. Tres limitacoes devem ser destacadas. Primeiro, a cobertura e estimada em 10-20 por cento do desempenho de compute IA agregado global (marco de 2025), com heterogeneidade significativa entre empresas e tipos de chip: aproximadamente 20-37 por cento dos NVIDIA H100, 12 por cento dos A100, mas menos de 4 por cento dos TPU Google e uma fracao desprezivel dos chips customizados de AWS, Microsoft ou Meta.[7] Segundo, os sistemas chineses sao anonimizados (nomes removidos, valores arredondados a um algarismo significativo), limitando a precisao analitica para a China. Terceiro, a localizacao fisica de um cluster nao determina o acesso: muitos clusters sao acessiveis via servicos cloud a partir de outros paises.",
        "Complementamos esses dados com o Working Paper da OCDE de Lehdonvirta, Wu, Hawkins et al. (outubro de 2025), que desenvolve uma metodologia para estimar a disponibilidade de compute IA em cloud publico por pais, contando regioes cloud dos principais fornecedores equipadas com aceleradores IA (A100, H100, GB200) em 39 economias.[8] Essa abordagem complementar distingue compute instalado (Epoch AI) de compute acessivel (OCDE), uma distincao crucial para o CACI.",
    ]),
    ("2.3 Construcao dos cenarios", [
        "A construcao dos cenarios segue um protocolo de quatro etapas, inspirado na metodologia da matriz 2x2 (Schwartz, 1991; van der Heijden, 2004) e adaptado ao contexto geoestrategico da IA.[9]",
        "Etapa 1 - Identificacao das forcas motrizes. Os elementos predeterminados (cuja evolucao e razoavelmente previsivel independentemente do cenario) incluem: o crescimento continuo da demanda mundial de compute IA; o aumento estrutural do consumo de energia dos data centers; a dependencia europeia das fundicoes asiaticas e americanas para chips de ponta; a concentracao do cloud global em torno de tres hyperscalers americanos; e o aumento exponencial dos custos de treinamento de modelos de fronteira.",
        "As incertezas criticas (cuja evolucao depende de escolhas politicas, reacoes estrategicas ou disrupcoes tecnologicas) sao agrupadas em duas dimensoes. Dimensao 1 - Intensidade do protecionismo tecnologico americano: esta dimensao cobre um espectro que vai da manutencao das restricoes atuais ao endurecimento agressivo (cotas GPU para a UE, restricoes a APIs e modelos, priorizacao explicita de entregas a empresas americanas). Dimensao 2 - Capacidade de resposta europeia: esta dimensao cobre um espectro que vai de uma postura passiva (adaptacao marginal, aceitacao da dependencia) a uma resposta ativa (Compute Zones com energia derrogada, AI Factories aceleradas, SMR nucleares para data centers, parcerias alternativas Japao-Coreia-Taiwan, revisao do AI Act).",
        "Etapa 2 - Matriz 2x2 e geracao de cenarios. O cruzamento das duas dimensoes de incerteza gera uma matriz de quatro cenarios. A escolha de quatro em vez de tres cenarios e deliberada. Schwartz (1991) e os praticantes do metodo Shell recomendam nunca construir tres cenarios, pois a mente humana tende a tratar o cenario mediano como o mais provavel, reduzindo a utilidade do exercicio.[10] A matriz 2x2 forca o analista a explorar quadrantes extremos - precisamente onde se desenrolam as rupturas estrategicas.",
        "Etapa 3 - Desenvolvimento narrativo e quantificacao. Cada cenario e desenvolvido segundo um protocolo padronizado composto de tres componentes: uma narrativa estrategica descrevendo a sequencia plausivel de eventos entre 2026 e 2030; uma quantificacao de metricas-chave calibradas em dados empiricos 2020-2026 e projetadas conforme os pressupostos do cenario; e indicadores antecipados (leading indicators) que permitem identificar, a partir de 2026-2027, para qual cenario a realidade esta convergindo.",
        "Etapa 4 - Analise de sensibilidade e robustez. Para cada recomendacao formulada no Capitulo VII, avaliamos sua robustez ao longo dos quatro cenarios. Uma recomendacao e considerada robusta se produz resultados positivos ou neutros em pelo menos tres dos quatro cenarios.",
    ]),
    ("2.4 Metricas-chave e indicador original: o CACI", [
        "Definimos seis metricas a serem calculadas ou estimadas no diagnostico empirico (Capitulo III), e em seguida projetadas em cada cenario (Capitulo V). Em conjunto, formam um painel da divergencia EUA/UE em IA: M1 (compute gap, razao EUA/UE de FLOPs IA instalados normalizada pelo PIB), M2 (custo relativo do FLOP para treinamento), M3 (dependencia cloud, parcela de cargas IA UE em infraestrutura EUA), M4 (produtividade IA setorial), M5 (restricao energetica, razao demanda/capacidade dos data centers) e M6 (deslocalizacoes IA).",
    ]),
    ("2.4.1 Fundamentos teoricos do CACI", [
        "Ancoragem na literatura. A construcao de um indicador composto de competitividade IA centrado no compute responde a uma necessidade identificada por varios fluxos de pesquisa convergentes. Desde 2023-2024, a literatura academica e institucional enfatiza cada vez mais que a capacidade computacional se tornou o fator de producao mais discriminante para a IA de fronteira (Sevilla et al., 2022; Epoch AI, 2025; Pilz et al., 2025). Os controles de exportacao americanos (BIS, outubro de 2022; atualizacoes 2023 e 2025) colocam explicitamente o compute avancado no cerne da competicao geopolitica, enquanto Hawkins, Lehdonvirta e Wu (2025) introduzem o conceito de soberania do compute como dimensao estruturante da autonomia estrategica.[11]",
        "No entanto, os indices existentes de competitividade IA nao colocam o compute no centro de sua construcao. O AI Preparedness Index do FMI (Cazzaniga et al., 2024), cobrindo 174 paises, agrega quatro dimensoes (infraestrutura digital, capital humano, inovacao/integracao economica, regulacao/etica) sem medir diretamente a capacidade de compute instalada.[12] O Global AI Index da Tortoise Media (2024), classificando 83 paises em 122 indicadores em tres pilares (Implementation, Innovation, Investment), inclui um componente de infraestrutura/supercomputacao, mas o subsume em um indice aditivo ponderado onde o compute e apenas um fator entre muitos.[13] O Stanford AI Index (2025) fornece dados ricos sobre tendencias de compute, mas nao propoe nenhum indice composto. Nenhum desses instrumentos formaliza o mecanismo pelo qual o compute, ajustado por seu custo energetico e relacionado a capacidade de absorcao de uma economia, determina a competitividade IA.",
        "O CACI visa preencher essa lacuna propondo um indicador parcimonioso, mas teoricamente fundamentado, que captura a interacao multiplicativa entre quatro fatores: compute acessivel, forca de trabalho relevante, condicoes de acesso geopolitico e custo energetico.",
    ]),
    ("2.4.2 Definicao formal", [
        "O CACI para uma regiao r no periodo t e definido em dois modos complementares. Em Power Mode (absoluto): CACI(r,t) = F(r,t)^0,40 x L(r,t)^0,20 x R(r,t)^0,15 / E(r,t)^0,25. Em Intensity Mode (por unidade de riqueza nacional): CACI(r,t) = F(r,t)^0,40 x L(r,t)^0,20 x R(r,t)^0,15 / (E(r,t)^0,25 x PIB(r,t)).",
        "F(r,t) - capacidade de compute IA instalada e acessivel na regiao r no periodo t. Agregamos o conjunto Epoch AI GPU Clusters por pais, em equivalentes H100 (a metrica utilizada pelo Epoch AI), complementado por estimativas OCDE de compute cloud acessivel. F admite duas especificacoes utilizadas em paralelo: F_total (todos os clusters localizados no territorio independentemente da propriedade, definindo o CACI fisico) e F_dom (clusters possuidos e operados por atores domesticos, definindo o CACI soberano).",
        "L(r,t) - populacao em idade ativa com competencias IA (proxy: graduados STEM mais formacoes IA certificadas mais densidade de competencias IA pelo LinkedIn Economic Graph). Este fator captura a capacidade de absorcao no sentido de Cohen e Levinthal (1990): compute abundante sem capital humano para explora-lo nao produz competitividade.[14]",
        "R(r,t) - indice de acesso geopolitico (entre 0 e 1) capturando a posicao do pais no regime de controles de exportacao. R = 1 para os Estados Unidos (acesso ilimitado ao compute de fronteira), R = 0,95 para aliados Tier-1 proximos (Reino Unido), R = 0,9 para a UE e a Franca (Tier-1 com restricoes seletivas), R = 0,85 para a India (Tier-2), R = 0,7 para a China (fortemente restrito).",
        "E(r,t) - custo energetico efetivo medio do compute na regiao r, em USD por MWh, ajustado pelos Power Purchase Agreements dos hyperscalers. Os valores de referencia do painel para o snapshot de abril de 2026 sao EUA 85, China 92, Franca 115, Alemanha 140, Reino Unido 190 USD por MWh. Estes numeros correspondem aos precos ajustados PPA efetivamente enfrentados pelos data centers, nao aos tarifas Eurostat industriais brutos (que sao tipicamente mais altos na Europa). Os custos efetivos UE neste escopo PPA-ajustado sao 1,4 a 1,7 vezes os niveis EUA.[15]",
        "PIB(r,t) - produto interno bruto da regiao r (Banco Mundial, Eurostat). Utilizado apenas em Intensity Mode para normalizar entre economias de tamanhos muito diferentes.",
        "Justificativa da forma multiplicativa. A escolha de uma agregacao multiplicativa (geometrica) em vez de aditiva (aritmetica) e deliberada e repousa em tres argumentos. Primeiro, a teoria das General Purpose Technologies (Bresnahan e Trajtenberg, 1995) postula uma forte complementaridade entre os insumos de inovacao: compute sem energia acessivel ou sem capital humano qualificado nao produz ganhos de competitividade, justificando uma forma onde a fraqueza em um fator penaliza o conjunto. Segundo, o OECD/JRC Handbook (2008, p. 33) recomenda a agregacao geometrica quando os componentes nao sao perfeitamente substituiveis: ao contrario da media aritmetica, a media geometrica nao permite que um escore muito alto em uma dimensao compense plenamente um escore muito baixo em outra. Terceiro, trabalhos recentes sobre construcao de indices IA (Koronakos, Kritikos e Sotiros, 2024; analise do GAII Tortoise via integral de Choquet) confirmam que as dimensoes de competitividade IA apresentam interacoes (complementaridades e redundancias) que tornam problematica a agregacao linear simples.[16]",
        f"Interpretacao economica. Em forma logaritmica, o CACI em Intensity Mode pode ser escrito ln(CACI) = 0,40 ln(F) + 0,20 ln(L) + 0,15 ln(R) - 0,25 ln(E) - ln(PIB). Esta transformacao lineariza a relacao e permite interpretar cada coeficiente como uma elasticidade. O indicador e concebido para a comparacao bilateral: a razao CACI(EUA)/CACI(UE) mede a vantagem competitiva relativa. O snapshot de abril de 2026 produz uma razao de {fmt_fr(us_eu_caci, 2)}:1 em Power Mode (F_total) - este e o numero principal deste estudo.",
    ]),
    ("2.4.3 Protocolo de calibracao e fontes de dados", [
        "A calibracao do CACI segue um protocolo de quatro etapas consistente com as recomendacoes do OECD/JRC Handbook (2008): identificacao e coleta de dados brutos; tratamento de valores ausentes e normalizacao; agregacao; e analise de sensibilidade.",
        "F(r,t) - Capacidade de compute IA instalada. A estimacao segue uma abordagem bottom-up. A camada principal vem do conjunto Epoch AI GPU Clusters (versao abril de 2026), fornecendo o desempenho em FLOP/s 16 bits e em equivalentes H100 agregado por pais para 2019-2026. As parcelas nacionais de compute operacional no snapshot de abril de 2026 sao: EUA 76,9 por cento, China 12,8 por cento, UE (13 economias principais) 4,4 por cento, Noruega 1,1 por cento, Japao 1,0 por cento.[17] Quando todos os clusters reportados sao incluidos (operacionais mais planejados), a parcela EUA passa para cerca de 50 por cento, com o Golfo (EAU, Arabia Saudita) e a Coreia capturando uma grande parcela da construcao projetada 2027-2030. A camada complementar vem da OCDE (Lehdonvirta et al., 2025), catalogando regioes cloud com aceleradores IA em 39 economias, capturando a dimensao de acessibilidade que o compute fisicamente instalado nao reflete plenamente.",
        "Procedimento de agregacao de F. Para cada par pais-ano, extraimos do conjunto Epoch AI a soma dos equivalentes H100 de todos os clusters localizados no pais, filtrando os sistemas confirmados (certeza maior ou igual a Likely); aplicamos um fator de extrapolacao para corrigir a sub-cobertura do conjunto (10 a 20 por cento do compute global), usando as estimativas de cobertura setorial do Epoch AI por tipo de chip; e adicionamos as capacidades cloud acessiveis estimadas via metodologia OCDE para paises onde o compute cloud estrangeiro representa uma parcela significativa da capacidade efetiva. Os dados brutos e o codigo de calculo (Python/pandas) estao documentados no apendice metodologico e no painel publico.[18]",
        "E(r,t) - Custo energetico. Os valores de referencia do CSV do painel utilizados ao longo deste estudo, em USD por MWh e ajustados PPA, sao: EUA 85, China 92, Franca 115, Alemanha 140, Reino Unido 190, India 88. Estes numeros incorporam tarifas negociadas para grandes consumidores (PPAs hyperscalers), usando estimativas IEA (2025, Energy and AI) do mix energetico dos data centers por regiao. Sao sistematicamente menores do que as tarifas Eurostat industriais brutas na Europa, porque os PPAs hyperscalers e o sourcing direto nuclear/renovavel reduzem materialmente o custo efetivo enfrentado pelos data centers IA. O Federal Reserve Board (outubro de 2025) documenta uma correlacao negativa significativa entre custos energeticos e adocao IA no nivel das empresas europeias.[19]",
        "PIB(r,t) e L(r,t). O PIB esta disponivel no Banco Mundial (World Development Indicators) e no Eurostat. O snapshot de abril de 2026 do painel utiliza valores de PIB em trilhoes de USD: EUA 29,3; China 18,7; UE 18,9; Franca 3,16; Alemanha 4,68; Reino Unido 3,6. O proxy de capital humano IA L(r,t) combina tres sub-indicadores: numero de graduados STEM (OCDE, Education at a Glance); densidade de competencias IA medida pelo LinkedIn Economic Graph (perfis com competencias IA em relacao a populacao ativa); e certificacoes IA (estimativas baseadas em programas cloud certificados AWS/Google/Microsoft por pais). O snapshot do painel utiliza, em milhoes de trabalhadores qualificados em IA: EUA 3,5; China 4,8; UE 3,1; Franca 1,5; Alemanha 1,9; Reino Unido 1,8. Reconhecemos um vies favorecendo paises de lingua inglesa e economias onde o LinkedIn e dominante, e documentamos o efeito desse vies na analise de sensibilidade (secao 2.4.5).[20]",
        "Normalizacao do CACI e selecao do benchmark. O CACI bruto produz valores absolutos heterogeneos cuja interpretacao direta e dificil. Em conformidade com a pratica padrao (Nardo et al., 2008; Saisana e Tarantola, 2002), aplicamos uma normalizacao min-max ao lider: o CACI normalizado de cada pais e expresso como percentagem do CACI bruto dos Estados Unidos, de modo que CACI_norm(EUA) = 100 por construcao.",
        "Esta escolha de normalizacao e metodologicamente fundamentada em tres argumentos. Primeiro, os Estados Unidos constituem um benchmark natural dada sua posicao dominante em todos os quatro fatores do CACI: 76,9 por cento do compute IA operacional global (Epoch AI, abril de 2026), os custos energeticos industriais efetivos mais baixos entre as principais economias da OCDE (85 USD/MWh, ajustado PPA, contra 115 na Franca e 140 na Alemanha - CSV do painel, abril de 2026), o maior PIB nominal e um ecossistema de capital humano IA inigualavel. Segundo, a normalizacao ao lider e a convencao estabelecida para indices de competitividade relativa: o RCA de Balassa (1965), o GCI do WEF (Schwab, 2019) e o Competitive Industrial Performance Index da UNIDO procedem todos por razao ao lider setorial ou nacional. Terceiro, esta normalizacao elimina ambiguidades interpretativas dos valores absolutos e permite comparacoes temporais: um CACI da Franca passando de 25 para 35 entre 2024 e 2028 significa recuperacao parcial, mesmo se os valores brutos de ambos os paises tiverem aumentado.",
        f"A consequencia metodologica mais importante desta normalizacao e a revelacao de uma diferenca estrutural que os indices tradicionais ocultam. O AI Preparedness Index do FMI (Cazzaniga et al., 2024) atribui escores de 0,85 (EUA) e 0,74 (Franca), uma razao de 1,15:1, sugerindo quase paridade. O Global AI Index da Tortoise produz diferencas similares. No entanto, o CACI normalizado em Power Mode coloca a Franca em cerca de 25 e a Alemanha em cerca de 5 em uma escala onde os Estados Unidos = 100, ou seja, razoes respectivas de 4:1 e 19:1. A razao principal EUA/UE (UE agregada como 13 economias principais) e de {fmt_fr(us_eu_caci, 2)}:1. Esta divergencia nao e contraditoria: resulta do fato de que os indices multidimensionais agregam linearmente dimensoes normativas (regulacao, etica, publicacoes) que compensam aritmeticamente os deficits em fatores materiais (compute, energia). O CACI, ao isolar esses fatores materiais, expoe o abismo fisico que as metricas compostas em media apagam.",
        "Esta propriedade do CACI - revelar a verdadeira magnitude do compute gap - constitui a contribuicao analitica central do indice e a tese empirica fundamental deste estudo. A diferenca CACI nao e um artefato de normalizacao: e o resultado cientifico.",
    ]),
    ("2.4.4 Posicionamento em relacao aos indices existentes de competitividade IA", [
        "O AI Preparedness Index (AIPI) do FMI cobre 174 paises (2023) e agrega quatro pilares: infraestrutura digital, capital humano e politicas do mercado de trabalho, inovacao e integracao economica, regulacao e etica. O AIPI nao mede o compute IA instalado e nao pondera pelo custo energetico. Sua correlacao esperada com o CACI e positiva, mas imperfeita: paises bem classificados no AIPI (Singapura, Dinamarca, Paises Baixos) nao sao necessariamente aqueles com o compute efetivo por unidade de PIB mais alto.[21]",
        "O Global AI Index (GAII) da Tortoise Media classifica 83 paises em 122 indicadores em tres pilares (Implementation, Innovation, Investment). Inclui um componente de infraestrutura/supercomputacao, mas o agrega linearmente com pesos subjetivos (reconhecido pela Tortoise como uma limitacao). O GAII e mais amplo e mais multidimensional do que o CACI, mas precisamente por ser amplo, dilui o sinal do compute em um composto de muitos indicadores. Como mostram Koronakos et al. (2024), a subjetividade das ponderacoes do GAII pode inverter classificacoes de paises dependendo dos cenarios de ponderacao escolhidos.",
        "O Stanford AI Index (2025) constitui a referencia mais completa em termos de dados brutos (modelos notaveis por pais, investimento, publicacoes, patentes, tendencias de compute). Nao propoe um indice composto, mas fornece as series temporais usadas por muitos outros indices. Os dados publicos do Stanford AI Index estao disponiveis via uma pasta Google Drive aberta.[22]",
        "Valor agregado especifico do CACI. O CACI distingue-se por quatro propriedades: coloca o compute no centro em vez de na periferia do indicador, refletindo o papel agora dominante do compute na IA de fronteira; integra explicitamente o custo energetico como gargalo, consistente com as conclusoes da IEA (2025); usa uma agregacao multiplicativa teoricamente fundamentada em vez de aditiva; e e parcimonioso (quatro variaveis), tornando-o transparente e reproduzivel, ao custo de menor abrangencia.",
    ]),
    ("2.4.5 Limitacoes do CACI e analise de sensibilidade", [
        "Primeira limitacao - opacidade de F(r,t). A medicao do compute instalado depende de dados privados incompletos. O conjunto Epoch AI cobre apenas 10-20 por cento do compute global, com cobertura desigual entre setores e empresas. A atribuicao nacional e ela mesma discutivel: uma parcela crescente do compute e detida por alguns hyperscalers privados operando globalmente. Mitigacao: documentamos sistematicamente as margens de incerteza, apresentamos intervalos em vez de valores pontuais e verificamos a estabilidade dos resultados quando F varia em mais ou menos 30 por cento. A decomposicao Fisico/Soberano trata parcialmente o problema de atribuicao.",
        "Segunda limitacao - heterogeneidade qualitativa do compute. O CACI agrega FLOPs sem distinguir geracoes de GPU (um H200 nao equivale a um A100 em eficiencia energetica e desempenho real). Mitigacao: o conjunto Epoch AI fornece o desempenho em equivalentes H100, oferecendo uma normalizacao parcial. Propomos um fator de ponderacao por geracao de GPU no apendice e mostramos que o ajuste modifica marginalmente as classificacoes.",
        "Terceira limitacao - o proxy de capital humano L(r,t). A combinacao STEM mais LinkedIn mais certificacoes apresenta um vies favorecendo os paises de lingua inglesa. Este vies provavelmente subestima a China e algumas economias asiaticas. Mitigacao: replicamos a analise usando o sub-indice Human Capital do AIPI do FMI como proxy alternativo e mostramos a sensibilidade das classificacoes a esta escolha.",
        "Quarta limitacao - estaticidade do indicador. O CACI mede um estado em um momento dado, nao uma dinamica. E por isso que o calculamos em varios anos (2022, 2024, 2026) e o projetamos em cada cenario, permitindo o acompanhamento da trajetoria e a avaliacao da evolucao da diferenca.",
        "Quinta limitacao - endogeneidade. Os paises que ganham produtividade IA investem maciamente em compute, criando um risco de causalidade reversa: o CACI poderia capturar a consequencia em vez da causa da competitividade. Este estudo, em seu marco de tese, nao instrumentaliza formalmente essa relacao (sem variaveis instrumentais ou estrategia GMM). No entanto, observamos que o choque exogeno das regras BIS de outubro de 2022 oferece um quase-experimento natural que poderia fundamentar uma estrategia de identificacao causal em Difference-in-Differences. Identificamos o tratamento formal dessa endogeneidade (DiD, IV ou GMM Arellano-Bond) como uma via de pesquisa prioritaria para qualquer extensao publicavel deste trabalho.[23]",
        "Sexta limitacao - vies de normalizacao para pequenas economias. A analise comparativa do CACI face aos indices AIPI do FMI e GAII da Tortoise revela uma anomalia instrutiva: certos paises com baixo PIB e forca de trabalho IA reduzida obtem escores CACI desproporcionalmente altos. A Africa do Sul e emblematica: classificada em ultimo ou penultimo lugar nos indices AIPI e Tortoise, pode aparecer como lider global no CACI em Intensity Mode quando a normalizacao pelo PIB amplifica um efeito de pequeno denominador. Este fenomeno e um caso classico de vies de normalizacao identificado na literatura de indicadores compostos. O OECD/JRC Handbook (Nardo et al., 2008, pp. 27-29) adverte que a normalizacao por PIB ou populacao pode produzir resultados enganosos para pequenas economias. Nossa mitigacao e tripla: um limiar de massa critica em F (excluindo economias abaixo de 10 000 equivalentes H100); um fator de escala alfa(r) = min(1, F(r) / F_mediano) para testes de sensibilidade; e uma classificacao dupla (Power Mode para potencia absoluta, Intensity Mode para densidade), com a ressalva explicita de que as comparacoes bilaterais neste estudo se concentram em EUA, UE, Franca e China.[24]",
    ]),
    ("2.4.6 Numeros reproduziveis - snapshot de abril de 2026", [
        "Para permitir ao leitor refazer o calculo a partir dos CSVs do painel publico, expomos aqui as entradas e saidas utilizadas ao longo deste estudo. As quatro variaveis de entrada sao provenientes da pasta de dados do painel (energy_prices.csv, gdp_data.csv, workforce_data.csv, gpu_clusters.csv) em https://mo0ogly.github.io/America-First-IA/dashboard/data/.",
        "Entradas (snapshot de abril de 2026). EUA: F_total = 39,65 milhoes de equiv. H100, L = 3,5 milhoes, R = 1,00, E = 85 USD/MWh, PIB = 29,3 trilhoes USD. UE (13 economias principais agregadas): F_total = 2,62 milhoes de equiv. H100, L = 3,1 milhoes, R = 0,90, E = 135 USD/MWh, PIB = 18,9 trilhoes USD. Franca: F_total = 2,44 milhoes de equiv. H100, L = 1,5 milhoes, R = 0,90, E = 115 USD/MWh, PIB = 3,16 trilhoes USD. Alemanha: F_total = 0,05 milhao de equiv. H100, L = 1,9 milhoes, R = 0,90, E = 140 USD/MWh, PIB = 4,68 trilhoes USD. China: F_total = 0,40 milhao de equiv. H100, L = 4,8 milhoes, R = 0,70, E = 92 USD/MWh, PIB = 18,7 trilhoes USD.",
        f"Saidas (Power Mode F_total, normalizadas em EUA = 100). EUA = 100,0; UE(13) = {fmt_fr(caci_scores['EU'], 1)}; Franca = 25,3; India = 22,2; China = 15,7; Reino Unido = 7,0; Alemanha = 5,4. EAU = 55,7 em CACI Fisico mas apenas 6,0 em CACI Soberano (F_dom restrito a clusters detidos por entidades domesticas G42, MGX, Mubadala, Khazna, TII), ilustrando a decomposicao Fisico/Soberano: 99,6 por cento do F_total dos EAU e detido por atores US-side (Stargate UAE, Microsoft, OpenAI). Razoes principais: EUA/UE(13) = {fmt_fr(us_eu_caci, 2)}:1; EUA/Franca = 3,96:1; EUA/Alemanha = 18,59:1; EUA/China = 7,46:1.",
        f"Estas saidas sao reproduzidas ao vivo no painel publico. Qualquer atualizacao subsequente do painel (novos clusters, revisoes energeticas, atualizacao do PIB) atualizara mecanicamente os numeros; a metodologia nao muda. A razao de {fmt_fr(us_eu_caci, 2)}:1 em Power Mode F_total reportada na capa e ao longo deste estudo corresponde exatamente a este snapshot.",
    ]),
    ("2.5 Escopo e delimitacoes", [
        "Escopo geografico. A analise concentra-se na relacao bilateral Estados Unidos / Uniao Europeia, com foco especifico na Franca. A China e tratada como variavel contextual (alvo principal dos controles de exportacao americanos, fator de pressao sobre as capacidades de producao de chips), mas nao e objeto de analise aprofundada. Japao, Coreia do Sul e Taiwan aparecem como atores da cadeia de suprimentos de semicondutores.",
        "Escopo temporal. O diagnostico cobre 2020-2026, os cenarios cobrem 2026-2030. O horizonte 2030 e escolhido porque corresponde a convergencia de varios prazos: projecoes da AIE para a energia dos data centers, maturidade esperada do EU Chips Act, objetivos da SNIA Franca 2030 e chegada potencial dos primeiros SMR nucleares operacionais.",
        "Escopo tecnologico. O estudo cobre a IA de fronteira (modelos fundacionais, intensivos em compute) e seus pre-requisitos materiais (GPU/ASIC, data centers, energia). Integra a robotica IA como fator amplificador da demanda energetica. Nao trata da IA embarcada edge (smartphones, IoT), exceto na medida em que constitui um objetivo especifico da SNIA francesa.",
    ]),
    ("2.6 Limitacoes metodologicas gerais", [
        "Incerteza politica radical. O protecionismo tecnologico depende de decisoes politicas discricionarias com previsibilidade estruturalmente baixa. Uma mudanca de administracao americana em 2028, um acordo comercial UE-EUA inesperado ou uma escalada do conflito EUA-China poderia invalidar certos pressupostos. E precisamente por isso que propomos quatro cenarios em vez de uma trajetoria unica.",
        "Disrupcoes tecnologicas. O episodio DeepSeek (janeiro de 2025), em que um modelo chines alcancou desempenho proximo da fronteira com orcamento de treinamento substancialmente reduzido, ilustra a possibilidade de avancos de eficiencia que alterariam os termos do problema. A IEA (2025, Energy and AI) dedica um estudo de caso ao DeepSeek e conclui que mesmo com melhorias significativas de eficiencia, o crescimento da demanda absorve os ganhos (efeito rebote de Jevons).[25]",
        "Opacidade dos dados de compute. O numero exato de GPUs implantadas por hyperscaler, a distribuicao geografica precisa dos data centers e os volumes de GPUs exportados por regiao sao dados parcial ou totalmente confidenciais. Nossas estimativas de compute instalado carregam margens de erro significativas, que documentamos sistematicamente.",
        "Vies das fontes de consultoria. Como observado na secao 2.2, as fontes industriais carregam um vies sistematico de otimismo. Mitigamos isso por meio da triangulacao, mas nao podemos elimina-lo inteiramente.",
        "Estas limitacoes nao comprometem a validade da analise. O metodo dos cenarios e precisamente concebido para funcionar em ambientes de alta incerteza, onde o objetivo nao e a previsao, mas a exploracao estruturada das possibilidades. Como observa Schwartz, os cenarios nao sao previsoes; sao historias plausiveis que ajudam a pensar.[26] Nossa contribuicao reside no rigor do enquadramento, na explicitacao dos pressupostos, na transparencia das fontes de dados e na originalidade do indicador CACI, em vez de na precisao das projecoes numericas.",
    ]),
]

PT.notes = EN.notes  # bibliographic references identical


# ===========================================================================
# Main
# ===========================================================================

def main() -> None:
    """Build the three Chapter II .docx files."""
    out_dir = Path(__file__).parent
    for lp in (EN, FR, PT):
        build(lp, out_dir)
    log.info("Done.")


if __name__ == "__main__":
    main()
