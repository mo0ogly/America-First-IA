"""

import sys
from pathlib import Path
sys.path.append(str(Path(__file__).resolve().parents[1]))
import caci_data_helper
m = caci_data_helper.get_metrics()
us_share = m['us_compute_share']
us_eu_caci = m['us_eu_caci_power_ratio']
us_eu_raw = m['us_eu_raw_ratio']
eu_sov = m['eu_sovereignty_ratio']
caci_scores = {k: v['caci_power_phys'] for k, v in m['country_results'].items()}

def fmt_fr(val, decimals=1):
    return f"{val:.{decimals}f}".replace(".", ",")

def fmt_en(val, decimals=1):
    return f"{val:.{decimals}f}"

Chapter V - Prospective Scenarios 2026-2030 - trilingual generator.

Generates the Chapter V .docx for the doctoral study
"AI for Americans First" in English, French and Brazilian Portuguese.

All metrics are aligned with the April 2026 dashboard snapshot:
    - Cover banner: {fmt_en(us_share, 1)}% / 1.59x / {fmt_en(us_eu_caci, 2)}:1
    - CACI Power Mode: {fmt_en(us_eu_caci, 2)}:1
    - Scenario M1 (compute) and M6 (CACI) rebased on the live dataset.

Author: Fabrice Pizzi (Universite Paris-Sorbonne, M2 Economic Intelligence).
Build: python3 generate_chapter5_trilingual.py
"""

from __future__ import annotations

import sys
from pathlib import Path
sys.path.append(str(Path(__file__).resolve().parents[1]))
import caci_data_helper
m = caci_data_helper.get_metrics()
us_share = m['us_compute_share']
us_eu_caci = m['us_eu_caci_power_ratio']
us_eu_raw = m['us_eu_raw_ratio']
eu_sov = m['eu_sovereignty_ratio']
caci_scores = {k: v['caci_power_phys'] for k, v in m['country_results'].items()}

def fmt_fr(val, decimals=1):
    return f"{val:.{decimals}f}".replace(".", ",")

def fmt_en(val, decimals=1):
    return f"{val:.{decimals}f}"


import logging
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
log = logging.getLogger("chapter5_gen")


# ---------------------------------------------------------------------------
# Visual identity (shared with Chapters I, II, III, IV)
# ---------------------------------------------------------------------------

NAVY = RGBColor(0x1A, 0x27, 0x44)
GOLD = RGBColor(0xB8, 0x92, 0x2F)
GREY = RGBColor(0x55, 0x55, 0x55)


@dataclass
class LangPack:
    """Container for one language version of Chapter V."""

    code: str
    filename: str
    cover_subtitle: str
    cover_title: str
    cover_blurb: str
    cover_chip_lines: list[str]
    cover_meta: str
    cover_keywords_label: str
    cover_keywords: str
    chapter_label: str
    chapter_title: str
    chapter_intro: str
    sections: list[tuple[str, list[str]]] = field(default_factory=list)
    table_blocks: list[tuple[str, str, list[list[str]]]] = field(default_factory=list)
    notes_label: str = "Notes"
    notes: list[str] = field(default_factory=list)
    license_block: list[str] = field(default_factory=list)
    page_footer: str = ""


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def set_run(run, *, font="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text, *, style=None, align=None, space_after=6, **run_kwargs):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run(run, **run_kwargs)
    return p


def add_heading(doc, text, level):
    sizes = {1: 22, 2: 16, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, font="Calibri", size=sizes.get(level, 11),
            bold=True, color=NAVY)
    return p


def add_cover(doc, lp: LangPack):
    banner = {
        "EN": "RESEARCH STUDY - FEBRUARY 2026",
        "FR": "ETUDE DE RECHERCHE - FEVRIER 2026",
        "PT-BR": "ESTUDO DE PESQUISA - FEVEREIRO DE 2026",
    }.get(lp.code, "RESEARCH STUDY - FEBRUARY 2026")
    add_paragraph(doc, "", space_after=0)
    add_paragraph(doc, banner,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, bold=True, color=GREY, space_after=4)
    add_paragraph(doc, lp.cover_title,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=26, bold=True, color=NAVY, space_after=4)
    add_paragraph(doc, lp.cover_subtitle,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, italic=True, color=GREY, space_after=4)
    add_paragraph(doc, lp.cover_blurb,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=11, italic=True, color=GREY, space_after=18)

    table = doc.add_table(rows=1, cols=3)
    for i, line in enumerate(lp.cover_chip_lines):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        set_run(run, size=11, bold=True, color=NAVY)

    add_paragraph(doc, "", space_after=8)
    add_paragraph(doc, "Fabrice Pizzi",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, bold=True, color=NAVY, space_after=2)
    add_paragraph(doc, "Universite Paris-Sorbonne",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, color=GREY, space_after=2)
    master_line = {
        "EN": "Master 2 Economic Intelligence - Intelligence Warfare",
        "FR": "Master 2 Intelligence Economique - Intelligence Warfare",
        "PT-BR": "Mestrado 2 Inteligencia Economica - Intelligence Warfare",
    }.get(lp.code, "Master 2 Economic Intelligence - Intelligence Warfare")
    add_paragraph(doc, master_line,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, italic=True, color=GREY, space_after=14)
    add_paragraph(doc, lp.cover_meta,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, italic=True, color=GREY, space_after=10)
    add_paragraph(doc, f"{lp.cover_keywords_label}: {lp.cover_keywords}",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=9, italic=True, color=GREY, space_after=24)


def add_chapter_header(doc, lp: LangPack):
    add_paragraph(doc, lp.chapter_label,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=11, bold=True, color=GOLD, space_after=2)
    add_paragraph(doc, lp.chapter_title,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=18, bold=True, color=NAVY, space_after=12)
    if lp.chapter_intro:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(lp.chapter_intro)
        set_run(run, size=11, italic=True, color=GREY)


def render_sections(doc, lp: LangPack):
    for heading, paragraphs in lp.sections:
        # Determine level based on numbering
        first = heading.split()[0]
        level = 2 if first.count(".") == 1 else 3
        add_heading(doc, heading, level)
        for para in paragraphs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.25
            run = p.add_run(para)
            set_run(run, size=11, color=RGBColor(0x20, 0x20, 0x20))


def render_tables(doc, lp: LangPack):
    for caption, source, rows in lp.table_blocks:
        add_paragraph(doc, caption,
                      size=10, bold=True, color=NAVY, space_after=2)
        add_paragraph(doc, source,
                      size=9, italic=True, color=GREY, space_after=4)
        n_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=n_cols)
        table.style = "Light Grid Accent 1"
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                cell = table.rows[i].cells[j]
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(cell_text)
                set_run(run, size=10,
                        bold=(i == 0),
                        color=NAVY if i == 0 else RGBColor(0x20, 0x20, 0x20))
        doc.add_paragraph()


def render_notes(doc, lp: LangPack):
    add_heading(doc, lp.notes_label, 2)
    for i, note in enumerate(lp.notes, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run_id = p.add_run(f"{i}. ")
        set_run(run_id, size=9, bold=True, color=GOLD)
        run_txt = p.add_run(note)
        set_run(run_txt, size=9, color=GREY)


def render_license(doc, lp: LangPack):
    doc.add_paragraph()
    for line in lp.license_block:
        add_paragraph(doc, line,
                      align=WD_ALIGN_PARAGRAPH.LEFT,
                      size=8, italic=True, color=GREY, space_after=2)
    add_paragraph(doc, lp.page_footer,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=8, italic=True, color=GREY, space_after=0)


def build(lp: LangPack, out_dir: Path) -> Path:
    log.info("Building Chapter V [%s] -> %s", lp.code, lp.filename)
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    add_cover(doc, lp)
    add_chapter_header(doc, lp)
    render_sections(doc, lp)
    render_tables(doc, lp)
    render_notes(doc, lp)
    render_license(doc, lp)

    out = out_dir / lp.filename
    doc.save(out)
    log.info("Saved %s", out)
    return out


# ===========================================================================
# Content - English
# ===========================================================================

EN = LangPack(
    code="EN",
    filename="Chapter_V_Prospective_Scenarios_2026_2030_EN.docx",
    cover_title="AI FOR AMERICANS FIRST",
    cover_subtitle="AI Protectionism, Energy and Semiconductors: US/Europe Divergence Trajectories 2024-2030",
    cover_blurb="Integrated Geostrategic and Economic Analysis - Chapter V",
    cover_chip_lines=[
        f"{fmt_en(us_share, 1)}% global operational AI compute = USA",
        "1.59x energy cost EU/US (PPA-adjusted)",
        f"{fmt_en(us_eu_caci, 2)}:1 CACI ratio US/EU (Power Mode)",
    ],
    cover_meta="Paris - February 2026  |  7 chapters  |  4 prospective scenarios  |  3 geographic zones",
    cover_keywords_label="Keywords",
    cover_keywords=("artificial intelligence, technology protectionism, semiconductors, "
                    "export controls, sovereign compute, AI geopolitics, France, "
                    "United States, China"),
    chapter_label="CHAPTER V",
    chapter_title="Prospective Scenarios 2026-2030",
    chapter_intro=(
        "This chapter constitutes the heart of the original contribution of this study. "
        "By applying the methodological protocol defined in Chapter II (2x2 matrix, "
        "six divergence metrics, CACI calibration), we construct four scenarios for the "
        "evolution of the transatlantic relationship in AI, energy, and semiconductors "
        "for the 2026-2030 period. Each scenario is determined by the combination of two "
        "critical uncertainties identified in Chapter III: the degree of US protectionism "
        "and the European strategic response capacity. We then evaluate each scenario "
        "on its six metrics, before synthesising the tipping point conditions between trajectories."
    ),
    notes_label="Notes",
    license_block=[
        "License and Disclaimer. This work, 'AI for Americans First,' is made available under the terms of the Creative Commons Attribution - NonCommercial - ShareAlike 4.0 International License (CC BY-NC-SA 4.0).",
        "You are free to share and adapt the material for non-commercial purposes, provided you give appropriate credit to Fabrice Pizzi (Universite Paris-Sorbonne) and distribute your contributions under the same license. This document is provided for educational and research purposes only.",
        "Public dashboard: https://mo0ogly.github.io/America-First-IA/dashboard/",
        "Repository: https://github.com/mo0ogly/America-First-IA",
    ],
    page_footer="AI for Americans First - Fabrice Pizzi - Chapter V",
)

EN.sections = [
    ("5.1 Predetermined elements: what will not change", [
        "In accordance with the Schwartz method (1991), we distinguish predetermined elements (quasi-certain trends by the 2030 horizon) from critical uncertainties (factors whose evolution depends on political decisions not yet taken). Four predetermined elements structure all scenarios.",
        "EP1 - Exponential growth in AI compute demand. Semiconductor sales doubled in two years (2023-2025), the power of installed AI chips doubles every seven months (Epoch AI), and no sign of slowing down is observable as of February 2026. Even under the hypothesis of a deceleration in scaling laws (Chinchilla saturation), the diffusion of AI towards inference, robotics, and autonomous agents will maintain a strongly growing compute demand.[1]",
        f"EP2 - Persistent concentration of compute in the United States. The US/EU ratio of {fmt_en(us_eu_raw, 1)}:1 in raw installed compute (Chapter III, April 2026 dashboard snapshot: 2,759,968 vs 156,632 PFLOP/s), resulting in a CACI Power Mode ratio of {fmt_en(us_eu_caci, 2)}:1 once weighted by the geometric formula F^0.40 x L^0.20 x R^0.15 / E^0.25, reflects investment decisions taken in 2022-2025 whose effects materialise until 2028-2029 (data center construction delays: 18-36 months). Even an immediate political reversal would not alter the installed stock before the end of the decade.",
        "EP3 - Growing energy tension. Global data center consumption, estimated at 415 TWh in 2024, will reach 800-950 TWh by 2030 according to IEA projections (Chapter III). The asymmetry in energy costs (US 1.4-1.7x cheaper after PPA correction, 2-3x cheaper on unadjusted industrial Eurostat tariffs) will persist, unless massive investment in European nuclear power, whose deployment delays (SMR: 5-7 years for the first reactors) exceed the 2030 horizon.[2]",
        "EP4 - Section 232 regulatory framework in place. Proclamation 11002 of January 14, 2026, is a legal fait accompli. Unlike IEEPA tariffs (annulled by the Supreme Court on February 20, 2026[3]), Section 232 tariffs rest on a confirmed legal basis. The Secretary of Commerce's report on the data center semiconductor market is expected by July 1, 2026, and may recommend an extension or modification of tariffs. Regardless of the direction taken, the legal instrument will remain available.[4]",
    ]),
    ("5.2 Critical uncertainties and 2x2 matrix", []),
    ("5.2.1 Axis 1: Degree of US protectionism", [
        "The first uncertainty concerns the evolution of American policy between two poles. The moderate pole corresponds to maintaining the January 2026 status quo: a 25 percent tariff limited to re-exported advanced chips, broad domestic exemptions, a US-EU trade agreement capping semiconductor tariffs at 15 percent, and no significant extension to cloud or models. The EU core (France, Germany) remains in the 'trusted partners' category. The aggressive pole assumes an extension after the July 2026 report: tariffs extended to derived semiconductors and equipment, GPU quotas for the EU (including France), restrictive conditions for access to frontier AI cloud, and the use of compute as a trade negotiation lever (compute-for-concessions).[5]",
    ]),
    ("5.2.2 Axis 2: European strategic response capacity", [
        "The second uncertainty concerns the EU's ability to deploy a coherent and rapid response. The reactive pole corresponds to fragmented national responses, dispersed investments, an AI Act creating additional compliance costs, and slow deployment of AI Factories/Gigafactories (bureaucratic delays, 24+ month authorisations). The proactive pole assumes the accelerated implementation of the AI Continent program (19 AI Factories plus up to 5 Gigafactories of 100,000+ GPUs), the adoption of Special Compute Zones (180-day authorisation), the effective mobilisation of the InvestAI fund (20 billion EUR), and the pooling of French nuclear capacity as a competitive advantage.[6]",
    ]),
    ("5.2.3 Matrix and scenario naming", [
        "The intersection of these two axes produces four scenarios. Reactive EU + moderate US protectionism gives Scenario A (Reinforced Status Quo, slow drift towards dependence). Reactive EU + aggressive US gives Scenario B (Digital Fracture, structural European decoupling). Proactive EU + moderate US gives Scenario C (Asymmetric Partnership, junior Western technological partner). Proactive EU + aggressive US gives Scenario D (Contested Sovereignty, race for autonomy under pressure).",
    ]),
    ("5.3 Scenario A - Reinforced Status Quo (Moderate protectionism + Reactive EU)", []),
    ("5.3.1 Narrative", [
        "Following the July 2026 report, the Secretary of Commerce recommends maintaining the 25 percent tariff on re-exported advanced chips but without extending it significantly. The August 2025 US-EU trade agreement is respected: semiconductor tariffs for the EU remain capped at 15 percent.[7] The EU, reassured by this status quo, slows the deployment of its own initiatives. EuroHPC AI Factories struggle to reach nominal capacity (authorisation delays, inter-state coordination). Gigafactories are postponed to 2029-2030. The InvestAI fund is partially mobilised (8-10 billion EUR out of 20). European companies continue to rely heavily on US cloud, whose performance and costs remain unbeatable.",
    ]),
    ("5.3.2 Metric trajectories", [
        f"M1 - Compute ratio (US/EU installed GPUs): goes from {fmt_en(us_eu_raw, 1)}:1 raw (2025) to 18-22:1 raw (2030) on operational installed compute. The gap widens slightly as US investments accelerate (Stargate, xAI mega-clusters, Meta) while the EU only adds the 19 AI Factories (max 25,000 GPUs each, approx. 475,000 public GPUs, an order of magnitude below a single US hyperscaler).[8]",
        "M2 - FLOP cost gap (EU/US): remains in the 2.4-3.2x range. The absence of aggressive tariffs on the EU maintains access to US cloud at prices close to current levels, but European energy costs continue to weigh in.",
        "M3 - US cloud share in European AI spending: goes from 70 percent (2024) to 72-75 percent (2030). European providers (OVHcloud, Deutsche Telekom) maintain their 15 percent in the sovereignty segment but do not gain ground on generative AI services.",
        "M4 - AI productivity (percent/year): US +2.5-3.0; EU +1.0-1.5. The EU realises part of the AI potential via downstream applications (SAP, Siemens, fintech), but slow adoption and compute deficit cap the gains.",
        "M5 - Energy dependence (data center TWh): EU approx. 115 TWh in 2030 (+65 percent vs 2024). French nuclear power absorbs part of the demand, but the absence of Special Compute Zones delays the network connection of new data centers.",
        f"M6 - CACI(US)/CACI(EU): goes from {fmt_en(us_eu_caci, 2)}:1 (April 2026) to 4-5:1 (2030). The gap widens moderately as the F factor (compute) accumulates on the US side while E costs (energy) weigh on the EU side.",
    ]),
    ("5.3.3 Consequences for France", [
        "This scenario is the most probable in the short term (estimated probability: 40-50 percent). It is also the most insidious: the absence of a visible shock demobilises European actors, while dependence deepens structurally. French companies benefit from access to US cloud for AI adoption (BNP Paribas, Airbus, TotalEnergies via AWS/Azure), but this adoption reinforces the lock-in described in Chapter IV. The AI productivity deficit relative to the United States (-1.0 to -1.5 points per year) accumulates over five years, widening the competitiveness gap by 5 to 8 GDP points.",
    ]),
    ("5.4 Scenario B - Digital Fracture (Aggressive protectionism + Reactive EU)", []),
    ("5.4.1 Narrative", [
        "The July 2026 report leads to a significant extension. The Secretary of Commerce recommends tariffs extended to semiconductor equipment and derived products, with a tariff offset program reserved for companies investing in American production.[9] The 15 percent EU agreement is revised upwards or accompanied by restrictive conditions (volume quotas on advanced GPUs, reciprocity requirements on the AI Act). Simultaneously, access to frontier AI cloud is made conditional for non-US entities (API access limitations to frontier models, weight restrictions). The EU, fragmented, fails to formulate a coherent response: member states divide between accommodation (Nordic countries, Netherlands) and confrontation (France, Italy).",
    ]),
    ("5.4.2 Metric trajectories", [
        f"M1 - Compute ratio: goes from {fmt_en(us_eu_raw, 1)}:1 raw (2025) to 25-35:1 raw (2030). GPU quotas limit European imports at the moment demand explodes. AI Factory projects are compromised by the inability to procure Nvidia/AMD GPUs at the planned volumes.",
        "M2 - FLOP cost gap: jumps to 4-6x. Extended tariffs, combined with quotas and energy asymmetry, massively increase European compute costs. French companies face a 3x to 5x surcharge for model training.",
        "M3 - US cloud share: paradoxically rises to 78-82 percent. Failing a credible local alternative, European companies wanting access to frontier AI must go through US hyperscalers, at the price conditions they dictate. Sovereign services (OVHcloud, Scaleway) lack the hardware to offer competitive GenAI services.",
        "M4 - AI productivity: US +2.5-3.5; EU +0.3-0.8. European AI potential is severely constrained. The McKinsey Global Institute estimates that with slow adoption, European productivity would not exceed 0.3 percent, close to stagnation.[10]",
        "M5 - Energy dependence: EU approx. 95 TWh only (2030), not by virtue but by default - the lack of GPUs limits data center construction. Ironically, the compute constraint mitigates the energy constraint.",
        f"M6 - CACI ratio: explodes from {fmt_en(us_eu_caci, 2)}:1 (April 2026) to 6-8:1 (2030). This is the scenario where the gap is largest, with all three CACI factors deteriorating simultaneously on the European side: F capped by quotas, E inflated by tariffs, L weakened by an accelerated brain drain to the United States.",
    ]),
    ("5.4.3 Consequences for France", [
        "This scenario (estimated probability: 15-20 percent) represents the worst-case. France suffers structural technological decoupling: compute-intensive projects (Mistral foundation models, Exotec robotics, Dassault simulations) are relocated to the United States or depend on increasingly costly US cloud access. Time-to-market for French AI solutions lengthens by 25 to 40 percent. SMEs, unable to absorb surcharges, renounce frontier AI and opt for degraded solutions (smaller open-source models, local inference). The cumulative productivity gap with the United States reaches 10 to 15 points over five years.",
    ]),
    ("5.5 Scenario C - Asymmetric Partnership (Moderate protectionism + Proactive EU)", []),
    ("5.5.1 Narrative", [
        "US protectionism remains moderate (as in A), but the EU exploits this window to accelerate its own investments. AI Factories are deployed on schedule (2026-2027), the first 100,000+ GPU Gigafactories are ordered end-2026 and delivered in 2028.[11] France plays a central role thanks to its nuclear fleet (65-70 percent of electricity mix, competitive marginal cost), and Special Compute Zones are designated on former industrial sites with heavy network connections.[12] However, the EU de facto accepts a status as a junior technological partner: it uses Nvidia/AMD GPUs (no European AI ASIC design champion), depends on TSMC/Samsung/Intel foundries for production, and its foundation models remain one step below US leaders.",
    ]),
    ("5.5.2 Metric trajectories", [
        f"M1 - Compute ratio: drops from {fmt_en(us_eu_raw, 1)}:1 raw (2025) to 8-10:1 raw (2030) on installed compute. Gigafactories and private investment (InvestAI plus industrial co-investments) add 1-2 million H100 equivalents in Europe, narrowing the gap without closing it.",
        "M2 - FLOP cost gap: drops to 1.5-2.0x. French nuclear power and Gigafactory economies of scale compress energy and infrastructure costs, although a residual gap persists (absence of proprietary GPU design).",
        "M3 - US cloud share: drops slightly to 60-65 percent. European sovereign services gain market share in regulated segments (defense, health, finance), while US cloud retains the majority of commercial workloads. The market segments into sovereign and performance.",
        "M4 - AI productivity: US +2.5-3.0; EU +1.8-2.5. The EU reaches 60-80 percent of its theoretical potential thanks to sufficient local compute for large-scale adoption of downstream applications, even if frontier model training remains dependent on US hardware.",
        "M5 - Energy: EU approx. 140 TWh (2030). Demand is higher than in A because European compute increases, but nuclear and planned SMRs absorb most of it. RTE France confirms the feasibility of +10 GW subject to network investments.",
        f"M6 - CACI ratio: drops from {fmt_en(us_eu_caci, 2)}:1 (April 2026) to 2.0-2.5:1 (2030). This is the most favourable realistically achievable scenario at the 2030 horizon. The F factor improves significantly, E benefits from nuclear, but L remains slightly lower (the US AI ecosystem being more attractive for top talent).",
    ]),
    ("5.5.3 Consequences for France", [
        "This scenario (estimated probability: 15-20 percent) is the most favourable for France in the short-to-medium term. France becomes the EU's AI energy hub thanks to its nuclear fleet, attracting data center and Gigafactory investments. French companies gain access to competitive local compute for inference and fine-tuning, reducing dependence on US cloud for standard uses. Mistral and French startups can train specialised models locally. However, frontier model training remains dependent on US hardware, and strategic autonomy is partial: France is sovereign in application, but not in the creation of fundamental technologies.",
    ]),
    ("5.6 Scenario D - Contested Sovereignty (Aggressive protectionism + Proactive EU)", []),
    ("5.6.1 Narrative", [
        "US protectionism intensifies (as in B), but the EU responds with determination. The American threat becomes the political catalyst for an unprecedented European industrial mobilisation since the AIRBUS project of the 1970s. The AI Continent program is accelerated and extended: the 5 Gigafactories are ordered as an emergency, France announces 20 GW of nuclear capacity dedicated to AI data centers by 2032 (combining fleet extension, new EPR2s and SMRs), the DARE project (European RISC-V) is escalated to design AI accelerators reducing dependence on Nvidia.[13] Simultaneously, the EU negotiates alternative technological alliances (Japan, South Korea, Taiwan) to secure GPU and foundry supplies.",
    ]),
    ("5.6.2 Metric trajectories", [
        f"M1 - Compute ratio: evolves from {fmt_en(us_eu_raw, 1)}:1 raw (2025) to 12-15:1 raw (2030). The EU invests massively but starts from very far. US quotas slow imports, but alternative alliances and local production (Gigafactories using Samsung/Intel GPUs as alternatives to Nvidia) partially compensate.",
        "M2 - FLOP cost gap: 2.5-4.0x initially (2027, peak tariff shock), then progressive reduction towards 1.8-2.5x (2030) as Gigafactories scale up and GPU alternatives mature.",
        "M3 - US cloud share: drops to 50-55 percent (2030), the most pronounced decline of the four scenarios. Geopolitical defiance and US restrictions push European companies towards sovereign alternatives, even if imperfect. US hyperscalers lose ground in regulated segments.",
        "M4 - AI productivity: US +2.5-3.5; EU +1.2-2.0. The EU traverses a productivity trough in 2027-2028 (transition period where US restrictions bite but European investments are not yet operational), then a partial catch-up from 2029.",
        "M5 - Energy: EU approx. 150-160 TWh (2030). This is the most energy-intensive scenario for the EU, with massive local data center construction creating huge demand. French nuclear power becomes a continental strategic asset, but network pressure is at its maximum.",
        f"M6 - CACI ratio: follows a U-shaped trajectory: degradation from {fmt_en(us_eu_caci, 2)}:1 to 8-12:1 in 2027-2028 (peak shock), then improvement to 4-7:1 by 2030. The result depends heavily on European execution speed: each year of delay on Gigafactories prolongs the period of maximum vulnerability.",
    ]),
    ("5.6.3 Consequences for France", [
        "This scenario (estimated probability: 15-20 percent) is the most ambitious and riskiest. It places France at the heart of an unprecedented European technological sovereignty effort. Massive nuclear investments (SMR, fleet extension) become a top-tier geopolitical issue. The DARE/RISC-V project could, if successful, constitute the first credible European alternative to Nvidia GPUs for AI, but on a 5-7 year horizon, well beyond 2030. In the short term (2026-2028), France traverses a period of maximum vulnerability where surcharges and shortages degrade competitiveness, before a catch-up conditional on infrastructure deployment speed.",
    ]),
    ("5.7 Comparative synthesis and tipping points", []),
    ("5.7.1 Metric synthesis table", [
        f"Table 11 below consolidates the trajectory of the six divergence metrics for the four scenarios by the 2030 horizon, anchored on the April 2026 dashboard snapshot (US/EU raw operational compute {fmt_en(us_eu_raw, 1)}:1, CACI Power Mode {fmt_en(us_eu_caci, 2)}:1).",
    ]),
    ("5.7.2 Tipping points between scenarios", [
        "The real trajectory will likely follow a hybrid path between these scenarios. Three tipping points determine possible transitions.",
        "First tipping point: the July 2026 Commerce report. This report will determine whether US protectionism extends (shift to B or D) or remains targeted (maintenance in A or C). Indicators to watch: evolution of the US semiconductor trade deficit, CHIPS Act fab filling rates (Intel, TSMC Arizona, Samsung Taylor), domestic political pressure (2026 midterms). Phase 1 negotiation results (report due April 14, 2026) will be an early signal.[14]",
        "Second tipping point: EU Gigafactory deployment speed. The Commission expects the first Gigafactories to be operational in 2027-2028. If this schedule is held, the EU shifts towards proactive scenarios (C or D). If authorisation, financing, or hardware supply delays push deliveries to 2029-2030, the EU remains in reactive mode (A or B). The CFG proposal for Special Compute Zones (180-day authorisation vs 24+ months currently) is the key acceleration factor.[15]",
        "Third tipping point: the French decision on nuclear for AI. France possesses a unique asset in Europe: a nuclear fleet providing 65-70 percent of electricity at a globally competitive marginal cost. The decision to dedicate significant capacity (10-20 GW) to AI data centers, via fleet extension, new EPR2s and SMRs, will determine whether France becomes Europe's AI energy hub or loses this position to others (Scandinavia with hydro, Eastern Europe with low land costs). This tipping point is specifically French and determines France's position within the European scenarios.[16]",
    ]),
    ("5.7.3 The convergence point: 2028", [
        "The four scenarios converge on a common critical point in 2028. This is the year when: (i) compute demand will exceed installed capacity in Europe, creating hardware bottlenecks (even under moderate protectionism); (ii) the full effects of extended tariffs (if adopted) will be felt; (iii) Gigafactories, if deployed on time, will start producing significant local compute; (iv) data center energy demand will saturate network connection capacity in several member states. The year 2028 constitutes the moment of truth when Europe will discover if it is on an A/B trajectory (growing dependence) or C/D (catch-up started). Decisions taken in 2026-2027 (US Commerce report, Gigafactories, French nuclear) will be irreversibly engaged.",
    ]),
    ("5.8 Origins of the tipping point: juridico-technical foundations", [
        "The 'Great Decoupling' scenario is not built ex nihilo. It is the logical projection, at the 2028 horizon, of a control architecture whose foundational layers are already operational in 2026. Identifying these layers is a matter of academic rigour: distinguishing documented trends from extrapolated projections.",
    ]),
    ("5.8.1 The legal layer: extraterritoriality as a structural instrument", [
        "The first layer is legal and largely precedes AI policy. The CLOUD Act (Clarifying Lawful Overseas Use of Data Act, 2018) establishes that any cloud provider subject to US jurisdiction is required to produce data regardless of where such communications, records, or other information are stored.[17] This law, confirmed by federal case law (United States v. Microsoft), creates a fundamental dissociation between the physical location of data and its legal nationality.",
        "The immediate consequence for AI compute is radical: an H100 cluster physically located in Dubai, Singapore, or an AWS eu-west-1 center in Ireland remains legally American. If ordered by the US government, the operator (AWS, Azure, Google Cloud) is legally bound to comply, regardless of the client's will or local legislation. Microsoft acknowledged before a French court in 2024 that it could not guarantee data sovereignty for European clients in the event of a legally grounded US injunction.[18]",
        "This pre-existing legal architecture is the substrate on which the Cloud-Nationality Pivot grafts: the Cloud Sovereignty Mandates projected for 2028 do not create a new power. They activate and systematise an existing jurisdictional power by extending it to the operational compute layer.",
    ]),
    ("5.8.2 The technical layer: from location verification to cluster throttling", [
        "The second layer is technical. The US AI Action Plan of July 23, 2025, explicitly introduces the concept of 'location verification features' applied to advanced AI chips. Michael Kratsios, OSTP director at the White House, confirmed discussions on software or hardware modifications to enable better location-tracking, explicitly included in the plan.[19]",
        "This announcement is not rhetorical. The BIS already has, since the Framework for Artificial Intelligence Diffusion (January 2025), a compute-tiering control architecture for destination countries (Tier 1-2-3) and compute caps per entity and country.[20] The BIS Affiliates Rule, suspended for one year in November 2025 but maintained in principle, stipulates that affiliation to a parent entity in a restricted country is sufficient to deny access to advanced compute, regardless of physical location.[21]",
        "The technological trajectory towards 2028 is therefore: (i) chips equipped with location verification mechanisms, (ii) automatic reporting to the BIS in case of deviation, (iii) the ability to suspend access or throttle performance via export licence. This is not science fiction: it is the extension to compute of a principle already applied to software (OFAC sanctions, service access freeze).",
        "Critical limitation: Throttling production clusters is technically complex and potentially disruptive for operators themselves. NVIDIA does not currently have a remote deactivation mechanism for H100/B200 GPUs in production. Such a mechanism would require significant architectural modifications to firmware and remote attestation protocols (via TPM or equivalent). The 2028 scenario is therefore conditional on the effective implementation of these modifications over 24-36 months.",
    ]),
    ("5.9 The Cloud-Nationality Pivot mechanism", []),
    ("5.9.1 The trigger: Cloud Sovereignty Mandates as extension of export controls", [
        "The 2028 tipping point scenario assumes that the United States crosses a qualitative threshold: moving from hardware access control (BIS chip controls) to operational compute service access control (Cloud Sovereignty Mandates). This threshold is not an arbitrary rupture; it responds to a structural flaw identified in the current control regime.",
        "This flaw is documented: despite BIS restrictions, investigations revealed that approx. 1 billion USD of Nvidia chips reached China via third countries (Malaysia, UAE, Singapore) in the first months of 2025.[22] The AI Action Plan's response (location verification and cluster monitoring) is the first step towards continuous post-export control.",
        "The plausible trigger in 2028 is an executive order extending Framework for AI Diffusion obligations to the cloud layer. It would impose a 'Data Residency and Jurisdiction Compliance' certification on all US hyperscalers operating advanced offshore clusters, with revocation of compute access on US soil as the enforcement mechanism.",
    ]),
    ("5.9.2 Dissociating Physical Factor / Sovereign Factor in CACI", [
        "This is where the scenario produces its most significant impact on the CACI model. Under the Cloud-Nationality Pivot, the compute variable F(r) decomposes into two components: F(r) = F_phys(r) x F_sov(r), where F_phys is physically installed compute and F_sov is the operational sovereignty factor (fraction of F_phys outside US jurisdiction).",
        "Note the distinction between this dynamic 2028 F_sov and the static Sovereign CACI introduced in Chapter I (Fig 1.8). The Chapter I Sovereign CACI captures who owns the compute today (April 2026 snapshot). The 2028 F_sov captures who controls compute under a hypothetical Mandate regime, which depends on cloud workload share, not installed capacity. The two metrics agree where compute is owned by domestic operators; they diverge where local clusters are owned by US-side operators (UAE 99.6%, EU cloud workloads majorly on AWS/Azure/GCP).",
        "Calibrated estimations for 2028, under the Mandate activation hypothesis, are presented in Table 12.",
    ]),
    ("5.10 Emergence of jurisdictional AI blocs (2028-2030)", [
        "The Great Decoupling does not produce a binary US/non-US world. It produces fragmentation into blocs of variable intensity, according to each zone's capacity to develop credible sovereign compute. Four blocs emerge.",
    ]),
    ("5.10.1 The Extended American Bloc (American AI Alliance)", [
        "The July 23, 2025 AI Action Plan explicitly lays the foundations for an American AI Alliance: exporting the full US tech stack (hardware, models, software, standards) to aligned allies in exchange for aligned export controls.[23] The strategy is 'carrot and stick': aligned allies access frontier chips and models without additional restrictions; those who refuse are exposed to Foreign Direct Product Rule mechanisms and secondary tariffs.",
        "Confirmed Tier 1 members: US, UK, Canada, Australia, Japan, South Korea, Netherlands, Germany, France (subject to export control alignment). For these countries, effective F_sov increases: their entities access US hyperscalers without restriction, and sovereign compute in development (EU Gigafactories) receives preferential treatment.",
    ]),
    ("5.10.2 The Eurasian Sovereign Bloc", [
        "China constitutes the only complete example of a pre-existing sovereign bloc. With an F_sov of approx. 0.98 and a national cloud ecosystem (Alibaba, Tencent, Huawei) operating outside US jurisdiction, Cloud Sovereignty Mandates have no direct traction. The Chinese constraint remains the shortage of advanced chips, but the American bloc cannot throttle a Huawei Ascend 910B cluster.",
        "Post-2028 dynamics: China holds the only large-scale sovereign compute outside the American bloc. Countries seeking to emancipate themselves from Mandates face a binary choice: conditional American compute or Chinese compute under other forms of dependence.",
    ]),
    ("5.10.3 Digital Non-aligned: an untenable position", [
        "India, Brazil, Southeast Asia, and Gulf countries (without special US treaties) constitute a digital non-alignment bloc. Their position is structurally uncomfortable: too dependent on US hyperscalers to shift to sovereignty, insufficiently integrated into the US alliance to escape restrictions in case of geopolitical disagreement.",
        "The UAE case illustrates this with quantitative force. Fig 1.8 documented that 99.6% of UAE F_total is owned by US-side actors (Stargate UAE, Microsoft, OpenAI), collapsing Sovereign CACI from a Physical 55.7 to just 6.0. Dubai invested massively to become an AI hub, notably via AWS/Microsoft/G42 agreements. Yet, G42 was already pressured in 2024 to break ties with Chinese entities as a condition for advanced chip access.[24] Under Mandates, this pressure becomes systemic: Gulf compute, physically present but legally American, would become a permanent negotiation lever.",
    ]),
    ("5.10.4 The European Bloc: between alliance and autonomy", [
        "Europe occupies an intermediate and evolving position. Legally Tier 1, the UE nevertheless maintains a strategic autonomy ambition that the American Alliance does not fully satisfy.",
        "The Cloud and AI Development Act (CADA), expected Q1 2026, attempts to resolve this by defining an 'EU Sovereignty Level' that would structurally exclude CLOUD Act-subject providers from sensitive public procurement.[25] The October 2025 Cloud Sovereignty Framework defines three assurance levels, with SOV-3 requiring the provider to be beyond the reach of non-European extraterritorial legislation.[26]",
        "This architecture is under construction, but its timing is problematic: CADA will be at best operational in 2027-2028, precisely when US Mandates could be activated. The vulnerability window is maximal.",
        f"Note from Chapter I: the EU is largely sovereign on installed compute ({fmt_en(eu_sov, 1)}% of F_total is EU-owned). The vulnerability is not on installed F but on the cloud workload layer (the compute actually used by EU firms, majorly on AWS/Azure/GCP). CADA targets exactly this layer.",
    ]),
    ("5.11 Transversal impacts on scenarios A-D", [
        "The Cloud-Nationality Pivot superimposes on the four scenarios, modifying their conclusions non-linearly. It adds a third dimension: the degree of installed compute autonomy. Table 13 synthesises the impact.",
    ]),
    ("5.12 Implications for France: the question of genuinely sovereign compute", []),
    ("5.12.1 'Sovereignty washing' as a systemic risk", [
        "The term 'sovereignty washing' designates the practice of US hyperscalers marketing sovereign cloud offers by implanting data centers on European soil, while remaining subject to the CLOUD Act.[27] Microsoft acknowledged in its own documentation that it cannot guarantee sovereignty for European clients in case of a legally grounded US injunction.[28]",
    ]),
    ("5.12.2 The French nuclear advantage in the new equation", [
        "The sovereign reinterpretation of CACI paradoxically reinforces the French strategic asset. If F(r) = F_phys x F_sov, the optimal strategy is not just to increase F_phys (attracting hyperscalers) but F_sov (developing compute independent of US jurisdictions).",
        "French nuclear power creates a top-tier competitive advantage: Special Compute Zones adjoined to carbon-free, competitive nuclear electricity, hosting Gigafactories operated by European entities (OVHcloud, Scaleway, Mistral AI, IONOS), would produce compute with an F_sov close to 1, genuinely beyond the reach of Cloud Sovereignty Mandates.",
    ]),
    ("5.12.3 The DARE/RISC-V project: from ambition to strategic necessity", [
        "Under Scenario D and especially under Cloud Sovereignty Mandates, the DARE project (Digital Autonomy with RISC-V in Europe, 2025) shifts from a long-term ambition to a strategic necessity.[29] As long as Europe depends exclusively on NVIDIA/AMD GPUs, US control over these architectures creates a residual vulnerability even in European-operated Gigafactories. Genuine strategic autonomy requires the ability to design independent accelerators, an horizon DARE situates at 2030-2032.",
    ]),
    ("5.13 Synthesis: the fourth tipping point", [
        "The Great Decoupling is not inevitable. It represents a conditional systemic risk. To the three tipping points in 5.7.2, a fourth is added.",
        "Fourth tipping point: implementation of location verification features in advanced GPUs. If the BIS and US Commerce Department successfully deploy remote attestation mechanisms in H100/B200 chips by end-2026-2027, the technical substrate for the Cloud-Nationality Pivot will be in place. The question will shift from technical feasibility to political decision.",
        "For Europe and France, a controlled decoupling strategy rests on three pillars: (i) accelerating sovereign compute deployment (EU Gigafactories) to increase F_sov before Mandates; (ii) securing Tier 1 status in the American alliance; (iii) investing in architectural alternatives (DARE/RISC-V, Huawei Ascend for non-sensitive workloads) to reduce medium-term dependence.",
    ]),
]

EN.table_blocks = [
    ("Table 10. 2x2 matrix of prospective scenarios 2026-2030.",
     "Source: author's construction, Schwartz methodology (1991).",
     [
         ["", "Reactive EU response", "Proactive EU response"],
         ["Moderate US protectionism",
          "Scenario A - Reinforced Status Quo (slow drift towards dependence)",
          "Scenario C - Asymmetric Partnership (junior Western technological partner)"],
         ["Aggressive US protectionism",
          "Scenario B - Digital Fracture (structural European decoupling)",
          "Scenario D - Contested Sovereignty (race for autonomy under pressure)"],
     ]),
    ("Table 11. Comparative synthesis of the four scenarios on the six divergence metrics (2030 horizon).",
     f"Source: author's construction; baseline April 2026 snapshot (US/EU raw operational compute {fmt_en(us_eu_raw, 1)}:1, CACI Power Mode {fmt_en(us_eu_caci, 2)}:1).",
     [
         ["Metric (2030)", "A - Status Quo", "B - Fracture", "C - Partnership", "D - Sovereignty"],
         ["M1 Raw compute ratio US/EU (operational)", "18-22:1", "25-35:1", "8-10:1", "12-15:1"],
         ["M2 FLOP cost gap", "2.4-3.2x", "4-6x", "1.5-2.0x", "1.8-2.5x"],
         ["M3 US cloud share (pct)", "72-75", "78-82", "60-65", "50-55"],
         ["M4 EU productivity (pct/yr)", "+1.0-1.5", "+0.3-0.8", "+1.8-2.5", "+1.2-2.0"],
         ["M5 EU energy (TWh)", "~115", "~95", "~140", "~155"],
         ["M6 CACI Power Mode ratio", "4-5:1", "6-8:1", "2.0-2.5:1", "4-7:1 (post-trough)"],
         ["Estimated probability", "40-50 pct", "15-20 pct", "15-20 pct", "15-20 pct"],
     ]),
    ("Table 12. Estimation of F_sov factor by jurisdiction and CACI impact under Cloud Sovereignty Mandates activation (2028).",
     "Source: author's construction; Synergy Research Group (2025), Statista (2025), and Chapter I Fig 1.8 baseline.",
     [
         ["Jurisdiction", "F_phys cloud-US share", "Estimated F_sov", "Current CACI (phys baseline)", "Post-Mandate CACI impact"],
         ["United States", "~5 pct (domestic cloud unaffected)", "1.00", "100 (reference)", "100 (unchanged)"],
         ["EU (France, Germany)", "~77 pct (AWS/Azure/GCP dominance)", "0.28", "28.9 (Power Mode)", "30-50 pct collapse on workloads"],
         ["UAE (Dubai hub)", "~88 pct (US hyperscaler dominance)", "0.12", "55.7 phys / 6.0 souv", "60-80 pct collapse - illusory hub"],
         ["Singapore", "~82 pct (US hyperscaler dominance)", "0.18", "high - APAC hub", "55-75 pct collapse"],
         ["China", "~2 pct (Alibaba/Tencent/Huawei Cloud)", "0.98", "15.7 (chip shortage caps)", "Unchanged - sovereignty effective"],
         ["India", "~60 pct (AWS/Azure + local)", "0.40", "22.2 (Power Mode)", "Moderate collapse - intermediate pos"],
     ]),
    ("Table 13. Impacts of the Cloud-Nationality Pivot (Cloud Sovereignty Mandates 2028) on the four 2x2 matrix scenarios.",
     "Source: author's construction.",
     [
         ["Scenario", "Without Mandates", "With Mandates 2028", "EU CACI Impact", "Strategic Reading"],
         ["A - Status Quo", "Slow dependence, CACI 4-5:1",
          "Partial activation - hyperscalers cooperate without major restriction",
          "Moderate degradation 15-25 pct on workloads",
          "More stable but security illusion revealed"],
         ["B - Fracture", "Structural decoupling, CACI 6-8:1",
          "Maximal activation - conditional US cloud and restricted chips simultaneously",
          "Double scissor effect: rare chips + conditional compute. CACI ratio potentially > 8:1",
          "Absolute worst case - chips + cloud combo"],
         ["C - Partnership", "Partial catch-up, CACI 2.0-2.5:1",
          "Sovereign Gigafactories absorb shock if EU F_sov rises to 0.45-0.55",
          "Limited impact if deployed on time - Gigafactories = sovereign hedge",
          "Best resilience - prior investment proves its value"],
         ["D - Sovereignty", "Catch-up under pressure, CACI 4-7:1 post-trough",
          "Mandates become political catalyst - accelerate EU deployment and JP/KR/TW alliances",
          "Accelerated U-curve - 2028-29 trough, faster catch-up",
          "Paradoxical: Mandates can accelerate EU sovereignty if response is fast enough"],
     ]),
]

EN.notes = [
    "Epoch AI (Jan 2026), Trends in AI Hardware and Compute. Doubling every 7 months combines 1.6x/yr quantity and 1.6x/yr performance. Even a slowdown to 12 months implies quadrupling by 2030.",
    "IEA (Apr 2025), Energy and AI, Paris. 800-950 TWh projections correspond to median/high scenarios. Energy ratio (1.4-1.7x PPA-adjusted) derived from dashboard: USA 85, EU 135 USD/MWh.",
    "US Supreme Court (Feb 20, 2026), Learning Resources Inc. v. Trump. 6-3 decision: 'IEEPA does not authorize the President to impose tariffs.'",
    "White House (Jan 14, 2026), Proclamation 11002, sec (2): Secretary update on data center semi market due July 1, 2026.",
    "Tax Foundation (Feb 2026). Aug 2025 US-EU agreement caps semi tariffs for EU at 15 percent, but Proclamation 11002 allows broader tariffs after Phase 1.",
    "European Commission (2025), AI Continent Action Plan. Target: triple EU data center capacity in 5-7 years. 19 AI Factories, 5 Gigafactories. InvestAI: 20 bn EUR. CFG (Oct 2025) Special Compute Zones.",
    "Tax Foundation (Feb 2026), op. cit.",
    "EuroHPC JU (2025). 19 AI Factories approx 475,000 public GPUs, vs single xAI Colossus (200k H100 extensible).",
    "Proclamation 11002, tariff offset program section.",
    "McKinsey Global Institute (May 2024). 0.3 percent corresponds to slow adoption scenario.",
    "EU Council (Dec 2025). Provisional schedule: operational Gigafactories in 2027-2028.",
    "CFG (Oct 2025), 'Tripling the EU Data Centre Stock with Special AI Compute Zones.'",
    "EuroHPC JU (Mar 2025), DARE project (Digital Autonomy with RISC-V in Europe).",
    "Proclamation 11002, sec (2): April 14, 2026 negotiation report due.",
    "CFG (Oct 2025). EU DC authorisation currently 24+ months (vs 6-12 US). SCZ would reduce to 180 days.",
    "RTE (2024), Futurs energetiques 2050. France +10 GW AI demand by 2030. Nuclear 65-70% of mix.",
    "CLOUD Act, Pub. L. 115-141 (Mar 23, 2018), Title III, sec 103(a). 18 U.S.C. sec 2713.",
    "Microsoft France, Tribunal judiciaire de Paris (2024). Cited in The Register (Dec 22, 2025).",
    "Michael Kratsios (OSTP), Seoul (Aug 2025), cited in TechResearchOnline.",
    "BIS, Framework for Artificial Intelligence Diffusion, Fed. Reg. vol 90, no 10 (Jan 15, 2025).",
    "BIS, 'Suspension of the Affiliates Rule for One Year' (Nov 10, 2025).",
    "Financial Times (July 25, 2025), '$1bn Nvidia chips smuggled to China'.",
    "White House, EO on Promoting the Export of the American AI Technology Stack (July 23, 2025).",
    "US Dept of Commerce, G42/Microsoft agreement (2024). Confirmed by Sec Gina Raimondo.",
    "European Commission, 2026 Work Programme, CADA, expected Q1 2026.",
    "European Commission, Cloud Sovereignty Framework (Oct 2025), SOV-1 to SOV-3.",
    "Cristina Caffarra, Eurostack Foundation, cited in The Register (Dec 22, 2025).",
    "See note 18.",
    "EuroHPC JU, DARE project (Digital Autonomy with RISC-V in Europe), launched Mar 2025.",
]


# ===========================================================================
# Content - French (Original)
# ===========================================================================

FR = LangPack(
    code="FR",
    filename="Chapitre_V_Scenarios_Prospectifs_2026_2030_FR.docx",
    cover_title="AI FOR AMERICANS FIRST",
    cover_subtitle="Protectionnisme IA, Energie et Semi-conducteurs : Trajectoires de divergence US/Europe 2024-2030",
    cover_blurb="Analyse geostrategique et economique integree - Chapitre V",
    cover_chip_lines=[
        f"{fmt_fr(us_share, 1)} pct du compute IA operationnel mondial = USA",
        "1,59x cout energie EU/US (ajuste-PPA)",
        f"{fmt_fr(us_eu_caci, 2)}:1 ratio CACI US/EU (Power Mode)",
    ],
    cover_meta="Paris - fevrier 2026  |  7 chapitres  |  4 scenarios prospectifs  |  3 zones geographiques",
    cover_keywords_label="Mots-cles",
    cover_keywords=("intelligence artificielle, protectionnisme technologique, semi-conducteurs, "
                    "controles a l'exportation, compute souverain, geopolitique IA, France, "
                    "Etats-Unis, Chine"),
    chapter_label="CHAPITRE V",
    chapter_title="Scenarios prospectifs 2026-2030",
    chapter_intro=(
        "Ce chapitre constitue le coeur de la contribution originale de cette etude. "
        "En appliquant le protocole methodologique defini au chapitre II (matrice 2x2, "
        "six metriques de divergence, calibration CACI), nous construisons quatre "
        "scenarios d'evolution de la relation transatlantique en IA, energie et "
        "semi-conducteurs pour la periode 2026-2030. Chaque scenario est determine par "
        "la combinaison de deux incertitudes critiques identifiees au chapitre III : "
        "le degre de protectionnisme americain et la capacite de reponse strategique "
        "europeenne. Nous evaluons ensuite chaque scenario sur ses six metriques, avant "
        "de synthetiser les conditions de bascule entre trajectoires."
    ),
    notes_label="Notes",
    license_block=[
        "Licence et avertissement. Cette oeuvre, 'AI for Americans First', est mise a disposition selon les termes de la licence Creative Commons Attribution - Pas d'Utilisation Commerciale - Partage dans les Memes Conditions 4.0 International (CC BY-NC-SA 4.0).",
        "Vous etes libre de partager et adapter le materiel a des fins non commerciales, a condition d'attribuer correctement l'oeuvre a Fabrice Pizzi (Universite Paris-Sorbonne) et de distribuer vos contributions sous la meme licence. Ce document est fourni a des fins educatives et de recherche uniquement.",
        "Tableau de bord public : https://mo0ogly.github.io/America-First-IA/dashboard/",
        "Depot : https://github.com/mo0ogly/America-First-IA",
    ],
    page_footer="AI for Americans First - Fabrice Pizzi - Chapitre V",
)

# Load FR sections and tables from the provided reference (truncated here for brevity but fully implemented in build)
# Note: I'll include the actual data to ensure the builder works.
FR.sections = [
    ("5.1 Elements predetermines : ce qui ne changera pas", [
        "Conformement a la methode Schwartz (1991), nous distinguons les elements predetermines (tendances quasi-certaines a l'horizon 2030) des incertitudes critiques (facteurs dont l'evolution depend de decisions politiques non encore prises). Quatre elements predetermines structurent l'ensemble des scenarios.",
        "EP1 - Croissance exponentielle de la demande de compute IA. Les ventes de semi-conducteurs ont double en deux ans (2023-2025), la puissance des puces IA installees double tous les sept mois (Epoch AI), et aucun signe de ralentissement n'est observable au moment de fevrier 2026. Meme dans l'hypothese d'une deceleration des scaling laws (saturation Chinchilla), la diffusion de l'IA vers l'inference, la robotique et les agents autonomes maintiendra une demande de compute fortement croissante.[1]",
        f"EP2 - Concentration persistante du compute aux Etats-Unis. Le ratio US/UE de {fmt_fr(us_eu_raw, 1)}:1 en compute installe brut (chapitre III, snapshot du tableau de bord d'avril 2026 : 2 759 968 vs 156 632 PFLOP/s), se traduisant par un ratio CACI Power Mode de {fmt_fr(us_eu_caci, 2)}:1 une fois pondere par la formule geometrique F^0,40 x L^0,20 x R^0,15 / E^0,25, reflete des decisions d'investissement prises en 2022-2025 dont les effets se materialisent jusqu'en 2028-2029 (delais de construction des centres de donnees : 18-36 mois). Meme un revirement politique immediat n'alterait pas le stock installe avant la fin de la decennie.",
        "EP3 - Tension energetique croissante. La consommation mondiale des centres de donnees, estimee a 415 TWh en 2024, atteindra 800-950 TWh d'ici 2030 selon les projections AIE (chapitre III). L'asymetrie des couts energetiques (US 1,4-1,7x moins chers apres correction PPA, 2-3x moins chers sur les tarifs Eurostat industriels non ajustes) persistera, sauf investissement massif dans le nucleaire europeen, dont les delais de deploiement (SMR : 5-7 ans pour les premiers reacteurs) depassent l'horizon 2030.[2]",
        "EP4 - Cadre reglementaire Section 232 en place. La Proclamation 11002 du 14 janvier 2026 est un fait accompli juridique. Contrairement aux tarifs IEEPA (annules par la Cour supreme le 20 fevrier 2026[3]), les tarifs Section 232 reposent sur une base legale confirmee. Le rapport du Secretaire au Commerce sur le marche des semi-conducteurs pour centres de donnees est attendu d'ici le 1er juillet 2026, et peut recommander une extension ou modification des tarifs. Quelle que soit la direction prise, l'instrument legal restera disponible.[4]",
    ]),
    ("5.2 Incertitudes critiques et matrice 2x2", []),
    ("5.2.1 Axe 1 : Degre de protectionnisme americain", [
        "La premiere incertitude porte sur l'evolution de la politique americaine entre deux poles. Le pole modere correspond au maintien du statu quo de janvier 2026 : tarif de 25 pour cent limite aux puces avancees re-exportees, exemptions domestiques larges, accord commercial UE plafonnant les tarifs sur semi-conducteurs a 15 pour cent, et aucune extension significative au cloud ou aux modeles. Le coeur UE (France, Allemagne) reste dans la categorie partenaires de confiance. Le pole agressif suppose une extension apres le rapport de juillet 2026 : tarifs etendus aux semi-conducteurs derives et equipements, quotas GPU pour l'UE (incluant la France), conditions restrictives pour l'acces au cloud IA de pointe, et utilisation du compute comme levier de negociation commerciale (compute-for-concessions).[5]",
    ]),
    ("5.2.2 Axe 2 : Capacite de reponse strategique europeenne", [
        "La seconde incertitude porte sur la capacite de l'UE a deployer une reponse coherente et rapide. Le pole reactif correspond a des reponses nationales fragmentees, des investissements disperses, un AI Act creant des couts de conformite supplementaires, et un deploiement lent des AI Factories/Gigafactories (delais bureaucratiques, autorisations 24+ mois). Le pole proactif suppose la mise en oeuvre acceleree du programme AI Continent (19 AI Factories plus jusqu'a 5 Gigafactories de 100 000+ GPU), l'adoption de Special Compute Zones (autorisation en 180 jours), la mobilisation effective du fonds InvestAI (20 milliards EUR), et la mise en commun de la capacite nucleaire francaise comme avantage competitif.[6]",
    ]),
    ("5.2.3 Matrice et nommage des scenarios", [
        "L'intersection de ces deux axes produit quatre scenarios. UE reactive + protectionnisme US modere donne le scenario A (Statu quo renforce, derive lente vers la dependance). UE reactive + US agressif donne le scenario B (Fracture numerique, decouplage europeen structurel). UE proactive + US modere donne le scenario C (Partenariat asymetrique, partenaire technologique junior occidental). UE proactive + US agressif donne le scenario D (Souverainete contestee, course a l'autonomie sous pression).",
    ]),
    ("5.3 Scenario A - Statu quo renforce (Protectionnisme modere + UE reactive)", []),
    ("5.3.1 Recit", [
        "Apres le rapport de juillet 2026, le Secretaire au Commerce recommande de maintenir le tarif de 25 pour cent sur les puces avancees re-exportees mais sans l'etendre significativement. L'accord commercial US-UE d'aout 2025 est respecte : les tarifs sur semi-conducteurs pour l'UE restent plafonnes a 15 pour cent.[7] L'UE, rassuree par ce statu quo, ralentit le deploiement de ses propres initiatives. Les AI Factories EuroHPC peinent a atteindre leur capacite nominale (delais d'autorisation, coordination inter-Etats). Les Gigafactories sont reportees a 2029-2030. Le fonds InvestAI est partiellement mobilise (8-10 milliards EUR sur 20). Les entreprises europeennes continuent de s'appuyer fortement sur le cloud US, dont la performance et les couts restent imbattables.",
    ]),
    ("5.3.2 Trajectoire des metriques", [
        f"M1 - Ratio compute (GPU installes US/UE) : passe de {fmt_fr(us_eu_raw, 1)}:1 brut (2025) a 18-22:1 brut (2030) sur le compute installe operationnel. L'ecart se creuse legerement a mesure que les investissements US s'accelerent (Stargate, mega-clusters xAI, Meta) tandis que l'UE n'ajoute que les 19 AI Factories (25 000 GPU max chacune, soit environ 475 000 GPU publics, un ordre de grandeur en dessous d'un seul hyperscaler US).[8]",
        "M2 - Ecart cout du FLOP (UE/US) : reste dans la fourchette 2,4-3,2x. L'absence de tarifs agressifs sur l'UE maintient l'acces au cloud US a des prix proches des niveaux actuels, mais les couts energetiques europeens continuent de peser.",
        "M3 - Part cloud US dans les depenses IA europeennes : passe de 70 pour cent (2024) a 72-75 pour cent (2030). Les fournisseurs europeens (OVHcloud, Deutsche Telekom) conservent leurs 15 pour cent sur le segment souverainete mais ne gagnent pas de terrain sur les services IA generative.",
        "M4 - Productivite IA (pct/an) : US +2,5-3,0 ; UE +1,0-1,5. L'UE realise une partie du potentiel IA via les applications aval (SAP, Siemens, fintech), mais l'adoption lente et le deficit de compute plafonnent les gains.",
        "M5 - Dependance energetique (TWh centres de donnees) : UE environ 115 TWh en 2030 (+65 pour cent vs 2024). Le nucleaire francais absorbe une partie de la demande, mais l'absence de Special Compute Zones retarde la connexion au reseau de nouveaux centres de donnees.",
        f"M6 - CACI(US)/CACI(UE) : passe de {fmt_fr(us_eu_caci, 2)}:1 (avril 2026) a 4-5:1 (2030). L'ecart se creuse moderement a mesure que le facteur F (compute) s'accumule cote US tandis que les couts E (energie) pesent cote UE.",
    ]),
    ("5.3.3 Consequences pour la France", [
        "Ce scenario est le plus probable a court terme (probabilite estimee : 40-50 pour cent). Il est aussi le plus insidieux : l'absence de choc visible demobilise les acteurs europeens, tandis que la dependance se creuse structurellement. Les entreprises francaises beneficient de l'acces au cloud US pour l'adoption IA (BNP Paribas, Airbus, TotalEnergies via AWS/Azure), mais cette adoption renforce le verrouillage decrit au chapitre IV. Le deficit de productivite IA par rapport aux Etats-Unis (-1,0 a -1,5 points par an) s'accumule sur cinq ans, creusant l'ecart de competitivite de 5 a 8 points de PIB.",
    ]),
    ("5.4 Scenario B - Fracture numerique (Protectionnisme agressif + UE reactive)", []),
    ("5.4.1 Recit", [
        "Le rapport de juillet 2026 conduit a une extension significative. Le Secretaire au Commerce recommande des tarifs etendus aux equipements semi-conducteurs et produits derives, avec un tariff offset program reserve aux entreprises investissant dans la production americaine.[9] L'accord UE a 15 pour cent est revise a la hausse, ou accompagne de conditions restrictives (quotas de volume sur les GPU avances, exigences de reciprocite sur l'AI Act). Simultanement, l'acces au cloud IA de pointe est rendu conditionnel pour les entites non americaines (limitations d'acces aux API des modeles de frontiere, restrictions sur les poids). L'UE, fragmentee, ne parvient pas a formuler une reponse coherente : les Etats membres se divisent entre accommodation (pays nordiques, Pays-Bas) et confrontation (France, Italie).",
    ]),
    ("5.4.2 Trajectoire des metriques", [
        f"M1 - Ratio compute : passe de {fmt_fr(us_eu_raw, 1)}:1 brut (2025) a 25-35:1 brut (2030). Les quotas GPU limitent les importations europeennes au moment ou la demande explose. Les projets AI Factory sont compromis par l'incapacite a se procurer des GPU Nvidia/AMD aux volumes prevus.",
        "M2 - Ecart cout du FLOP : bondit a 4-6x. Les tarifs etendus, combines aux quotas et a l'asymetrie energetique, augmentent massivement les couts du compute europeen. Les entreprises francaises font face a une surcharge de 3x a 5x pour l'entrainement de modeles.",
        "M3 - Part cloud US : paradoxalement, monte a 78-82 pour cent. Faute d'alternative locale credible, les entreprises europeennes voulant acceder a l'IA de pointe doivent passer par les hyperscalers US, aux conditions tarifaires qu'ils dictent. Les services souverains (OVHcloud, Scaleway) manquent du materiel pour offrir des services GenAI competitifs.",
        "M4 - Productivite IA : US +2,5-3,5 ; UE +0,3-0,8. Le potentiel IA europeen est severement contraint. Le McKinsey Global Institute estime qu'avec une adoption lente, la productivite europeenne ne depasserait pas 0,3 pour cent, proche de la stagnation.[10]",
        "M5 - Dependance energetique : UE environ 95 TWh seulement (2030), non par vertu mais par defaut - le manque de GPU limite la construction des centres de donnees. Ironiquement, la contrainte de compute attenue la contrainte energetique.",
        f"M6 - Ratio CACI : explose de {fmt_fr(us_eu_caci, 2)}:1 (avril 2026) a 6-8:1 (2030). C'est le scenario ou l'ecart est le plus important, avec les trois facteurs CACI se deteriorant simultanement cote europeen : F plafonne par les quotas, E gonfle par les tarifs, L affaibli par un brain drain accelere vers les Etats-Unis.",
    ]),
    ("5.4.3 Consequences pour la France", [
        "Ce scenario (probabilite estimee : 15-20 pour cent) represente le pire des cas. La France subit un decouplage technologique structurel : les projets compute-intensifs (modeles de fondation Mistral, robotique Comau/Exotec, simulations Dassault) sont relocalises aux Etats-Unis ou dependent d'un acces au cloud US de plus en plus couteux. Le time-to-market des solutions IA francaises s'allonge de 25 a 40 pour cent. Les PME industrielles, incapables d'absorber les surcharges, renoncent a l'IA de pointe et optent pour des solutions degradees (modeles open-source plus petits, inference locale). L'ecart de productivite cumule avec les Etats-Unis atteint 10 a 15 points sur cinq ans.",
    ]),
    ("5.5 Scenario C - Partenariat asymetrique (Protectionnisme modere + UE proactive)", []),
    ("5.5.1 Recit", [
        "Le protectionnisme US reste modere (comme en A), mais l'UE exploite cette fenetre pour accelerer ses propres investissements. Les AI Factories sont deployees dans les delais (2026-2027), les premieres Gigafactories de 100 000+ GPU sont commandees fin 2026 et livrees en 2028.[11] La France joue un role central grace a son parc nucleaire (65-70 pour cent du mix electrique, cout marginal competitif), et des Special Compute Zones sont designees sur d'anciens sites industriels avec connexions reseau lourdes.[12] Toutefois, l'UE accepte de facto un statut de partenaire technologique junior : elle utilise des GPU Nvidia/AMD (pas de champion europeen en design d'ASIC IA), depend des fonderies TSMC/Samsung/Intel pour la production, et ses modeles de fondation restent un cran en dessous des leaders US.",
    ]),
    ("5.5.2 Trajectoire des metriques", [
        f"M1 - Ratio compute : descend de {fmt_fr(us_eu_raw, 1)}:1 brut (2025) a 8-10:1 brut (2030) sur le compute installe. Les Gigafactories et l'investissement prive (InvestAI plus co-investissements industriels) ajoutent 1-2 millions d'equivalents H100 en Europe, reduisant l'ecart sans le combler.",
        "M2 - Ecart cout du FLOP : descend a 1,5-2,0x. Le nucleaire francais et les economies d'echelle des Gigafactories compriment les couts d'energie et d'infrastructure, bien qu'un ecart residuel persiste (absence de design GPU proprietaire).",
        "M3 - Part cloud US : descend legerement a 60-65 pour cent. Les services souverains europeens gagnent des parts de marche sur les segments regules (defense, sante, finance), tandis que le cloud US conserve la majorite des charges commerciales. Le marche se segmente en souverain et performance.",
        "M4 - Productivite IA : US +2,5-3,0 ; UE +1,8-2,5. L'UE atteint 60-80 pour cent de son potentiel theorique grace a un compute local suffisant pour l'adoption a grande echelle d'applications aval, meme si l'entrainement des modeles de frontiere reste dependant du materiel US.",
        "M5 - Energie : UE environ 140 TWh (2030). La demande est plus elevee qu'en A car le compute europeen augmente, mais le nucleaire et les SMR planifies absorbent l'essentiel. RTE France confirme la faisabilite de +10 GW sous reserve d'investissements reseau.",
        f"M6 - Ratio CACI : descend de {fmt_fr(us_eu_caci, 2)}:1 (avril 2026) a 2,0-2,5:1 (2030). C'est le scenario le plus favorable realistement atteignable a l'horizon 2030. Le facteur F s'ameliore significativement, E beneficie du nucleaire, mais L reste legerement inferieur (l'ecosysteme IA US plus attractif pour les meilleurs talents).",
    ]),
    ("5.5.3 Consequences pour la France", [
        "Ce scenario (probabilite estimee : 15-20 pour cent) est le plus favorable pour la France a court-moyen terme. La France devient le hub energetique IA de l'UE grace a son parc nucleaire, attirant les investissements en centres de donnees et Gigafactories. Les entreprises francaises gagnent un acces a un compute local competitif pour l'inference et le fine-tuning, reduisant la dependance au cloud US pour les usages standard. Mistral et les startups francaises peuvent entrainer des modeles specialises localement. Toutefois, l'entrainement des modeles de frontiere reste dependant du materiel US, et l'autonomie strategique est partielle : la France est souveraine en application, mais pas dans la creation des technologies fondamentales.",
    ]),
    ("5.6 Scenario D - Souverainete contestee (Protectionnisme agressif + UE proactive)", []),
    ("5.6.1 Recit", [
        "Le protectionnisme US s'intensifie (comme en B), mais l'UE repond avec determination. La menace americaine devient le catalyseur politique d'une mobilisation industrielle europeenne sans precedent depuis le projet AIRBUS des annees 1970. Le programme AI Continent est accelere et etendu : les 5 Gigafactories sont commandees en urgence, la France annonce 20 GW de capacite nucleaire dediee aux centres de donnees IA d'ici 2032 (combinant extension du parc existant et SMR), le projet DARE (RISC-V europeen) est escalade pour concevoir des accelerateurs IA reduisant la dependance a Nvidia.[13] Simultanement, l'UE negocie des alliances technologiques alternatives (Japon, Coree du Sud, Taiwan) pour securiser l'approvisionnement en GPU et fonderies.",
    ]),
    ("5.6.2 Trajectoire des metriques", [
        f"M1 - Ratio compute : evolue de {fmt_fr(us_eu_raw, 1)}:1 brut (2025) a 12-15:1 brut (2030) sur le compute installe. L'UE investit massivement mais part de tres loin. Les quotas US ralentissent les importations, mais les alliances alternatives et la production locale (Gigafactories utilisant des GPU Samsung/Intel comme alternatives a Nvidia) compensent partiellement.",
        "M2 - Ecart cout du FLOP : 2,5-4,0x initialement (2027, pic du choc tarifaire), puis reduction progressive vers 1,8-2,5x (2030) a mesure que les Gigafactories montent en cadence et que les alternatives GPU murissent.",
        "M3 - Part cloud US : descend a 50-55 pour cent (2030), le declin le plus prononce des quatre scenarios. La defiance geopolitique et les restrictions US poussent les entreprises europeennes vers les alternatives souveraines, meme imparfaites. Les hyperscalers US perdent du terrain sur les segments regules.",
        "M4 - Productivite IA : US +2,5-3,5 ; UE +1,2-2,0. L'UE traverse un creux de productivite en 2027-2028 (periode de transition ou les restrictions US mordent mais les investissements europeens ne sont pas encore operationnels), puis un rattrapage partiel a partir de 2029.",
        "M5 - Energie : UE environ 150-160 TWh (2030). C'est le scenario le plus energivore pour l'UE, la construction massive de centres de donnees locaux creant une demande enorme. Le nucleaire francais devient un actif strategique continental, mais la pression sur le reseau est maximale.",
        f"M6 - Ratio CACI : suit une trajectoire en U : degradation de {fmt_fr(us_eu_caci, 2)}:1 a 8-12:1 en 2027-2028 (pic du choc), puis amelioration vers 4-7:1 d'ici 2030. Le resultat depend fortement de la vitesse d'execution europeenne : chaque annee de retard sur les Gigafactories prolonge la periode de vulnerabilite maximale.",
    ]),
    ("5.6.3 Consequences pour la France", [
        "Ce scenario (probabilite estimee : 15-20 pour cent) est le plus ambitieux et le plus risque. Il place la France au coeur d'un effort de souverainete technologique europeen sans precedent. Les investissements nucleaires massifs (SMR, extension du parc) deviennent un enjeu geopolitique de premier ordre. Le projet DARE/RISC-V pourrait, en cas de succes, constituer la premiere alternative europeenne credible aux GPU Nvidia pour l'IA, mais sur un horizon de 5-7 ans, bien au-dela de 2030. A court terme (2026-2028), la France traverse une periode de vulnerabilite maximale ou les surcharges et penuries degradent la competitivite, avant un rattrapage conditionnel a la vitesse de deploiement de l'infrastructure.",
    ]),
    ("5.7 Synthese comparative et conditions de bascule", []),
    ("5.7.1 Tableau de synthese des metriques", [
        f"Le Tableau 11 ci-dessous consolide la trajectoire des six metriques de divergence pour les quatre scenarios a l'horizon 2030, ancre sur le snapshot du tableau de bord d'avril 2026 (compute brut operationnel US/UE {fmt_fr(us_eu_raw, 1)}:1, CACI Power Mode {fmt_fr(us_eu_caci, 2)}:1).",
    ]),
    ("5.7.2 Conditions de bascule entre scenarios", [
        "La trajectoire reelle suivra probablement un chemin hybride entre ces scenarios. Trois points de bascule determinent les transitions possibles.",
        "Premier point de bascule : le rapport Commerce de juillet 2026. Ce rapport determinera si le protectionnisme US s'etend (basculement vers B ou D) ou reste cible (maintien en A ou C). Indicateurs a surveiller : evolution du deficit commercial americain en semi-conducteurs, taux de remplissage des fabs CHIPS Act (Intel, TSMC Arizona, Samsung Taylor), pression politique interieure (midterms 2026). Le resultat des negociations de Phase 1 (rapport du a 14 avril 2026) sera un signal precoce.[14]",
        "Deuxieme point de bascule : la vitesse de deploiement des Gigafactories UE. La Commission prevoit les premieres Gigafactories operationnelles en 2027-2028. Si ce calendrier est tenu, l'UE bascule vers les scenarios proactifs (C ou D). Si les delais d'autorisation, de financement ou d'approvisionnement materiel repoussent les livraisons a 2029-2030, l'UE reste en mode reactif (A ou B). La proposition CFG des Special Compute Zones (autorisation en 180 jours vs 24+ mois actuellement) est le facteur d'acceleration cle.[15]",
        "Troisieme point de bascule : la decision francaise sur le nucleaire pour l'IA. La France possede un atout unique en Europe : un parc nucleaire fournissant 65-70 pour cent de l'electricite, avec un cout marginal globalement competitif. La decision de dedier une capacite significative (10-20 GW) aux centres de donnees IA, via l'extension du parc, les nouveaux EPR2 et les SMR, determinera si la France devient le hub energetique IA de l'Europe ou cede cette position a d'autres (Scandinavie avec l'hydroelectrique, Europe de l'Est avec des couts fonciers bas). Ce point de bascule est specifiquement francais et determine la position de la France au sein des scenarios europeens.[16]",
    ]),
    ("5.7.3 Le point de convergence : 2028", [
        "Les quatre scenarios convergent sur un point critique commun en 2028. C'est l'annee ou : (i) la demande de compute depassera la capacite installee en Europe, creant des goulots d'etranglement materiels (meme sous protectionnisme modere) ; (ii) les premiers effets des tarifs etendus (s'ils sont adoptes) seront pleinement ressentis ; (iii) les Gigafactories, si elles sont deployees a temps, commenceront a produire un compute local significatif ; (iv) la demande energetique des centres de donnees saturera la capacite de connexion au reseau dans plusieurs Etats membres. L'annee 2028 constitue donc le moment de verite ou l'Europe decouvrira si elle est sur trajectoire A/B (dependance croissante) ou C/D (rattrapage commence). Les decisions prises en 2026-2027 (rapport Commerce US, Gigafactories, nucleaire francais) seront irreversiblement engagees.",
    ]),
    ("5.8 Origines du point de bascule : fondations juridico-techniques deja en place", [
        "Le scenario du Grand Decouplage n'est pas construit ex nihilo. Il est la projection logique, a l'horizon 2028, d'une architecture de controle dont les couches fondationnelles sont deja operationnelles en 2026. Identifier ces couches releve de la rigueur academique : distinguer ce qui releve de tendances documentees de ce qui constitue une projection extrapolee.",
    ]),
    ("5.8.1 La couche legale : l'extraterritorialite comme instrument structurel", [
        "La premiere couche est legale et precede largement la politique IA. Le CLOUD Act (Clarifying Lawful Overseas Use of Data Act, 2018) etablit que tout fournisseur de cloud soumis a la juridiction US est tenu de produire les donnees independamment du lieu ou ces communications, enregistrements ou autres informations sont stockes.[17] Cette loi, confirmee par la jurisprudence federale (United States v. Microsoft, resolu par l'adoption du CLOUD Act avant que la Cour supreme ne statue), cree une dissociation fondamentale entre la localisation physique des donnees et leur nationalite legale.",
        "La consequence immediate pour le compute IA est radicale : un cluster H100 physiquement localise a Dubai, Singapour ou dans un centre AWS eu-west-1 en Irlande reste legalement americain. En cas d'ordre du gouvernement US, l'operateur (AWS, Azure, Google Cloud) est legalement tenu de se conformer, independamment de la volonte du client ou de la legislation locale. Microsoft a reconnu devant un tribunal francais en 2024 ne pas pouvoir garantir la souverainete des donnees pour les clients europeens en cas d'injonction US fondee en droit.[18]",
        "Cette architecture legale preexistante est le substrat sur lequel se greffe le Pivot Cloud-Nationalite : les Cloud Sovereignty Mandates projetes pour 2028 ne creent pas un pouvoir nouveau. Ils activent et systematisent un pouvoir juridictionnel deja existant en l'etendant a la couche du compute operationnel.",
    ]),
    ("5.8.2 La couche technique : de la verification de localisation au throttling de cluster", [
        "La seconde couche est technique. L'AI Action Plan US du 23 juillet 2025 introduit explicitement le concept de fonctionnalites de verification de localisation appliquees aux puces IA avancees. Michael Kratsios, directeur de l'OSTP de la Maison-Blanche, a confirme qu'il y a une discussion sur les types de modifications logicielles ou physiques que l'on pourrait apporter aux puces elles-memes pour faire un meilleur location-tracking, ce qui a ete explicitement inclus dans le plan.[19]",
        "Cette annonce n'est pas rhetorique. Le BIS dispose deja, depuis le Framework for Artificial Intelligence Diffusion (janvier 2025, abroge en mai puis remplace par des mesures de guidance), d'une architecture de controle compute-tiering pour les pays de destination (Tier 1-2-3) et de plafonds de compute par entite et par pays.[20] La BIS Affiliates Rule, suspendue pour un an en novembre 2025 mais maintenue en principe, stipule que l'affiliation d'une societe a une entite mere dans un pays restreint suffit a lui refuser l'acces au compute avance, independamment de la localisation physique du cluster.[21]",
        "La trajectoire technologique vers 2028 est donc : (i) des puces equipees de mecanismes de verification de localisation (logicielle ou materielle), (ii) des systemes de reporting automatique au BIS en cas de deviation, (iii) la capacite a suspendre l'acces ou throttler la performance via licence d'exportation, un cluster operant sous licence d'exportation US potentiellement soumis a des restrictions operationnelles par decision administrative. Ce n'est pas de la science-fiction : c'est l'extension au compute d'un principe deja applique au logiciel (sanctions OFAC sur les licences logicielles, gel d'acces aux services cloud pour les entites listees).",
        "Limitation critique de ce mecanisme. Une tension technique reelle doit etre identifiee ici : throttler des clusters de production est techniquement complexe et potentiellement perturbant pour les operateurs eux-memes. NVIDIA ne dispose pas actuellement de mecanisme de desactivation a distance pour les GPU H100/B200 en production. Un tel mecanisme exigerait une modification architecturale significative des firmware et des protocoles d'attestation a distance (via TPM ou equivalent). Le scenario 2028 est donc conditionnel a la mise en oeuvre effective de ces modifications, hypothese plausible sur 24-36 mois d'effort industriel, mais pas une certitude.",
    ]),
    ("5.9 Le mecanisme du Pivot Cloud-Nationalite", []),
    ("5.9.1 Le declencheur : Cloud Sovereignty Mandates comme extension des controles a l'exportation", [
        "Le scenario du point de bascule 2028 suppose que les Etats-Unis franchissent un palier qualitatif : passer du controle d'acces materiel (controles BIS sur les puces) au controle d'acces aux services de compute operationnel (Cloud Sovereignty Mandates). Ce palier n'est pas une rupture arbitraire, il repond a une faille structurelle identifiee dans le regime de controle actuel.",
        "Cette faille est documentee : malgre les restrictions BIS sur les exportations de H100/A100, des enquetes ont revele qu'environ 1 milliard USD de puces Nvidia ont ete acheminees vers la Chine en contournant les controles a l'exportation via des pays tiers (Malaisie, EAU, Singapour) au cours des seuls premiers mois de 2025.[22] La reponse de l'AI Action Plan (fonctionnalites de verification de localisation et monitoring de cluster) constitue le premier pas vers un controle continu post-exportation.",
        "Le declencheur plausible en 2028 est un executive order etendant les obligations du Framework for AI Diffusion a la couche cloud. Sa structure imposerait une certification Data Residency and Jurisdiction Compliance a tous les hyperscalers US operant des clusters avances offshore, avec revocation de l'acces au compute sur sol US comme mecanisme de mise en conformite, et le BIS se reservant le droit de throttler ou suspendre la performance des clusters autorises en cas d'irregularite.",
    ]),
    ("5.9.2 La dissociation Facteur Physique / Facteur Souverain dans le CACI", [
        "C'est ici que le scenario produit son impact le plus analytiquement significatif sur le modele CACI developpe aux chapitres I a IV. Le modele actuel integre le compute F(r) comme mesure de la capacite physiquement installee dans la region r. Sous activation du Pivot Cloud-Nationalite, cependant, la variable se decompose en deux composantes distinctes : F(r) = F_phys(r) x F_sov(r), ou F_phys est le compute physiquement installe dans la juridiction et F_sov est le facteur de souverainete operationnelle (fraction de F_phys hors juridiction US et donc insensible aux Cloud Sovereignty Mandates).",
        "Notez la distinction entre ce F_sov dynamique 2028 et le CACI souverain statique introduit au chapitre I (Fig 1.8). Le CACI souverain du chapitre I a ete calcule en filtrant les clusters Epoch AI par Owner : il capture qui detient le compute installe aujourd'hui (snapshot avril 2026). Le F_sov 2028 capture qui controle le compute sous un regime hypothetique de Cloud Sovereignty Mandates, qui depend de la part des charges hyperscaler, pas de la capacite installee. Les deux metriques s'accordent la ou le compute est detenu par des operateurs domestiques (US, Chine, France domestique Fluidstack/Sesterce) ; elles divergent fortement pour les juridictions ou les clusters domestiquement localises sont detenus par des operateurs US-side (EAU 99,6 pour cent, charges cloud UE majoritairement sur AWS/Azure/GCP).",
        "Les estimations calibrees pour 2028, sous l'hypothese d'activation des Cloud Sovereignty Mandates, sont presentees dans le Tableau 12 ci-dessous.",
    ]),
    ("5.10 L'emergence de blocs IA juridictionnels (2028-2030)", [
        "Le Grand Decouplage ne produit pas un monde binaire US/non-US. Il produit une fragmentation en blocs d'intensite variable, selon la capacite de chaque zone a developper un compute souverain credible. Quatre blocs emergent avec des caracteristiques distinctes.",
    ]),
    ("5.10.1 Le bloc americain etendu (American AI Alliance)", [
        "L'AI Action Plan du 23 juillet 2025 pose explicitement les bases d'une American AI Alliance : exportation de la pile technologique US complete (materiel, modeles, logiciels, standards) aux allies disposes, en echange de l'adoption de controles a l'exportation alignes.[23] La strategie est explicitement decrite comme carrot and stick : les allies alignes accedent aux puces avancees et aux modeles de frontiere sans restrictions supplementaires ; ceux qui refusent sont exposes aux mecanismes Foreign Direct Product Rule et aux tarifs secondaires.",
        "Membres du bloc americain etendu (Tier 1 confirmes) : Etats-Unis, Royaume-Uni, Canada, Australie, Japon, Coree du Sud, Pays-Bas, Allemagne, France (sous reserve d'alignement sur les controles a l'exportation). Pour ces pays, le F_sov effectif augmente : leurs entites accedent aux hyperscalers US sans restriction, et le compute souverain en developpement (Gigafactories UE pour l'Europe) recoit un traitement preferentiel. Le CACI de ces pays n'est pas degrade par les Mandates, il peut meme beneficier d'un effet d'alliance.",
    ]),
    ("5.10.2 Le bloc souverain eurasien", [
        "La Chine constitue le seul exemple complet de bloc souverain preexistant. Avec un F_sov d'environ 0,98 et un ecosysteme cloud national (Alibaba Cloud, Tencent Cloud, Huawei Cloud) operant hors de la juridiction US, les Cloud Sovereignty Mandates n'ont aucune traction directe. La contrainte chinoise reste la penurie de puces avancees (les controles a l'exportation 2022-2025 ont limite l'acces aux GPU H100/A100/B200), mais le bloc americain ne peut pas throttler un cluster Huawei Ascend 910B.",
        "La dynamique post-2028 : la Chine detient le seul compute souverain a grande echelle hors du bloc americain. Les pays cherchant a s'emanciper des Cloud Sovereignty Mandates se retrouvent structurellement face a une alternative binaire : compute americain conditionnel ou compute chinois sous d'autres formes de dependance. Cette contrainte binaire est l'impact geopolitique le plus profond du Grand Decouplage.",
    ]),
    ("5.10.3 Les non-alignes numeriques : une position intenable", [
        "L'Inde, le Bresil, l'Asie du Sud-Est et les pays du Golfe (en l'absence de traites speciaux avec les Etats-Unis) constituent un bloc de non-alignement numerique. Leur position est structurellement inconfortable : too dependants des hyperscalers US pour basculer vers la souverainete, insuffisamment integres a l'alliance americaine pour echapper aux restrictions en cas de desaccord geopolitique.",
        "Le cas des EAU illustre cette fragilite avec une force quantitative. La Fig 1.8 du chapitre I a documente que 99,6 pour cent du F_total des EAU (22,9 millions d'equivalents H100) est detenu par des acteurs US-side (Stargate UAE, Microsoft, OpenAI), faisant s'effondrer le CACI souverain d'un Physique 55,7 a seulement 6,0. Dubai a investi massivement depuis 2022 pour devenir un hub IA regional, notamment via des accords avec AWS, Microsoft et G42. Pourtant, G42 a deja ete soumis a une intense pression US in 2024 pour rompre ses liens avec des entites chinoises, condition imposee par Washington pour l'acces aux puces avancees.[24] Sous Cloud Sovereignty Mandates, cette pression deviendrait systemique : le compute des hubs du Golfe, physiquement present mais legalement americain, deviendrait un levier de negociation permanent.",
    ]),
    ("5.10.4 Le bloc europeen : entre alliance et autonomie", [
        "L'Europe occupe une position intermediaire et evolutive. Legalement Tier 1 (France, Allemagne, Pays-Bas, etc. sont explicitement dans la presomption d'approbation BIS pour les puces avancees), l'UE maintient neanmoins une ambition d'autonomie strategique que l'Alliance americaine ne satisfait pas pleinement.",
        "Le Cloud and AI Development Act (CADA), dont la proposition formelle est attendue au T1 2026, tente de repondre a ce dilemme en definissant un EU Sovereignty Level qui exclurait structurellement les fournisseurs soumis au CLOUD Act des marches publics sensibles.[25] La Commission europeenne a publie en octobre 2025 un Cloud Sovereignty Framework definissant trois niveaux d'assurance (SOV-1 a SOV-3), avec SOV-3 exigeant que le fournisseur soit hors d'atteinte de toute legislation extraterritoriale non europeenne.[26]",
        "Cette architecture legislative est en construction, mais son calendrier est problematique : le CADA sera au mieux operationnel en 2027-2028, precisement quand les Cloud Sovereignty Mandates US pourraient etre actives. La fenetre de vulnerabilite est maximale entre 2028 et 2030.",
        f"Une nuance importante du chapitre I : l'UE est largement souveraine sur le compute installe ({fmt_fr(eu_sov, 1)} pour cent du F_total est detenu par des operateurs UE). La fenetre de vulnerabilite n'est donc pas sur le F installe mais sur la couche des charges cloud (le compute reellement utilise par les entreprises UE, majoritairement heberge sur AWS/Azure/GCP). Le CADA cible exactement cette couche.",
    ]),
    ("5.11 Impacts transversaux sur les scenarios A-D", [
        "Le Pivot Cloud-Nationalite se superpose aux quatre scenarios, modifiant leurs conclusions de maniere non-lineaire. Il n'invalide pas la matrice 2x2 mais ajoute une troisieme dimension : le degre d'autonomie du compute installe. Le Tableau 13 synthetise l'impact.",
    ]),
    ("5.12 Implications pour la France : la question du compute reellement souverain", []),
    ("5.12.1 Le sovereignty washing comme risque systemique", [
        "Le terme sovereignty washing, popularise par Cristina Caffarra (Eurostack Foundation), designe la pratique des hyperscalers US qui commercialisent des offres sovereign cloud en implantant des centres de donnees sur sol europeen, tout en restant soumis au CLOUD Act.[27] Microsoft a reconnu dans sa propre documentation commerciale ne pas pouvoir garantir la souverainete pour les clients europeens en cas d'injonction US fondee en droit.[28]",
    ]),
    ("5.12.2 L'avantage nucleaire francais dans la nouvelle equation", [
        "La reinterpretation souverainiste du CACI renforce paradoxalement l'atout strategique francais. Si F(r) se decompose en F_phys x F_sov, alors la strategie optimale de la France n'est pas seulement d'augmenter F_phys (attirer plus de centres de donnees hyperscaler) mais d'augmenter F_sov (developper du compute independant des juridictions US).",
        "L'energie nucleaire francaise cree ici un avantage competitif de premier ordre : des Special Compute Zones adossees a une electricite nucleaire decarbonee et economiquement competitive, hebergeant des Gigafactories operees par des entites europeennes (OVHcloud, Scaleway, Mistral AI, IONOS), produiraient un compute avec un F_sov proche de 1, le seul compute genuinement hors d'atteinte des Cloud Sovereignty Mandates.",
    ]),
    ("5.12.3 Le projet DARE/RISC-V : de l'ambition a la necessite strategique", [
        "Dans le cadre du scenario D et a fortiori sous Cloud Sovereignty Mandates, le projet DARE (Digital Autonomy with RISC-V in Europe, EuroHPC JU, 2025) change de statut : il n'est plus une ambition de long terme mais une necessite strategique.[29] Tant que l'Europe depend exclusivement des GPU NVIDIA/AMD pour son compute IA, le controle US sur ces architectures cree une vulnerabilite residuelle meme dans les Gigafactories operees par des entites europeennes, une mise a jour firmware imposee par NVIDIA dans le cadre du programme location verification pourrait theoriquement degrader la performance des clusters europeens.",
    ]),
    ("5.13 Synthese : le quatrieme point de bascule et les conditions d'un decouplage maitrise", [
        "Le Grand Decouplage n'est pas inevitable. Il represente un risque systemique conditionnel dont l'activation depend des choix politiques US et de la vitesse de reponse europeenne. Aux trois points de bascule identifies en 5.7.2, un quatrieme s'ajoute.",
        "Quatrieme point de bascule : la mise en oeuvre des fonctionnalites de verification de localisation dans les GPU avances. Si le BIS et le Department of Commerce US, conformement a l'AI Action Plan de juillet 2025, parviennent a deployer des mecanismes d'attestation a distance dans les puces H100/B200/GB300 d'ici fin 2026-2027, le substrat technique du Pivot Cloud-Nationalite sera en place. La question ne sera plus de savoir si les Cloud Sovereignty Mandates sont techniquement faisables, mais uniquement s'ils sont politiquement decides. Ce point de bascule devrait etre surveille des 2026 : les premiers appels d'offres BIS/NIST sur les standards d'attestation a distance pour puces IA constitueront le signal precoce.",
        "Pour l'Europe et la France, la strategie d'un decouplage maitrise repose sur trois piliers interdependants : (i) accelerer le deploiement du compute souverain (Gigafactories UE sous juridiction europeenne) pour augmenter F_sov avant l'activation des Mandates ; (ii) securiser le statut Tier 1 dans l'alliance americaine pour maintenir un acces sans restriction aux puces avancees, en acceptant la coordination sur les controles a l'exportation ; (iii) investir dans les alternatives architecturales (DARE/RISC-V, Huawei Ascend comme alternative court terme pour les charges non sensibles) pour reduire la dependance aux GPU US a moyen terme.",
    ]),
]

FR.table_blocks = [
    ("Tableau 10. Matrice 2x2 des scenarios prospectifs 2026-2030.",
     "Source : construction de l'auteur, methodologie Schwartz (1991).",
     [
         ["", "Reponse UE reactive", "Reponse UE proactive"],
         ["Protectionnisme US modere",
          "Scenario A - Statu quo renforce (derive lente vers la dependance)",
          "Scenario C - Partenariat asymetrique (partenaire technologique junior occidental)"],
         ["Protectionnisme US agressif",
          "Scenario B - Fracture numerique (decouplage europeen structurel)",
          "Scenario D - Souverainete contestee (course a l'autonomie sous pression)"],
     ]),
    ("Tableau 11. Synthese comparative des quatre scenarios sur les six metriques de divergence (horizon 2030).",
     f"Source : construction de l'auteur ; baseline snapshot avril 2026 (compute brut operationnel US/UE {fmt_fr(us_eu_raw, 1)}:1, CACI Power Mode {fmt_fr(us_eu_caci, 2)}:1).",
     [
         ["Metrique (2030)", "A - Statu quo", "B - Fracture", "C - Partenariat", "D - Souverainete"],
         ["M1 Ratio compute brut US/UE (operationnel)", "18-22:1", "25-35:1", "8-10:1", "12-15:1"],
         ["M2 Ecart cout du FLOP", "2,4-3,2x", "4-6x", "1,5-2,0x", "1,8-2,5x"],
         ["M3 Part cloud US (pct)", "72-75", "78-82", "60-65", "50-55"],
         ["M4 Productivite UE (pct/an)", "+1,0-1,5", "+0,3-0,8", "+1,8-2,5", "+1,2-2,0"],
         ["M5 Energie UE (TWh)", "~115", "~95", "~140", "~155"],
         ["M6 Ratio CACI Power Mode", "4-5:1", "6-8:1", "2,0-2,5:1", "4-7:1 (post-creux)"],
         ["Probabilite estimee", "40-50 pct", "15-20 pct", "15-20 pct", "15-20 pct"],
     ]),
    ("Tableau 12. Estimation du facteur F_sov par juridiction et impact CACI sous activation des Cloud Sovereignty Mandates (2028).",
     "Source : construction de l'auteur ; Synergy Research Group (2025), Statista Enterprise Cloud (2025), et chapitre I Fig 1.8 pour le baseline souverain sur compute installe.",
     [
         ["Juridiction", "F_phys part cloud-US", "F_sov estime", "CACI actuel (baseline phys)", "Impact CACI post-Mandate"],
         ["Etats-Unis", "~5 pct (cloud domestique non affecte)", "1,00", "100 (reference)", "100 (inchange)"],
         ["UE (France, Allemagne)", "~77 pct (AWS/Azure/GCP dominent les charges UE)", "0,28", "28,9 (Power Mode)", "Effondrement 30-50 pct sur les charges"],
         ["EAU (hub Dubai)", "~88 pct (hyperscalers US dominants)", "0,12", "55,7 phys / 6,0 souv sur installe", "Effondrement 60-80 pct - hub illusoire"],
         ["Singapour", "~82 pct (hyperscalers US dominants)", "0,18", "eleve - hub APAC", "Effondrement 55-75 pct"],
         ["Chine", "~2 pct (Alibaba/Tencent/Huawei Cloud)", "0,98", "15,7 (penurie de puces plafonne)", "Inchange - souverainete deja effective"],
         ["Inde", "~60 pct (AWS/Azure + locaux)", "0,40", "22,2 (Power Mode)", "Effondrement modere - position intermediaire"],
     ]),
    ("Tableau 13. Impacts du Pivot Cloud-Nationalite (Cloud Sovereignty Mandates 2028) sur les quatre scenarios de la matrice 2x2.",
     "Source : construction de l'auteur.",
     [
         ["Scenario", "Sans Mandates", "Avec Mandates 2028", "Impact CACI UE", "Lecture strategique"],
         ["A - Statu quo", "Dependance lente, CACI 4-5:1",
          "Activation partielle - hyperscalers cooperent sans restriction majeure",
          "Degradation moderee 15-25 pct sur les charges",
          "Plus stable mais illusion de securite revelee"],
         ["B - Fracture", "Decouplage structurel, CACI 6-8:1",
          "Activation maximale - cloud US conditionnel et puces restreintes simultanement",
          "Effet ciseau double : puces rares + compute conditionnel. Ratio CACI potentiellement > 8:1",
          "Pire cas absolu - combo puces et cloud"],
         ["C - Partenariat asymetrique", "Rattrapage partiel, CACI 2,0-2,5:1",
          "Gigafactories souveraines absorbent le choc si F_sov UE monte a 0,45-0,55",
          "Impact limite si deploye dans les delais - Gigafactories = couverture souveraine",
          "Meilleure resilience - investissement prealable prouve sa valeur"],
         ["D - Souverainete contestee", "Rattrapage sous pression, CACI 4-7:1 post-creux",
          "Mandates deviennent catalyseur politique - accelerent deploiement UE et alliances JP/KR/TW",
          "Courbe en U acceleree - creux 2028-2029, rattrapage plus rapide",
          "Paradoxal : les Mandates peuvent accelerer la souverainete UE si la reponse est assez rapide"],
     ]),
]

FR.notes = EN.notes


# ===========================================================================
# Content - Brazilian Portuguese
# ===========================================================================

PT = LangPack(
    code="PT-BR",
    filename="Capitulo_V_Cenarios_Prospectivos_2026_2030_PT-BR.docx",
    cover_title="AI FOR AMERICANS FIRST",
    cover_subtitle="Protecionismo de IA, Energia e Semicondutores: Trajetorias de Divergencia EUA/Europa 2024-2030",
    cover_blurb="Analise geoestrategica e economica integrada - Capitulo V",
    cover_chip_lines=[
        f"{fmt_fr(us_share, 1)} pct do compute IA operacional mundial = EUA",
        "1,59x custo de energia UE/EUA (ajustado-PPA)",
        f"{fmt_fr(us_eu_caci, 2)}:1 razao CACI EUA/UE (Power Mode)",
    ],
    cover_meta="Paris - fevereiro de 2026  |  7 capitulos  |  4 cenarios prospectivos  |  3 zonas geograficas",
    cover_keywords_label="Palavras-chave",
    cover_keywords=("inteligencia artificial, protecionismo tecnologico, semicondutores, "
                    "controles de exportacao, compute soberano, geopolitica da IA, Franca, "
                    "Estados Unidos, China"),
    chapter_label="CAPITULO V",
    chapter_title="Cenarios prospectivos 2026-2030",
    chapter_intro=(
        "Este capitulo constitui o coracao da contribuicao original deste estudo. "
        "Ao aplicar o protocolo metodologico definido no Capitulo II (matriz 2x2, "
        "seis metricas de divergencia, calibracao CACI), construimos quatro cenarios "
        "de evolucao da relacao transatlantica em IA, energia e semicondutores para "
        "o periodo 2026-2030. Cada cenario e determinado pela combinacao de duas "
        "incertezas criticas identificadas no Capitulo III: o grau de protecionismo "
        "americano e a capacidade de resposta estrategica europeia. Avaliamos entao "
        "cada cenario em suas seis metricas, antes de sintetizar as condicoes de "
        "bascula entre trajetorias."
    ),
    notes_label="Notas",
    license_block=[
        "Licenca e isencao de responsabilidade. Esta obra, 'AI for Americans First', e disponibilizada sob os termos da Licenca Creative Commons Atribuicao - NaoComercial - CompartilhaIgual 4.0 Internacional (CC BY-NC-SA 4.0).",
        "Voce e livre para compartilhar e adaptar o material para fins nao comerciais, desde que credite adequadamente a obra a Fabrice Pizzi (Universite Paris-Sorbonne) e distribua suas contribuicoes sob a mesma licenca. Este documento e fornecido apenas para fins educacionais e de pesquisa.",
        "Painel publico: https://mo0ogly.github.io/America-First-IA/dashboard/",
        "Repositorio: https://github.com/mo0ogly/America-First-IA",
    ],
    page_footer="AI for Americans First - Fabrice Pizzi - Capitulo V",
)

PT.sections = [
    ("5.1 Elementos predeterminados: o que nao mudara", [
        "De acordo com o metodo Schwartz (1991), distinguimos os elementos predeterminados (tendencias quase certas no horizonte 2030) das incertezas criticas (fatores cuja evolucao depende de decisoes politicas ainda nao tomadas). Quatro elementos predeterminados estruturam todos os cenarios.",
        "EP1 - Crescimento exponencial da demanda de compute IA. As vendas de semicondutores dobraram em dois anos (2023-2025), a potencia dos chips de IA instalados dobra a cada sete meses (Epoch AI), e nenhum sinal de desaceleracao e observavel em fevereiro de 2026. Mesmo sob a hipotese de uma desaceleracao nas scaling laws (saturacao Chinchilla), a difusao da IA em direcao a inferencia, robotica e agentes autonomos mantera uma demanda de compute em forte crescimento.[1]",
        f"EP2 - Concentracao persistente do compute nos Estados Unidos. A razao EUA/UE de {fmt_fr(us_eu_raw, 1)}:1 em compute instalado bruto (Capitulo III, snapshot do painel de abril de 2026: 2.759.968 vs 156.632 PFLOP/s), resultando em uma razao CACI Power Mode de {fmt_fr(us_eu_caci, 2)}:1 uma vez ponderada pela formula geometrica F^0,40 x L^0,20 x R^0,15 / E^0,25, reflete decisoes de investimento tomadas em 2022-2025 cujos efeitos se materializam ate 2028-2029 (atrasos na construcao de data centers: 18-36 meses). Mesmo uma reversao politica imediata nao alteraria o estoque instalado antes do final da decada.",
        "EP3 - Tensao energetica crescente. O consumo global dos data centers, estimado em 415 TWh em 2024, atingira 800-950 TWh ate 2030, de acordo com as projecoes da IEA (Capitulo III). A assimetria nos custos de energia (EUA 1,4-1,7x mais baratos apos correcao PPA, 2-3x mais baratos nas tarifas Eurostat industriais nao ajustadas) persistira, a menos que haja um investimento massivo na energia nuclear europeia, cujos atrasos de implantacao (SMR: 5-7 anos para os primeiros reatores) excedem o horizonte de 2030.[2]",
        "EP4 - Estrutura regulatoria da Secao 232 em vigor. A Proclamacao 11002 de 14 de janeiro de 2026 e um fato consumado juridico. Ao contrario das tarifas IEEPA (anuladas pela Suprema Corte em 20 de fevereiro de 2026[3]), as tarifas da Secao 232 repousam sobre uma base legal confirmada. O relatorio do Secretario do Comercio sobre o mercado de semicondutores para data centers e esperado para 1 de julho de 2026, e pode recomendar uma extensao ou modificacao das tarifas. Independentemente da direcao tomada, o instrumento legal permanecera disponivel.[4]",
    ]),
    ("5.2 Incertezas criticas e matriz 2x2", []),
    ("5.2.1 Eixo 1: Grau de protecionismo americano", [
        "A primeira incerteza diz respeito a evolucao da politica americana entre dois polos. O polo moderado corresponde a manutencao do status quo de janeiro de 2026: tarifa de 25 por cento limitada a chips avancados reexportados, isencoes domesticas amplas, acordo comercial EUA-UE limitando as tarifas de semicondutores a 15 por cento, e nenhuma extensao significativa para nuvem ou modelos. O cerne da UE (Franca, Alemanha) permanece na categoria de 'parceiros confiaveis'. O polo agressivo pressupoe uma extensao apos o relatorio de julho de 2026: tarifas estendidas a semicondutores derivados e equipamentos, quotas de GPU para a UE (incluindo a Franca), condicoes restritivas para o acesso a nuvem de IA de fronteira, e o uso de compute como alavanca de negociacao comercial (compute-for-concessions).[5]",
    ]),
    ("5.2.2 Eixo 2: Capacidade de resposta estrategica europeia", [
        "A segunda incerteza diz respeito a capacidade da UE de implantar uma resposta coerente e rapida. O polo reativo corresponde a respostas nacionais fragmentadas, investimentos dispersos, um AI Act criando custos de conformidade adicionais e implantacao lenta de AI Factories/Gigafactories (atrasos burocraticos, autorizacoes de mais de 24 meses). O polo proativo pressupoe a implementacao acelerada do programa AI Continent (19 AI Factories mais ate 5 Gigafactories de 100.000+ GPUs), a adocao de Special Compute Zones (autorizacao em 180 dias), a mobilizacao efetiva do fundo InvestAI (20 bilhoes de EUR) e a uniao da capacidade nuclear francesa como vantagem competitiva.[6]",
    ]),
    ("5.2.3 Matriz e nomeacao dos cenarios", [
        "A interseccao desses dois eixos produz quatro cenarios. UE reativa + protecionismo moderado dos EUA resulta no Cenario A (Status Quo Reforcado, deriva lenta para a dependencia). UE reativa + EUA agressivos resulta no Cenario B (Fratura Digital, desacoplamento europeu estrutural). UE proativa + protecionismo moderado dos EUA resulta no Cenario C (Parceria Assimetrica, parceiro tecnologico junior ocidental). UE proativa + EUA agressivos resulta no Cenario D (Soberania Contestada, corrida pela autonomia sob pressao).",
    ]),
    ("5.3 Cenario A - Status Quo Reforcado (Protecionismo moderado + UE reativa)", []),
    ("5.3.1 Narrativa", [
        "Apos o relatorio de julho de 2026, o Secretario do Comercio recomenda manter a tarifa de 25 por cento sobre chips avancados reexportados, mas sem estende-la significativamente. O acordo comercial EUA-UE de agosto de 2025 e respeitado: as tarifas de semicondutores para a UE permanecem limitadas a 15 por cento.[7] A UE, tranquilizada por esse status quo, retarda a implantacao de suas proprias iniciativas. As AI Factories do EuroHPC lutam para atingir a capacidade nominal (atrasos nas autorizacoes, coordenacao interestadual). As Gigafactories sao adiadas para 2029-2030. O fundo InvestAI e parcialmente mobilizado (8-10 bilhoes de EUR de um total de 20). As empresas europeias continuam a depender fortemente da nuvem dos EUA, cujo desempenho e custos permanecem imbativeis.",
    ]),
    ("5.3.2 Trajetoria das metricas", [
        f"M1 - Razao de compute (GPUs instaladas EUA/UE): passa de {fmt_fr(us_eu_raw, 1)}:1 bruto (2025) para 18-22:1 bruto (2030) no compute instalado operacional. O gap aumenta ligeiramente a medida que os investimentos nos EUA aceleram (Stargate, mega-clusters xAI, Meta), enquanto a UE adiciona apenas as 19 AI Factories (max 25.000 GPUs cada, aprox. 475.000 GPUs publicas, uma ordem de grandeza abaixo de um unico hyperscaler dos EUA).[8]",
        "M2 - Diferenca de custo do FLOP (UE/EUA): permanece na faixa de 2,4-3,2x. A ausencia de tarifas agressivas sobre a UE mantem o acesso a nuvem dos EUA a precos proximos aos niveis atuais, mas os custos de energia europeus continuam a pesar.",
        "M3 - Parcela da nuvem dos EUA nos gastos de IA europeus: passa de 70 por cento (2024) para 72-75 por cento (2030). Os provedores europeus (OVHcloud, Deutsche Telekom) mantem seus 15 por cento no segmento de soberania, mas nao ganham terreno nos servicos de IA generativa.",
        "M4 - Produtividade da IA (por cento/ano): EUA +2,5-3,0; UE +1,0-1,5. A UE realiza parte do potencial de IA via aplicacoes a jusante (SAP, Siemens, fintech), mas a adocao lenta e o deficit de compute limitam os ganhos.",
        "M5 - Dependencia energetica (data center TWh): UE aprox. 115 TWh em 2030 (+65 por cento vs 2024). A energia nuclear francesa absorve parte da demanda, mas a ausencia de Special Compute Zones atrasa a conexao a rede de novos data centers.",
        f"M6 - CACI(EUA)/CACI(UE): passa de {fmt_fr(us_eu_caci, 2)}:1 (abril de 2026) para 4-5:1 (2030). O gap aumenta moderadamente a medida que o fator F (compute) se acumula no lado dos EUA, enquanto os custos E (energia) pesam no lado da UE.",
    ]),
    ("5.3.3 Consequencias para a Franca", [
        "Este cenario e o mais provavel no curto prazo (probabilidade estimada: 40-50 por cento). E tambem o mais insidioso: a ausencia de um choque visivel desmobiliza os atores europeus, enquanto a dependencia se aprofunda estruturalmente. As empresas francesas beneficiam-se do acesso a nuvem dos EUA para a adocao de IA (BNP Paribas, Airbus, TotalEnergies via AWS/Azure), mas essa adocao reforca o bloqueio descrito no Capitulo IV. O deficit de produtividade da IA em relacao aos Estados Unidos (-1,0 a -1,5 pontos por ano) acumula-se ao longo de cinco anos, ampliando a lacuna de competitividade em 5 a 8 pontos do PIB.",
    ]),
    ("5.4 Cenario B - Fratura Digital (Protecionismo agressivo + UE reativa)", []),
    ("5.4.1 Narrativa", [
        "O relatorio de julho de 2026 leva a uma extensao significativa. O Secretario do Comercio recomenda tarifas estendidas a equipamentos de semicondutores e produtos derivados, com um programa de compensacao tarifaria reservado para empresas que investem na producao americana.[9] O acordo de 15 por cento da UE e revisado para cima ou acompanhado de condicoes restritivas (quotas de volume em GPUs avancadas, requisitos de reciprocidade no AI Act). Simultaneamente, o acesso a nuvem de IA de fronteira e tornado condicional para entidades nao americanas (limitacoes de acesso a API para modelos de fronteira, restricoes de peso). A UE, fragmentada, nao consegue formular uma resposta coerente: os Estados-membros dividem-se entre acomodacao (paises nordicos, Holanda) e confronto (Franca, Italia).",
    ]),
    ("5.4.2 Trajetoria das metricas", [
        f"M1 - Razao de compute: passa de {fmt_fr(us_eu_raw, 1)}:1 bruto (2025) para 25-35:1 bruto (2030). As quotas de GPU limitam as importacoes europeias no momento em que a demanda explode. Os projetos de AI Factory sao comprometidos pela incapacidade de adquirir GPUs Nvidia/AMD nos volumes planejados.",
        "M2 - Diferenca de custo do FLOP: salta para 4-6x. Tarifas estendidas, combinadas com quotas e assimetria energetica, aumentam massivamente os custos de compute europeus. As empresas francesas enfrentam uma sobretaxa de 3x a 5x para treinamento de modelos.",
        "M3 - Parcela da nuvem dos EUA: paradoxalmente sobe para 78-82 por cento. Na falta de uma alternativa local credivel, as empresas europeias que desejam acesso a IA de fronteira devem passar pelos hyperscalers dos EUA, nas condicoes de preco que eles ditam. Servicos soberanos (OVHcloud, Scaleway) carecem de hardware para oferecer servicos competitivos de GenAI.",
        "M4 - Produtividade da IA: EUA +2,5-3,5; UE +0,3-0,8. O potencial de IA europeu e severamente restrito. O McKinsey Global Institute estima que, com adocao lenta, a produtividade europeia nao passaria de 0,3 por cento, proximo a estagnacao.[10]",
        "M5 - Dependencia energetica: UE aprox. 95 TWh apenas (2030), nao por virtude, mas por padrao - a falta de GPUs limita a construcao de data centers. Ironicamente, a restricao de compute mitiga a restricao de energia.",
        f"M6 - Razao CACI: explode de {fmt_fr(us_eu_caci, 2)}:1 (abril de 2026) para 6-8:1 (2030). Este e o cenario onde o gap e maior, com todos os tres fatores CACI deteriorando-se simultaneamente no lado europeu: F limitado por quotas, E inflado por tarifas, L enfraquecido por uma fuga de cerebros acelerada para os Estados Unidos.",
    ]),
    ("5.4.3 Consequencias para a Franca", [
        "Este cenario (probabilidade estimada: 15-20 por cento) representa o pior caso. A Franca sofre um desacoplamento tecnologico estrutural: projetos intensivos em compute (modelos de fundacao Mistral, robotica Exotec, simulacoes Dassault) sao realocados para os Estados Unidos ou dependem de um acesso cada vez mais caro a nuvem dos EUA. O time-to-market das solucoes de IA francesas aumenta de 25 para 40 por cento. As PMEs, incapazes de absorver sobretaxas, renunciam a IA de fronteira e optam por solucoes degradadas (modelos de codigo aberto menores, inferencia local). O gap de produtividade acumulado com os Estados Unidos atinge 10 a 15 pontos em cinco anos.",
    ]),
    ("5.5 Cenario C - Parceria Assimetrica (Protecionismo moderado + UE proativa)", []),
    ("5.5.1 Narrativa", [
        "O protecionismo dos EUA permanece moderado (como em A), mas a UE explora esta janela para acelerar seus proprios investimentos. As AI Factories sao implantadas no cronograma (2026-2027), as primeiras Gigafactories de 100.000+ GPUs sao encomendadas no final de 2026 e entregues em 2028.[11] A Franca desempenha um papel central gracas a sua frota nuclear (65-70 por cento da matriz eletrica, custo marginal competitivo), e Special Compute Zones sao designadas em antigos locais industriais com conexoes de rede pesadas.[12] No entanto, a UE aceita de facto um status de parceiro tecnologico junior: utiliza GPUs Nvidia/AMD (nenhum campeao europeu de design de ASIC IA), depende de fundicoes TSMC/Samsung/Intel para producao, e seus modelos de fundacao permanecem um degrau abaixo dos lideres dos EUA.",
    ]),
    ("5.5.2 Trajetoria das metricas", [
        f"M1 - Razao de compute: cai de {fmt_fr(us_eu_raw, 1)}:1 bruto (2025) para 8-10:1 bruto (2030) no compute instalado. Gigafactories e investimento privado (InvestAI mais co-investimentos industriais) adicionam 1-2 milhoes de equivalentes H100 na Europa, reduzindo o gap sem fecha-lo.",
        "M2 - Diferenca de custo do FLOP: cai para 1,5-2,0x. A energia nuclear francesa e as economias de escala das Gigafactories comprimem os custos de energia e infraestrutura, embora persista um gap residual (ausencia de design de GPU proprietario).",
        "M3 - Parcela da nuvem dos EUA: cai ligeiramente para 60-65 por cento. Servicos soberanos europeus ganham participacao de mercado em segmentos regulamentados (defesa, saude, financas), enquanto a nuvem dos EUA mantem a maioria das cargas comerciais. O mercado segmenta-se em soberano e desempenho.",
        "M4 - Produtividade da IA: EUA +2,5-3,0; UE +1,8-2,5. A UE atinge 60-80 por cento de seu potencial teorico gracas ao compute local suficiente para a adocao em larga escala de aplicacoes a jusante, mesmo que o treinamento de modelos de fronteira permaneca dependente do hardware dos EUA.",
        "M5 - Energia: UE aprox. 140 TWh (2030). A demanda e maior do que em A porque o compute europeu aumenta, mas a energia nuclear e os SMRs planejados absorvem a maior parte. A RTE France confirma a viabilidade de +10 GW sujeitos a investimentos na rede.",
        f"M6 - Razao CACI: cai de {fmt_fr(us_eu_caci, 2)}:1 (abril de 2026) para 2,0-2,5:1 (2030). Este e o cenario mais favoravel realisticamente alcancavel no horizonte de 2030. O fator F melhora significativamente, E beneficia-se da energia nuclear, mas L permanece um pouco abaixo (o ecossistema de IA dos EUA sendo mais atraente para os melhores talentos).",
    ]),
    ("5.5.3 Consequencias para a Franca", [
        "Este cenario (probabilidade estimada: 15-20 por cento) e o mais favoravel para a Franca no curto a medio prazo. A Franca torna-se o hub de energia de IA da UE gracas a sua frota nuclear, atraindo investimentos em data centers e Gigafactories. As empresas francesas ganham acesso a compute local competitivo para inferencia e fine-tuning, reduzindo a dependencia da nuvem dos EUA para usos padrao. Mistral e startups francesas podem treinar modelos especializados localmente. No entanto, o treinamento de modelos de fronteira permanece dependente do hardware dos EUA, e a autonomia estrategica e parcial: a Franca e soberana na aplicacao, mas nao na criacao de tecnologias fundamentais.",
    ]),
    ("5.6 Cenario D - Soberania Contestada (Protecionismo agressivo + UE proativa)", []),
    ("5.6.1 Narrativa", [
        "O protecionismo dos EUA intensifica-se (como em B), mas a UE responde com determinacao. A ameaca americana torna-se o catalisador politico para uma mobilizacao industrial europeia sem precedentes desde o projeto AIRBUS da decada de 1970. O programa AI Continent e acelerado e estendido: as 5 Gigafactories sao encomendadas em carater de emergencia, a Franca anuncia 20 GW de capacidade nuclear dedicada a data centers de IA ate 2032 (combinando extensao da frota, novos EPR2s e SMRs), o projeto DARE (RISC-V europeu) e escalado para projetar aceleradores de IA reduzindo a dependencia da Nvidia.[13] Simultaneamente, a UE negocia aliancas tecnologicas alternativas (Japao, Coreia do Sul, Taiwan) para garantir suprimentos de GPU e fundicao.",
    ]),
    ("5.6.2 Trajetoria das metricas", [
        f"M1 - Razao de compute: evolui de {fmt_fr(us_eu_raw, 1)}:1 bruto (2025) para 12-15:1 bruto (2030). A UE investe massivamente, mas comeca de muito longe. As quotas dos EUA retardam as importacoes, mas aliancas alternativas e producao local (Gigafactories usando GPUs Samsung/Intel como alternativas a Nvidia) compensam parcialmente.",
        "M2 - Diferenca de custo do FLOP: 2,5-4,0x inicialmente (2027, pico do choque tarifario), depois reducao progressiva para 1,8-2,5x (2030) a medida que as Gigafactories escalam e as alternativas de GPU amadurecem.",
        "M3 - Parcela da nuvem dos EUA: cai para 50-55 por cento (2030), o declinio mais acentuado dos quatro cenarios. O desafio geopolitico e as restricoes dos EUA empurram as empresas europeias para alternativas soberanas, mesmo que imperfeitas. Hyperscalers dos EUA perdem terreno em segmentos regulamentados.",
        "M4 - Produtividade da IA: EUA +2,5-3,5; UE +1,2-2,0. A UE atravessa um vale de produtividade em 2027-2028 (periodo de transicao onde as restricoes dos EUA mordem, mas os investimentos europeus ainda nao estao operacionais), depois uma recuperacao parcial a partir de 2029.",
        "M5 - Energia: UE aprox. 150-160 TWh (2030). Este e o cenario mais intensivo em energia para a UE, com a construcao massiva de data centers locais criando uma demanda enorme. A energia nuclear francesa torna-se um ativo estrategico continental, mas a pressao na rede esta no maximo.",
        f"M6 - Razao CACI: segue uma trajetoria em forma de U: degradacao de {fmt_fr(us_eu_caci, 2)}:1 para 8-12:1 em 2027-2028 (pico do choque), depois melhora para 4-7:1 ate 2030. O resultado depende fortemente da velocidade de execucao europeia: cada ano de atraso nas Gigafactories prolonga o periodo de vulnerabilidade maxima.",
    ]),
    ("5.6.3 Consequencias para a Franca", [
        "Este cenario (probabilidade estimada: 15-20 por cento) e o mais ambicioso e arriscado. Coloca a Franca no coracao de um esforco sem precedentes de soberania tecnologica europeia. Investimentos nucleares massivos (SMR, extensao da frota) tornam-se uma questao geopolitica de primeira ordem. O projeto DARE/RISC-V poderia, se bem-sucedido, constituir a primeira alternativa europeia credivel as GPUs Nvidia para IA, mas em um horizonte de 5 a 7 anos, bem alem de 2030. No curto prazo (2026-2028), a Franca atravessa um periodo de vulnerabilidade maxima onde sobretaxas e escassez degradam a competitividade, antes de uma recuperacao condicional a velocidade de implantacao da infraestrutura.",
    ]),
    ("5.7 Sintese comparativa e pontos de inflexao", []),
    ("5.7.1 Tabela de sintese das metricas", [
        f"A Tabela 11 abaixo consolida a trajetoria das seis metricas de divergencia para os quatro cenarios no horizonte de 2030, ancorada no snapshot do painel de abril de 2026 (compute operacional bruto EUA/UE {fmt_fr(us_eu_raw, 1)}:1, CACI Power Mode {fmt_fr(us_eu_caci, 2)}:1).",
    ]),
    ("5.7.2 Pontos de inflexao entre os cenarios", [
        "A trajetoria real provavelmente seguira um caminho hibrido entre esses cenarios. Tres pontos de inflexao determinam as transicoes possiveis.",
        "Primeiro ponto de inflexao: o relatorio de Comercio de julho de 2026. Este relatorio determinara se o protecionismo dos EUA se estendera (mudanca para B ou D) ou permanecera direcionado (manutencao em A ou C). Indicadores a serem observados: evolucao do deficit comercial de semicondutores dos EUA, taxas de preenchimento de fabricas do CHIPS Act (Intel, TSMC Arizona, Samsung Taylor), pressao politica interna (midterms de 2026). Os resultados das negociacoes da Fase 1 (relatorio previsto para 14 de abril de 2026) serao um sinal precoce.[14]",
        "Segundo ponto de inflexao: velocidade de implantacao da Gigafactory da UE. A Comissao espera que as primeiras Gigafactories estejam operacionais em 2027-2028. Se este cronograma for cumprido, a UE mudara para cenarios proativos (C ou D). Se os atrasos de autorizacao, financiamento ou fornecimento de hardware empurrarem as entregas para 2029-2030, a UE permanecera em modo reativo (A ou B). A proposta da CFG para Special Compute Zones (autorizacao em 180 dias vs 24+ meses atualmente) e o fator chave de aceleracao.[15]",
        "Terceiro ponto de inflexao: a decisao francesa sobre nuclear para IA. A Franca possui um ativo unico na Europa: uma frota nuclear que fornece 65-70 por cento da eletricidade a um custo marginal globalmente competitivo. A decisao de dedicar uma capacidade significativa (10-20 GW) a data centers de IA, via extensao da frota, novos EPR2s e SMRs, determinara se a Franca se tornara o hub de energia de IA da Europa ou perdera esta posicao para outros (Escandinavia com hidreletrica, Europa Oriental com baixos custos de terra). Este ponto de inflexao e especificamente frances e determina a posicao da Franca dentro dos cenarios europeus.[16]",
    ]),
    ("5.7.3 O ponto de convergencia: 2028", [
        "Os quatro cenarios convergem em um ponto critico comum em 2028. Este e o ano em que: (i) a demanda por compute excedera a capacidade instalada na Europa, criando gargalos de hardware (mesmo sob protecionismo moderado); (ii) os efeitos totais das tarifas estendidas (se adotadas) serao sentidos; (iii) as Gigafactories, se implantadas a tempo, comecarao a produzir compute local significativo; (iv) a demanda de energia dos data centers saturara a capacidade de conexao da rede em varios Estados-membros. O ano de 2028 constitui o momento da verdade em que a Europa descobrira se esta em uma trajetoria A/B (dependencia crescente) ou C/D (recuperacao iniciada). Decisoes tomadas em 2026-2027 (relatorio de Comercio dos EUA, Gigafactories, nuclear frances) estarao irreversivelmente engajadas.",
    ]),
    ("5.8 Origens do ponto de inflexao: fundamentos juridico-tecnicos", [
        "O cenario do 'Grande Desacoplamento' nao e construido ex nihilo. E a projecao logica, no horizonte de 2028, de uma arquitetura de controle cujas camadas fundamentais ja estao operacionais em 2026. Identificar essas camadas e uma questao de rigor academico: distinguir tendencias documentadas de projecoes extrapoladas.",
    ]),
    ("5.8.1 A camada legal: extraterritorialidade como instrumento estrutural", [
        "A primeira camada e legal e precede amplamente a politica de IA. O CLOUD Act (Clarifying Lawful Overseas Use of Data Act, 2018) estabelece que qualquer provedor de nuvem sujeito a jurisdicao dos EUA e obrigado a produzir dados, independentemente de onde tais comunicacoes, registros ou outras informacoes sejam armazenados.[17] Esta lei, confirmada pela jurisprudencia federal (United States v. Microsoft), cria uma dissociacao fundamental entre a localizacao fisica dos dados e sua nacionalidade legal.",
        "A consequencia imediata para o compute IA e radical: um cluster H100 fisicamente localizado em Dubai, Cingapura ou em um centro AWS eu-west-1 na Irlanda permanece legalmente americano. Se ordenado pelo governo dos EUA, o operador (AWS, Azure, Google Cloud) e legalmente obrigado a cumprir, independentemente da vontade do cliente ou da legislacao local. A Microsoft reconheceu perante um tribunal frances em 2024 que nao poderia garantir a soberania dos dados para clientes europeus em caso de uma injuncao dos EUA juridicamente fundamentada.[18]",
        "Esta arquitetura legal preexistente e o substrato sobre o qual o Pivot de Nuvem-Nacionalidade se enxerta: os Cloud Sovereignty Mandates projetados para 2028 nao criam um novo poder. Eles ativam e sistematizam um poder jurisdicional existente, estendendo-o a camada de compute operacional.",
    ]),
    ("5.8.2 A camada tecnica: da verificacao de localizacao ao estrangulamento do cluster", [
        "A segunda camada e tecnica. O Plano de Acao de IA dos EUA de 23 de julho de 2025 introduz explicitamente o conceito de 'recursos de verificacao de localizacao' aplicados a chips de IA avancados. Michael Kratsios, diretor da OSTP na Casa Branca, confirmou discussoes sobre modificacoes de software ou hardware para permitir um melhor rastreamento de localizacao, explicitamente incluido no plano.[19]",
        "Este anuncio nao e retorico. O BIS ja possui, desde o Framework for Artificial Intelligence Diffusion (janeiro de 2025), uma arquitetura de controle de compute-tiering para paises de destino (Tier 1-2-3) e limites de compute por entidade e pais.[20] A BIS Affiliates Rule, suspensa por um ano em novembro de 2025, mas mantida em principio, estipula que a afiliacao a uma entidade controladora em um pais restrito e suficiente para negar o acesso a compute avancado, independentemente da localizacao fisica.[21]",
        "A trajetoria tecnologica em direcao a 2028 e, portanto: (i) chips equipados com mecanismos de verificacao de localizacao, (ii) relatorios automaticos ao BIS em caso de desvio, (iii) a capacidade de suspender o acesso ou estrangular o desempenho via licenca de exportacao. Isso nao e ficcao cientifica: e a extensao ao compute de um principio ja aplicado ao software (sancoes da OFAC, congelamento de acesso ao servico).",
        "Limitacao critica: Estrangular clusters de producao e tecnicamente complexo e potencialmente perturbador para os proprios operadores. A NVIDIA nao possui atualmente um mecanismo de desativacao remota para GPUs H100/B200 em producao. Tal mecanismo exigiria modificacoes arquiteturais significativas no firmware e nos protocolos de atestacao remota (via TPM ou equivalente). O cenario de 2028 e, portanto, condicional a implementacao efetiva dessas modificacoes ao longo de 24 a 36 meses.",
    ]),
    ("5.9 O mecanismo de Pivot de Nuvem-Nacionalidade", []),
    ("5.9.1 O gatilho: Cloud Sovereignty Mandates como extensao dos controles de exportacao", [
        "O cenario do ponto de inflexao de 2028 assume que os Estados Unidos cruzam um limiar qualitativo: passando do controle de acesso ao hardware (controles de chip do BIS) para o controle de acesso ao servico de compute operacional (Cloud Sovereignty Mandates). Este limiar nao e uma ruptura arbitraria; ele responde a uma falha estrutural identificada no atual regime de controle.",
        "Esta falha esta documentada: apesar das restricoes do BIS, investigacoes revelaram que aprox. 1 bilhao de USD de chips Nvidia chegaram a China via terceiros paises (Malasia, Emirados Arabes Unidos, Cingapura) nos primeiros meses de 2025.[22] A resposta do Plano de Acao de IA (verificacao de localizacao e monitoramento de cluster) e o primeiro passo para o controle continuo pos-exportacao.",
        "O gatilho plausivel em 2028 e uma ordem executiva estendendo as obrigacoes do Framework for AI Diffusion para a camada de nuvem. Imporia uma certificacao de 'Residencia de Dados e Conformidade de Jurisdicao' a todos os hyperscalers dos EUA que operam clusters offshore avancados, com a revogacao do acesso ao compute em solo americano como mecanismo de execucao.",
    ]),
    ("5.9.2 Dissociando Fator Fisico / Fator Soberano no CACI", [
        "E aqui que o cenario produz seu impacto mais significativo no modelo CACI. Sob o Pivot de Nuvem-Nacionalidade, a variavel de compute F(r) decompoe-se em dois componentes: F(r) = F_phys(r) x F_sov(r), onde F_phys e o compute fisicamente instalado e F_sov e o fator de soberania operacional (fracao de F_phys fora da jurisdicao dos EUA).",
        "Observe a distincao entre este F_sov dinamico de 2028 e o CACI Soberano estatico introduzido no Capitulo I (Fig 1.8). O CACI Soberano do Capitulo I captura quem e o dono do compute hoje (snapshot de abril de 2026). O F_sov de 2028 captura quem controla o compute sob um regime hipotetico de Mandatos, que depende da parcela da carga de trabalho na nuvem, nao da capacidade instalada. Os dois indicadores concordam onde o compute e de propriedade de operadores domesticos; divergem onde os clusters locais sao de propriedade de operadores do lado dos EUA (EAU 99,6%, cargas de trabalho em nuvem da UE majoritariamente em AWS/Azure/GCP).",
        "Estimativas calibradas para 2028, sob a hipotese de ativacao do Mandato, sao apresentadas na Tabela 12.",
    ]),
    ("5.10 Surgimento de blocos de IA jurisdicionais (2028-2030)", [
        "O Grande Desacoplamento nao produz um mundo binario EUA/nao-EUA. Produz fragmentacao em blocos de intensidade variavel, de acordo com a capacidade de cada zona de desenvolver compute soberano credivel. Surgem quatro blocos.",
    ]),
    ("5.10.1 O Bloco Americano Estendido (American AI Alliance)", [
        "O Plano de Acao de IA de 23 de julho de 2025 estabelece explicitamente as bases para uma American AI Alliance: exportar a pilha tecnologica completa dos EUA (hardware, modelos, software, padroes) para aliados alinhados em troca de controles de exportacao alinhados.[23] A estrategia e 'cenoura e pau': aliados alinhados acessam chips e modelos de fronteira sem restricoes adicionais; aqueles que recusam sao expostos aos mecanismos da Foreign Direct Product Rule e a tarifas secundarias.",
        "Membros Tier 1 confirmados: EUA, Reino Unido, Canada, Australia, Japao, Coreia do Sul, Holanda, Alemanha, Franca (sujeito ao alinhamento do controle de exportacao). Para esses paises, o F_sov efetivo aumenta: suas entidades acessam os hyperscalers dos EUA sem restricao, e o compute soberano em desenvolvimento (Gigafactories da UE) recebe tratamento preferencial.",
    ]),
    ("5.10.2 O Bloco Soberano Eurasiano", [
        "A China constitui o unico exemplo completo de um bloco soberano preexistente. Com um F_sov de aprox. 0,98 e um ecossistema de nuvem nacional (Alibaba, Tencent, Huawei) operando fora da jurisdicao dos EUA, os Cloud Sovereignty Mandates nao tem tracao direta. A restricao chinesa continua sendo a escassez de chips avancados, mas o bloco americano nao pode estrangular um cluster Huawei Ascend 910B.",
        "Dinamica pos-2028: a China detem o unico compute soberano em larga escala fora do bloco americano. Os paises que buscam se emancipar dos Mandatos enfrentam uma escolha binaria: compute americano condicional ou compute chines sob outras formas de dependencia.",
    ]),
    ("5.10.3 Nao alinhados digitais: uma posicao insustentavel", [
        "India, Brasil, Sudeste Asiatico e paises do Golfo (sem tratados especiais com os EUA) constituem um bloco de nao alinhamento digital. Sua posicao e estruturalmente desconfortavel: dependentes demais dos hyperscalers dos EUA para mudar para a soberania, insuficientemente integrados na alianca dos EUA para escapar de restricoes em caso de desacordo geopolitico.",
        "O caso dos Emirados Arabes Unidos ilustra isso com forca quantitativa. A Fig 1.8 documentou que 99,6% do F_total dos EAU e de propriedade de atores do lado dos EUA (Stargate UAE, Microsoft, OpenAI), colapsando o CACI Soberano de um Fisico de 55,7 para apenas 6,0. Dubai investiu massivamente para se tornar um hub de IA, notadamente via acordos com AWS/Microsoft/G42. No entanto, a G42 ja foi pressionada em 2024 a romper lacos com entidades chinesas como condicao para o acesso a chips avancados.[24] Sob Mandatos, essa pressao torna-se sistemica: o compute dos hubs do Golfe, fisicamente presente, mas legalmente americano, tornar-se-ia uma alavanca de negociacao permanente.",
    ]),
    ("5.10.4 O Bloco Europeu: entre alianca e autonomia", [
        "A Europa ocupa uma posicao intermediaria e em evolucao. Legalmente Tier 1, a UE mantém, no entanto, uma ambicao de autonomia estrategica que a Alianca Americana nao satisfaz plenamente.",
        "O Cloud and AI Development Act (CADA), esperado para o primeiro trimestre de 2026, tenta resolver isso definindo um 'Nivel de Soberania da UE' que excluiria estruturalmente os provedores sujeitos ao CLOUD Act das compras publicas sensiveis.[25] O Cloud Sovereignty Framework de outubro de 2025 define tres niveis de garantia, com o SOV-3 exigindo que o provedor esteja alem do alcance da legislacao extraterritorial nao europeia.[26]",
        "Esta arquitetura esta em construcao, mas seu tempo e problematico: o CADA estara, na melhor das hipoteses, operacional em 2027-2028, precisamente quando os Mandatos dos EUA poderao ser ativados. A janela de vulnerabilidade e maxima.",
        f"Nota do Capitulo I: a UE e amplamente soberana em compute instalado ({fmt_fr(eu_sov, 1)}% do F_total e de propriedade da UE). A vulnerabilidade nao e no F instalado, mas na camada de carga de trabalho em nuvem (o compute realmente usado por empresas da UE, majoritariamente em AWS/Azure/GCP). O CADA visa exatamente esta camada.",
    ]),
    ("5.11 Impactos transversais nos cenarios A-D", [
        "O Pivot de Nuvem-Nacionalidade sobrepoe-se aos quatro cenarios, modificando suas conclusoes de forma nao linear. Adiciona uma terceira dimensao: o grau de autonomia do compute instalado. A Tabela 13 sintetiza o impacto.",
    ]),
    ("5.12 Implicacoes para a Franca: a questao do compute realmente soberano", []),
    ("5.12.1 'Sovereignty washing' como risco sistemico", [
        "O termo 'sovereignty washing' designa a pratica dos hyperscalers dos EUA de comercializar ofertas de nuvem soberanas implantando data centers em solo europeu, enquanto permanecem sujeitos ao CLOUD Act.[27] A Microsoft reconheceu em sua propria documentacao que nao pode garantir a soberania para clientes europeus em caso de uma injuncao dos EUA juridicamente fundamentada.[28]",
    ]),
    ("5.12.2 A vantagem nuclear francesa na nova equacao", [
        "A reinterpretacao soberana do CACI paradoxalmente reforca o ativo estrategico frances. Se F(r) = F_phys x F_sov, a estrategia ideal nao e apenas aumentar o F_phys (atraindo hyperscalers), mas o F_sov (desenvolvendo compute independente das jurisdicoes dos EUA).",
        "A energia nuclear francesa cria uma vantagem competitiva de primeira linha: Special Compute Zones anexas a eletricidade nuclear livre de carbono e competitiva, hospedando Gigafactories operadas por entidades europeias (OVHcloud, Scaleway, Mistral AI, IONOS), produziriam compute com um F_sov proximo de 1, genuinamente alem do alcance dos Cloud Sovereignty Mandates.",
    ]),
    ("5.12.3 O projeto DARE/RISC-V: de ambicao a necessidade estrategica", [
        "Sob o Cenario D e especialmente sob os Cloud Sovereignty Mandates, o projeto DARE (Digital Autonomy with RISC-V in Europe, 2025) muda de uma ambicao de longo prazo para uma necessidade estrategica.[29] Enquanto a Europa depender exclusivamente de GPUs NVIDIA/AMD, o controle dos EUA sobre essas arquiteturas cria uma vulnerabilidade residual mesmo em Gigafactories operadas por europeus. A verdadeira autonomia estrategica requer a capacidade de projetar aceleradores independentes, um horizonte que o DARE situa em 2030-2032.",
    ]),
    ("5.13 Sintese: o quarto ponto de inflexao", [
        "O Grande Desacoplamento nao e inevitavel. Representa um risco sistemico condicional. Aos tres pontos de inflexao em 5.7.2, um quarto e adicionado.",
        "Quarto ponto de inflexao: implementacao de recursos de verificacao de localizacao em GPUs avancadas. Se o BIS e o Departamento de Comercio dos EUA implantarem com sucesso mecanismos de atestacao remota em chips H100/B200 ate o final de 2026-2027, o substrato tecnico para o Pivot de Nuvem-Nacionalidade estara pronto. A questao passara da viabilidade tecnica para a decisao politica.",
        "Para a Europa e a Franca, uma estrategia de desacoplamento controlado repousa em tres pilares: (i) acelerar a implantacao de compute soberano (Gigafactories da UE) para aumentar o F_sov antes dos Mandatos; (ii) garantir o status Tier 1 na alianca americana; (iii) investir em alternativas arquitetonicas (DARE/RISC-V, Huawei Ascend para cargas de trabalho nao sensiveis) para reduzir a dependencia de medio prazo.",
    ]),
]

PT.table_blocks = [
    ("Tabela 10. Matriz 2x2 de cenarios prospectivos 2026-2030.",
     "Fonte: construcao do autor, metodologia Schwartz (1991).",
     [
         ["", "Resposta reativa da UE", "Resposta proativa da UE"],
         ["Protecionismo moderado dos EUA",
          "Cenario A - Status Quo Reforcado (deriva lenta para a dependencia)",
          "Cenario C - Parceria Assimetrica (parceiro tecnologico junior ocidental)"],
         ["Protecionismo agressivo dos EUA",
          "Cenario B - Fratura Digital (desacoplamento europeu estrutural)",
          "Cenario D - Soberania Contestada (corrida pela autonomia sob pressao)"],
     ]),
    ("Tabela 11. Sintese comparativa dos quatro cenarios nas seis metricas de divergencia (horizonte 2030).",
     f"Fonte: construcao do autor; snapshot de abril de 2026 (compute operacional bruto EUA/UE {fmt_fr(us_eu_raw, 1)}:1, CACI Power Mode {fmt_fr(us_eu_caci, 2)}:1).",
     [
         ["Metrica (2030)", "A - Status Quo", "B - Fratura", "C - Parceria", "D - Soberania"],
         ["M1 Razao de compute bruto EUA/UE (operacional)", "18-22:1", "25-35:1", "8-10:1", "12-15:1"],
         ["M2 Diferenca de custo do FLOP", "2,4-3,2x", "4-6x", "1,5-2,0x", "1,8-2,5x"],
         ["M3 Parcela da nuvem dos EUA (pct)", "72-75", "78-82", "60-65", "50-55"],
         ["M4 Produtividade da UE (pct/ano)", "+1,0-1,5", "+0,3-0,8", "+1,8-2,5", "+1,2-2,0"],
         ["M5 Energia da UE (TWh)", "~115", "~95", "~140", "~155"],
         ["M6 Razao CACI Power Mode", "4-5:1", "6-8:1", "2,0-2,5:1", "4-7:1 (pos-vale)"],
         ["Probabilidade estimada", "40-50 pct", "15-20 pct", "15-20 pct", "15-20 pct"],
     ]),
    ("Tabela 12. Estimativa do fator F_sov por jurisdicao e impacto no CACI sob ativacao dos Cloud Sovereignty Mandates (2028).",
     "Fonte: construcao do autor; Synergy Research Group (2025), Statista (2025), e Capitulo I Fig 1.8 baseline.",
     [
         ["Jurisdicao", "Parcela F_phys nuvem-EUA", "F_sov estimado", "CACI atual (base fisico)", "Impacto CACI pos-Mandato"],
         ["Estados Unidos", "~5 pct (nuvem domestica nao afetada)", "1,00", "100 (referencia)", "100 (inalterado)"],
         ["UE (Franca, Alemanha)", "~77 pct (dominancia AWS/Azure/GCP)", "0,28", "28,9 (Power Mode)", "Colapso de 30-50 pct nas cargas"],
         ["EAU (hub Dubai)", "~88 pct (dominancia hyperscaler EUA)", "0,12", "55,7 phys / 6,0 souv", "Colapso de 60-80 pct - hub ilusorio"],
         ["Cingapura", "~82 pct (dominancia hyperscaler EUA)", "0,18", "alto - hub APAC", "Colapso de 55-75 pct"],
         ["China", "~2 pct (Alibaba/Tencent/Huawei Cloud)", "0,98", "15,7 (escassez de chips)", "Inalterado - soberania efetiva"],
         ["India", "~60 pct (AWS/Azure + local)", "0,40", "22,2 (Power Mode)", "Colapso moderado - posicao interm."],
     ]),
    ("Tabela 13. Impactos do Pivot de Nuvem-Nacionalidade (Cloud Sovereignty Mandates 2028) nos quatro cenarios da matriz 2x2.",
     "Fonte: construcao do autor.",
     [
         ["Cenario", "Sem Mandatos", "Com Mandatos 2028", "Impacto no CACI da UE", "Leitura Estrategica"],
         ["A - Status Quo", "Dependencia lenta, CACI 4-5:1",
          "Ativacao parcial - hyperscalers cooperam sem restricao majoritaria",
          "Degradacao moderada 15-25 pct nas cargas",
          "Mais estavel, mas ilusao de seguranca revelada"],
         ["B - Fratura", "Desacoplamento estrutural, CACI 6-8:1",
          "Ativacao maxima - nuvem dos EUA condicional e chips restritos simultaneamente",
          "Efeito tesoura duplo: chips raros + compute condicional. Razao CACI potencialmente > 8:1",
          "Pior caso absoluto - combo chips + nuvem"],
         ["C - Parceria", "Recuperacao parcial, CACI 2,0-2,5:1",
          "Gigafactories soberanas absorvem o choque se o F_sov da UE subir para 0,45-0,55",
          "Impacto limitado se implantado no prazo - Gigafactories = hedge soberano",
          "Melhor resiliencia - investimento anterior prova seu valor"],
         ["D - Soberania", "Recuperacao sob pressao, CACI 4-7:1 pos-vale",
          "Mandatos tornam-se catalisador politico - aceleram implantacao da UE e aliancas JP/KR/TW",
          "Curva em U acelerada - vale 2028-29, recuperacao mais rapida",
          "Paradoxal: Mandatos podem acelerar soberania da UE se a resposta for rapida o suficiente"],
     ]),
]

PT.notes = EN.notes


# ===========================================================================
# Main
# ===========================================================================

def main() -> None:
    """Build the three Chapter V .docx files."""
    out_dir = Path(__file__).parent
    for lp in (EN, FR, PT):
        build(lp, out_dir)
    log.info("Done.")


if __name__ == "__main__":
    main()
