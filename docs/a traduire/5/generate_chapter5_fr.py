"""
Chapitre V - Scenarios Prospectifs 2026-2030 - generateur FR.

Genere le .docx du Chapitre V de la these "AI for Americans First"
en francais uniquement. Tous les chiffres sont alignes sur le
snapshot du tableau de bord d'avril 2026 :

    - Bandeau couverture : 76,9 pct / 1,59x / 3,46:1
    - EP2 (compute)      : US/UE 17,6:1 brut, 2 759 968 vs 156 632 PFLOP/s,
                           CACI Power Mode 3,46:1 (etait 16:1 / 2 763 554 vs 173 416)
    - EP3 (energie)      : 1,4-1,7x apres correction PPA (et 2-3x sur Eurostat brut)
    - Tous les M1/M6 des scenarios rebases sur 3,46:1 (avril 2026)
    - Section 5.7 Tableau 11 entierement rafraichi ; M1 du scenario C corrige
      (etait incoherent a 8-10:1, est maintenant 8-10:1 sur compute brut OK
      mais le CACI M6 descend bien a 2,0-2,5:1)
    - Section 5.9.2 Tableau 12 reutilise les valeurs Phys/Sov rigoureusement
      calculees (EAU 55,7 phys / 6,0 sov sur compute installe, UE 28,9 quasi
      pleinement souveraine sur le F installe ; le collapse se joue sur les
      charges cloud, pas sur le F installe)
    - Distinction explicite entre le CACI Souverain statique du chapitre I
      (ownership des clusters Epoch AI) et le F_sov dynamique 2028 (controle
      sous regime Cloud Sovereignty Mandate)

Auteur : Fabrice Pizzi (Universite Paris-Sorbonne, M2 Intelligence Economique).
Build  : python3 generate_chapter5_fr.py
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
log = logging.getLogger("chapitre5_fr")


NAVY = RGBColor(0x1A, 0x27, 0x44)
GOLD = RGBColor(0xB8, 0x92, 0x2F)
GREY = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x20, 0x20, 0x20)


# ---------------------------------------------------------------------------
# Style helpers (consistent with Chap I-IV generators)
# ---------------------------------------------------------------------------

def set_run(run, *, font="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text, *, align=None, space_after=6, **run_kwargs):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run(run, **run_kwargs)
    return p


def add_heading(doc, text, level):
    sizes = {1: 22, 2: 16, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, font="Calibri", size=sizes.get(level, 11),
            bold=True, color=NAVY)
    return p


def add_cover(doc):
    add_paragraph(doc, "", space_after=0)
    add_paragraph(doc, "ETUDE DE RECHERCHE - FEVRIER 2026",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, bold=True, color=GREY, space_after=4)
    add_paragraph(doc, "AI FOR AMERICANS FIRST",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=26, bold=True, color=NAVY, space_after=4)
    add_paragraph(doc,
                  "Protectionnisme IA, Energie et Semi-conducteurs : "
                  "Trajectoires de divergence US/Europe 2024-2030",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, italic=True, color=GREY, space_after=4)
    add_paragraph(doc,
                  "Analyse geostrategique et economique integree - Chapitre V",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=11, italic=True, color=GREY, space_after=18)

    table = doc.add_table(rows=1, cols=3)
    chips = [
        "76,9 pct du compute IA operationnel mondial = USA",
        "1,59x cout energie EU/US (ajuste-PPA)",
        "3,46:1 ratio CACI US/EU (Power Mode)",
    ]
    for i, line in enumerate(chips):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        set_run(run, size=11, bold=True, color=NAVY)

    add_paragraph(doc, "", space_after=8)
    add_paragraph(doc, "Fabrice Pizzi",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, bold=True, color=NAVY, space_after=2)
    add_paragraph(doc, "Universite Paris-Sorbonne",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, color=GREY, space_after=2)
    add_paragraph(doc,
                  "Master 2 Intelligence Economique - Intelligence Warfare",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, italic=True, color=GREY, space_after=14)
    add_paragraph(doc,
                  "Paris - fevrier 2026  |  7 chapitres  |  4 scenarios prospectifs  |  3 zones geographiques",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, italic=True, color=GREY, space_after=10)
    add_paragraph(doc,
                  "Mots-cles : intelligence artificielle, protectionnisme technologique, "
                  "semi-conducteurs, controles a l'exportation, compute souverain, "
                  "geopolitique IA, France, Etats-Unis, Chine",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=9, italic=True, color=GREY, space_after=24)


def add_chapter_header(doc):
    add_paragraph(doc, "CHAPITRE V",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=11, bold=True, color=GOLD, space_after=2)
    add_paragraph(doc, "Scenarios prospectifs 2026-2030",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=18, bold=True, color=NAVY, space_after=12)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(
        "Ce chapitre constitue le coeur de la contribution originale de cette etude. "
        "En appliquant le protocole methodologique defini au chapitre II (matrice 2x2, "
        "six metriques de divergence, calibration CACI), nous construisons quatre "
        "scenarios d'evolution de la relation transatlantique en IA, energie et "
        "semi-conducteurs pour la periode 2026-2030. Chaque scenario est determine par "
        "la combinaison de deux incertitudes critiques identifiees au chapitre III : "
        "le degre de protectionnisme americain et la capacite de reponse strategique "
        "europeenne. Nous evaluons ensuite chaque scenario sur ses six metriques, avant "
        "de synthetiser les conditions de bascule entre trajectoires."
    )
    set_run(run, size=11, italic=True, color=GREY)


def render_section(doc, title, paragraphs):
    """Render one section: title at the right level, then body paragraphs."""
    first = title.split()[0]
    level = 2 if first.count(".") == 1 else 3
    add_heading(doc, title, level)
    for para in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(para)
        set_run(run, size=11, color=DARK)


def render_table(doc, caption, source, rows):
    """Render one captioned table."""
    add_paragraph(doc, caption,
                  size=10, bold=True, color=NAVY, space_after=2)
    add_paragraph(doc, source,
                  size=9, italic=True, color=GREY, space_after=4)
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            set_run(run, size=10,
                    bold=(i == 0),
                    color=NAVY if i == 0 else DARK)
    doc.add_paragraph()


def render_notes(doc, notes):
    add_heading(doc, "Notes", 2)
    for i, note in enumerate(notes, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run_id = p.add_run(f"{i}. ")
        set_run(run_id, size=9, bold=True, color=GOLD)
        run_txt = p.add_run(note)
        set_run(run_txt, size=9, color=GREY)


def render_license(doc):
    doc.add_paragraph()
    license_lines = [
        "Licence et avertissement. Cette oeuvre, 'AI for Americans First', est mise a disposition selon les termes de la licence Creative Commons Attribution - Pas d'Utilisation Commerciale - Partage dans les Memes Conditions 4.0 International (CC BY-NC-SA 4.0).",
        "Vous etes libre de partager et adapter le materiel a des fins non commerciales, a condition d'attribuer correctement l'oeuvre a Fabrice Pizzi (Universite Paris-Sorbonne) et de distribuer vos contributions sous la meme licence. Ce document est fourni a des fins educatives et de recherche uniquement.",
        "Tableau de bord public : https://mo0ogly.github.io/America-First-IA/dashboard/",
        "Depot : https://github.com/mo0ogly/America-First-IA",
    ]
    for line in license_lines:
        add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT,
                      size=8, italic=True, color=GREY, space_after=2)
    add_paragraph(doc, "AI for Americans First - Fabrice Pizzi - Chapitre V",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=8, italic=True, color=GREY, space_after=0)


# ---------------------------------------------------------------------------
# Body content (FR)
# ---------------------------------------------------------------------------

SECTIONS: list[tuple[str, list[str]]] = [
    ("5.1 Elements predetermines : ce qui ne changera pas", [
        "Conformement a la methode Schwartz (1991), nous distinguons les elements predetermines (tendances quasi-certaines a l'horizon 2030) des incertitudes critiques (facteurs dont l'evolution depend de decisions politiques non encore prises). Quatre elements predetermines structurent l'ensemble des scenarios.",
        "EP1 - Croissance exponentielle de la demande de compute IA. Les ventes de semi-conducteurs ont double en deux ans (2023-2025), la puissance des puces IA installees double tous les sept mois (Epoch AI), et aucun signe de ralentissement n'est observable au moment de fevrier 2026. Meme dans l'hypothese d'une deceleration des scaling laws (saturation Chinchilla), la diffusion de l'IA vers l'inference, la robotique et les agents autonomes maintiendra une demande de compute fortement croissante.[1]",
        "EP2 - Concentration persistante du compute aux Etats-Unis. Le ratio US/UE de 17,6:1 en compute installe brut (chapitre III, snapshot du tableau de bord d'avril 2026 : 2 759 968 vs 156 632 PFLOP/s), se traduisant par un ratio CACI Power Mode de 3,46:1 une fois pondere par la formule geometrique F^0,40 x L^0,20 x R^0,15 / E^0,25, reflete des decisions d'investissement prises en 2022-2025 dont les effets se materialisent jusqu'en 2028-2029 (delais de construction des centres de donnees : 18-36 mois). Meme un revirement politique immediat n'alterait pas le stock installe avant la fin de la decennie.",
        "EP3 - Tension energetique croissante. La consommation mondiale des centres de donnees, estimee a 415 TWh en 2024, atteindra 800-950 TWh d'ici 2030 selon les projections AIE (chapitre III). L'asymetrie des couts energetiques (US 1,4-1,7x moins chers apres correction PPA, 2-3x moins chers sur les tarifs Eurostat industriels non ajustes) persistera, sauf investissement massif dans le nucleaire europeen, dont les delais de deploiement (SMR : 5-7 ans pour les premiers reacteurs) depassent l'horizon 2030.[2]",
        "EP4 - Cadre reglementaire Section 232 en place. La Proclamation 11002 du 14 janvier 2026 est un fait accompli juridique. Contrairement aux tarifs IEEPA (annules par la Cour supreme le 20 fevrier 2026[3]), les tarifs Section 232 reposent sur une base legale confirmee. Le rapport du Secretaire au Commerce sur le marche des semi-conducteurs pour centres de donnees est attendu d'ici le 1er juillet 2026, et peut recommander une extension ou modification des tarifs. Quelle que soit la direction prise, l'instrument legal restera disponible.[4]",
    ]),
    ("5.2 Incertitudes critiques et matrice 2x2", []),
    ("5.2.1 Axe 1 : Degre de protectionnisme americain", [
        "La premiere incertitude porte sur l'evolution de la politique americaine entre deux poles. Le pole modere correspond au maintien du statu quo de janvier 2026 : tarif de 25 pour cent limite aux puces avancees re-exportees, exemptions domestiques larges, accord commercial UE plafonnant les tarifs sur semi-conducteurs a 15 pour cent, et aucune extension significative au cloud ou aux modeles. Le coeur UE (France, Allemagne) reste dans la categorie partenaires de confiance. Le pole agressif suppose une extension apres le rapport de juillet 2026 : tarifs etendus aux semi-conducteurs derives et equipements, quotas GPU pour l'UE (incluant la France), conditions restrictives pour l'acces au cloud IA de pointe, et utilisation du compute comme levier de negociation commerciale (compute-for-concessions).[5]",
    ]),
    ("5.2.2 Axe 2 : Capacite de reponse strategique europeenne", [
        "La seconde incertitude porte sur la capacite de l'UE a deployer une reponse coherente et rapide. Le pole reactif correspond a des reponses nationales fragmentees, des investissements disperses, un AI Act creant des couts de conformite supplementaires, et un deploiement lent des AI Factories/Gigafactories (delais bureaucratiques, autorisations 24+ mois). Le pole proactif suppose la mise en oeuvre acceleree du programme AI Continent (19 AI Factories plus jusqu'a 5 Gigafactories de 100 000+ GPU), l'adoption de Special Compute Zones (autorisation en 180 jours), la mobilisation effective du fonds InvestAI (20 milliards EUR), et la mise en commun de la capacite nucleaire francaise comme avantage competitif.[6]",
    ]),
    ("5.2.3 Matrice et nommage des scenarios", [
        "L'intersection de ces deux axes produit quatre scenarios. UE reactive + protectionnisme US modere donne le scenario A (Statu quo renforce, derive lente vers la dependance). UE reactive + US agressif donne le scenario B (Fracture numerique, decouplage europeen structurel). UE proactive + US modere donne le scenario C (Partenariat asymetrique, partenaire technologique junior occidental). UE proactive + US agressif donne le scenario D (Souverainete contestee, course a l'autonomie sous pression).",
    ]),
    ("5.3 Scenario A - Statu quo renforce (Protectionnisme modere + UE reactive)", []),
    ("5.3.1 Recit", [
        "Apres le rapport de juillet 2026, le Secretaire au Commerce recommande de maintenir le tarif de 25 pour cent sur les puces avancees re-exportees mais sans l'etendre significativement. L'accord commercial US-UE d'aout 2025 est respecte : les tarifs sur semi-conducteurs pour l'UE restent plafonnes a 15 pour cent.[7] L'UE, rassuree par ce statu quo, ralentit le deploiement de ses propres initiatives. Les AI Factories EuroHPC peinent a atteindre leur capacite nominale (delais d'autorisation, coordination inter-Etats). Les Gigafactories sont reportees a 2029-2030. Le fonds InvestAI est partiellement mobilise (8-10 milliards EUR sur 20). Les entreprises europeennes continuent de s'appuyer fortement sur le cloud US, dont la performance et les couts restent imbattables.",
    ]),
    ("5.3.2 Trajectoire des metriques", [
        "M1 - Ratio compute (GPU installes US/UE) : passe de 17,6:1 brut (2025) a 18-22:1 brut (2030) sur le compute installe operationnel. L'ecart se creuse legerement a mesure que les investissements US s'accelerent (Stargate, mega-clusters xAI, Meta) tandis que l'UE n'ajoute que les 19 AI Factories (25 000 GPU max chacune, soit environ 475 000 GPU publics, un ordre de grandeur en dessous d'un seul hyperscaler US).[8]",
        "M2 - Ecart cout du FLOP (UE/US) : reste dans la fourchette 2,4-3,2x. L'absence de tarifs agressifs sur l'UE maintient l'acces au cloud US a des prix proches des niveaux actuels, mais les couts energetiques europeens continuent de peser.",
        "M3 - Part cloud US dans les depenses IA europeennes : passe de 70 pour cent (2024) a 72-75 pour cent (2030). Les fournisseurs europeens (OVHcloud, Deutsche Telekom) conservent leurs 15 pour cent sur le segment souverainete mais ne gagnent pas de terrain sur les services IA generative.",
        "M4 - Productivite IA (pct/an) : US +2,5-3,0 ; UE +1,0-1,5. L'UE realise une partie du potentiel IA via les applications aval (SAP, Siemens, fintech), mais l'adoption lente et le deficit de compute plafonnent les gains.",
        "M5 - Dependance energetique (TWh centres de donnees) : UE environ 115 TWh en 2030 (+65 pour cent vs 2024). Le nucleaire francais absorbe une partie de la demande, mais l'absence de Special Compute Zones retarde la connexion au reseau de nouveaux centres de donnees.",
        "M6 - CACI(US)/CACI(UE) : passe de 3,46:1 (avril 2026) a 4-5:1 (2030). L'ecart se creuse moderement a mesure que le facteur F (compute) s'accumule cote US tandis que les couts E (energie) pesent cote UE.",
    ]),
    ("5.3.3 Consequences pour la France", [
        "Ce scenario est le plus probable a court terme (probabilite estimee : 40-50 pour cent). Il est aussi le plus insidieux : l'absence de choc visible demobilise les acteurs europeens, tandis que la dependance se creuse structurellement. Les entreprises francaises beneficient de l'acces au cloud US pour l'adoption IA (BNP Paribas, Airbus, TotalEnergies via AWS/Azure), mais cette adoption renforce le verrouillage decrit au chapitre IV. Le deficit de productivite IA par rapport aux Etats-Unis (-1,0 a -1,5 points par an) s'accumule sur cinq ans, creusant l'ecart de competitivite de 5 a 8 points de PIB.",
    ]),
    ("5.4 Scenario B - Fracture numerique (Protectionnisme agressif + UE reactive)", []),
    ("5.4.1 Recit", [
        "Le rapport de juillet 2026 conduit a une extension significative. Le Secretaire au Commerce recommande des tarifs etendus aux equipements semi-conducteurs et produits derives, avec un tariff offset program reserve aux entreprises investissant dans la production americaine.[9] L'accord UE a 15 pour cent est revise a la hausse, ou accompagne de conditions restrictives (quotas de volume sur les GPU avances, exigences de reciprocite sur l'AI Act). Simultanement, l'acces au cloud IA de pointe est rendu conditionnel pour les entites non americaines (limitations d'acces aux API des modeles de frontiere, restrictions sur les poids). L'UE, fragmentee, ne parvient pas a formuler une reponse coherente : les Etats membres se divisent entre accommodation (pays nordiques, Pays-Bas) et confrontation (France, Italie).",
    ]),
    ("5.4.2 Trajectoire des metriques", [
        "M1 - Ratio compute : passe de 17,6:1 brut (2025) a 25-35:1 brut (2030). Les quotas GPU limitent les importations europeennes au moment ou la demande explose. Les projets AI Factory sont compromis par l'incapacite a se procurer des GPU Nvidia/AMD aux volumes prevus.",
        "M2 - Ecart cout du FLOP : bondit a 4-6x. Les tarifs etendus, combines aux quotas et a l'asymetrie energetique, augmentent massivement les couts du compute europeen. Les entreprises francaises font face a une surcharge de 3x a 5x pour l'entrainement de modeles.",
        "M3 - Part cloud US : paradoxalement, monte a 78-82 pour cent. Faute d'alternative locale credible, les entreprises europeennes voulant acceder a l'IA de pointe doivent passer par les hyperscalers US, aux conditions tarifaires qu'ils dictent. Les services souverains (OVHcloud, Scaleway) manquent du materiel pour offrir des services GenAI competitifs.",
        "M4 - Productivite IA : US +2,5-3,5 ; UE +0,3-0,8. Le potentiel IA europeen est severement contraint. Le McKinsey Global Institute estime qu'avec une adoption lente, la productivite europeenne ne depasserait pas 0,3 pour cent, proche de la stagnation.[10]",
        "M5 - Dependance energetique : UE environ 95 TWh seulement (2030), non par vertu mais par defaut - le manque de GPU limite la construction des centres de donnees. Ironiquement, la contrainte de compute attenue la contrainte energetique.",
        "M6 - Ratio CACI : explose de 3,46:1 (avril 2026) a 6-8:1 (2030). C'est le scenario ou l'ecart est le plus important, avec les trois facteurs CACI se deteriorant simultanement cote europeen : F plafonne par les quotas, E gonfle par les tarifs, L affaibli par un brain drain accelere vers les Etats-Unis.",
    ]),
    ("5.4.3 Consequences pour la France", [
        "Ce scenario (probabilite estimee : 15-20 pour cent) represente le pire des cas. La France subit un decouplage technologique structurel : les projets compute-intensifs (modeles de fondation Mistral, robotique Comau/Exotec, simulations Dassault) sont relocalises aux Etats-Unis ou dependent d'un acces au cloud US de plus en plus couteux. Le time-to-market des solutions IA francaises s'allonge de 25 a 40 pour cent. Les PME industrielles, incapables d'absorber les surcharges, renoncent a l'IA de pointe et optent pour des solutions degradees (modeles open-source plus petits, inference locale). L'ecart de productivite cumule avec les Etats-Unis atteint 10 a 15 points sur cinq ans.",
    ]),
    ("5.5 Scenario C - Partenariat asymetrique (Protectionnisme modere + UE proactive)", []),
    ("5.5.1 Recit", [
        "Le protectionnisme US reste modere (comme en A), mais l'UE exploite cette fenetre pour accelerer ses propres investissements. Les AI Factories sont deployees dans les delais (2026-2027), les premieres Gigafactories de 100 000+ GPU sont commandees fin 2026 et livrees en 2028.[11] La France joue un role central grace a son parc nucleaire (65-70 pour cent du mix electrique, cout marginal competitif), et des Special Compute Zones sont designees sur d'anciens sites industriels avec connexions reseau lourdes.[12] Toutefois, l'UE accepte de facto un statut de partenaire technologique junior : elle utilise des GPU Nvidia/AMD (pas de champion europeen en design d'ASIC IA), depend des fonderies TSMC/Samsung/Intel pour la production, et ses modeles de fondation restent un cran en dessous des leaders US.",
    ]),
    ("5.5.2 Trajectoire des metriques", [
        "M1 - Ratio compute : descend de 17,6:1 brut (2025) a 8-10:1 brut (2030) sur le compute installe. Les Gigafactories et l'investissement prive (InvestAI plus co-investissements industriels) ajoutent 1-2 millions d'equivalents H100 en Europe, reduisant l'ecart sans le combler.",
        "M2 - Ecart cout du FLOP : descend a 1,5-2,0x. Le nucleaire francais et les economies d'echelle des Gigafactories compriment les couts d'energie et d'infrastructure, bien qu'un ecart residuel persiste (absence de design GPU proprietaire).",
        "M3 - Part cloud US : descend legerement a 60-65 pour cent. Les services souverains europeens gagnent des parts de marche sur les segments regules (defense, sante, finance), tandis que le cloud US conserve la majorite des charges commerciales. Le marche se segmente en souverain et performance.",
        "M4 - Productivite IA : US +2,5-3,0 ; UE +1,8-2,5. L'UE atteint 60-80 pour cent de son potentiel theorique grace a un compute local suffisant pour l'adoption a grande echelle d'applications aval, meme si l'entrainement des modeles de frontiere reste dependant du materiel US.",
        "M5 - Energie : UE environ 140 TWh (2030). La demande est plus elevee qu'en A car le compute europeen augmente, mais le nucleaire et les SMR planifies absorbent l'essentiel. RTE France confirme la faisabilite de +10 GW sous reserve d'investissements reseau.",
        "M6 - Ratio CACI : descend de 3,46:1 (avril 2026) a 2,0-2,5:1 (2030). C'est le scenario le plus favorable realistement atteignable a l'horizon 2030. Le facteur F s'ameliore significativement, E beneficie du nucleaire, mais L reste legerement inferieur (l'ecosysteme IA US plus attractif pour les meilleurs talents).",
    ]),
    ("5.5.3 Consequences pour la France", [
        "Ce scenario (probabilite estimee : 15-20 pour cent) est le plus favorable pour la France a court-moyen terme. La France devient le hub energetique IA de l'UE grace a son parc nucleaire, attirant les investissements en centres de donnees et Gigafactories. Les entreprises francaises gagnent un acces a un compute local competitif pour l'inference et le fine-tuning, reduisant la dependance au cloud US pour les usages standard. Mistral et les startups francaises peuvent entrainer des modeles specialises localement. Toutefois, l'entrainement des modeles de frontiere reste dependant du materiel US, et l'autonomie strategique est partielle : la France est souveraine en application, mais pas dans la creation des technologies fondamentales.",
    ]),
    ("5.6 Scenario D - Souverainete contestee (Protectionnisme agressif + UE proactive)", []),
    ("5.6.1 Recit", [
        "Le protectionnisme US s'intensifie (comme en B), mais l'UE repond avec determination. La menace americaine devient le catalyseur politique d'une mobilisation industrielle europeenne sans precedent depuis le projet AIRBUS des annees 1970. Le programme AI Continent est accelere et etendu : les 5 Gigafactories sont commandees en urgence, la France annonce 20 GW de capacite nucleaire dediee aux centres de donnees IA d'ici 2032 (combinant extension du parc existant et SMR), le projet DARE (RISC-V europeen) est escalade pour concevoir des accelerateurs IA reduisant la dependance a Nvidia.[13] Simultanement, l'UE negocie des alliances technologiques alternatives (Japon, Coree du Sud, Taiwan) pour securiser l'approvisionnement en GPU et fonderies.",
    ]),
    ("5.6.2 Trajectoire des metriques", [
        "M1 - Ratio compute : evolue de 17,6:1 brut (2025) a 12-15:1 brut (2030) sur le compute installe. L'UE investit massivement mais part de tres loin. Les quotas US ralentissent les importations, mais les alliances alternatives et la production locale (Gigafactories utilisant des GPU Samsung/Intel comme alternatives a Nvidia) compensent partiellement.",
        "M2 - Ecart cout du FLOP : 2,5-4,0x initialement (2027, pic du choc tarifaire), puis reduction progressive vers 1,8-2,5x (2030) a mesure que les Gigafactories montent en cadence et que les alternatives GPU murissent.",
        "M3 - Part cloud US : descend a 50-55 pour cent (2030), le declin le plus prononce des quatre scenarios. La defiance geopolitique et les restrictions US poussent les entreprises europeennes vers les alternatives souveraines, meme imparfaites. Les hyperscalers US perdent du terrain sur les segments regules.",
        "M4 - Productivite IA : US +2,5-3,5 ; UE +1,2-2,0. L'UE traverse un creux de productivite en 2027-2028 (periode de transition ou les restrictions US mordent mais les investissements europeens ne sont pas encore operationnels), puis un rattrapage partiel a partir de 2029.",
        "M5 - Energie : UE environ 150-160 TWh (2030). C'est le scenario le plus energivore pour l'UE, la construction massive de centres de donnees locaux creant une demande enorme. Le nucleaire francais devient un actif strategique continental, mais la pression sur le reseau est maximale.",
        "M6 - Ratio CACI : suit une trajectoire en U : degradation de 3,46:1 a 8-12:1 en 2027-2028 (pic du choc), puis amelioration vers 4-7:1 d'ici 2030. Le resultat depend fortement de la vitesse d'execution europeenne : chaque annee de retard sur les Gigafactories prolonge la periode de vulnerabilite maximale.",
    ]),
    ("5.6.3 Consequences pour la France", [
        "Ce scenario (probabilite estimee : 15-20 pour cent) est le plus ambitieux et le plus risque. Il place la France au coeur d'un effort de souverainete technologique europeen sans precedent. Les investissements nucleaires massifs (SMR, extension du parc) deviennent un enjeu geopolitique de premier ordre. Le projet DARE/RISC-V pourrait, en cas de succes, constituer la premiere alternative europeenne credible aux GPU Nvidia pour l'IA, mais sur un horizon de 5-7 ans, bien au-dela de 2030. A court terme (2026-2028), la France traverse une periode de vulnerabilite maximale ou les surcharges et penuries degradent la competitivite, avant un rattrapage conditionnel a la vitesse de deploiement de l'infrastructure.",
    ]),
    ("5.7 Synthese comparative et conditions de bascule", []),
    ("5.7.1 Tableau de synthese des metriques", [
        "Le Tableau 11 ci-dessous consolide la trajectoire des six metriques de divergence pour les quatre scenarios a l'horizon 2030, ancre sur le snapshot du tableau de bord d'avril 2026 (compute brut operationnel US/UE 17,6:1, CACI Power Mode 3,46:1).",
    ]),
    ("5.7.2 Conditions de bascule entre scenarios", [
        "La trajectoire reelle suivra probablement un chemin hybride entre ces scenarios. Trois points de bascule determinent les transitions possibles.",
        "Premier point de bascule : le rapport Commerce de juillet 2026. Ce rapport determinera si le protectionnisme US s'etend (basculement vers B ou D) ou reste cible (maintien en A ou C). Indicateurs a surveiller : evolution du deficit commercial americain en semi-conducteurs, taux de remplissage des fabs CHIPS Act (Intel, TSMC Arizona, Samsung Taylor), pression politique interieure (midterms 2026). Le resultat des negociations de Phase 1 (rapport du a 14 avril 2026) sera un signal precoce.[14]",
        "Deuxieme point de bascule : la vitesse de deploiement des Gigafactories UE. La Commission prevoit les premieres Gigafactories operationnelles en 2027-2028. Si ce calendrier est tenu, l'UE bascule vers les scenarios proactifs (C ou D). Si les delais d'autorisation, de financement ou d'approvisionnement materiel repoussent les livraisons a 2029-2030, l'UE reste en mode reactif (A ou B). La proposition CFG des Special Compute Zones (autorisation en 180 jours vs 24+ mois actuellement) est le facteur d'acceleration cle.[15]",
        "Troisieme point de bascule : la decision francaise sur le nucleaire pour l'IA. La France possede un atout unique en Europe : un parc nucleaire fournissant 65-70 pour cent de l'electricite, avec un cout marginal globalement competitif. La decision de dedier une capacite significative (10-20 GW) aux centres de donnees IA, via l'extension du parc, les nouveaux EPR2 et les SMR, determinera si la France devient le hub energetique IA de l'Europe ou cede cette position a d'autres (Scandinavie avec l'hydroelectrique, Europe de l'Est avec des couts fonciers bas). Ce point de bascule est specifiquement francais et determine la position de la France au sein des scenarios europeens.[16]",
    ]),
    ("5.7.3 Le point de convergence : 2028", [
        "Les quatre scenarios convergent sur un point critique commun en 2028. C'est l'annee ou : (i) la demande de compute depassera la capacite installee en Europe, creant des goulots d'etranglement materiels (meme sous protectionnisme modere) ; (ii) les premiers effets des tarifs etendus (s'ils sont adoptes) seront pleinement ressentis ; (iii) les Gigafactories, si elles sont deployees a temps, commenceront a produire un compute local significatif ; (iv) la demande energetique des centres de donnees saturera la capacite de connexion au reseau dans plusieurs Etats membres. L'annee 2028 constitue donc le moment de verite ou l'Europe decouvrira si elle est sur trajectoire A/B (dependance croissante) ou C/D (rattrapage commence). Les decisions prises en 2026-2027 (rapport Commerce US, Gigafactories, nucleaire francais) seront irreversiblement engagees.",
    ]),
    ("5.8 Origines du point de bascule : fondations juridico-techniques deja en place", [
        "Le scenario du Grand Decouplage n'est pas construit ex nihilo. Il est la projection logique, a l'horizon 2028, d'une architecture de controle dont les couches fondationnelles sont deja operationnelles en 2026. Identifier ces couches releve de la rigueur academique : distinguer ce qui releve de tendances documentees de ce qui constitue une projection extrapolee.",
    ]),
    ("5.8.1 La couche legale : l'extraterritorialite comme instrument structurel", [
        "La premiere couche est legale et precede largement la politique IA. Le CLOUD Act (Clarifying Lawful Overseas Use of Data Act, 2018) etablit que tout fournisseur de cloud soumis a la juridiction US est tenu de produire les donnees independamment du lieu ou ces communications, enregistrements ou autres informations sont stockes.[17] Cette loi, confirmee par la jurisprudence federale (United States v. Microsoft, resolu par l'adoption du CLOUD Act avant que la Cour supreme ne statue), cree une dissociation fondamentale entre la localisation physique des donnees et leur nationalite legale.",
        "La consequence immediate pour le compute IA est radicale : un cluster H100 physiquement localise a Dubai, Singapour ou dans un centre AWS eu-west-1 en Irlande reste legalement americain. En cas d'ordre du gouvernement US, l'operateur (AWS, Azure, Google Cloud) est legalement tenu de se conformer, independamment de la volonte du client ou de la legislation locale. Microsoft a reconnu devant un tribunal francais en 2024 ne pas pouvoir garantir la souverainete des donnees pour les clients europeens en cas d'injonction US fondee en droit.[18]",
        "Cette architecture legale preexistante est le substrat sur lequel se greffe le Pivot Cloud-Nationalite : les Cloud Sovereignty Mandates projetes pour 2028 ne creent pas un pouvoir nouveau. Ils activent et systematisent un pouvoir juridictionnel deja existant en l'etendant a la couche du compute operationnel.",
    ]),
    ("5.8.2 La couche technique : de la verification de localisation au throttling de cluster", [
        "La seconde couche est technique. L'AI Action Plan US du 23 juillet 2025 introduit explicitement le concept de fonctionnalites de verification de localisation appliquees aux puces IA avancees. Michael Kratsios, directeur de l'OSTP de la Maison-Blanche, a confirme qu'il y a une discussion sur les types de modifications logicielles ou physiques que l'on pourrait apporter aux puces elles-memes pour faire un meilleur location-tracking, ce qui a ete explicitement inclus dans le plan.[19]",
        "Cette annonce n'est pas rhetorique. Le BIS dispose deja, depuis le Framework for Artificial Intelligence Diffusion (janvier 2025, abroge en mai puis remplace par des mesures de guidance), d'une architecture de controle compute-tiering pour les pays de destination (Tier 1-2-3) et de plafonds de compute par entite et par pays.[20] La BIS Affiliates Rule, suspendue pour un an en novembre 2025 mais maintenue en principe, stipule que l'affiliation d'une societe a une entite mere dans un pays restreint suffit a lui refuser l'acces au compute avance, independamment de la localisation physique du cluster.[21]",
        "La trajectoire technologique vers 2028 est donc : (i) des puces equipees de mecanismes de verification de localisation (logicielle ou materielle), (ii) des systemes de reporting automatique au BIS en cas de deviation, (iii) la capacite a suspendre l'acces ou throttler la performance via licence d'exportation, un cluster operant sous licence d'exportation US potentiellement soumis a des restrictions operationnelles par decision administrative. Ce n'est pas de la science-fiction : c'est l'extension au compute d'un principe deja applique au logiciel (sanctions OFAC sur les licences logicielles, gel d'acces aux services cloud pour les entites listees).",
        "Limitation critique de ce mecanisme. Une tension technique reelle doit etre identifiee ici : throttler des clusters de production est techniquement complexe et potentiellement perturbant pour les operateurs eux-memes. NVIDIA ne dispose pas actuellement de mecanisme de desactivation a distance pour les GPU H100/B200 en production. Un tel mecanisme exigerait une modification architecturale significative des firmware et des protocoles d'attestation a distance (via TPM ou equivalent). Le scenario 2028 est donc conditionnel a la mise en oeuvre effective de ces modifications, hypothese plausible sur 24-36 mois d'effort industriel, mais pas une certitude.",
    ]),
    ("5.9 Le mecanisme du Pivot Cloud-Nationalite", []),
    ("5.9.1 Le declencheur : Cloud Sovereignty Mandates comme extension des controles a l'exportation", [
        "Le scenario du point de bascule 2028 suppose que les Etats-Unis franchissent un palier qualitatif : passer du controle d'acces materiel (controles BIS sur les puces) au controle d'acces aux services de compute operationnel (Cloud Sovereignty Mandates). Ce palier n'est pas une rupture arbitraire, il repond a une faille structurelle identifiee dans le regime de controle actuel.",
        "Cette faille est documentee : malgre les restrictions BIS sur les exportations de H100/A100, des enquetes ont revele qu'environ 1 milliard USD de puces Nvidia ont ete acheminees vers la Chine en contournant les controles a l'exportation via des pays tiers (Malaisie, EAU, Singapour) au cours des seuls premiers mois de 2025.[22] La reponse de l'AI Action Plan (fonctionnalites de verification de localisation et monitoring de cluster) constitue le premier pas vers un controle continu post-exportation.",
        "Le declencheur plausible en 2028 est un executive order etendant les obligations du Framework for AI Diffusion a la couche cloud. Sa structure imposerait une certification Data Residency and Jurisdiction Compliance a tous les hyperscalers US operant des clusters avances offshore, avec revocation de l'acces au compute sur sol US comme mecanisme de mise en conformite, et le BIS se reservant le droit de throttler ou suspendre la performance des clusters autorises en cas d'irregularite.",
    ]),
    ("5.9.2 La dissociation Facteur Physique / Facteur Souverain dans le CACI", [
        "C'est ici que le scenario produit son impact le plus analytiquement significatif sur le modele CACI developpe aux chapitres I a IV. Le modele actuel integre le compute F(r) comme mesure de la capacite physiquement installee dans la region r. Sous activation du Pivot Cloud-Nationalite, cependant, la variable se decompose en deux composantes distinctes : F(r) = F_phys(r) x F_sov(r), ou F_phys est le compute physiquement installe dans la juridiction et F_sov est le facteur de souverainete operationnelle (fraction de F_phys hors juridiction US et donc insensible aux Cloud Sovereignty Mandates).",
        "Notez la distinction entre ce F_sov dynamique 2028 et le CACI souverain statique introduit au chapitre I (Fig 1.8). Le CACI souverain du chapitre I a ete calcule en filtrant les clusters Epoch AI par Owner : il capture qui detient le compute installe aujourd'hui (snapshot avril 2026). Le F_sov 2028 capture qui controle le compute sous un regime hypothetique de Cloud Sovereignty Mandates, qui depend de la part des charges hyperscaler, pas de la capacite installee. Les deux metriques s'accordent la ou le compute est detenu par des operateurs domestiques (US, Chine, France domestique Fluidstack/Sesterce) ; elles divergent fortement pour les juridictions ou les clusters domestiquement localises sont detenus par des operateurs US-side (EAU 99,6 pour cent, charges cloud UE majoritairement sur AWS/Azure/GCP).",
        "Les estimations calibrees pour 2028, sous l'hypothese d'activation des Cloud Sovereignty Mandates, sont presentees dans le Tableau 12 ci-dessous.",
    ]),
    ("5.10 L'emergence de blocs IA juridictionnels (2028-2030)", [
        "Le Grand Decouplage ne produit pas un monde binaire US/non-US. Il produit une fragmentation en blocs d'intensite variable, selon la capacite de chaque zone a developper un compute souverain credible. Quatre blocs emergent avec des caracteristiques distinctes.",
    ]),
    ("5.10.1 Le bloc americain etendu (American AI Alliance)", [
        "L'AI Action Plan du 23 juillet 2025 pose explicitement les bases d'une American AI Alliance : exportation de la pile technologique US complete (materiel, modeles, logiciels, standards) aux allies disposes, en echange de l'adoption de controles a l'exportation alignes.[23] La strategie est explicitement decrite comme carrot and stick : les allies alignes accedent aux puces avancees et aux modeles de frontiere sans restrictions supplementaires ; ceux qui refusent sont exposes aux mecanismes Foreign Direct Product Rule et aux tarifs secondaires.",
        "Membres du bloc americain etendu (Tier 1 confirmes) : Etats-Unis, Royaume-Uni, Canada, Australie, Japon, Coree du Sud, Pays-Bas, Allemagne, France (sous reserve d'alignement sur les controles a l'exportation). Pour ces pays, le F_sov effectif augmente : leurs entites accedent aux hyperscalers US sans restriction, et le compute souverain en developpement (Gigafactories UE pour l'Europe) recoit un traitement preferentiel. Le CACI de ces pays n'est pas degrade par les Mandates, il peut meme beneficier d'un effet d'alliance.",
    ]),
    ("5.10.2 Le bloc souverain eurasien", [
        "La Chine constitue le seul exemple complet de bloc souverain preexistant. Avec un F_sov d'environ 0,98 et un ecosysteme cloud national (Alibaba Cloud, Tencent Cloud, Huawei Cloud) operant hors de la juridiction US, les Cloud Sovereignty Mandates n'ont aucune traction directe. La contrainte chinoise reste la penurie de puces avancees (les controles a l'exportation 2022-2025 ont limite l'acces aux GPU H100/A100/B200), mais le bloc americain ne peut pas throttler un cluster Huawei Ascend 910B.",
        "La dynamique post-2028 : la Chine detient le seul compute souverain a grande echelle hors du bloc americain. Les pays cherchant a s'emanciper des Cloud Sovereignty Mandates se retrouvent structurellement face a une alternative binaire : compute americain conditionnel ou compute chinois sous d'autres formes de dependance. Cette contrainte binaire est l'impact geopolitique le plus profond du Grand Decouplage.",
    ]),
    ("5.10.3 Les non-alignes numeriques : une position intenable", [
        "L'Inde, le Bresil, l'Asie du Sud-Est et les pays du Golfe (en l'absence de traites speciaux avec les Etats-Unis) constituent un bloc de non-alignement numerique. Leur position est structurellement inconfortable : trop dependants des hyperscalers US pour basculer vers la souverainete, insuffisamment integres a l'alliance americaine pour echapper aux restrictions en cas de desaccord geopolitique.",
        "Le cas des EAU illustre cette fragilite avec une force quantitative. La Fig 1.8 du chapitre I a documente que 99,6 pour cent du F_total des EAU (22,9 millions d'equivalents H100) est detenu par des acteurs US-side (Stargate UAE, Microsoft, OpenAI), faisant s'effondrer le CACI souverain d'un Physique 55,7 a seulement 6,0. Dubai a investi massivement depuis 2022 pour devenir un hub IA regional, notamment via des accords avec AWS, Microsoft et G42. Pourtant, G42 a deja ete soumis a une intense pression US en 2024 pour rompre ses liens avec des entites chinoises, condition imposee par Washington pour l'acces aux puces avancees.[24] Sous Cloud Sovereignty Mandates, cette pression deviendrait systemique : le compute des hubs du Golfe, physiquement present mais legalement americain, deviendrait un levier de negociation permanent.",
    ]),
    ("5.10.4 Le bloc europeen : entre alliance et autonomie", [
        "L'Europe occupe une position intermediaire et evolutive. Legalement Tier 1 (France, Allemagne, Pays-Bas, etc. sont explicitement dans la presomption d'approbation BIS pour les puces avancees), l'UE maintient neanmoins une ambition d'autonomie strategique que l'Alliance americaine ne satisfait pas pleinement.",
        "Le Cloud and AI Development Act (CADA), dont la proposition formelle est attendue au T1 2026, tente de repondre a ce dilemme en definissant un EU Sovereignty Level qui exclurait structurellement les fournisseurs soumis au CLOUD Act des marches publics sensibles.[25] La Commission europeenne a publie en octobre 2025 un Cloud Sovereignty Framework definissant trois niveaux d'assurance (SOV-1 a SOV-3), avec SOV-3 exigeant que le fournisseur soit hors d'atteinte de toute legislation extraterritoriale non europeenne.[26]",
        "Cette architecture legislative est en construction, mais son calendrier est problematique : le CADA sera au mieux operationnel en 2027-2028, precisement quand les Cloud Sovereignty Mandates US pourraient etre actives. La fenetre de vulnerabilite est maximale entre 2028 et 2030.",
        "Une nuance importante du chapitre I : l'UE est largement souveraine sur le compute installe (99,2 pour cent du F_total est detenu par des operateurs UE). La fenetre de vulnerabilite n'est donc pas sur le F installe mais sur la couche des charges cloud (le compute reellement utilise par les entreprises UE, majoritairement heberge sur AWS/Azure/GCP). Le CADA cible exactement cette couche.",
    ]),
    ("5.11 Impacts transversaux sur les scenarios A-D", [
        "Le Pivot Cloud-Nationalite se superpose aux quatre scenarios, modifiant leurs conclusions de maniere non-lineaire. Il n'invalide pas la matrice 2x2 mais ajoute une troisieme dimension : le degre d'autonomie du compute installe. Le Tableau 13 synthetise l'impact.",
    ]),
    ("5.12 Implications pour la France : la question du compute reellement souverain", []),
    ("5.12.1 Le sovereignty washing comme risque systemique", [
        "Le terme sovereignty washing, popularise par Cristina Caffarra (Eurostack Foundation), designe la pratique des hyperscalers US qui commercialisent des offres sovereign cloud en implantant des centres de donnees sur sol europeen, tout en restant soumis au CLOUD Act.[27] Microsoft a reconnu dans sa propre documentation commerciale ne pas pouvoir garantir la souverainete pour les clients europeens en cas d'injonction US fondee en droit.[28]",
        "Le Cloud Sovereignty Framework publie par la Commission en octobre 2025 commence a formaliser cette distinction. Le niveau SOV-3 exclut explicitement les entites soumises a une legislation extraterritoriale non europeenne. Mais si elle est maintenue dans le texte final du CADA, cette exigence excluerait de facto AWS, Azure et GCP des marches publics les plus sensibles, decision politiquement conflictuelle vis-a-vis des Etats-Unis.",
    ]),
    ("5.12.2 L'avantage nucleaire francais dans la nouvelle equation", [
        "La reinterpretation souverainiste du CACI renforce paradoxalement l'atout strategique francais. Si F(r) se decompose en F_phys x F_sov, alors la strategie optimale de la France n'est pas seulement d'augmenter F_phys (attirer plus de centres de donnees hyperscaler) mais d'augmenter F_sov (developper du compute independant des juridictions US).",
        "L'energie nucleaire francaise cree ici un avantage competitif de premier ordre : des Special Compute Zones adossees a une electricite nucleaire decarbonee et economiquement competitive, hebergeant des Gigafactories operees par des entites europeennes (OVHcloud, Scaleway, Mistral AI, IONOS), produiraient un compute avec un F_sov proche de 1, le seul compute genuinement hors d'atteinte des Cloud Sovereignty Mandates.",
        "La decision-cle de la France en 2026-2027 n'est donc plus simplement combien de GPU mais combien de GPU sous juridiction francaise. Ces deux metriques peuvent diverger considerablement si la politique d'attraction des investissements reste indifferente a la nationalite legale des operateurs.",
    ]),
    ("5.12.3 Le projet DARE/RISC-V : de l'ambition a la necessite strategique", [
        "Dans le cadre du scenario D et a fortiori sous Cloud Sovereignty Mandates, le projet DARE (Digital Autonomy with RISC-V in Europe, EuroHPC JU, 2025) change de statut : il n'est plus une ambition de long terme mais une necessite strategique.[29] Tant que l'Europe depend exclusivement des GPU NVIDIA/AMD pour son compute IA, le controle US sur ces architectures cree une vulnerabilite residuelle meme dans les Gigafactories operees par des entites europeennes, une mise a jour firmware imposee par NVIDIA dans le cadre du programme location verification pourrait theoriquement degrader la performance des clusters europeens.",
        "Ce risque est speculatif mais pas deraisonnable : il illustre la profondeur de la dependance architecturale. L'autonomie strategique authentique en compute IA exige in fine la capacite de concevoir des accelerateurs independants, horizon que le projet DARE situe a 2030-2032, bien au-dela du point de bascule 2028.",
    ]),
    ("5.13 Synthese : le quatrieme point de bascule et les conditions d'un decouplage maitrise", [
        "Le Grand Decouplage n'est pas inevitable. Il represente un risque systemique conditionnel dont l'activation depend des choix politiques US et de la vitesse de reponse europeenne. Aux trois points de bascule identifies en 5.7.2, un quatrieme s'ajoute.",
        "Quatrieme point de bascule : la mise en oeuvre des fonctionnalites de verification de localisation dans les GPU avances. Si le BIS et le Department of Commerce US, conformement a l'AI Action Plan de juillet 2025, parviennent a deployer des mecanismes d'attestation a distance dans les puces H100/B200/GB300 d'ici fin 2026-2027, le substrat technique du Pivot Cloud-Nationalite sera en place. La question ne sera plus de savoir si les Cloud Sovereignty Mandates sont techniquement faisables, mais uniquement s'ils sont politiquement decides. Ce point de bascule devrait etre surveille des 2026 : les premiers appels d'offres BIS/NIST sur les standards d'attestation a distance pour puces IA constitueront le signal precoce.",
        "Pour l'Europe et la France, la strategie d'un decouplage maitrise repose sur trois piliers interdependants : (i) accelerer le deploiement du compute souverain (Gigafactories UE sous juridiction europeenne) pour augmenter F_sov avant l'activation des Mandates ; (ii) securiser le statut Tier 1 dans l'alliance americaine pour maintenir un acces sans restriction aux puces avancees, en acceptant la coordination sur les controles a l'exportation ; (iii) investir dans les alternatives architecturales (DARE/RISC-V, Huawei Ascend comme alternative court terme pour les charges non sensibles) pour reduire la dependance aux GPU US a moyen terme.",
        "L'annee 2028 constitue le moment de verite non seulement pour les quatre scenarios initiaux mais pour une question plus fondamentale : l'Europe sera-t-elle capable de maintenir son acces au compute IA de pointe dans un monde ou la nationalite legale du compute prime sur sa localisation physique ? Les decisions prises en 2026-2027 (CADA, Gigafactories, energie nucleaire, alignement sur les controles a l'exportation) determineront si le Grand Decouplage est pour l'Europe une menace existentielle ou une opportunite de consolider son autonomie strategique.",
    ]),
]


TABLES = [
    ("Tableau 10. Matrice 2x2 des scenarios prospectifs 2026-2030.",
     "Source : construction de l'auteur, methodologie Schwartz (1991).",
     [
         ["", "Reponse UE reactive", "Reponse UE proactive"],
         ["Protectionnisme US modere",
          "Scenario A - Statu quo renforce (derive lente vers la dependance)",
          "Scenario C - Partenariat asymetrique (partenaire technologique junior occidental)"],
         ["Protectionnisme US agressif",
          "Scenario B - Fracture numerique (decouplage europeen structurel)",
          "Scenario D - Souverainete contestee (course a l'autonomie sous pression)"],
     ]),
    ("Tableau 11. Synthese comparative des quatre scenarios sur les six metriques de divergence (horizon 2030).",
     "Source : construction de l'auteur ; baseline snapshot avril 2026 (compute brut operationnel US/UE 17,6:1, CACI Power Mode 3,46:1).",
     [
         ["Metrique (2030)", "A - Statu quo", "B - Fracture", "C - Partenariat", "D - Souverainete"],
         ["M1 Ratio compute brut US/UE (operationnel)", "18-22:1", "25-35:1", "8-10:1", "12-15:1"],
         ["M2 Ecart cout du FLOP", "2,4-3,2x", "4-6x", "1,5-2,0x", "1,8-2,5x"],
         ["M3 Part cloud US (pct)", "72-75", "78-82", "60-65", "50-55"],
         ["M4 Productivite UE (pct/an)", "+1,0-1,5", "+0,3-0,8", "+1,8-2,5", "+1,2-2,0"],
         ["M5 Energie UE (TWh)", "~115", "~95", "~140", "~155"],
         ["M6 Ratio CACI Power Mode", "4-5:1", "6-8:1", "2,0-2,5:1", "4-7:1 (post-creux)"],
         ["Probabilite estimee", "40-50 pct", "15-20 pct", "15-20 pct", "15-20 pct"],
     ]),
    ("Tableau 12. Estimation du facteur F_sov par juridiction et impact CACI sous activation des Cloud Sovereignty Mandates (2028).",
     "Source : construction de l'auteur ; Synergy Research Group (2025), Statista Enterprise Cloud (2025), et chapitre I Fig 1.8 pour le baseline souverain sur compute installe.",
     [
         ["Juridiction", "F_phys part cloud-US", "F_sov estime", "CACI actuel (baseline phys)", "Impact CACI post-Mandate"],
         ["Etats-Unis", "~5 pct (cloud domestique non affecte)", "1,00", "100 (reference)", "100 (inchange)"],
         ["UE (France, Allemagne)", "~77 pct (AWS/Azure/GCP dominent les charges UE)", "0,28", "28,9 (Power Mode)", "Effondrement 30-50 pct sur les charges"],
         ["EAU (hub Dubai)", "~88 pct (hyperscalers US dominants)", "0,12", "55,7 phys / 6,0 souv sur installe", "Effondrement 60-80 pct - hub illusoire"],
         ["Singapour", "~82 pct (hyperscalers US dominants)", "0,18", "eleve - hub APAC", "Effondrement 55-75 pct"],
         ["Chine", "~2 pct (Alibaba/Tencent/Huawei Cloud)", "0,98", "15,7 (penurie de puces plafonne)", "Inchange - souverainete deja effective"],
         ["Inde", "~60 pct (AWS/Azure + locaux)", "0,40", "22,2 (Power Mode)", "Effondrement modere - position intermediaire"],
     ]),
    ("Tableau 13. Impacts du Pivot Cloud-Nationalite (Cloud Sovereignty Mandates 2028) sur les quatre scenarios de la matrice 2x2.",
     "Source : construction de l'auteur.",
     [
         ["Scenario", "Sans Mandates", "Avec Mandates 2028", "Impact CACI UE", "Lecture strategique"],
         ["A - Statu quo", "Dependance lente, CACI 4-5:1",
          "Activation partielle - hyperscalers cooperent sans restriction majeure",
          "Degradation moderee 15-25 pct sur les charges",
          "Plus stable mais illusion de securite revelee"],
         ["B - Fracture", "Decouplage structurel, CACI 6-8:1",
          "Activation maximale - cloud US conditionnel et puces restreintes simultanement",
          "Effet ciseau double : puces rares + compute conditionnel. Ratio CACI potentiellement > 8:1",
          "Pire cas absolu - combo puces et cloud"],
         ["C - Partenariat asymetrique", "Rattrapage partiel, CACI 2,0-2,5:1",
          "Gigafactories souveraines absorbent le choc si F_sov UE monte a 0,45-0,55",
          "Impact limite si deploye dans les delais - Gigafactories = couverture souveraine",
          "Meilleure resilience - investissement prealable prouve sa valeur"],
         ["D - Souverainete contestee", "Rattrapage sous pression, CACI 4-7:1 post-creux",
          "Mandates deviennent catalyseur politique - accelerent deploiement UE et alliances JP/KR/TW",
          "Courbe en U acceleree - creux 2028-2029, rattrapage plus rapide",
          "Paradoxal : les Mandates peuvent accelerer la souverainete UE si la reponse est assez rapide"],
     ]),
]


NOTES = [
    "Epoch AI (janvier 2026), Trends in AI Hardware and Compute. Le doublement tous les 7 mois de la production de puces IA combine 1,6x/an en quantite et 1,6x/an en performance par puce. Meme un ralentissement a 12 mois impliquerait un quadruplement d'ici 2030.",
    "AIE (avril 2025), Energy and AI, Paris. Les projections 800-950 TWh correspondent aux scenarios median et haut de l'AIE. Le ratio energetique (1,4-1,7x apres correction PPA, 2-3x sur Eurostat brut) derive des valeurs de reference du tableau de bord public : USA 85, Chine 92, France 115, Allemagne 140, Royaume-Uni 190, UE 135 USD/MWh, toutes ajustees-PPA (chapitre II §2.4.6).",
    "Cour supreme des Etats-Unis (20 fevrier 2026), Learning Resources Inc. v. Trump et V.O.S. Selections v. United States, decision 6-3 : 'IEEPA does not authorize the President to impose tariffs.' Voir Tax Foundation (2026), Tariff Tracker.",
    "Maison-Blanche (14 janvier 2026), Proclamation 11002, section (2) : 'By July 1, 2026, the Secretary shall provide me with an update on the market for semiconductors used in United States data centers.'",
    "Tax Foundation (fevrier 2026), op. cit. L'accord US-UE (aout 2025) plafonne les tarifs sur semi-conducteurs a 15 pour cent pour l'UE, mais la Proclamation 11002 prevoit explicitement des tarifs plus larges possibles apres la Phase 1.",
    "Commission europeenne (2025), AI Continent Action Plan. Objectif : tripler la capacite des centres de donnees UE en 5-7 ans. 19 AI Factories selectionnees, jusqu'a 5 Gigafactories prevues. Fonds InvestAI : 20 milliards EUR. CFG (octobre 2025), 'Special Compute Zones' : autorisation en 180 jours.",
    "Tax Foundation (fevrier 2026), op. cit. L'accord US-UE d'aout 2025 inclut un plafond de 15 pour cent sur les tarifs sur semi-conducteurs pour l'UE.",
    "EuroHPC JU (2025). Les 19 AI Factories prevoient jusqu'a 25 000 GPU chacune (sites standards). Meme a pleine capacite, cela represente environ 475 000 GPU, moins que le seul cluster xAI Colossus (200 000 GPU H100, extensible). Segler Consulting (juin 2025) estime la capacite publique totale UE a environ 57 000 accelerateurs en 2025.",
    "Proclamation 11002, section sur le tariff offset program. Snell & Wilmer (fevrier 2026), 'The Continued Utilization of Tariffs to Control the Semiconductor Industry.'",
    "McKinsey Global Institute (mai 2024), op. cit. Le chiffre de 0,3 pour cent correspond au scenario d'adoption lente, proche du niveau actuel de croissance de la productivite en Europe occidentale.",
    "Conseil de l'UE (decembre 2025), adoption de la position sur le reglement amende pour les Gigafactories IA. Le calendrier provisoire place les premiers appels d'offres fin 2025 et les premieres installations operationnelles en 2027-2028.",
    "CFG (octobre 2025), 'Tripling the EU Data Centre Stock with Special AI Compute Zones.' La proposition prone la reutilisation de sites industriels desaffectes avec connexions reseau lourdes.",
    "EuroHPC JU (mars 2025), projet DARE (Digital Autonomy with RISC-V in Europe), programme de 6 ans. Trois projets de processeurs par des entreprises distinctes.",
    "Proclamation 11002, section (2) : le Secretaire au Commerce et l'USTR doivent fournir un rapport sur l'etat des negociations dans les 90 jours, soit le 14 avril 2026.",
    "CFG (octobre 2025), op. cit. Le delai moyen d'autorisation pour un centre de donnees dans l'UE est actuellement de 24+ mois (vs 6-12 mois aux Etats-Unis). La proposition SCZ reduirait ce delai a 180 jours via un guichet unique.",
    "RTE (2024), Futurs energetiques 2050, scenario N03. La France prevoit +10 GW de demande pour les centres de donnees d'ici 2030. L'avantage nucleaire francais (65-70 pour cent du mix) est unique en Europe.",
    "Clarifying Lawful Overseas Use of Data (CLOUD) Act, Pub. L. 115-141 (23 mars 2018), Title III, section 103(a). Disposition codifiee a 18 U.S.C. section 2713.",
    "Microsoft France, Tribunal judiciaire de Paris (2024) - reconnaissance que la societe ne peut garantir la souverainete des donnees pour les clients europeens en cas d'injonction US fondee en droit. Cite dans The Register (22 decembre 2025).",
    "Michael Kratsios, directeur OSTP, declarations publiques au APEC Digital and AI Ministerial Meeting, Seoul (aout 2025), cite dans TechResearchOnline (5 aout 2025), 'US AI Chip Tracking Plan.'",
    "Bureau of Industry and Security, U.S. Department of Commerce, Framework for Artificial Intelligence Diffusion, Federal Register vol. 90, no. 10 (15 janvier 2025), ECCN 3A090.a - plafonds de compute de 50 000 equivalents H100 par pays Tier 2, 1 700 equivalents H100 par entite.",
    "BIS, 'Suspension of the Affiliates Rule for One Year' (10 novembre 2025). La regle, bien que suspendue, conserve son architecture de controle et peut etre reactivee par decision administrative.",
    "Financial Times (25 juillet 2025), 'Nvidia AI chips worth $1bn smuggled to China after Trump export controls' - rapport sur les circuits de contournement via la Malaisie, les EAU et Singapour dans les 6 premiers mois de 2025.",
    "Maison-Blanche, Executive Order on Promoting the Export of the American AI Technology Stack (23 juillet 2025). L'AI Action Plan detaille la strategie carrot and stick.",
    "US Department of Commerce, accord G42/Microsoft (2024) - G42 (Abu Dhabi) a accepte de se desinvestir de ses partenariats avec des entites chinoises comme condition d'acces aux puces NVIDIA avancees. Confirme par les declarations de la Secretaire au Commerce Gina Raimondo.",
    "Commission europeenne, Programme de travail 2026, Cloud and AI Development Act (CADA), prevu T1 2026.",
    "Commission europeenne, Cloud Sovereignty Framework (octobre 2025), niveaux SOV-1 a SOV-3, publie dans le cadre du Cloud III Dynamic Purchasing System.",
    "Cristina Caffarra, Eurostack Foundation, citee dans The Register (22 decembre 2025) : 'A company subject to the extraterritorial laws of the United States cannot be considered sovereign for Europe.'",
    "Voir note 18.",
    "EuroHPC JU, projet DARE (Digital Autonomy with RISC-V in Europe), lance en mars 2025, programme de 6 ans.",
]


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def build(out_dir: Path) -> Path:
    """Build the FR Chapter V .docx and return its path."""
    log.info("Building Chapitre V [FR] -> Chapitre_V_Scenarios_Prospectifs_2026_2030_FR.docx")
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    add_cover(doc)
    add_chapter_header(doc)

    for title, paragraphs in SECTIONS:
        render_section(doc, title, paragraphs)

    for caption, source, rows in TABLES:
        render_table(doc, caption, source, rows)

    render_notes(doc, NOTES)
    render_license(doc)

    out = out_dir / "Chapitre_V_Scenarios_Prospectifs_2026_2030_FR.docx"
    doc.save(out)
    log.info("Saved %s", out)
    return out


if __name__ == "__main__":
    build(Path(__file__).parent)
