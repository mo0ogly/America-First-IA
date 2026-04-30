"""
Chapter I — Introduction & Theoretical Framework — trilingual generator.

Generates the Chapter I .docx for the doctoral study
"AI for Americans First — AI Protectionism, Energy and Semiconductors:
US/Europe Divergence Trajectories 2024-2030"
in English, French and Brazilian Portuguese.

Key revisions vs previous version:
    1. CACI Power Mode formula explicitly defined in section 1.2
       (geometric weighted index F^0.40 x L^0.20 x R^0.15 / E^0.25,
       weights 40/25/20/15).
    2. Distinction between Physical CACI (F_total) and Sovereign CACI
       (F domestic only) introduced in section 1.2.
    3. Section 1.4: footnote pointing to public dashboard and working paper.
    4. Headline figures aligned with the dashboard live data
       (US share of operational compute = 76.9%, EU electricity / US ratio
       = 1.59, US/EU CACI Power ratio = 3.46:1).

Author: Fabrice Pizzi (Universite Paris-Sorbonne, M2 Economic Intelligence).
Build: python3 generate_chapter1_trilingual.py
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
log = logging.getLogger("chapter1_gen")


# ---------------------------------------------------------------------------
# Shared visual identity
# ---------------------------------------------------------------------------

NAVY = RGBColor(0x1A, 0x27, 0x44)
GOLD = RGBColor(0xB8, 0x92, 0x2F)
GREY = RGBColor(0x55, 0x55, 0x55)


@dataclass
class LangPack:
    """Container for one language version of Chapter I."""

    code: str                       # "EN", "FR", "PT-BR"
    filename: str                   # output .docx filename
    cover_subtitle: str             # below the main title
    cover_title: str                # AI for Americans First (translated)
    cover_blurb: str                # short tagline under the title
    cover_chip_lines: list[str]     # 3 KPIs on the cover
    cover_meta: str                 # "Paris - February 2026 / 7 chapters..."
    cover_keywords_label: str
    cover_keywords: str
    chapter_label: str              # "CHAPTER I", "CHAPITRE I", "CAPITULO I"
    chapter_title: str              # "Introduction and Theoretical Framework"
    sections: list[tuple[str, list[str]]] = field(default_factory=list)
    notes_label: str = "Notes"
    notes: list[str] = field(default_factory=list)
    license_block: list[str] = field(default_factory=list)
    page_footer: str = ""


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------


def set_run(run, *, font="Calibri", size=11, bold=False, italic=False, color=None):
    """Apply consistent typography to a docx run."""
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text, *, style=None, align=None, space_after=6, **run_kwargs):
    """Add a paragraph and return it (text may be empty for spacers)."""
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run(run, **run_kwargs)
    return p


def add_heading(doc, text, level):
    """Custom heading rendering (avoids built-in colored heading styles)."""
    sizes = {1: 22, 2: 16, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, font="Calibri", size=sizes.get(level, 11),
            bold=True, color=NAVY)
    return p


def add_cover(doc, lp: LangPack):
    """Render the cover block (title, KPIs, metadata)."""
    add_paragraph(doc, "", space_after=0)
    banner = {
        "EN": "RESEARCH STUDY - FEBRUARY 2026",
        "FR": "ETUDE DE RECHERCHE - FEVRIER 2026",
        "PT-BR": "ESTUDO DE PESQUISA - FEVEREIRO DE 2026",
    }.get(lp.code, "RESEARCH STUDY - FEBRUARY 2026")
    add_paragraph(doc, banner,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, bold=True, color=GREY, space_after=4)
    add_paragraph(doc, lp.cover_title,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=26, bold=True, color=NAVY, space_after=4)
    add_paragraph(doc, lp.cover_subtitle,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, italic=True, color=GREY, space_after=4)
    add_paragraph(doc, lp.cover_blurb,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=11, italic=True, color=GREY, space_after=18)

    # KPI band as a 3-column table
    table = doc.add_table(rows=1, cols=3)
    table.autofit = True
    for i, line in enumerate(lp.cover_chip_lines):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        set_run(run, size=11, bold=True, color=NAVY)

    add_paragraph(doc, "", space_after=8)
    add_paragraph(doc, "Fabrice Pizzi",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, bold=True, color=NAVY, space_after=2)
    add_paragraph(doc, "Universite Paris-Sorbonne",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, color=GREY, space_after=2)
    master_line = {
        "EN": "Master 2 Economic Intelligence - Intelligence Warfare",
        "FR": "Master 2 Intelligence Economique - Intelligence Warfare",
        "PT-BR": "Mestrado 2 Inteligencia Economica - Intelligence Warfare",
    }.get(lp.code, "Master 2 Economic Intelligence - Intelligence Warfare")
    add_paragraph(doc, master_line,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, italic=True, color=GREY, space_after=14)
    add_paragraph(doc, lp.cover_meta,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, italic=True, color=GREY, space_after=10)
    add_paragraph(doc, f"{lp.cover_keywords_label}: {lp.cover_keywords}",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=9, italic=True, color=GREY, space_after=24)


def add_chapter_header(doc, lp: LangPack):
    """Render the chapter header strip below the cover."""
    add_paragraph(doc, lp.chapter_label,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=11, bold=True, color=GOLD, space_after=2)
    add_paragraph(doc, lp.chapter_title,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=18, bold=True, color=NAVY, space_after=18)


def render_sections(doc, lp: LangPack):
    """Render the body sections from (heading, [paragraph, ...]) tuples."""
    for heading, paragraphs in lp.sections:
        # Detect heading level from numbering
        if heading.startswith(("1.1 ", "1.2 ", "1.3 ", "1.4 ", "1.5 ")):
            add_heading(doc, heading, 2)
        elif heading.startswith(("1.3.1", "1.3.2", "1.3.3", "1.3.4")):
            add_heading(doc, heading, 3)
        else:
            add_heading(doc, heading, 3)
        for para in paragraphs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.25
            run = p.add_run(para)
            set_run(run, size=11, color=RGBColor(0x20, 0x20, 0x20))


def render_notes(doc, lp: LangPack):
    """Render the footnotes block at the end of the chapter."""
    add_heading(doc, lp.notes_label, 2)
    for i, note in enumerate(lp.notes, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run_id = p.add_run(f"{i}. ")
        set_run(run_id, size=9, bold=True, color=GOLD)
        run_txt = p.add_run(note)
        set_run(run_txt, size=9, color=GREY)


def render_license(doc, lp: LangPack):
    """Render the license disclaimer block."""
    doc.add_paragraph()
    for line in lp.license_block:
        add_paragraph(doc, line,
                      align=WD_ALIGN_PARAGRAPH.LEFT,
                      size=8, italic=True, color=GREY, space_after=2)
    add_paragraph(doc, lp.page_footer,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=8, italic=True, color=GREY, space_after=0)


def build(lp: LangPack, out_dir: Path) -> Path:
    """Build the .docx file for one language pack."""
    log.info("Building Chapter I [%s] -> %s", lp.code, lp.filename)
    doc = Document()
    # Page setup A4
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    add_cover(doc, lp)
    add_chapter_header(doc, lp)
    render_sections(doc, lp)
    render_notes(doc, lp)
    render_license(doc, lp)

    out = out_dir / lp.filename
    doc.save(out)
    log.info("Saved %s", out)
    return out


# ===========================================================================
# Content - English
# ===========================================================================

EN = LangPack(
    code="EN",
    filename="Chapter_I_Introduction_Cadrage_Theorique_EN.docx",
    cover_title="AI FOR AMERICANS FIRST",
    cover_subtitle="AI Protectionism, Energy and Semiconductors: US/Europe Divergence Trajectories 2024-2030",
    cover_blurb="Integrated Geostrategic and Economic Analysis - Chapter I",
    cover_chip_lines=[
        "76.9% global operational AI compute = USA",
        "1.59x energy cost EU/US",
        "3.46:1 CACI ratio US/EU (Power Mode)",
    ],
    cover_meta="Paris - February 2026  |  7 chapters  |  4 prospective scenarios  |  3 geographic zones",
    cover_keywords_label="Keywords",
    cover_keywords=("artificial intelligence, technology protectionism, semiconductors, "
                    "export controls, sovereign compute, AI geopolitics, France, "
                    "United States, China"),
    chapter_label="CHAPTER I",
    chapter_title="Introduction and Theoretical Framework",
    notes_label="Notes",
    license_block=[
        "License and Disclaimer. This work, 'AI for Americans First,' is made available under the terms of the America-First-IA Creative Commons Attribution - NonCommercial - ShareAlike 4.0 International License (CC BY-NC-SA 4.0).",
        "You are free to share and adapt the material for non-commercial purposes, provided you give appropriate credit to Fabrice Pizzi (Universite Paris-Sorbonne) and distribute your contributions under the same license. This document is provided for educational and research purposes only.",
        "Public dashboard: https://mo0ogly.github.io/America-First-IA/dashboard/",
        "Repository: https://github.com/mo0ogly/America-First-IA",
    ],
    page_footer="AI for Americans First - Fabrice Pizzi - Chapter I",
)

EN.sections = [
    ("1.1 Research Question", [
        "Artificial intelligence is reshaping the foundations of global economic competitiveness. Since the launch of ChatGPT in November 2022 and the spectacular acceleration of investments in foundation models, generative AI has established itself as a cross-cutting transformative factor of the economy, simultaneously affecting finance, industry, healthcare, transportation and services. Yet this transformation rests on a precise material substrate: considerable computing capacity, powered by cutting-edge semiconductors and abundant electrical energy. Mastery of this triptych - compute, chips, energy - has become a geostrategic issue of the first order.",
        "In this context, the United States has progressively erected a control regime over access to frontier AI technologies. As early as October 2022, the Bureau of Industry and Security (BIS) of the Department of Commerce imposed restrictions on the export of advanced GPUs to China. In January 2025, the Biden administration extended these controls to more than 120 countries via the AI Diffusion Rule, creating a tier-based segmentation system that conditions access to the most powerful AI chips on the degree of geopolitical alignment with Washington. The Trump administration, taking office in January 2025, replaced this framework with a more explicitly competitive approach, culminating in July 2025 with the publication of America's AI Action Plan, then in January 2026 with the imposition of 25 percent tariffs on certain advanced AI semiconductors (Nvidia H200, AMD MI325X) under Section 232.[1]",
        "These measures, officially motivated by national security imperatives, de facto produce a structural competitive advantage for American companies: they benefit from unlimited access to frontier compute, while actors from other regions - including European allies - see their capacities capped, made more expensive, or conditioned. We are thus witnessing the emergence of a new type of technology protectionism, where the tax is not merely tariff-based but also regulatory, logistical and strategic.",
        "This study poses the following question: to what extent does US technology protectionism on AI - export controls, tariffs, domestic compute prioritization - create a structural competitiveness divergence between the United States and Europe, and what are the measurable consequences for France by 2030?",
        "This research question breaks down into three sub-questions articulating the empirical, prospective and normative dimensions of the analysis:",
        "(a) What is the current AI computing capacity gap (compute gap) between the United States and the European Union, and how does this gap evolve under different scenarios of American trade and technology policy?",
        "(b) How does the asymmetry of compute access translate into sectoral productivity, model training costs, and market share in AI services?",
        "(c) Do European energy constraints and the rise of AI robotics amplify the divergence, and can France leverage its nuclear advantage to mitigate this structural disadvantage?",
        "The originality of this study lies in its integrated approach. The existing literature treats separately export controls (Carnegie Endowment, CSIS, Hudson Institute), data center energy projections (IEA), semiconductor market dynamics (McKinsey, SIA, Deloitte), compute sovereignty (Hawkins, Lehdonvirta and Wu, 2025), and AI-related competitive barriers (Bruegel, OECD). No work proposes the complete causal trajectory we seek to establish here: US protectionism then restriction of European compute then productivity divergence then strategic dependence, with energy and robotics as amplifying factors.",
    ]),
    ("1.2 Operational Definitions", [
        "The analysis deployed in this study rests on a set of concepts that require rigorous definition, as their usage varies across disciplinary contexts.",
        "Frontier compute. We define frontier compute as the installed computing capacity in the form of AI accelerators - GPUs (Nvidia A100, H100, H200, B200), ASICs (Google TPU), or specialized circuits - deployed in industrial-scale data centers. This capacity is measured in FLOPs (floating-point operations per second) aggregated at the national or regional level, in H100-equivalents (the metric used by the Epoch AI dataset), or in GW of IT load. Frontier compute is the fundamental input for the development and deployment of frontier AI: without it, it is impossible to train competitive foundation models or operate large-scale inference services.",
        "AI technology protectionism. AI technology protectionism designates the set of state measures - export controls, tariffs, quotas, export licenses, logistical prioritization, restrictions on model weights and APIs - that create an asymmetry of access to compute, models, and AI services between geographic regions. This concept extends the classical notion of trade protectionism by integrating the intangible dimension (software, models, cloud services) and the infrastructural dimension (energy, chip production capacity). Unlike traditional tariff barriers, AI technology protectionism can operate through non-tariff channels - for example, prioritizing GPU deliveries to domestic companies in a context of global shortage.",
        "Compute gap. The compute gap measures the ratio of effectively available AI computing capacity between two regions. We define it as the ratio between installed and accessible AI FLOPs (or H100-equivalents) for economic actors in region A and those in region B, normalized by active population or GDP. A compute gap of x15 between the United States and the EU means that, per unit of GDP, American actors have fifteen times more AI computing power than their European counterparts. This ratio constitutes a synthetic indicator of structural advantage in AI.",
        "Compute sovereignty. We adopt the definition proposed by Hawkins, Lehdonvirta and Wu (2025), who decompose compute sovereignty into three levels: (1) the quantity of AI compute physically present on national territory, (2) the nationality of the companies owning the data centers, and (3) the nationality of the accelerator suppliers whose chips power these data centers.[2] A country can have significant computing capacity on its soil while being dependent on foreign operators and chip suppliers, which limits its real sovereignty. This concept is essential for understanding why the presence of AWS or Azure data centers in Europe does not, in itself, constitute a form of European sovereignty over compute.",
        "Geopolitical vendor lock-in. We propose the concept of geopolitical vendor lock-in to designate the structural dependence of an economic ecosystem on foreign technology providers whose access can be restricted, made more expensive, or conditioned by political decision of a third-party government. This concept extends the classical notion of vendor lock-in in IT (cloud provider migration costs, for example) by adding a geopolitical dimension: the risk that access to critical infrastructure may be used as a negotiating lever between states. The Starlink-Ukraine episode of March 2025, where American control of a communication system was perceived as an instrument of pressure, strikingly illustrates this type of risk.[3]",
        "Compute-Adjusted Competitiveness Index (CACI). We construct, as the central original indicator of this study, a Compute-Adjusted Competitiveness Index combining the four dimensions identified above into a weighted geometric composite. In its absolute formulation, called Power Mode, the indicator is defined as CACI = F^0.40 x L^0.20 x R^0.15 / E^0.25, where F is the installed AI compute (H100-equivalents, total or sovereign), L is the relevant labor force (millions of AI-skilled workers), R is a geopolitical access index (0 to 1), and E is the average industrial electricity price (USD per MWh). The four exponents (40, 20, 15, 25) sum to one and reflect the relative importance attributed to compute, labor, geopolitical access and energy cost. A second formulation, called Intensity Mode, divides the result by GDP to obtain a competitiveness density per unit of national wealth. The full methodological derivation is presented in Chapter II and the indicator is updated in real time on the public dashboard.[a]",
        "Two variants of the indicator are used in parallel. Physical CACI integrates all compute physically present on the territory (F_total), regardless of ownership. Sovereign CACI restricts F to compute owned and operated by domestic actors (F_dom), thereby capturing the actual sovereign autonomy. The gap between the two variants measures the structural dependence of a region on foreign hyperscalers. For the European Union, this gap is the most direct quantitative measure of the strategic vulnerability that this study seeks to characterize.",
    ]),
    ("1.3 Theoretical Framework", [
        "Our analysis is grounded in four complementary theoretical currents, which provide the conceptual tools necessary to articulate the technological, economic and geopolitical dimensions of the phenomenon under study.",
    ]),
    ("1.3.1 AI as a General Purpose Technology", [
        "The first theoretical anchor is the General Purpose Technologies (GPT) theory formalized by Bresnahan and Trajtenberg (1995). A GPT is a technology characterized by three properties: its pervasiveness (it is used as an input in numerous downstream sectors), its inherent potential for technical improvement, and its innovation complementarities (R&D productivity in user sectors increases as a consequence of GPT improvement).[4] The model predicts that GPTs generate increasing returns to scale and that their diffusion throughout the economy is a source of generalized productivity gains. However, Bresnahan and Trajtenberg also emphasize that a decentralized economy may have difficulty fully exploiting a GPT potential, as market transactions between the GPT producer and its users may lead to too little, too late innovation.",
        "Brynjolfsson, Rock and Syverson (2019) applied this framework to contemporary AI, demonstrating that AI, and in particular machine learning, satisfies Bresnahan and Trajtenberg three criteria to be qualified as a GPT.[5] Their productivity J-curve model explains why productivity gains linked to a GPT may initially be invisible in statistics: firms must first invest massively in intangible assets (reorganization, training, process reengineering) before reaping the benefits.",
        "This framework is fundamental to our analysis because it implies that early and massive access to AI compute - that is, the capacity to invest in the GPT at large scale from the earliest deployment phases - generates cumulative advantages that are difficult to reverse. Innovation complementarities create a path dependence dynamic: actors who access abundant compute early develop superior models, capture usage data, generate revenues they reinvest in compute, and thus widen a gap that self-reinforces over time. Any policy restricting compute access for a given region therefore has non-linear consequences: it does not merely delay adoption, it structurally compromises it.",
    ]),
    ("1.3.2 Weaponized Interdependence and Control of Global Networks", [
        "The second theoretical anchor is the weaponized interdependence theory developed by Farrell and Newman (2019). These authors demonstrate that global economic networks, far from creating symmetrical relations of mutual interdependence as classical liberal theory posited, tend to produce asymmetric structures in which certain nodes (hubs) become far more connected than others. States exercising political jurisdiction over these central nodes can instrumentalize them for coercive purposes, via two mechanisms: the panopticon effect (strategic information collection) and the chokepoint effect (capacity to cut or restrict flows).[6]",
        "The application to the AI value chain is remarkably pertinent. The United States controls several critical chokepoints: AI accelerator design (Nvidia holds over 80 percent of the GPU market for data centers), cloud infrastructure (AWS, Azure and Google Cloud represent approximately 70 percent of the global market), and the most advanced foundation models (OpenAI, Anthropic, Google DeepMind). Control of these chokepoints allows the US government to modulate global access to AI compute as a geopolitical lever, in exactly the same way that the SWIFT system was used as a lever in the financial domain. Farrell and Newman themselves recognized, in an update published in Foreign Affairs in December 2025, that semiconductors and AI had become a major application terrain for their theory, the Trump administration having explicitly used export controls on AI chips as bargaining currency in negotiations with China.[7]",
    ]),
    ("1.3.3 Concentration Economics and Innovation Rents", [
        "The third theoretical anchor mobilizes the literature on digital market concentration and barriers to entry in the AI ecosystem. Martens (2024), in a policy brief for Bruegel, demonstrates that the training costs of foundation models grow exponentially, constituting an insurmountable barrier to entry for most actors.[8] He estimates that a compute farm on the order of a trillion dollars is conceivable in the medium term, an investment threshold completely beyond the reach of public funding and the vast majority of companies. Only the GAMMAN (Google, Apple, Meta, Microsoft, Amazon, Nvidia) have the resources to access it.",
        "The OECD (2025) confirms this analysis in its report on competition in AI infrastructure, identifying barriers to entry at multiple levels of the supply chain: extremely high capital requirements, massive economies of scale, switching costs between providers, and absence of interoperability standards.[9] The Federal Reserve Board (October 2025), in a comparative analysis of AI competitiveness in advanced economies, shows that the United States concentrates more than 75 percent of global venture capital investment in generative AI, and that Europe lags significantly not only in investment but also in enterprise adoption, with high energy costs constituting an additional brake.[10]",
        "This concentration dynamic has a direct consequence for our analysis: US technology protectionism does not merely limit European access to compute, it reinforces the position of dominant actors who are precisely those benefiting from domestic exemption. It is a dual advantage mechanism: reduction of constraints for US companies, increase of constraints for their competitors.",
    ]),
    ("1.3.4 Digital Sovereignty and European Strategic Autonomy", [
        "The fourth theoretical anchor mobilizes the European current of digital sovereignty, analyzed notably by Mugge (2024) in the Journal of European Public Policy. Mugge identifies three fundamental tensions in Europe AI sovereignty ambition: does sovereignty oppose the EU to other AI powers, or citizens to large platforms? Does it aim at economic competitiveness or rights protection? And who really benefits from it - European champions or the entire ecosystem?[11] These tensions are precisely those that American protectionism exacerbates: by compressing European technological sovereignty space, it forces the EU to arbitrate between these contradictory objectives in a context of urgency.",
        "Hawkins, Lehdonvirta and Wu (2025) provide a decisive empirical contribution by measuring compute sovereignty through the infrastructure of the nine main global cloud providers, which represent approximately 70 percent of the global market. Their results reveal that the degree of sovereignty varies considerably depending on the level of analysis (territorial, corporate, or hardware), and that most European countries exhibit a sovereignty deficit at at least two of these three levels.[12] McKinsey (December 2025) estimates the European sovereign AI opportunity at 480 billion euros annually by 2030, contingent on a scenario of strong technological sovereignty and high AI adoption.[13]",
    ]),
    ("1.4 Targeted Literature Review and Gap Identification", [
        "The state of the literature in February 2026 reveals a rapidly forming field of research, but one that remains highly fragmented. We identify five main corpora, each covering one dimension of our research question.",
        "Corpus 1: Export controls and AI geopolitics. The Carnegie Endowment for International Peace provides the most detailed analysis of US AI export control policy. Winter-Levy and Phillips-Robins (May 2025) described the Biden AI Diffusion Rule as a compromise between three objectives - control, promotion and geopolitical leverage - and analyzed replacement options under Trump.[14] CSIS, the Hudson Institute and Pillsbury Law have documented the legal mechanisms and recent developments, notably the Section 232 tariffs of January 2026. Contrary Research (November 2025) offers a particularly rich analysis of the underlying temporal dilemma: if AGI is five years away, export controls strengthen American dominance; if it is ten years or more, they accelerate Chinese technological autonomization.[15]",
        "Corpus 2: Energy and data center infrastructure. The IEA special report Energy and AI (April 2025) constitutes the global reference for data center energy projections. It establishes that global data center electricity consumption will reach 945 TWh in 2030 in the base scenario, compared to 415 TWh in 2024, with AI as the primary driver of this growth.[16] In the United States, data centers will consume more electricity than all energy-intensive industry combined by 2030. In Europe, growth will be plus 45 TWh (plus 70 percent), with a risk of delay on approximately 20 percent of projects linked to electricity grid constraints. IEA-4E published in parallel a critical review of estimation models, highlighting the magnitude of uncertainties (2030 projections vary by a factor of 40 across studies).[17]",
        "Corpus 3: Semiconductor market. McKinsey (January 2026) significantly revised upward its estimates of the semiconductor market size by integrating captive designers (Apple, Amazon, Tesla) and fabless operators whose value does not appear in traditional statistics. Their base estimate goes from 775 billion USD in 2024 to 1,600 billion USD in 2030, a CAGR of 13 percent.[18] SIA reports record sales of 627.6 billion USD in 2024, while Deloitte (February 2026) anticipates that generative AI chips alone will represent nearly half of the sector revenue in 2026.[19]",
        "Corpus 4: European AI sovereignty and competitiveness. Several recent publications document the European deficit. The European Parliament (2025) finds that only 11 percent of small European enterprises use AI, compared to 58 percent of small American businesses.[20] Accenture (November 2025) reports that 62 percent of European organizations are seeking sovereign solutions in the face of geopolitical uncertainty, but only 19 percent see it as a competitive advantage. The majority (48 percent) is motivated by regulatory compliance obligations, suggesting a defensive rather than strategic approach.[21] The Draghi report (September 2024) on European competitiveness had already identified the digital investment deficit as a structural factor in Europe decline.",
        "Corpus 5: Competitive barriers and concentration. Bruegel, the OECD and the Federal Reserve Board converge on the observation of growing concentration of the AI ecosystem around a small number of actors. The CERRE report (June 2025) on competition policy for cloud and AI identifies migration barriers, cloud credit practices likely to create lock-in, and growing dependence of small AI developers on hyperscalers for access to accelerated compute.[22] The structural collaboration between AI startups and Big Tech, illustrated by the Mistral-Microsoft agreement, testifies to this dependence.",
        "The gap this study aims to fill. Examination of these five corpora reveals an uncovered analytical space: no study proposes an integrated trajectory linking US technology protectionism to EU competitiveness divergence via the compute-energy-semiconductor triptych. Existing works treat each dimension in isolation - export controls without energy, energy without semiconductors, semiconductors without productivity. Moreover, the specific angle of AI robotics as an amplifier of energy demand is virtually absent from the literature. It is precisely this gap that our study intends to fill, by proposing a unified analytical framework and an original indicator - the Compute-Adjusted Competitiveness Index (CACI) - enabling measurement and projection of this divergence. The indicator and the underlying dataset are publicly accessible and updated in real time on the project dashboard.[a]",
    ]),
    ("1.5 Report Structure", [
        "The report is organized in eight chapters. After this introduction, Chapter II presents the methodology, detailing the multi-scenario approach, data sources, CACI construction and methodological limitations. Chapter III establishes the 2020-2026 empirical diagnosis, covering energy trajectories, the semiconductor market, installed compute and the export controls chronology. Chapter IV analyzes the mechanisms of US competitive advantage. Chapter V presents four 2026-2030 scenarios structured around two axes of uncertainty. Chapter VI details the consequences for France and Europe, differentiated by actor type and sector. Chapter VII formulates strategic recommendations at three time horizons. Chapter VIII concludes by synthesizing the study contributions and identifying avenues for future research.",
    ]),
]

EN.notes = [
    "Pillsbury Law (2026), 'Trump Admin Targets Advanced AI Semiconductors, Defers Broader Tariffs,' January 15, 2026. The Section 232 proclamation imposes a 25 percent tariff on logic integrated circuits meeting specific technical parameters (TTP > 14,000, DRAM bandwidth > 4,500 GB/s), covering notably Nvidia H200 and AMD MI325X destined for re-export.",
    "Hawkins, Z.J., Lehdonvirta, V. and Wu, B. (2025), 'AI Compute Sovereignty: Infrastructure Control Across Territories, Cloud Providers and Accelerators,' SSRN, June 2025, https://ssrn.com/abstract=5312977",
    "Carnegie Endowment for International Peace (2025), 'The EU AI Power Play: Between Deregulation and Innovation,' May 2025. See also the March 2025 revelation on the use of Starlink as a pressure lever on Ukraine.",
    "Bresnahan, T.F. and Trajtenberg, M. (1995), 'General Purpose Technologies Engines of Growth?,' Journal of Econometrics, 65(1), pp. 83-108.",
    "Brynjolfsson, E., Rock, D. and Syverson, C. (2019), 'Artificial Intelligence and the Modern Productivity Paradox: A Clash of Expectations and Statistics,' in Agrawal, Gans and Goldfarb (eds.), The Economics of Artificial Intelligence, University of Chicago Press, pp. 23-57. See also Brynjolfsson, Rock and Syverson (2021), 'The Productivity J-Curve,' American Economic Journal: Macroeconomics, 13(1), pp. 268-320.",
    "Farrell, H. and Newman, A.L. (2019), 'Weaponized Interdependence: How Global Economic Networks Shape State Coercion,' International Security, 44(1), pp. 42-79. doi:10.1162/isec_a_00351",
    "Farrell, H. and Newman, A.L. (2025), 'The Weaponized World Economy: Surviving the New Age of Economic Coercion,' Foreign Affairs, December 2025.",
    "Martens, B. (2024), 'Why artificial intelligence is creating fundamental challenges for competition policy,' Bruegel Policy Brief 16/2024.",
    "OECD (2025), 'Competition in Artificial Intelligence Infrastructure,' https://www.oecd.org/en/publications/competition-in-artificial-intelligence-infrastructure_623d1874-en.html",
    "Federal Reserve Board (2025), 'The State of AI Competition in Advanced Economies,' FEDS Notes, October 6, 2025.",
    "Mugge, D. (2024), 'EU AI sovereignty: for whom, to what end, and to whose benefit?,' Journal of European Public Policy, 31(8), pp. 2200-2225. doi:10.1080/13501763.2024.2318475",
    "Hawkins, Lehdonvirta and Wu (2025), op. cit.",
    "McKinsey (2025), 'Accelerating Europe AI Adoption: The Role of Sovereign AI Capabilities,' December 2025.",
    "Winter-Levy, S. and Phillips-Robins, A. (2025), 'The Trump Administration May Be About to Repeal the AI Diffusion Rule. Here Is What It Should Do Next,' Carnegie Endowment, May 8, 2025.",
    "Contrary Research (2025), 'Deep Dive: Export Controls and the AI Race,' November 6, 2025. https://research.contrary.com/report/drawing-geopolitical-boundaries",
    "IEA (2025), 'Energy and AI,' special report, April 10, 2025. https://www.iea.org/reports/energy-and-ai",
    "IEA-4E (2025), 'Data Centre Energy Use: Critical Review of Models and Results.'",
    "McKinsey (2026), 'Hiding in Plain Sight: The Underestimated Size of the Semiconductor Industry,' January 2026.",
    "Deloitte (2026), '2026 Semiconductor Industry Outlook,' February 2026. SIA (2025), 2024 sales = 627.6 billion USD (WSTS).",
    "European Parliament (2025), 'Making Europe an AI Continent,' EPRS BRI(2025)775923. EU large enterprise AI adoption: 41 percent vs 11 percent for small ones. US small business adoption: 58 percent (US Chamber of Commerce, August 2025).",
    "Accenture (2025), 'Europe Seeking Greater AI Sovereignty,' November 3, 2025. Survey of 1,928 organizations in 28 countries.",
    "CERRE / Meyers, Z. (2025), 'A Competition Policy for Cloud and AI,' Issue Paper, June 2025.",
    "[a] Public dashboard: https://mo0ogly.github.io/America-First-IA/dashboard/. Working Paper: https://mo0ogly.github.io/America-First-IA/pdf/Working_Paper_CACI_AI_Competitiveness.pdf. The CACI Power Mode formula F^0.40 x L^0.20 x R^0.15 / E^0.25 reproduces the live calculation; the headline ratio of 3.46:1 between USA and EU (aggregated) is computed on the live dataset as of April 2026.",
]


# ===========================================================================
# Content - French
# ===========================================================================

FR = LangPack(
    code="FR",
    filename="Chapitre_I_Introduction_Cadrage_Theorique_FR.docx",
    cover_title="AI FOR AMERICANS FIRST",
    cover_subtitle="Protectionnisme IA, Energie et Semi-conducteurs : Trajectoires de divergence US/Europe 2024-2030",
    cover_blurb="Analyse geostrategique et economique integree - Chapitre I",
    cover_chip_lines=[
        "76,9 pct du compute IA operationnel mondial = USA",
        "1,59x cout energie EU/US",
        "3,46:1 ratio CACI US/EU (Power Mode)",
    ],
    cover_meta="Paris - fevrier 2026  |  7 chapitres  |  4 scenarios prospectifs  |  3 zones geographiques",
    cover_keywords_label="Mots-cles",
    cover_keywords=("intelligence artificielle, protectionnisme technologique, semi-conducteurs, "
                    "controles a l'exportation, compute souverain, geopolitique IA, France, "
                    "Etats-Unis, Chine"),
    chapter_label="CHAPITRE I",
    chapter_title="Introduction et cadrage theorique",
    notes_label="Notes",
    license_block=[
        "Licence et avertissement. Cette oeuvre, 'AI for Americans First', est mise a disposition selon les termes de la licence Creative Commons Attribution - Pas d'Utilisation Commerciale - Partage dans les Memes Conditions 4.0 International (CC BY-NC-SA 4.0) du projet America-First-IA.",
        "Vous etes libre de partager et adapter le materiel a des fins non commerciales, a condition d'attribuer correctement l'oeuvre a Fabrice Pizzi (Universite Paris-Sorbonne) et de distribuer vos contributions sous la meme licence. Ce document est fourni a des fins educatives et de recherche uniquement.",
        "Tableau de bord public : https://mo0ogly.github.io/America-First-IA/dashboard/",
        "Depot : https://github.com/mo0ogly/America-First-IA",
    ],
    page_footer="AI for Americans First - Fabrice Pizzi - Chapitre I",
)

FR.sections = [
    ("1.1 Question de recherche", [
        "L'intelligence artificielle reconfigure les fondations de la competitivite economique mondiale. Depuis le lancement de ChatGPT en novembre 2022 et l'acceleration spectaculaire des investissements dans les modeles de fondation, l'IA generative s'est imposee comme un facteur transformateur transversal de l'economie, affectant simultanement la finance, l'industrie, la sante, les transports et les services. Or cette transformation repose sur un substrat materiel precis : une capacite de calcul considerable, alimentee par des semi-conducteurs de pointe et une energie electrique abondante. La maitrise de ce triptyque - compute, puces, energie - est devenue un enjeu geostrategique de premier ordre.",
        "Dans ce contexte, les Etats-Unis ont progressivement erige un regime de controle sur l'acces aux technologies IA de frontiere. Des octobre 2022, le Bureau of Industry and Security (BIS) du Department of Commerce a impose des restrictions a l'exportation des GPU avances vers la Chine. En janvier 2025, l'administration Biden a etendu ces controles a plus de 120 pays via l'AI Diffusion Rule, creant un systeme de segmentation par paliers conditionnant l'acces aux puces IA les plus puissantes au degre d'alignement geopolitique avec Washington. L'administration Trump, entree en fonction en janvier 2025, a remplace ce cadre par une approche plus explicitement competitive, culminant en juillet 2025 avec la publication du America's AI Action Plan, puis en janvier 2026 avec l'imposition de tarifs de 25 pour cent sur certains semi-conducteurs IA avances (Nvidia H200, AMD MI325X) au titre de la Section 232.[1]",
        "Ces mesures, officiellement motivees par des imperatifs de securite nationale, produisent de facto un avantage competitif structurel pour les entreprises americaines : elles beneficient d'un acces illimite au compute de frontiere, tandis que les acteurs d'autres regions - y compris les allies europeens - voient leurs capacites plafonnees, rencheries ou conditionnees. Nous assistons ainsi a l'emergence d'un nouveau type de protectionnisme technologique, ou la taxe n'est pas seulement tarifaire mais egalement reglementaire, logistique et strategique.",
        "Cette etude pose la question suivante : dans quelle mesure le protectionnisme technologique americain sur l'IA - controles a l'exportation, tarifs, priorisation du compute domestique - cree-t-il une divergence structurelle de competitivite entre les Etats-Unis et l'Europe, et quelles en sont les consequences mesurables pour la France a l'horizon 2030 ?",
        "Cette question de recherche se decompose en trois sous-questions articulant les dimensions empirique, prospective et normative de l'analyse :",
        "(a) Quel est l'ecart actuel de capacite de calcul IA (compute gap) entre les Etats-Unis et l'Union europeenne, et comment cet ecart evolue-t-il selon differents scenarios de politique commerciale et technologique americaine ?",
        "(b) Comment l'asymetrie d'acces au compute se traduit-elle en productivite sectorielle, en couts d'entrainement des modeles et en parts de marche dans les services IA ?",
        "(c) Les contraintes energetiques europeennes et la montee de la robotique IA amplifient-elles la divergence, et la France peut-elle mobiliser son avantage nucleaire pour mitiger ce desavantage structurel ?",
        "L'originalite de cette etude reside dans son approche integree. La litterature existante traite separement les controles a l'exportation (Carnegie Endowment, CSIS, Hudson Institute), les projections energetiques des centres de donnees (AIE), la dynamique du marche des semi-conducteurs (McKinsey, SIA, Deloitte), la souverainete du compute (Hawkins, Lehdonvirta et Wu, 2025) et les barrieres concurrentielles liees a l'IA (Bruegel, OCDE). Aucun travail ne propose la trajectoire causale complete que nous cherchons a etablir ici : protectionnisme americain puis restriction du compute europeen puis divergence de productivite puis dependance strategique, avec l'energie et la robotique comme facteurs amplificateurs.",
    ]),
    ("1.2 Definitions operationnelles", [
        "L'analyse deployee dans cette etude repose sur un ensemble de concepts qui requierent une definition rigoureuse, leur usage variant selon les contextes disciplinaires.",
        "Compute de frontiere. Nous definissons le compute de frontiere comme la capacite de calcul installee sous forme d'accelerateurs IA - GPU (Nvidia A100, H100, H200, B200), ASIC (Google TPU) ou circuits specialises - deployes dans des centres de donnees a l'echelle industrielle. Cette capacite se mesure en FLOPs (operations en virgule flottante par seconde) agreges au niveau national ou regional, en H100-equivalents (la metrique utilisee par le jeu de donnees Epoch AI), ou en GW de charge IT. Le compute de frontiere constitue l'intrant fondamental du developpement et du deploiement de l'IA de frontiere : sans lui, il est impossible d'entrainer des modeles de fondation competitifs ou d'operer des services d'inference a grande echelle.",
        "Protectionnisme technologique IA. Le protectionnisme technologique IA designe l'ensemble des mesures etatiques - controles a l'exportation, tarifs, quotas, licences d'exportation, priorisation logistique, restrictions sur les poids de modeles et les API - qui creent une asymetrie d'acces au compute, aux modeles et aux services IA entre regions geographiques. Ce concept etend la notion classique de protectionnisme commercial en integrant la dimension immaterielle (logiciels, modeles, services cloud) et la dimension infrastructurelle (energie, capacites de production de puces). A la difference des barrieres tarifaires traditionnelles, le protectionnisme technologique IA peut operer par des canaux non tarifaires - par exemple en priorisant les livraisons de GPU aux entreprises domestiques en contexte de penurie mondiale.",
        "Compute gap. Le compute gap mesure le ratio de capacite de calcul IA effectivement disponible entre deux regions. Nous le definissons comme le ratio entre les FLOPs IA installes et accessibles (ou H100-equivalents) pour les acteurs economiques de la region A et ceux de la region B, normalise par la population active ou le PIB. Un compute gap de x15 entre les Etats-Unis et l'UE signifie que, par unite de PIB, les acteurs americains disposent de quinze fois plus de puissance de calcul IA que leurs homologues europeens. Ce ratio constitue un indicateur synthetique de l'avantage structurel en IA.",
        "Souverainete du compute. Nous adoptons la definition proposee par Hawkins, Lehdonvirta et Wu (2025), qui decomposent la souverainete du compute en trois niveaux : (1) la quantite de compute IA physiquement presente sur le territoire national, (2) la nationalite des entreprises proprietaires des centres de donnees, et (3) la nationalite des fournisseurs d'accelerateurs dont les puces alimentent ces centres.[2] Un pays peut disposer d'une capacite de calcul significative sur son sol tout en etant dependant d'operateurs et de fournisseurs etrangers, ce qui limite sa souverainete reelle. Ce concept est essentiel pour comprendre pourquoi la presence de centres AWS ou Azure en Europe ne constitue pas, en soi, une forme de souverainete europeenne sur le compute.",
        "Vendor lock-in geopolitique. Nous proposons le concept de vendor lock-in geopolitique pour designer la dependance structurelle d'un ecosysteme economique a des fournisseurs technologiques etrangers dont l'acces peut etre restreint, rencheri ou conditionne par decision politique d'un gouvernement tiers. Ce concept etend la notion classique de vendor lock-in informatique (couts de migration entre fournisseurs cloud, par exemple) en y ajoutant une dimension geopolitique : le risque que l'acces a une infrastructure critique soit utilise comme levier de negociation entre Etats. L'episode Starlink-Ukraine de mars 2025, ou le controle americain d'un systeme de communication a ete percu comme un instrument de pression, illustre de facon saisissante ce type de risque.[3]",
        "Compute-Adjusted Competitiveness Index (CACI). Nous construisons, comme indicateur original central de cette etude, un Compute-Adjusted Competitiveness Index combinant en un composite geometrique pondere les quatre dimensions identifiees ci-dessus. Dans sa formulation absolue, dite Power Mode, l'indicateur s'ecrit CACI = F^0,40 x L^0,20 x R^0,15 / E^0,25, ou F est le compute IA installe (H100-equivalents, total ou souverain), L est la main-d'oeuvre pertinente (millions de travailleurs qualifies en IA), R est un indice d'acces geopolitique (0 a 1) et E est le prix moyen de l'electricite industrielle (USD par MWh). Les quatre exposants (40, 20, 15, 25) somment a un et refletent l'importance relative attribuee au compute, au travail, a l'acces geopolitique et au cout de l'energie. Une seconde formulation, dite Intensity Mode, divise le resultat par le PIB pour obtenir une densite de competitivite par unite de richesse nationale. La derivation methodologique complete est presentee au chapitre II et l'indicateur est mis a jour en temps reel sur le tableau de bord public.[a]",
        "Deux variantes de l'indicateur sont utilisees en parallele. Le CACI physique integre tout le compute physiquement present sur le territoire (F_total), independamment de la propriete. Le CACI souverain restreint F au compute possede et opere par des acteurs domestiques (F_dom), capturant ainsi l'autonomie souveraine effective. L'ecart entre les deux variantes mesure la dependance structurelle d'une region aux hyperscalers etrangers. Pour l'Union europeenne, cet ecart constitue la mesure quantitative la plus directe de la vulnerabilite strategique que cette etude cherche a caracteriser.",
    ]),
    ("1.3 Cadre theorique", [
        "Notre analyse s'enracine dans quatre courants theoriques complementaires, qui fournissent les outils conceptuels necessaires pour articuler les dimensions technologique, economique et geopolitique du phenomene etudie.",
    ]),
    ("1.3.1 L'IA comme General Purpose Technology", [
        "Le premier ancrage theorique est la theorie des General Purpose Technologies (GPT) formalisee par Bresnahan et Trajtenberg (1995). Une GPT est une technologie caracterisee par trois proprietes : sa pervasivite (elle est utilisee comme intrant dans de nombreux secteurs en aval), son potentiel inherent d'amelioration technique, et ses complementarites d'innovation (la productivite R&D dans les secteurs utilisateurs croit en consequence de l'amelioration de la GPT).[4] Le modele predit que les GPT generent des rendements d'echelle croissants et que leur diffusion dans l'economie est source de gains de productivite generalises. Toutefois, Bresnahan et Trajtenberg soulignent egalement qu'une economie decentralisee peut peiner a exploiter pleinement le potentiel d'une GPT, les transactions de marche entre le producteur de GPT et ses utilisateurs pouvant conduire a une innovation insuffisante et tardive.",
        "Brynjolfsson, Rock et Syverson (2019) ont applique ce cadre a l'IA contemporaine, demontrant que l'IA, et en particulier le machine learning, satisfait les trois criteres de Bresnahan et Trajtenberg pour etre qualifiee de GPT.[5] Leur modele de productivity J-curve explique pourquoi les gains de productivite lies a une GPT peuvent etre initialement invisibles dans les statistiques : les entreprises doivent d'abord investir massivement dans des actifs intangibles (reorganisation, formation, reingenierie des processus) avant d'en recolter les benefices.",
        "Ce cadre est fondamental pour notre analyse car il implique que l'acces precoce et massif au compute IA - c'est-a-dire la capacite a investir dans la GPT a grande echelle des les premieres phases de deploiement - genere des avantages cumulatifs difficilement reversibles. Les complementarites d'innovation creent une dynamique de path dependence : les acteurs qui accedent tot a un compute abondant developpent des modeles superieurs, captent des donnees d'usage, generent des revenus qu'ils reinvestissent dans le compute, et creusent ainsi un ecart qui s'auto-renforce dans le temps. Toute politique restreignant l'acces au compute pour une region donnee a donc des consequences non lineaires : elle ne fait pas que retarder l'adoption, elle la compromet structurellement.",
    ]),
    ("1.3.2 Interdependance armee et controle des reseaux mondiaux", [
        "Le deuxieme ancrage theorique est la theorie de la weaponized interdependence developpee par Farrell et Newman (2019). Ces auteurs demontrent que les reseaux economiques mondiaux, loin de creer des relations symetriques d'interdependance mutuelle comme le postulait la theorie liberale classique, tendent a produire des structures asymetriques dans lesquelles certains noeuds (hubs) deviennent bien plus connectes que d'autres. Les Etats exercant une juridiction politique sur ces noeuds centraux peuvent les instrumentaliser a des fins coercitives, via deux mecanismes : l'effet panopticon (collecte d'informations strategiques) et l'effet chokepoint (capacite de couper ou de restreindre les flux).[6]",
        "L'application a la chaine de valeur IA est remarquablement pertinente. Les Etats-Unis controlent plusieurs chokepoints critiques : la conception des accelerateurs IA (Nvidia detient plus de 80 pour cent du marche GPU pour centres de donnees), l'infrastructure cloud (AWS, Azure et Google Cloud representent environ 70 pour cent du marche mondial) et les modeles de fondation les plus avances (OpenAI, Anthropic, Google DeepMind). Le controle de ces chokepoints permet au gouvernement americain de moduler l'acces mondial au compute IA comme un levier geopolitique, exactement comme le systeme SWIFT a ete utilise comme levier dans le domaine financier. Farrell et Newman ont eux-memes reconnu, dans une mise a jour publiee dans Foreign Affairs en decembre 2025, que les semi-conducteurs et l'IA etaient devenus un terrain d'application majeur pour leur theorie, l'administration Trump ayant explicitement utilise les controles a l'exportation sur les puces IA comme monnaie d'echange dans les negociations avec la Chine.[7]",
    ]),
    ("1.3.3 Economie de la concentration et rentes d'innovation", [
        "Le troisieme ancrage theorique mobilise la litterature sur la concentration des marches numeriques et les barrieres a l'entree dans l'ecosysteme IA. Martens (2024), dans un policy brief pour Bruegel, demontre que les couts d'entrainement des modeles de fondation croissent exponentiellement, constituant une barriere a l'entree insurmontable pour la plupart des acteurs.[8] Il estime qu'une ferme de calcul de l'ordre du trillion de dollars est concevable a moyen terme, un seuil d'investissement totalement hors de portee du financement public et de la grande majorite des entreprises. Seules les GAMMAN (Google, Apple, Meta, Microsoft, Amazon, Nvidia) disposent des ressources pour y acceder.",
        "L'OCDE (2025) confirme cette analyse dans son rapport sur la concurrence dans l'infrastructure IA, identifiant des barrieres a l'entree a plusieurs niveaux de la chaine d'approvisionnement : exigences capitalistiques extremement elevees, economies d'echelle massives, couts de switching entre fournisseurs et absence de standards d'interoperabilite.[9] La Federal Reserve Board (octobre 2025), dans une analyse comparative de la competitivite IA dans les economies avancees, montre que les Etats-Unis concentrent plus de 75 pour cent de l'investissement mondial en venture capital dans l'IA generative, et que l'Europe accuse un retard significatif non seulement en investissement mais aussi en adoption d'entreprise, les couts energetiques eleves constituant un frein supplementaire.[10]",
        "Cette dynamique de concentration a une consequence directe pour notre analyse : le protectionnisme technologique americain ne se contente pas de limiter l'acces europeen au compute, il renforce la position des acteurs dominants qui sont precisement ceux qui beneficient de l'exemption domestique. C'est un mecanisme de double avantage : reduction des contraintes pour les entreprises americaines, augmentation des contraintes pour leurs concurrentes.",
    ]),
    ("1.3.4 Souverainete numerique et autonomie strategique europeenne", [
        "Le quatrieme ancrage theorique mobilise le courant europeen de la souverainete numerique, analyse notamment par Mugge (2024) dans le Journal of European Public Policy. Mugge identifie trois tensions fondamentales dans l'ambition de souverainete IA europeenne : la souverainete oppose-t-elle l'UE aux autres puissances IA, ou les citoyens aux grandes plateformes ? Vise-t-elle la competitivite economique ou la protection des droits ? Et qui en beneficie reellement - les champions europeens ou l'ecosysteme entier ?[11] Ces tensions sont precisement celles que le protectionnisme americain exacerbe : en comprimant l'espace de souverainete technologique europeen, il force l'UE a arbitrer entre ces objectifs contradictoires dans un contexte d'urgence.",
        "Hawkins, Lehdonvirta et Wu (2025) apportent une contribution empirique decisive en mesurant la souverainete du compute a travers l'infrastructure des neuf principaux fournisseurs cloud mondiaux, qui representent environ 70 pour cent du marche mondial. Leurs resultats revelent que le degre de souverainete varie considerablement selon le niveau d'analyse (territorial, corporate ou hardware), et que la plupart des pays europeens presentent un deficit de souverainete sur au moins deux de ces trois niveaux.[12] McKinsey (decembre 2025) estime l'opportunite IA souveraine europeenne a 480 milliards d'euros par an a l'horizon 2030, conditionnelle a un scenario de souverainete technologique forte combine a une adoption IA elevee.[13]",
    ]),
    ("1.4 Revue de litterature ciblee et identification du gap", [
        "L'etat de la litterature en fevrier 2026 revele un champ de recherche en formation rapide, mais qui demeure tres fragmente. Nous identifions cinq corpus principaux, chacun couvrant une dimension de notre question de recherche.",
        "Corpus 1 : Controles a l'exportation et geopolitique IA. Le Carnegie Endowment for International Peace fournit l'analyse la plus detaillee de la politique americaine de controle a l'exportation IA. Winter-Levy et Phillips-Robins (mai 2025) ont decrit l'AI Diffusion Rule de Biden comme un compromis entre trois objectifs - controle, promotion et levier geopolitique - et analyse les options de remplacement sous Trump.[14] Le CSIS, le Hudson Institute et Pillsbury Law ont documente les mecanismes legaux et les developpements recents, notamment les tarifs Section 232 de janvier 2026. Contrary Research (novembre 2025) propose une analyse particulierement riche du dilemme temporel sous-jacent : si l'AGI est a cinq ans, les controles a l'exportation renforcent la domination americaine ; si elle est a dix ans ou plus, ils accelerent l'autonomisation technologique chinoise.[15]",
        "Corpus 2 : Energie et infrastructure des centres de donnees. Le rapport special Energy and AI de l'AIE (avril 2025) constitue la reference mondiale en matiere de projections energetiques pour les centres de donnees. Il etablit que la consommation electrique mondiale des centres de donnees atteindra 945 TWh en 2030 dans le scenario central, contre 415 TWh en 2024, l'IA etant le principal moteur de cette croissance.[16] Aux Etats-Unis, les centres de donnees consommeront plus d'electricite que toute l'industrie energo-intensive combinee a l'horizon 2030. En Europe, la croissance sera de plus 45 TWh (plus 70 pour cent), avec un risque de retard sur environ 20 pour cent des projets lie aux contraintes du reseau electrique. L'AIE-4E a publie en parallele une revue critique des modeles d'estimation, soulignant l'ampleur des incertitudes (les projections 2030 varient d'un facteur 40 selon les etudes).[17]",
        "Corpus 3 : Marche des semi-conducteurs. McKinsey (janvier 2026) a revise significativement a la hausse ses estimations de la taille du marche des semi-conducteurs en integrant les concepteurs captifs (Apple, Amazon, Tesla) et les operateurs fabless dont la valeur n'apparait pas dans les statistiques traditionnelles. Leur estimation centrale passe de 775 milliards USD en 2024 a 1 600 milliards USD en 2030, soit un TCAM de 13 pour cent.[18] La SIA rapporte des ventes record de 627,6 milliards USD en 2024, tandis que Deloitte (fevrier 2026) anticipe que les puces IA generatives representeront a elles seules pres de la moitie du chiffre d'affaires du secteur en 2026.[19]",
        "Corpus 4 : Souverainete et competitivite IA europeennes. Plusieurs publications recentes documentent le deficit europeen. Le Parlement europeen (2025) constate que seules 11 pour cent des petites entreprises europeennes utilisent l'IA, contre 58 pour cent des petites entreprises americaines.[20] Accenture (novembre 2025) rapporte que 62 pour cent des organisations europeennes recherchent des solutions souveraines face a l'incertitude geopolitique, mais seulement 19 pour cent y voient un avantage competitif. La majorite (48 pour cent) est motivee par des obligations de conformite reglementaire, suggerant une approche defensive plutot que strategique.[21] Le rapport Draghi (septembre 2024) sur la competitivite europeenne avait deja identifie le deficit d'investissement numerique comme un facteur structurel du declin europeen.",
        "Corpus 5 : Barrieres concurrentielles et concentration. Bruegel, l'OCDE et la Federal Reserve Board convergent sur le constat d'une concentration croissante de l'ecosysteme IA autour d'un petit nombre d'acteurs. Le rapport CERRE (juin 2025) sur la politique de concurrence pour le cloud et l'IA identifie des barrieres a la migration, des pratiques de cloud credits susceptibles de creer du lock-in, et une dependance croissante des petits developpeurs IA aux hyperscalers pour l'acces au compute accelere.[22] La collaboration structurelle entre startups IA et Big Tech, illustree par l'accord Mistral-Microsoft, temoigne de cette dependance.",
        "Le gap que cette etude entend combler. L'examen de ces cinq corpus revele un espace analytique non couvert : aucune etude ne propose une trajectoire integree liant le protectionnisme technologique americain a la divergence de competitivite UE via le triptyque compute-energie-semi-conducteurs. Les travaux existants traitent chaque dimension isolement - controles a l'exportation sans energie, energie sans semi-conducteurs, semi-conducteurs sans productivite. Par ailleurs, l'angle specifique de la robotique IA comme amplificateur de la demande energetique est quasiment absent de la litterature. C'est precisement ce gap que notre etude entend combler, en proposant un cadre analytique unifie et un indicateur original - le Compute-Adjusted Competitiveness Index (CACI) - permettant la mesure et la projection de cette divergence. L'indicateur et le jeu de donnees sous-jacent sont accessibles publiquement et mis a jour en temps reel sur le tableau de bord du projet.[a]",
    ]),
    ("1.5 Structure du rapport", [
        "Le rapport est organise en huit chapitres. Apres cette introduction, le chapitre II presente la methodologie, detaillant l'approche multi-scenarios, les sources de donnees, la construction du CACI et les limites methodologiques. Le chapitre III etablit le diagnostic empirique 2020-2026, couvrant les trajectoires energetiques, le marche des semi-conducteurs, le compute installe et la chronologie des controles a l'exportation. Le chapitre IV analyse les mecanismes de l'avantage competitif americain. Le chapitre V presente quatre scenarios 2026-2030 structures autour de deux axes d'incertitude. Le chapitre VI detaille les consequences pour la France et l'Europe, differenciees par type d'acteur et par secteur. Le chapitre VII formule des recommandations strategiques a trois horizons temporels. Le chapitre VIII conclut en synthetisant les apports de l'etude et en identifiant des pistes de recherche futures.",
    ]),
]

FR.notes = EN.notes  # bibliographic references identical


# ===========================================================================
# Content - Brazilian Portuguese
# ===========================================================================

PT = LangPack(
    code="PT-BR",
    filename="Capitulo_I_Introduction_PT-BR.docx",
    cover_title="AI FOR AMERICANS FIRST",
    cover_subtitle="Protecionismo de IA, Energia e Semicondutores: Trajetorias de Divergencia EUA/Europa 2024-2030",
    cover_blurb="Analise geoestrategica e economica integrada - Capitulo I",
    cover_chip_lines=[
        "76,9 pct do compute IA operacional mundial = EUA",
        "1,59x custo de energia UE/EUA",
        "3,46:1 razao CACI EUA/UE (Power Mode)",
    ],
    cover_meta="Paris - fevereiro de 2026  |  7 capitulos  |  4 cenarios prospectivos  |  3 zonas geograficas",
    cover_keywords_label="Palavras-chave",
    cover_keywords=("inteligencia artificial, protecionismo tecnologico, semicondutores, "
                    "controles de exportacao, compute soberano, geopolitica da IA, Franca, "
                    "Estados Unidos, China"),
    chapter_label="CAPITULO I",
    chapter_title="Introducao e enquadramento teorico",
    notes_label="Notas",
    license_block=[
        "Licenca e isencao de responsabilidade. Esta obra, 'AI for Americans First', e disponibilizada nos termos da Licenca Creative Commons Atribuicao - NaoComercial - CompartilhaIgual 4.0 Internacional (CC BY-NC-SA 4.0) do projeto America-First-IA.",
        "Voce e livre para compartilhar e adaptar o material para fins nao comerciais, desde que credite adequadamente a obra a Fabrice Pizzi (Universite Paris-Sorbonne) e distribua suas contribuicoes sob a mesma licenca. Este documento e fornecido apenas para fins educacionais e de pesquisa.",
        "Painel publico: https://mo0ogly.github.io/America-First-IA/dashboard/",
        "Repositorio: https://github.com/mo0ogly/America-First-IA",
    ],
    page_footer="AI for Americans First - Fabrice Pizzi - Capitulo I",
)

PT.sections = [
    ("1.1 Pergunta de pesquisa", [
        "A inteligencia artificial esta reconfigurando os fundamentos da competitividade economica global. Desde o lancamento do ChatGPT em novembro de 2022 e a aceleracao espetacular dos investimentos em modelos fundacionais, a IA generativa estabeleceu-se como um fator transformador transversal da economia, afetando simultaneamente as financas, a industria, a saude, os transportes e os servicos. Contudo, essa transformacao repousa sobre um substrato material preciso: capacidade computacional consideravel, alimentada por semicondutores de ponta e energia eletrica abundante. O dominio desse triptico - compute, chips, energia - tornou-se uma questao geoestrategica de primeira ordem.",
        "Nesse contexto, os Estados Unidos erigiram progressivamente um regime de controle sobre o acesso as tecnologias de IA de fronteira. Ja em outubro de 2022, o Bureau of Industry and Security (BIS) do Department of Commerce impos restricoes a exportacao de GPUs avancadas para a China. Em janeiro de 2025, a administracao Biden estendeu esses controles a mais de 120 paises por meio da AI Diffusion Rule, criando um sistema de segmentacao por niveis que condiciona o acesso aos chips de IA mais potentes ao grau de alinhamento geopolitico com Washington. A administracao Trump, empossada em janeiro de 2025, substituiu esse marco por uma abordagem mais explicitamente competitiva, culminando em julho de 2025 com a publicacao do America's AI Action Plan, e em janeiro de 2026 com a imposicao de tarifas de 25 por cento sobre certos semicondutores avancados de IA (Nvidia H200, AMD MI325X) ao abrigo da Secao 232.[1]",
        "Essas medidas, oficialmente motivadas por imperativos de seguranca nacional, produzem de facto uma vantagem competitiva estrutural para as empresas americanas: elas se beneficiam de acesso ilimitado ao compute de fronteira, enquanto atores de outras regioes - incluindo aliados europeus - veem suas capacidades limitadas, encarecidas ou condicionadas. Assistimos assim ao surgimento de um novo tipo de protecionismo tecnologico, em que o imposto nao e apenas tarifario, mas tambem regulatorio, logistico e estrategico.",
        "Este estudo coloca a seguinte questao: em que medida o protecionismo tecnologico americano em IA - controles de exportacao, tarifas, priorizacao do compute domestico - cria uma divergencia estrutural de competitividade entre os Estados Unidos e a Europa, e quais sao as consequencias mensuraveis para a Franca ate 2030?",
        "Esta questao de pesquisa decompoe-se em tres subquestoes que articulam as dimensoes empirica, prospectiva e normativa da analise:",
        "(a) Qual e a lacuna atual de capacidade computacional de IA (compute gap) entre os Estados Unidos e a Uniao Europeia, e como essa lacuna evolui sob diferentes cenarios de politica comercial e tecnologica americana?",
        "(b) Como a assimetria de acesso ao compute se traduz em produtividade setorial, custos de treinamento de modelos e participacao de mercado em servicos de IA?",
        "(c) As restricoes energeticas europeias e a ascensao da robotica de IA amplificam a divergencia, e a Franca pode mobilizar sua vantagem nuclear para mitigar essa desvantagem estrutural?",
        "A originalidade deste estudo reside em sua abordagem integrada. A literatura existente trata separadamente os controles de exportacao (Carnegie Endowment, CSIS, Hudson Institute), as projecoes energeticas dos data centers (AIE), a dinamica do mercado de semicondutores (McKinsey, SIA, Deloitte), a soberania do compute (Hawkins, Lehdonvirta e Wu, 2025) e as barreiras competitivas relacionadas a IA (Bruegel, OCDE). Nenhum trabalho propoe a trajetoria causal completa que buscamos estabelecer aqui: protecionismo americano e em seguida restricao do compute europeu e em seguida divergencia de produtividade e em seguida dependencia estrategica, com energia e robotica como fatores amplificadores.",
    ]),
    ("1.2 Definicoes operacionais", [
        "A analise desenvolvida neste estudo apoia-se em um conjunto de conceitos que requerem definicao rigorosa, pois seu uso varia entre contextos disciplinares.",
        "Compute de fronteira. Definimos compute de fronteira como a capacidade computacional instalada na forma de aceleradores de IA - GPUs (Nvidia A100, H100, H200, B200), ASICs (Google TPU) ou circuitos especializados - implantados em data centers em escala industrial. Essa capacidade e medida em FLOPs (operacoes de ponto flutuante por segundo) agregados em nivel nacional ou regional, em H100-equivalentes (a metrica utilizada pelo conjunto de dados Epoch AI), ou em GW de carga IT. O compute de fronteira constitui o insumo fundamental do desenvolvimento e da implantacao da IA de fronteira: sem ele, e impossivel treinar modelos fundacionais competitivos ou operar servicos de inferencia em larga escala.",
        "Protecionismo tecnologico em IA. O protecionismo tecnologico em IA designa o conjunto de medidas estatais - controles de exportacao, tarifas, cotas, licencas de exportacao, priorizacao logistica, restricoes sobre pesos de modelos e APIs - que criam uma assimetria de acesso ao compute, aos modelos e aos servicos de IA entre regioes geograficas. Este conceito estende a nocao classica de protecionismo comercial integrando a dimensao imaterial (software, modelos, servicos em nuvem) e a dimensao infraestrutural (energia, capacidades de producao de chips). Diferentemente das barreiras tarifarias tradicionais, o protecionismo tecnologico em IA pode operar por canais nao tarifarios - por exemplo, priorizando entregas de GPU a empresas domesticas em contexto de escassez global.",
        "Compute gap. O compute gap mede a razao da capacidade computacional de IA efetivamente disponivel entre duas regioes. Definimo-lo como a razao entre os FLOPs de IA instalados e acessiveis (ou H100-equivalentes) para os atores economicos da regiao A e os da regiao B, normalizada pela populacao ativa ou pelo PIB. Um compute gap de x15 entre os Estados Unidos e a UE significa que, por unidade de PIB, os atores americanos dispoem de quinze vezes mais poder computacional de IA do que seus homologos europeus. Essa razao constitui um indicador sintetico da vantagem estrutural em IA.",
        "Soberania do compute. Adotamos a definicao proposta por Hawkins, Lehdonvirta e Wu (2025), que decompoem a soberania do compute em tres niveis: (1) a quantidade de compute de IA fisicamente presente no territorio nacional, (2) a nacionalidade das empresas proprietarias dos data centers e (3) a nacionalidade dos fornecedores de aceleradores cujos chips alimentam esses data centers.[2] Um pais pode dispor de capacidade computacional significativa em seu solo enquanto e dependente de operadores e fornecedores estrangeiros, o que limita sua soberania real. Esse conceito e essencial para compreender por que a presenca de data centers AWS ou Azure na Europa nao constitui, em si, uma forma de soberania europeia sobre o compute.",
        "Vendor lock-in geopolitico. Propomos o conceito de vendor lock-in geopolitico para designar a dependencia estrutural de um ecossistema economico em relacao a fornecedores tecnologicos estrangeiros cujo acesso pode ser restringido, encarecido ou condicionado por decisao politica de um governo terceiro. Este conceito estende a nocao classica de vendor lock-in em TI (custos de migracao entre provedores de nuvem, por exemplo) acrescentando uma dimensao geopolitica: o risco de que o acesso a uma infraestrutura critica seja utilizado como alavanca de negociacao entre Estados. O episodio Starlink-Ucrania de marco de 2025, em que o controle americano de um sistema de comunicacao foi percebido como instrumento de pressao, ilustra de forma marcante esse tipo de risco.[3]",
        "Compute-Adjusted Competitiveness Index (CACI). Construimos, como indicador original central deste estudo, um Compute-Adjusted Competitiveness Index combinando em um composto geometrico ponderado as quatro dimensoes identificadas acima. Em sua formulacao absoluta, denominada Power Mode, o indicador e definido como CACI = F^0,40 x L^0,20 x R^0,15 / E^0,25, onde F e o compute de IA instalado (H100-equivalentes, total ou soberano), L e a forca de trabalho relevante (milhoes de trabalhadores qualificados em IA), R e um indice de acesso geopolitico (0 a 1), e E e o preco medio da eletricidade industrial (USD por MWh). Os quatro expoentes (40, 20, 15, 25) somam um e refletem a importancia relativa atribuida ao compute, ao trabalho, ao acesso geopolitico e ao custo da energia. Uma segunda formulacao, denominada Intensity Mode, divide o resultado pelo PIB para obter uma densidade de competitividade por unidade de riqueza nacional. A derivacao metodologica completa e apresentada no Capitulo II e o indicador e atualizado em tempo real no painel publico.[a]",
        "Duas variantes do indicador sao utilizadas em paralelo. O CACI fisico integra todo o compute fisicamente presente no territorio (F_total), independentemente da propriedade. O CACI soberano restringe F ao compute possuido e operado por atores domesticos (F_dom), capturando assim a autonomia soberana efetiva. A diferenca entre as duas variantes mede a dependencia estrutural de uma regiao em relacao aos hyperscalers estrangeiros. Para a Uniao Europeia, essa diferenca constitui a medida quantitativa mais direta da vulnerabilidade estrategica que este estudo busca caracterizar.",
    ]),
    ("1.3 Marco teorico", [
        "Nossa analise enraiza-se em quatro correntes teoricas complementares, que fornecem as ferramentas conceituais necessarias para articular as dimensoes tecnologica, economica e geopolitica do fenomeno estudado.",
    ]),
    ("1.3.1 A IA como General Purpose Technology", [
        "A primeira ancora teorica e a teoria das General Purpose Technologies (GPT) formalizada por Bresnahan e Trajtenberg (1995). Uma GPT e uma tecnologia caracterizada por tres propriedades: sua pervasividade (e utilizada como insumo em numerosos setores a jusante), seu potencial inerente de melhoria tecnica e suas complementaridades de inovacao (a produtividade de P&D nos setores usuarios cresce em consequencia da melhoria da GPT).[4] O modelo preve que as GPTs geram retornos crescentes de escala e que sua difusao na economia e fonte de ganhos generalizados de produtividade. Contudo, Bresnahan e Trajtenberg tambem ressaltam que uma economia descentralizada pode ter dificuldade em explorar plenamente o potencial de uma GPT, ja que as transacoes de mercado entre o produtor da GPT e seus usuarios podem levar a inovacao insuficiente e tardia.",
        "Brynjolfsson, Rock e Syverson (2019) aplicaram esse marco a IA contemporanea, demonstrando que a IA, e em particular o aprendizado de maquina, satisfaz os tres criterios de Bresnahan e Trajtenberg para ser qualificada como GPT.[5] Seu modelo de productivity J-curve explica por que os ganhos de produtividade ligados a uma GPT podem ser inicialmente invisiveis nas estatisticas: as empresas devem primeiro investir massivamente em ativos intangiveis (reorganizacao, formacao, reengenharia de processos) antes de colher os beneficios.",
        "Esse marco e fundamental para nossa analise porque implica que o acesso precoce e massivo ao compute de IA - ou seja, a capacidade de investir na GPT em larga escala desde as primeiras fases de implantacao - gera vantagens cumulativas dificilmente reversiveis. As complementaridades de inovacao criam uma dinamica de path dependence: os atores que acessam cedo um compute abundante desenvolvem modelos superiores, capturam dados de uso, geram receitas que reinvestem em compute e, assim, ampliam uma diferenca que se autorreforca ao longo do tempo. Toda politica que restrinja o acesso ao compute para uma determinada regiao tem, portanto, consequencias nao lineares: nao apenas atrasa a adocao, mas a compromete estruturalmente.",
    ]),
    ("1.3.2 Interdependencia armada e controle das redes globais", [
        "A segunda ancora teorica e a teoria da weaponized interdependence desenvolvida por Farrell e Newman (2019). Esses autores demonstram que as redes economicas globais, longe de criarem relacoes simetricas de interdependencia mutua como postulava a teoria liberal classica, tendem a produzir estruturas assimetricas em que certos nos (hubs) tornam-se muito mais conectados que outros. Os Estados que exercem jurisdicao politica sobre esses nos centrais podem instrumentaliza-los para fins coercitivos, por meio de dois mecanismos: o efeito panopticon (coleta de informacoes estrategicas) e o efeito chokepoint (capacidade de cortar ou restringir fluxos).[6]",
        "A aplicacao a cadeia de valor da IA e notavelmente pertinente. Os Estados Unidos controlam varios chokepoints criticos: o design de aceleradores de IA (a Nvidia detem mais de 80 por cento do mercado de GPUs para data centers), a infraestrutura em nuvem (AWS, Azure e Google Cloud representam aproximadamente 70 por cento do mercado global) e os modelos fundacionais mais avancados (OpenAI, Anthropic, Google DeepMind). O controle desses chokepoints permite ao governo americano modular o acesso global ao compute de IA como alavanca geopolitica, exatamente como o sistema SWIFT foi utilizado como alavanca no dominio financeiro. Os proprios Farrell e Newman reconheceram, em uma atualizacao publicada na Foreign Affairs em dezembro de 2025, que os semicondutores e a IA tinham se tornado um terreno de aplicacao maior para sua teoria, com a administracao Trump tendo utilizado explicitamente os controles de exportacao sobre chips de IA como moeda de troca nas negociacoes com a China.[7]",
    ]),
    ("1.3.3 Economia da concentracao e rendas de inovacao", [
        "A terceira ancora teorica mobiliza a literatura sobre concentracao de mercados digitais e barreiras a entrada no ecossistema de IA. Martens (2024), em um policy brief para o Bruegel, demonstra que os custos de treinamento de modelos fundacionais crescem exponencialmente, constituindo uma barreira a entrada intransponivel para a maioria dos atores.[8] Ele estima que uma fazenda de calculo da ordem de um trilhao de dolares e concebivel a medio prazo, um limiar de investimento totalmente fora do alcance do financiamento publico e da grande maioria das empresas. Apenas as GAMMAN (Google, Apple, Meta, Microsoft, Amazon, Nvidia) dispoem dos recursos para acede-lo.",
        "A OCDE (2025) confirma essa analise em seu relatorio sobre concorrencia em infraestrutura de IA, identificando barreiras a entrada em multiplos niveis da cadeia de suprimentos: requisitos de capital extremamente elevados, economias de escala massivas, custos de switching entre fornecedores e ausencia de padroes de interoperabilidade.[9] O Federal Reserve Board (outubro de 2025), em uma analise comparativa da competitividade de IA nas economias avancadas, mostra que os Estados Unidos concentram mais de 75 por cento do investimento global de venture capital em IA generativa, e que a Europa apresenta um atraso significativo nao apenas em investimento, mas tambem em adocao empresarial, com os altos custos energeticos constituindo um freio adicional.[10]",
        "Essa dinamica de concentracao tem uma consequencia direta para nossa analise: o protecionismo tecnologico americano nao apenas limita o acesso europeu ao compute, mas reforca a posicao dos atores dominantes que sao precisamente aqueles que se beneficiam da isencao domestica. Trata-se de um mecanismo de dupla vantagem: reducao das restricoes para as empresas americanas, aumento das restricoes para suas concorrentes.",
    ]),
    ("1.3.4 Soberania digital e autonomia estrategica europeia", [
        "A quarta ancora teorica mobiliza a corrente europeia da soberania digital, analisada notadamente por Mugge (2024) no Journal of European Public Policy. Mugge identifica tres tensoes fundamentais na ambicao de soberania de IA europeia: a soberania opoe a UE a outras potencias de IA, ou os cidadaos as grandes plataformas? Visa a competitividade economica ou a protecao de direitos? E quem realmente se beneficia dela - os campeoes europeus ou todo o ecossistema?[11] Essas tensoes sao precisamente aquelas que o protecionismo americano exacerba: ao comprimir o espaco de soberania tecnologica europeu, forca a UE a arbitrar entre esses objetivos contraditorios em um contexto de urgencia.",
        "Hawkins, Lehdonvirta e Wu (2025) trazem uma contribuicao empirica decisiva ao medir a soberania do compute por meio da infraestrutura dos nove principais provedores globais de nuvem, que representam aproximadamente 70 por cento do mercado global. Seus resultados revelam que o grau de soberania varia consideravelmente segundo o nivel de analise (territorial, corporate ou hardware) e que a maioria dos paises europeus apresenta um deficit de soberania em pelo menos dois desses tres niveis.[12] A McKinsey (dezembro de 2025) estima a oportunidade europeia de IA soberana em 480 bilhoes de euros por ano ate 2030, condicional a um cenario de soberania tecnologica forte combinado a alta adocao de IA.[13]",
    ]),
    ("1.4 Revisao da literatura focada e identificacao da lacuna", [
        "O estado da literatura em fevereiro de 2026 revela um campo de pesquisa em formacao rapida, mas que permanece muito fragmentado. Identificamos cinco corpus principais, cada um cobrindo uma dimensao da nossa pergunta de pesquisa.",
        "Corpus 1: Controles de exportacao e geopolitica de IA. O Carnegie Endowment for International Peace fornece a analise mais detalhada da politica americana de controle de exportacao de IA. Winter-Levy e Phillips-Robins (maio de 2025) descreveram a AI Diffusion Rule de Biden como um compromisso entre tres objetivos - controle, promocao e alavancagem geopolitica - e analisaram opcoes de substituicao sob Trump.[14] CSIS, Hudson Institute e Pillsbury Law documentaram os mecanismos legais e os desenvolvimentos recentes, notadamente as tarifas Secao 232 de janeiro de 2026. A Contrary Research (novembro de 2025) oferece uma analise particularmente rica do dilema temporal subjacente: se a AGI esta a cinco anos de distancia, os controles de exportacao reforcam a dominacao americana; se estiver a dez anos ou mais, eles aceleram a autonomizacao tecnologica chinesa.[15]",
        "Corpus 2: Energia e infraestrutura de data centers. O relatorio especial Energy and AI da AIE (abril de 2025) constitui a referencia global em projecoes energeticas para data centers. Estabelece que o consumo eletrico mundial dos data centers atingira 945 TWh em 2030 no cenario base, ante 415 TWh em 2024, com a IA como principal motor desse crescimento.[16] Nos Estados Unidos, os data centers consumirao mais eletricidade do que toda a industria intensiva em energia combinada ate 2030. Na Europa, o crescimento sera de mais 45 TWh (mais 70 por cento), com risco de atraso em aproximadamente 20 por cento dos projetos vinculado as restricoes da rede eletrica. A AIE-4E publicou paralelamente uma revisao critica dos modelos de estimativa, destacando a magnitude das incertezas (as projecoes para 2030 variam por um fator de 40 entre os estudos).[17]",
        "Corpus 3: Mercado de semicondutores. A McKinsey (janeiro de 2026) revisou significativamente para cima suas estimativas do tamanho do mercado de semicondutores ao integrar designers cativos (Apple, Amazon, Tesla) e operadores fabless cujo valor nao aparece nas estatisticas tradicionais. Sua estimativa central passa de 775 bilhoes de USD em 2024 para 1 600 bilhoes de USD em 2030, ou um TCAM de 13 por cento.[18] A SIA reporta vendas recordes de 627,6 bilhoes de USD em 2024, enquanto a Deloitte (fevereiro de 2026) antecipa que apenas os chips de IA generativa representarao quase metade da receita do setor em 2026.[19]",
        "Corpus 4: Soberania e competitividade de IA europeias. Varias publicacoes recentes documentam o deficit europeu. O Parlamento Europeu (2025) constata que apenas 11 por cento das pequenas empresas europeias usam IA, contra 58 por cento das pequenas empresas americanas.[20] A Accenture (novembro de 2025) reporta que 62 por cento das organizacoes europeias buscam solucoes soberanas frente a incerteza geopolitica, mas apenas 19 por cento veem nelas uma vantagem competitiva. A maioria (48 por cento) e motivada por obrigacoes de conformidade regulatoria, sugerindo uma abordagem defensiva em vez de estrategica.[21] O relatorio Draghi (setembro de 2024) sobre competitividade europeia ja havia identificado o deficit de investimento digital como fator estrutural do declinio europeu.",
        "Corpus 5: Barreiras competitivas e concentracao. Bruegel, OCDE e Federal Reserve Board convergem na constatacao de uma concentracao crescente do ecossistema de IA em torno de um pequeno numero de atores. O relatorio CERRE (junho de 2025) sobre politica de concorrencia para nuvem e IA identifica barreiras a migracao, praticas de creditos de nuvem suscetiveis de criar lock-in e dependencia crescente dos pequenos desenvolvedores de IA em relacao aos hyperscalers para acesso ao compute acelerado.[22] A colaboracao estrutural entre startups de IA e Big Tech, ilustrada pelo acordo Mistral-Microsoft, atesta essa dependencia.",
        "A lacuna que este estudo busca preencher. O exame desses cinco corpora revela um espaco analitico nao coberto: nenhum estudo propoe uma trajetoria integrada que ligue o protecionismo tecnologico americano a divergencia de competitividade da UE pelo triptico compute-energia-semicondutores. Os trabalhos existentes tratam cada dimensao isoladamente - controles de exportacao sem energia, energia sem semicondutores, semicondutores sem produtividade. Alem disso, o angulo especifico da robotica de IA como amplificadora da demanda energetica esta praticamente ausente da literatura. E precisamente essa lacuna que nosso estudo pretende preencher, propondo um marco analitico unificado e um indicador original - o Compute-Adjusted Competitiveness Index (CACI) - permitindo a medicao e a projecao dessa divergencia. O indicador e o conjunto de dados subjacente sao acessiveis publicamente e atualizados em tempo real no painel do projeto.[a]",
    ]),
    ("1.5 Estrutura do relatorio", [
        "O relatorio esta organizado em oito capitulos. Apos esta introducao, o Capitulo II apresenta a metodologia, detalhando a abordagem multi-cenario, as fontes de dados, a construcao do CACI e as limitacoes metodologicas. O Capitulo III estabelece o diagnostico empirico 2020-2026, cobrindo as trajetorias energeticas, o mercado de semicondutores, o compute instalado e a cronologia dos controles de exportacao. O Capitulo IV analisa os mecanismos da vantagem competitiva americana. O Capitulo V apresenta quatro cenarios 2026-2030 estruturados em torno de dois eixos de incerteza. O Capitulo VI detalha as consequencias para a Franca e a Europa, diferenciadas por tipo de ator e por setor. O Capitulo VII formula recomendacoes estrategicas em tres horizontes temporais. O Capitulo VIII conclui sintetizando as contribuicoes do estudo e identificando caminhos para pesquisas futuras.",
    ]),
]

PT.notes = EN.notes  # bibliographic references identical


# ===========================================================================
# Main
# ===========================================================================

def main() -> None:
    """Build the three Chapter I .docx files."""
    out_dir = Path(__file__).parent
    for lp in (EN, FR, PT):
        build(lp, out_dir)
    log.info("Done.")


if __name__ == "__main__":
    main()
