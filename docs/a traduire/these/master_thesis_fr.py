"""
Master thesis assembler - FR.

Combines the 11 French chapters into a single .docx deliverable with:
    1. Cover page (title, author, date, executive summary band)
    2. Auto-updating Table of Contents (TOC field, refresh on first open)
    3. The 11 chapters in order, separated by page breaks:
        I, II, III, IV, V, VI, VI bis, VI ter, VI quater, VII, Conclusion
    4. Each chapter content is appended with its existing internal headings
       and tables preserved.

Usage
-----
    python master_thesis_fr.py

Inputs (must exist in the same directory as this script):
    Chapitre_I_Introduction_Cadrage_Theorique_FR.docx
    Chapitre_II_Methodologie_FR.docx
    Chapitre_III_Diagnostic_Empirique_FR.docx
    Chapitre_IV_Mecanismes_Avantage_Concurrentiel_FR.docx
    Chapitre_V_Scenarios_Prospectifs_2026_2030_FR.docx
    Chapitre_VI_Consequences_France_Europe_FR.docx
    Chapitre_VI_bis_Amerique_Sud_Bresil_FR.docx
    Chapitre_VI_ter_Asie_FR.docx
    Chapitre_VI_quater_Afrique_FR.docx
    Chapitre_VII_Recommandations_FR.docx
    Conclusion_Generale_FR.docx

Output:
    These_AI_For_Americans_First_FR.docx

Auteur : Fabrice Pizzi (Universite Paris-Sorbonne, M2 Intelligence Economique).
"""

from __future__ import annotations

import copy
import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
log = logging.getLogger("master_thesis_fr")


# Visual identity (same palette as chapters)
NAVY = RGBColor(0x1A, 0x27, 0x44)
GOLD = RGBColor(0xB8, 0x92, 0x2F)
GREY = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x20, 0x20, 0x20)


# Order of chapters in the master document
CHAPTERS: list[tuple[str, str]] = [
    ("Chapitre_I_Introduction_Cadrage_Theorique_FR.docx",
     "Chapitre I - Introduction et cadrage theorique"),
    ("Chapitre_II_Methodologie_FR.docx",
     "Chapitre II - Methodologie"),
    ("Chapitre_III_Diagnostic_Empirique_FR.docx",
     "Chapitre III - Diagnostic empirique 2020-2026"),
    ("Chapitre_IV_Mecanismes_Avantage_Concurrentiel_FR.docx",
     "Chapitre IV - Mecanismes de l'avantage concurrentiel US"),
    ("Chapitre_V_Scenarios_Prospectifs_2026_2030_FR.docx",
     "Chapitre V - Scenarios prospectifs 2026-2030"),
    ("Chapitre_VI_Consequences_France_Europe_FR.docx",
     "Chapitre VI - Consequences pour la France et l'Europe"),
    ("Chapitre_VI_bis_Amerique_Sud_Bresil_FR.docx",
     "Chapitre VI bis - Consequences pour l'Amerique du Sud et le Bresil"),
    ("Chapitre_VI_ter_Asie_FR.docx",
     "Chapitre VI ter - Consequences pour l'Asie"),
    ("Chapitre_VI_quater_Afrique_FR.docx",
     "Chapitre VI quater - Consequences pour l'Afrique"),
    ("Chapitre_VII_Recommandations_FR.docx",
     "Chapitre VII - Recommandations strategiques"),
    ("Conclusion_Generale_FR.docx",
     "Conclusion generale"),
]


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _set_run(run, *, font="Calibri", size=11, bold=False, italic=False, color=None):
    """Apply consistent typography to a docx run."""
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _add_paragraph(doc, text, *, align=None, space_after=6, **run_kwargs):
    """Add a paragraph with one run."""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        _set_run(run, **run_kwargs)
    return p


def _add_page_break(doc) -> None:
    """Insert a hard page break."""
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _add_field(paragraph, instruction: str) -> None:
    """Insert a Word field (used for TOC / PAGE / NUMPAGES)."""
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


# ---------------------------------------------------------------------------
# Cover page + executive summary
# ---------------------------------------------------------------------------

def add_master_cover(doc) -> None:
    """Render the master thesis cover page."""
    _add_paragraph(doc, "", space_after=20)
    _add_paragraph(doc, "THESE DOCTORALE", align=WD_ALIGN_PARAGRAPH.CENTER,
                   size=11, bold=True, color=GREY, space_after=4)
    _add_paragraph(doc, "Universite Paris-Sorbonne",
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   size=11, italic=True, color=GREY, space_after=4)
    _add_paragraph(doc, "Master 2 Intelligence Economique - Intelligence Warfare",
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   size=11, italic=True, color=GREY, space_after=40)

    _add_paragraph(doc, "AI FOR AMERICANS FIRST",
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   size=32, bold=True, color=NAVY, space_after=10)
    _add_paragraph(doc,
                   "Protectionnisme IA, Energie et Semi-conducteurs :",
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   size=14, italic=True, color=NAVY, space_after=2)
    _add_paragraph(doc,
                   "Trajectoires de divergence US/Europe 2024-2030",
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   size=14, italic=True, color=NAVY, space_after=10)
    _add_paragraph(doc,
                   "Analyse geostrategique et economique integree",
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   size=12, italic=True, color=GREY, space_after=40)

    # Headline numbers band
    table = doc.add_table(rows=1, cols=3)
    chips = [
        "76,9 pct compute IA operationnel mondial = USA",
        "1,59x cout energie EU/US (ajuste-PPA)",
        "3,46:1 ratio CACI US/EU (Power Mode)",
    ]
    for i, line in enumerate(chips):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        _set_run(run, size=11, bold=True, color=NAVY)

    _add_paragraph(doc, "", space_after=40)

    _add_paragraph(doc, "Fabrice Pizzi",
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   size=18, bold=True, color=NAVY, space_after=4)
    _add_paragraph(doc, "(handle mo0ogly)",
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   size=10, italic=True, color=GREY, space_after=20)
    _add_paragraph(doc, "Paris - fevrier 2026",
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   size=12, color=GREY, space_after=40)

    _add_paragraph(doc,
                   "11 chapitres  |  4 scenarios prospectifs  |  4 zones geographiques  "
                   "(Europe, Amerique du Sud, Asie, Afrique)  |  108 figures",
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   size=10, italic=True, color=GREY, space_after=20)
    _add_paragraph(doc,
                   "Mots-cles : intelligence artificielle, protectionnisme technologique, "
                   "semi-conducteurs, controles a l'exportation, compute souverain, "
                   "geopolitique IA, France, Etats-Unis, Chine.",
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   size=10, italic=True, color=GREY, space_after=20)


# ---------------------------------------------------------------------------
# Executive summary
# ---------------------------------------------------------------------------

def add_executive_summary(doc) -> None:
    """Render a 1-page executive summary in front of the TOC."""
    _add_paragraph(doc, "Resume executif",
                   size=22, bold=True, color=NAVY, space_after=14)

    paragraphs = [
        "Cette these analyse l'emergence d'un regime protectionniste IA americain (decret implicite "
        "AI for Americans First) et ses consequences geostrategiques pour la France, l'Europe, et "
        "le reste du monde sur la periode 2024-2030. L'etude mobilise un appareil empirique original "
        "(tableau de bord public Epoch AI snapshot avril 2026) et propose un cadre theorique unifie "
        "(indice CACI - Compute-Adjusted Competitive Index) integrant FLOPs, capital humain, "
        "regulation et cout energetique selon une formule geometrique ponderee.",

        "Sur le snapshot avril 2026, les Etats-Unis concentrent 76,9 pour cent du compute IA "
        "operationnel mondial. Le ratio brut compute installe US/UE(13) atteint 17,6:1, traduit "
        "par la formule geometrique CACI Power Mode (F^0,40 x L^0,20 x R^0,15 / E^0,25) en un "
        "ratio US/UE de 3,46:1. La decomposition Phys/Sov rigoureusement calculee a partir du "
        "champ Owner d'Epoch AI revele un cas extreme : 99,6 pct du F_total emirati est detenu "
        "par des operateurs US-side, faisant chuter le CACI souverain emirati de 55,7 (Physique) "
        "a 6,0. L'UE, en revanche, est largement souveraine sur son F installe (99,2 pct). La "
        "vulnerabilite europeenne se situe non sur le compute installe mais sur la couche "
        "operationnelle des charges cloud (majoritairement hebergees sur AWS/Azure/GCP).",

        "Quatre scenarios prospectifs 2026-2030 sont developpes (Statu quo, Fracture, Partenariat "
        "asymetrique, Souverainete subie), avec un point de basculement identifie en 2028 a la "
        "convergence de la saturation compute/energie UE et de l'activation potentielle des "
        "Cloud Sovereignty Mandates US. La fenetre d'action strategique 2026-2028 est etroite : "
        "apres cette date, les positions se cristallisent autour de la baseline 17,6:1 brut / "
        "3,46:1 CACI Power Mode.",

        "Les recommandations strategiques s'articulent en cinq axes (compute, energie, alliances, "
        "regulation, talent) sur trois horizons (2026-2027, 2027-2029, 2029-2032). La France "
        "dispose d'avantages distinctifs (parc nucleaire 70 pct du mix, cout PPA 1,35x USA, "
        "Mistral AI valorise 11,7 milliards EUR, ASML partenaire industriel) qui en font le pilier "
        "credible d'une trajectoire d'autonomie strategique ciblee - ni alignement total (Japon : "
        "550 milliards USD investis aux US), ni confrontation (Chine : 246-300 EFLOP/s sous "
        "restriction), ni hesitation (Bresil : double bind US/Chine). L'enjeu n'est pas l'autarcie "
        "mais la capacite de choix.",
    ]
    for text in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.3
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        _set_run(run, size=11, color=DARK)


# ---------------------------------------------------------------------------
# Table of contents
# ---------------------------------------------------------------------------

def add_table_of_contents(doc, chapter_titles: list[str]) -> None:
    """Insert a pre-rendered table of contents listing the 11 chapters.

    A real Word TOC field is also embedded so users can refresh page
    numbers in Word/LibreOffice (Ctrl+A then F9). The pre-rendered
    listing guarantees a usable TOC even on read-only viewers that
    don't compute fields (e.g. PDF exports without field update).
    """
    _add_paragraph(doc, "Table des matieres",
                   size=22, bold=True, color=NAVY, space_after=14)

    # Pre-rendered listing (no page numbers because pagination depends
    # on rendering engine; refresh via Ctrl+A, F9 in Word for numbers).
    for i, title in enumerate(chapter_titles, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        run_n = p.add_run(f"{i:>2}.  ")
        _set_run(run_n, size=11, bold=True, color=GOLD)
        run_t = p.add_run(title)
        _set_run(run_t, size=11, color=NAVY)

    # Hint and dynamic TOC field for Word users
    _add_paragraph(doc, "", space_after=10)
    _add_paragraph(doc,
                   "Note : une table des matieres dynamique (avec sous-sections et numeros de "
                   "page) est integree ci-dessous. Sous Microsoft Word, faire Ctrl+A puis F9 pour "
                   "la mettre a jour. Sous LibreOffice, faire clic-droit sur la table et choisir "
                   "'Mettre a jour l'index'.",
                   size=9, italic=True, color=GREY, space_after=10)

    p = doc.add_paragraph()
    _add_field(p, r'TOC \o "1-2" \h \z \u')


# ---------------------------------------------------------------------------
# Chapter import
# ---------------------------------------------------------------------------

# Markers in the source covers that we want to skip when concatenating.
# Each chapter .docx starts with its own cover (label, title, ETUDE DE
# RECHERCHE, headline band, FabricePizzi, Paris fevrier 2026, Mots-cles).
# We detect the end of cover by looking for the chapter title repeated in
# bold at the start of the chapter body (e.g. "CHAPITRE I", "CHAPITRE V",
# "CONCLUSION GENERALE").
COVER_TERMINATORS = {
    "Chapitre_I_Introduction_Cadrage_Theorique_FR.docx": "CHAPITRE I",
    "Chapitre_II_Methodologie_FR.docx": "CHAPITRE II",
    "Chapitre_III_Diagnostic_Empirique_FR.docx": "CHAPITRE III",
    "Chapitre_IV_Mecanismes_Avantage_Concurrentiel_FR.docx": "CHAPITRE IV",
    "Chapitre_V_Scenarios_Prospectifs_2026_2030_FR.docx": "CHAPITRE V",
    "Chapitre_VI_Consequences_France_Europe_FR.docx": "CHAPITRE VI",
    "Chapitre_VI_bis_Amerique_Sud_Bresil_FR.docx": "CHAPITRE VI BIS",
    "Chapitre_VI_ter_Asie_FR.docx": "CHAPITRE VI TER",
    "Chapitre_VI_quater_Afrique_FR.docx": "CHAPITRE VI QUATER",
    "Chapitre_VII_Recommandations_FR.docx": "CHAPITRE VII",
    "Conclusion_Generale_FR.docx": "CONCLUSION GENERALE",
}


def _find_cover_end(source_doc, terminator: str) -> int:
    """Return the index of the body element where the chapter content starts."""
    body = source_doc.element.body
    for i, el in enumerate(body):
        text = "".join(t.text or "" for t in el.iter(qn("w:t")))
        if terminator.lower() in text.lower() and len(text) <= len(terminator) + 10:
            return i
    # Fallback: skip ~24 first paragraphs (typical cover length)
    return min(24, len(body) - 1)


def add_chapter_heading(doc, *, label: str, title: str) -> None:
    """Render a uniform chapter title as Heading 1 so it shows up in the TOC."""
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(label.upper())
    _set_run(run, font="Calibri", size=22, bold=True, color=NAVY)


def import_chapter_body(master_doc: Document, chapter_path: Path,
                        terminator: str) -> int:
    """Append the body elements of a chapter to the master document.

    Skips the source's own cover (which would duplicate the master cover
    style) and the source's license footer. Returns the number of XML
    elements imported, useful for logging.
    """
    src = Document(str(chapter_path))
    src_body = src.element.body
    cover_end = _find_cover_end(src, terminator)

    imported = 0
    target_body = master_doc.element.body
    # Insert before sectPr (the trailing section properties) to preserve
    # the master's section settings.
    sect_pr = target_body.find(qn("w:sectPr"))

    for el in list(src_body)[cover_end:]:
        # Skip the trailing sectPr from the source (we keep the master's)
        if el.tag == qn("w:sectPr"):
            continue

        # Skip license footer paragraphs (they all have a specific marker
        # text "Licence et avertissement" or a page-footer string).
        text = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
        if text.startswith("Licence et avertissement"):
            break  # Everything after is footer
        if "AI for Americans First - Fabrice Pizzi" in text:
            continue
        if text.startswith("Tableau de bord public") or text.startswith("Depot :"):
            continue
        if text.startswith("Vous etes libre"):
            continue

        # Deep copy the element to avoid sharing nodes between trees
        new_el = copy.deepcopy(el)
        if sect_pr is not None:
            target_body.insert(list(target_body).index(sect_pr), new_el)
        else:
            target_body.append(new_el)
        imported += 1

    return imported


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

def build_master(input_dir: Path, output_path: Path) -> Path:
    """Build the master thesis .docx by assembling the 11 FR chapters."""
    log.info("Initializing master document")
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Heading 1 / Heading 2 styles for TOC
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = NAVY

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = NAVY

    # 1. Cover page
    add_master_cover(doc)
    _add_page_break(doc)

    # 2. Executive summary
    add_executive_summary(doc)
    _add_page_break(doc)

    # 3. Table of contents (pre-rendered + auto-refreshing field)
    chapter_titles = [t for _, t in CHAPTERS]
    add_table_of_contents(doc, chapter_titles)
    _add_page_break(doc)

    # 4. Chapters in order
    total_imported = 0
    for filename, full_title in CHAPTERS:
        chapter_path = input_dir / filename
        if not chapter_path.exists():
            log.error("Missing chapter file: %s", chapter_path)
            raise FileNotFoundError(chapter_path)

        log.info("Importing %s", filename)
        # Add chapter Heading 1 so it appears in the TOC
        add_chapter_heading(doc, label=full_title, title=full_title)

        terminator = COVER_TERMINATORS[filename]
        n = import_chapter_body(doc, chapter_path, terminator)
        log.info("  -> %d elements imported", n)
        total_imported += n

        # Page break between chapters
        _add_page_break(doc)

    # 5. Final license/credits page
    _add_paragraph(doc, "Licence et credits",
                   size=18, bold=True, color=NAVY, space_after=12)
    _add_paragraph(doc,
                   "Cette these, AI for Americans First, est mise a disposition selon les termes "
                   "de la licence Creative Commons Attribution - Pas d'Utilisation Commerciale - "
                   "Partage dans les Memes Conditions 4.0 International (CC BY-NC-SA 4.0).",
                   size=10, italic=True, color=GREY, space_after=8)
    _add_paragraph(doc,
                   "Tableau de bord public : https://mo0ogly.github.io/America-First-IA/dashboard/",
                   size=10, italic=True, color=GREY, space_after=4)
    _add_paragraph(doc,
                   "Depot : https://github.com/mo0ogly/America-First-IA",
                   size=10, italic=True, color=GREY, space_after=4)
    _add_paragraph(doc,
                   "Auteur : Fabrice Pizzi (handle mo0ogly), Universite Paris-Sorbonne, "
                   "Master 2 Intelligence Economique - Intelligence Warfare. Paris, fevrier 2026.",
                   size=10, italic=True, color=GREY, space_after=4)

    log.info("Saving %s (%d body elements imported across %d chapters)",
             output_path, total_imported, len(CHAPTERS))
    doc.save(str(output_path))
    return output_path


def main() -> None:
    here = Path(__file__).parent.resolve()
    output = here / "These_AI_For_Americans_First_FR.docx"
    build_master(here, output)
    log.info("Master thesis built at: %s", output)


if __name__ == "__main__":
    main()
