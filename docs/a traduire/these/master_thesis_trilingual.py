"""
Master Thesis Assembler - Trilingual (EN, FR, PT-BR).

Industrializes the assembly of the 11 chapters (from folders 1-8) into 
three final trilingual deliverables.
"""

from __future__ import annotations

import sys
from pathlib import Path
sys.path.append(str(Path(__file__).resolve().parents[1]))
import caci_data_helper
m = caci_data_helper.get_metrics()
us_share = m['us_compute_share']
us_eu_caci = m['us_eu_caci_power_ratio']
us_eu_raw = m['us_eu_raw_ratio']
eu_sov = m['eu_sovereignty_ratio']
caci_scores = {k: v['caci_power_phys'] for k, v in m['country_results'].items()}

def fmt_fr(val, decimals=1):
    return f"{val:.{decimals}f}".replace(".", ",")

def fmt_en(val, decimals=1):
    return f"{val:.{decimals}f}"

import copy
import logging
from dataclasses import dataclass, field
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
log = logging.getLogger("master_thesis_trilingual")

# Visual identity
NAVY = RGBColor(0x1A, 0x27, 0x44)
GOLD = RGBColor(0xB8, 0x92, 0x2F)
GREY = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x20, 0x20, 0x20)

@dataclass
class MasterLangPack:
    code: str  # "EN", "FR", "PT-BR"
    filename: str
    doc_type: str
    university: str
    master_info: str
    title: str
    subtitle: str
    analysis_label: str
    chip_lines: list[str]
    author: str
    date_meta: str
    structure_info: str
    keywords_label: str
    keywords: str
    summary_title: str
    summary_paragraphs: list[str]
    toc_title: str
    toc_note: str
    license_title: str
    license_text: str
    chapters: list[tuple[str, str, str]]  # [(relative_path, filename, display_title)]

# ---------------------------------------------------------------------------
# Data Definitions
# ---------------------------------------------------------------------------

FR = MasterLangPack(
    code="FR",
    filename="These_AI_For_Americans_First_FR.docx",
    doc_type="THESE DOCTORALE",
    university="Universite Paris-Sorbonne",
    master_info="Master 2 Intelligence Economique - Intelligence Warfare",
    title="AI FOR AMERICANS FIRST",
    subtitle="Protectionnisme IA, Energie et Semi-conducteurs : Trajectoires de divergence US/Europe 2024-2030",
    analysis_label="Analyse geostrategique et economique integree",
    chip_lines=[
        f"{fmt_fr(us_share, 1)} pct compute IA operationnel mondial = USA",
        "1,59x cout energie EU/US (ajuste-PPA)",
        f"{fmt_fr(us_eu_caci, 2)}:1 ratio CACI US/EU (Power Mode)",
    ],
    author="Fabrice Pizzi",
    date_meta="Paris - fevrier 2026",
    structure_info="11 chapitres | 4 scenarios prospectifs | 4 zones geographiques | 108 figures",
    keywords_label="Mots-cles",
    keywords="intelligence artificielle, protectionnisme technologique, semi-conducteurs, controles a l'exportation, compute souverain, geopolitique IA, France, Etats-Unis, Chine.",
    summary_title="Resume executif",
    summary_paragraphs=[
        "Cette these analyse l'emergence d'un regime protectionniste IA americain (decret implicite 'AI for Americans First') et ses consequences geostrategiques pour la France, l'Europe, et le reste du monde sur la periode 2024-2030. L'etude mobilise un appareil empirique original (tableau de bord public snapshot avril 2026) et propose un cadre theorique unifie (indice CACI - Compute-Adjusted Competitive Index) integrant FLOPs, capital humain, regulation et cout energetique selon une formule geometrique ponderee.",
        f"Sur le snapshot avril 2026, les Etats-Unis concentrent {fmt_fr(us_share, 1)} pour cent du compute IA operationnel mondial. Le ratio brut compute installe US/UE(13) atteint {fmt_fr(us_eu_raw, 1)}:1, traduit par la formule geometrique CACI Power Mode (F^0,40 x L^0,20 x R^0,15 / E^0,25) en un ratio US/UE de {fmt_fr(us_eu_caci, 2)}:1. L'UE est largement souveraine sur son F installe ({fmt_fr(eu_sov, 1)} pct). La vulnerabilite europeenne se situe sur la couche operationnelle des charges cloud (majoritairement hebergees sur AWS/Azure/GCP).",
        "Quatre scenarios prospectifs 2026-2030 sont developpes (Statu quo, Fracture, Partenariat asymetrique, Souverainete subie), avec un point de basculement identifie en 2028 a la convergence de la saturation compute/energie UE et de l'activation potentielle des Cloud Sovereignty Mandates US. La fenetre d'action strategique 2026-2028 est etroite.",
        "Les recommandations strategiques s'articulent en cinq axes (compute, energie, alliances, regulation, talent) sur trois horizons. La France dispose d'avantages distinctifs (parc nucleaire, Mistral AI, ASML) qui en font le pilier credible d'une trajectoire d'autonomie strategique ciblee. L'enjeu n'est pas l'autarcie mais la capacite de choix."
    ],
    toc_title="Table des matieres",
    toc_note="Note : une table des matieres dynamique est integree ci-dessous. Faire Ctrl+A puis F9 sous Word pour mettre a jour.",
    license_title="Licence et credits",
    license_text="Cette these est mise a disposition selon les termes de la licence Creative Commons Attribution - Pas d'Utilisation Commerciale - Partage dans les Memes Conditions 4.0 International (CC BY-NC-SA 4.0).",
    chapters=[
        ("1", "Chapitre_I_Introduction_Cadrage_Theorique_FR.docx", "Chapitre I - Introduction et cadrage theorique"),
        ("2", "Chapitre_II_Methodologie_FR.docx", "Chapitre II - Methodologie"),
        ("3", "Chapitre_III_Diagnostic_Empirique_FR.docx", "Chapitre III - Diagnostic empirique 2020-2026"),
        ("4", "Chapitre_IV_Mecanismes_Avantage_Concurrentiel_FR.docx", "Chapitre IV - Mecanismes de l'avantage concurrentiel US"),
        ("5", "Chapitre_V_Scenarios_Prospectifs_2026_2030_FR.docx", "Chapitre V - Scenarios prospectifs 2026-2030"),
        ("6", "Chapitre_VI_Consequences_France_Europe_FR.docx", "Chapitre VI - Consequences pour la France et l'Europe"),
        ("6", "Chapitre_VI_bis_Amerique_Sud_Bresil_FR.docx", "Chapitre VI bis - Consequences pour l'Amerique du Sud et le Bresil"),
        ("6", "Chapitre_VI_ter_Asie_Chine_FR.docx", "Chapitre VI ter - Consequences pour l'Asie et la Chine"),
        ("6", "Chapitre_VI_quater_Afrique_FR.docx", "Chapitre VI quater - Consequences pour l'Afrique"),
        ("7", "Chapitre_VII_Recommandations_FR.docx", "Chapitre VII - Recommandations strategiques"),
        ("8", "Conclusion_Generale_FR.docx", "Conclusion generale"),
    ]
)

EN = MasterLangPack(
    code="EN",
    filename="Thesis_AI_For_Americans_First_EN.docx",
    doc_type="DOCTORAL THESIS",
    university="Paris-Sorbonne University",
    master_info="Master 2 Economic Intelligence - Intelligence Warfare",
    title="AI FOR AMERICANS FIRST",
    subtitle="AI Protectionism, Energy and Semiconductors: US/Europe Divergence Trajectories 2024-2030",
    analysis_label="Integrated Geostrategic and Economic Analysis",
    chip_lines=[
        f"{fmt_en(us_share, 1)}% of global operational AI compute = USA",
        "1.59x EU/US energy cost (PPP-adjusted)",
        f"{fmt_en(us_eu_caci, 2)}:1 US/EU CACI ratio (Power Mode)",
    ],
    author="Fabrice Pizzi",
    date_meta="Paris - February 2026",
    structure_info="11 Chapters | 4 Prospective Scenarios | 4 Geographic Zones | 108 Figures",
    keywords_label="Keywords",
    keywords="artificial intelligence, technological protectionism, semiconductors, export controls, sovereign compute, AI geopolitics, France, United States, China.",
    summary_title="Executive Summary",
    summary_paragraphs=[
        "This thesis analyzes the emergence of a US AI protectionist regime (the implicit 'AI for Americans First' doctrine) and its geostrategic consequences for France, Europe, and the rest of the world over the 2024-2030 period. The study employs an original empirical framework (public dashboard snapshot April 2026) and proposes a unified theoretical model: the Compute-Adjusted Competitive Index (CACI).",
        f"As of April 2026, the United States concentrates {fmt_en(us_share, 1)}% of global operational AI compute. The raw US/EU compute ratio stands at {fmt_en(us_eu_raw, 1)}:1, which the CACI Power Mode formula translates into a {fmt_en(us_eu_caci, 2)}:1 advantage. While the EU maintains high sovereignty over its physical install base ({fmt_en(eu_sov, 1)}%), its vulnerability lies in the operational cloud layer (hosted mainly on AWS/Azure/GCP).",
        "Four prospective scenarios for 2026-2030 are developed (Status Quo, Fracture, Asymmetric Partnership, Forced Sovereignty), identifying a tipping point in 2028 where EU compute/energy saturation meets the potential activation of US Cloud Sovereignty Mandates.",
        "Strategic recommendations focus on five axes: compute, energy, alliances, regulation, and talent. France's unique advantages (nuclear fleet, Mistral AI, ASML partnership) position it as a credible pillar for targeted strategic autonomy. The goal is not autarky but the capacity for sovereign choice."
    ],
    toc_title="Table of Contents",
    toc_note="Note: a dynamic Table of Contents is included below. Press Ctrl+A then F9 in Word to refresh.",
    license_title="License and Credits",
    license_text="This thesis is made available under the terms of the Creative Commons Attribution-NonCommercial-ShareAlike 4.0 International (CC BY-NC-SA 4.0) license.",
    chapters=[
        ("1", "Chapter_I_Introduction_Cadrage_Theorique_EN.docx", "Chapter I - Introduction and Theoretical Framework"),
        ("2", "Chapter_II_Methodologie_EN.docx", "Chapter II - Methodology"),
        ("3", "Chapter_III_Empirical_Diagnosis_EN.docx", "Chapter III - Empirical Diagnosis 2020-2026"),
        ("4", "Chapter_IV_Mechanisms_Competitive_Advantage_EN.docx", "Chapter IV - Mechanisms of US Competitive Advantage"),
        ("5", "Chapter_V_Prospective_Scenarios_2026_2030_EN.docx", "Chapter V - Prospective Scenarios 2026-2030"),
        ("6", "Chapter_VI_Consequences_France_Europe_EN.docx", "Chapter VI - Consequences for France and Europe"),
        ("6", "Chapter_VI_bis_Americas_Brazil_EN.docx", "Chapter VI bis - Consequences for South America and Brazil"),
        ("6", "Chapter_VI_ter_Asia_China_EN.docx", "Chapter VI ter - Consequences for Asia and China"),
        ("6", "Chapter_VI_quater_Africa_EN.docx", "Chapter VI quater - Consequences for Africa"),
        ("7", "Chapter_VII_Recommendations_EN.docx", "Chapter VII - Strategic Recommendations"),
        ("8", "General_Conclusion_EN.docx", "General Conclusion"),
    ]
)

PT = MasterLangPack(
    code="PT-BR",
    filename="Tese_AI_For_Americans_First_PT-BR.docx",
    doc_type="TESE DE DOUTORADO",
    university="Universidade Paris-Sorbonne",
    master_info="Mestrado 2 Inteligencia Economica - Intelligence Warfare",
    title="AI FOR AMERICANS FIRST",
    subtitle="Protecionismo de IA, Energia e Semicondutores: Trajetorias de Divergencia EUA/Europa 2024-2030",
    analysis_label="Analise Geoestrategica e Economica Integrada",
    chip_lines=[
        f"{fmt_fr(us_share, 1)}% do compute IA operacional mundial = EUA",
        "1,59x custo de energia UE/EUA (ajustado por PPP)",
        f"{fmt_fr(us_eu_caci, 2)}:1 razao CACI EUA/UE (Power Mode)",
    ],
    author="Fabrice Pizzi",
    date_meta="Paris - Fevereiro de 2026",
    structure_info="11 Capitulos | 4 Cenarios Prospectivos | 4 Zonas Geograficas | 108 Figuras",
    keywords_label="Palavras-chave",
    keywords="inteligencia artificial, protecionismo tecnologico, semicondutores, controles de exportacao, compute soberano, geopolitica da IA, Franca, Estados Unidos, China.",
    summary_title="Resumo Executivo",
    summary_paragraphs=[
        "Esta tese analisa a emergencia de um regime protecionista de IA nos EUA (doutrina implicita 'AI for Americans First') e suas consequencias geoestrategicas para a Franca, a Europa e o resto do mundo no periodo 2024-2030. O estudo utiliza um arcabouco empirico original e propoe um modelo teorico unificado: o Compute-Adjusted Competitive Index (CACI).",
        f"Em abril de 2026, os Estados Unidos concentram {fmt_fr(us_share, 1)}% do compute operacional global de IA. A razao bruta de compute instalado EUA/UE e de {fmt_fr(us_eu_raw, 1)}:1, que a formula CACI Power Mode traduz em uma vantagem de {fmt_fr(us_eu_caci, 2)}:1. Embora a UE mantenha alta soberania sobre sua base fisica instalada ({fmt_fr(eu_sov, 1)}%), sua vulnerabilidade reside na camada operacional de nuvem.",
        "Quatro cenarios prospectivos para 2026-2030 sao desenvolvidos, identificando um ponto de virada em 2028 com a convergencia da saturacao de compute/energia na UE e a potencial ativacao dos Mandatos de Soberania de Nuvem dos EUA.",
        "As recomendacoes estrategicas focam em cinco eixos: compute, energia, aliancas, regulacao e talento. As vantagens unicas da Franca posicionam-na como um pilar credivel para a autonomia estrategica direcionada. O objetivo nao e a autarquia, mas a capacidade de escolha soberana."
    ],
    toc_title="Sumario",
    toc_note="Nota: um sumario dinamico esta incluido abaixo. Pressione Ctrl+A e depois F9 no Word para atualizar.",
    license_title="Licenca e Creditos",
    license_text="Esta tese e disponibilizada sob os termos da licenca Creative Commons Atribuicao-NaoComercial-CompartilhaIgual 4.0 Internacional (CC BY-NC-SA 4.0).",
    chapters=[
        ("1", "Capitulo_I_Introduction_PT-BR.docx", "Capitulo I - Introducao e Enquadramento Teorico"),
        ("2", "Capitulo_II_Metodologia_PT-BR.docx", "Capitulo II - Metodologia"),
        ("3", "Capitulo_III_Diagnostico_Empirico_PT-BR.docx", "Capitulo III - Diagnostico Empirico 2020-2026"),
        ("4", "Capitulo_IV_Mecanismos_Vantagem_Competitiva_PT-BR.docx", "Capitulo IV - Mecanismos da Vantagem Competitiva dos EUA"),
        ("5", "Capitulo_V_Cenarios_Prospectivos_2026_2030_PT-BR.docx", "Capitulo V - Cenarios Prospectivos 2026-2030"),
        ("6", "Capitulo_VI_Consequencias_Franca_Europa_PT-BR.docx", "Capitulo VI - Consequencias para a Franca e a Europa"),
        ("6", "Capitulo_VI_bis_Americas_Brasil_PT-BR.docx", "Capitulo VI bis - Consequencias para a America do Sul e o Brasil"),
        ("6", "Capitulo_VI_ter_Asia_China_PT-BR.docx", "Capitulo VI ter - Consequencias para a Asia e a China"),
        ("6", "Capitulo_VI_quater_Africa_PT-BR.docx", "Capitulo VI quater - Consequencias para a Africa"),
        ("7", "Capitulo_VII_Recomendacoes_PT-BR.docx", "Capitulo VII - Recomendacoes Estrategicas"),
        ("8", "Conclusao_Geral_PT-BR.docx", "Conclusao Geral"),
    ]
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_run(run, *, font="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color

def _add_paragraph(doc, text, *, align=None, space_after=6, **run_kwargs):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        _set_run(run, **run_kwargs)
    return p

def _add_page_break(doc) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def _add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)

def add_master_cover(doc, lp: MasterLangPack) -> None:
    _add_paragraph(doc, "", space_after=20)
    _add_paragraph(doc, lp.doc_type, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, bold=True, color=GREY, space_after=4)
    _add_paragraph(doc, lp.university, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True, color=GREY, space_after=4)
    _add_paragraph(doc, lp.master_info, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True, color=GREY, space_after=40)
    _add_paragraph(doc, lp.title, align=WD_ALIGN_PARAGRAPH.CENTER, size=32, bold=True, color=NAVY, space_after=10)
    _add_paragraph(doc, lp.subtitle, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, italic=True, color=NAVY, space_after=10)
    _add_paragraph(doc, lp.analysis_label, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, italic=True, color=GREY, space_after=40)

    table = doc.add_table(rows=1, cols=3)
    for i, line in enumerate(lp.chip_lines):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run(para.add_run(line), size=11, bold=True, color=NAVY)

    _add_paragraph(doc, "", space_after=40)
    _add_paragraph(doc, lp.author, align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True, color=NAVY, space_after=4)
    _add_paragraph(doc, "(handle mo0ogly)", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, italic=True, color=GREY, space_after=20)
    _add_paragraph(doc, lp.date_meta, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, color=GREY, space_after=40)
    _add_paragraph(doc, lp.structure_info, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, italic=True, color=GREY, space_after=20)
    _add_paragraph(doc, f"{lp.keywords_label}: {lp.keywords}", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, italic=True, color=GREY, space_after=20)

def add_executive_summary(doc, lp: MasterLangPack) -> None:
    _add_paragraph(doc, lp.summary_title, size=22, bold=True, color=NAVY, space_after=14)
    for text in lp.summary_paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.3
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_run(p.add_run(text), size=11, color=DARK)

def add_table_of_contents(doc, lp: MasterLangPack) -> None:
    _add_paragraph(doc, lp.toc_title, size=22, bold=True, color=NAVY, space_after=14)
    for i, (_, _, title) in enumerate(lp.chapters, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        _set_run(p.add_run(f"{i:>2}.  "), size=11, bold=True, color=GOLD)
        _set_run(p.add_run(title), size=11, color=NAVY)
    _add_paragraph(doc, "", space_after=10)
    _add_paragraph(doc, lp.toc_note, size=9, italic=True, color=GREY, space_after=10)
    p = doc.add_paragraph()
    _add_field(p, r'TOC \o "1-2" \h \z \u')

def import_chapter_body(master_doc: Document, chapter_path: Path) -> int:
    src = Document(str(chapter_path))
    src_body = src.element.body
    
    # Skip cover: find the first bold paragraph with specific text
    cover_end = 0
    for i, el in enumerate(src_body):
        text = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
        # Typical chapter starts with "CHAPITRE I" or "CHAPTER I" or "CAPITULO I"
        if any(kw in text.upper() for kw in ["CHAPITRE ", "CHAPTER ", "CAPITULO ", "CONCLUSION GENERALE", "GENERAL CONCLUSION", "CONCLUSAO GERAL"]) and len(text) < 30:
            cover_end = i
            break
    if cover_end == 0: cover_end = 24 # fallback

    imported = 0
    target_body = master_doc.element.body
    sect_pr = target_body.find(qn("w:sectPr"))

    for el in list(src_body)[cover_end:]:
        if el.tag == qn("w:sectPr"): continue
        text = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
        if text.startswith("Licence et avertissement") or text.startswith("License and disclaimer") or text.startswith("Licenca e isencao"): break
        if "AI for Americans First - Fabrice Pizzi" in text: continue
        if text.startswith("Tableau de bord public") or text.startswith("Public dashboard") or text.startswith("Painel publico"): continue
        if text.startswith("Depot :") or text.startswith("Repository:"): continue
        if text.startswith("Vous etes libre") or text.startswith("You are free") or text.startswith("Voce e livre"): continue

        new_el = copy.deepcopy(el)
        if sect_pr is not None:
            target_body.insert(list(target_body).index(sect_pr), new_el)
        else:
            target_body.append(new_el)
        imported += 1
    return imported

def build_master(lp: MasterLangPack, base_dir: Path, output_dir: Path):
    log.info("Building Master Thesis [%s] -> %s", lp.code, lp.filename)
    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.2)
    section.top_margin = section.bottom_margin = Cm(2.0)

    # Styles
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"; h1.font.size = Pt(22); h1.font.bold = True; h1.font.color.rgb = NAVY
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"; h2.font.size = Pt(16); h2.font.bold = True; h2.font.color.rgb = NAVY

    add_master_cover(doc, lp)
    _add_page_break(doc)
    add_executive_summary(doc, lp)
    _add_page_break(doc)
    add_table_of_contents(doc, lp)
    _add_page_break(doc)

    for rel_path, filename, display_title in lp.chapters:
        chapter_path = base_dir / rel_path / filename
        log.info("  Importing %s", filename)
        
        p = doc.add_paragraph(style="Heading 1")
        _set_run(p.add_run(display_title.upper()), font="Calibri", size=22, bold=True, color=NAVY)
        
        n = import_chapter_body(doc, chapter_path)
        log.info("    -> %d elements", n)
        _add_page_break(doc)

    # Final page
    _add_paragraph(doc, lp.license_title, size=18, bold=True, color=NAVY, space_after=12)
    _add_paragraph(doc, lp.license_text, size=10, italic=True, color=GREY, space_after=8)
    
    out = output_dir / lp.filename
    doc.save(str(out))
    log.info("Saved %s", out)

def main():
    base_dir = Path(r"c:\Users\pizzif\Documents\GitHub\America-First-IA-main\docs\a traduire")
    output_dir = base_dir / "these"
    output_dir.mkdir(exist_ok=True)
    
    for lp in (FR, EN, PT):
        build_master(lp, base_dir, output_dir)

if __name__ == "__main__":
    main()
