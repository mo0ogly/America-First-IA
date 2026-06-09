"""

import sys
from pathlib import Path
sys.path.append(str(Path(__file__).resolve().parents[1]))
import caci_data_helper
m = caci_data_helper.get_metrics()
us_share = m['us_compute_share']
us_eu_caci = m['us_eu_caci_power_ratio']
us_eu_raw = m['us_eu_raw_ratio']
eu_sov = m['eu_sovereignty_ratio']
caci_scores = {k: v['caci_power_phys'] for k, v in m['country_results'].items()}

def fmt_fr(val, decimals=1):
    return f"{val:.{decimals}f}".replace(".", ",")

def fmt_en(val, decimals=1):
    return f"{val:.{decimals}f}"

Annexe C - Academic Note - Trilingual Generator (EN, FR, PT-BR).

Generates the .docx for Annex C in three languages.
Synthesis of the thesis arguments + Grand Decoupling 2028 addendum.
"""

from __future__ import annotations

import sys
from pathlib import Path
sys.path.append(str(Path(__file__).resolve().parents[1]))
import caci_data_helper
m = caci_data_helper.get_metrics()
us_share = m['us_compute_share']
us_eu_caci = m['us_eu_caci_power_ratio']
us_eu_raw = m['us_eu_raw_ratio']
eu_sov = m['eu_sovereignty_ratio']
caci_scores = {k: v['caci_power_phys'] for k, v in m['country_results'].items()}

def fmt_fr(val, decimals=1):
    return f"{val:.{decimals}f}".replace(".", ",")

def fmt_en(val, decimals=1):
    return f"{val:.{decimals}f}"


import logging
from dataclasses import dataclass
from pathlib import Path

from annexes_helpers import (
    add_chapter_header, add_cover, init_document,
    render_license, render_notes, render_section, render_table,
)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
log = logging.getLogger("annexe_c_trilingual")

@dataclass
class LangPack:
    code: str
    label: str
    title: str
    intro: str
    sections: list[tuple[str, list[str]]]
    tables: list[tuple[str, str, list[list[str]]]]
    notes: list[str]
    sources_line: str
    footer: str
    filename: str

# ---------------------------------------------------------------------------
# Content - ENGLISH
# ---------------------------------------------------------------------------

EN = LangPack(
    code="EN",
    label="ANNEX C - ACADEMIC NOTE",
    title="AI for Americans First: American AI Protectionism, Reconfiguration of the Global Technological Order and Consequences for France and Europe (2026-2030)",
    intro=(
        "This annex presents the Academic Synthesis Note of the dissertation. This study analyzes the mechanisms and "
        "consequences of American AI protectionism under the Trump 2.0 administration, integrating four dimensions: "
        "energy, semiconductors, compute, and regulation. Based on an empirical diagnosis (2020-2026), the construction "
        "of a Compute-Adjusted Competitiveness Index (CACI), and a 2x2 scenario matrix, the research demonstrates that "
        "the combination of tariffs (25%, Section 232) and export controls creates a measurable structural competitive "
        f"advantage (CACI US/EU ratio of {fmt_en(us_eu_caci, 2)}:1, Power mode, for a raw operational compute ratio of {fmt_en(us_eu_raw, 1)}:1, April 2026). "
        "The analysis reveals differentiated trajectories of dependence across regions (Europe, South America, Asia, Africa). "
        "A prospective addendum formalizes the 'Grand Decoupling 2028' scenario based on US Cloud Sovereignty Mandates, "
        "revealing that France and Europe must distinguish between physically installed capacity and operationally sovereign compute."
    ),
    sections=[
        ("C.1 Subject and Problem Statement", [
            f"AI has emerged since 2023 as the primary vector of economic innovation and geopolitical competition. Yet the AI value chain exhibits unprecedented concentration: as of the April 2026 public dashboard snapshot, the United States controls {fmt_en(us_share, 1)} percent of global operational AI compute, and five US hyperscalers plan 660-690 billion USD in capex for 2026 alone.",
            "In this context, the Trump 2.0 administration has transformed Biden-era export controls (2022-2025) into a hybrid protectionist regime. This study asks: to what extent does American AI protectionism create a measurable structural competitive advantage, and what are the differentiated consequences for France, Europe, and other world regions?",
        ]),
        ("C.2 Methodological Framework", [
            "The methodology rests on three pillars: (1) An empirical longitudinal diagnosis (2020-2026) of energy, semiconductors, and compute distributions; (2) The construction of the Compute-Adjusted Competitiveness Index (CACI) Power Mode formula: CACI = F^0.40 x L^0.20 x R^0.15 / E^0.25; (3) A 2x2 scenario matrix crossing US protectionism intensity with European strategic response.",
            "A prospective addendum extends the CACI model with the decomposition F(r) = F_phys(r) x F_sov(r), where F_sov measures the fraction of regional compute operated outside US jurisdiction. This reveals that physically installed capacity does not equal operationally sovereign capacity.",
        ]),
        ("C.3.1 A Three-Tiered Protectionism", [
            "The study identifies a cumulative three-level protectionist architecture: (1) Export controls segmenting the world into three access tiers; (2) 25% Section 232 tariffs on advanced AI semiconductors creating a direct cost differential; (3) Capitalistic gravity effects where massive capex concentration reinforces US supremacy without further regulation.",
            "A potential fourth tier looms for 2028: Cloud Sovereignty Mandates, transforming US hyperscalers operating offshore into conditional intermediaries of global compute.",
        ]),
        ("C.3.4 The Grand Decoupling 2028", [
            "This scenario models a qualitative shift: the transition from control of chip flows to control of the jurisdictional layer (operating sovereignty). The CLOUD Act (2018) already establishes that a cluster physically installed in Ireland remains legally American. The 'America's AI Action Plan' (2025) introduces location-verification features in chips for remote throttling.",
            "Decomposition reveals that while the US and China are fully sovereign (F_sov=1.00), the EU is largely sovereign on cluster ownership (F_sov=0.99) but highly dependent on the cloud workload layer (F_sov_workloads ~0.28).",
        ]),
    ],
    tables=[
        ("Table C.1. Summary of regional consequences of American AI protectionism.", "Source: Author's construction, calibrated on April 2026 snapshot.", [
            ["Region", "Tier", "Main Dynamic", "Strategic Asset", "Main Risk"],
            ["France / Europe", "1", "Dependence on US GPU + cloud; InvestAI response", "Nuclear energy (70% mix), Mistral, ASML, AI Act", "Geopolitical vendor lock-in; Low F_sov; CSM 2028"],
            ["Brazil / S. America", "2", "US-China competition arena; Scala/TikTok projects", "83% renewable energy mix; dynamic fintech market", "Triple fracture (N-S, E-W, intra-regional)"],
            ["China", "3", "Forced autonomy; alternative ecosystem (Huawei/DeepSeek)", "Domestic market 1.4B; 125B+ USD/yr investment", "2-3 gen GPU lag; tech isolation"],
        ]),
    ],
    notes=[
        f"Public dashboard snapshot April 2026: USA {fmt_en(us_share, 1)}% of global operational AI compute.",
        "CACI Power Mode: F^0.40 x L^0.20 x R^0.15 / E^0.25. Validated at beta=0.251, p<0.01.",
        "Section 232: 25% tariff on advanced AI semiconductors (Jan 15, 2026).",
        "CLOUD Act (2018): Jurisdictional control over data regardless of physical storage location.",
    ],
    sources_line="Main sources: IEA (2025-2026), World Bank (2025), Epoch AI, McKinsey, Synergy Research, White House BIS.",
    footer="AI for Americans First - Fabrice Pizzi - Annex C Academic Note",
    filename="Annex_C_Academic_Note_EN.docx"
)

# ---------------------------------------------------------------------------
# Content - FRENCH
# ---------------------------------------------------------------------------

FR = LangPack(
    code="FR",
    label="ANNEXE C - NOTE ACADEMIQUE",
    title="AI for Americans First : protectionnisme IA americain, recomposition de l'ordre technologique mondial et consequences pour la France et l'Europe (2026-2030)",
    intro=(
        "Cette annexe presente la Note Academique de synthese de la these. Cette etude analyse les mecanismes et les "
        "consequences du protectionnisme IA americain sous l'administration Trump 2.0, en integrant quatre dimensions : "
        "energie, semi-conducteurs, compute et regulation."
    ),
    sections=[
        ("C.1 Objet et problematique", [
            f"L'intelligence artificielle s'est imposee depuis 2023 comme le principal vecteur d'innovation economique et de competition geopolitique. Sur le snapshot avril 2026, les Etats-Unis controlent {fmt_fr(us_share, 1)} pour cent du compute IA operationnel mondial.",
        ]),
        ("C.2 Cadre methodologique", [
            "La methodologie repose sur trois piliers : diagnostic empirique 2020-2026, indice CACI (Compute-Adjusted Competitiveness Index) et matrice scenarielle 2x2.",
            "Un addendum prospectif formalise le 'Grand Decouplage 2028' fonde sur les Cloud Sovereignty Mandates americains.",
        ]),
    ],
    tables=[
        ("Tableau C.1. Synthese des consequences regionales.", "Source : construction de l'auteur, snapshot avril 2026.", [
            ["Region", "Tier", "Dynamique principale", "Atout strategique", "Risque principal"],
            ["France / Europe", "1", "Dependance GPU + cloud US ; reponse InvestAI", "Nucleaire (70% mix), Mistral, ASML, AI Act", "Vendor lock-in geopolitique ; F_sov bas"],
        ]),
    ],
    notes=[
        f"Snapshot avril 2026 : USA {fmt_fr(us_share, 1)}% du compute IA operationnel mondial.",
        "Indice CACI Power Mode valide econometriquement (beta = 0,251).",
    ],
    sources_line="Sources principales : AIE (2025-2026), Banque mondiale (2025), Epoch AI, McKinsey, Synergy Research.",
    footer="AI for Americans First - Fabrice Pizzi - Annexe C Note academique",
    filename="Annexe_C_Note_Academique_FR.docx"
)

# ---------------------------------------------------------------------------
# Content - PORTUGUESE (BRAZIL)
# ---------------------------------------------------------------------------

BR = LangPack(
    code="PT-BR",
    label="ANEXO C - NOTA ACADEMICA",
    title="AI for Americans First: Protecionismo de IA Americano, Recomposiçao da Ordem Tecnologica Mundial e Consequencias para a França e a Europa (2026-2030)",
    intro=(
        "Este anexo apresenta a Nota Academica de sintese da tese. Este estudo analisa os mecanismos e as consequencias "
        "do protecionismo de IA americano sob a administraçao Trump 2.0, integrando quatro dimensoes: energia, "
        "semicondutores, compute e regulaçao. Com base em um diagnostico empirico (2020-2026), na construçao de um "
        "Indice de Competitividade Ajustado ao Compute (CACI) e em uma matriz de cenarios 2x2, a pesquisa demonstra "
        "que a combinaçao de tarifas (25%, Seçao 232) e controles de exportaçao cria uma vantagem competitiva estrutural "
        f"mensuravel (razao CACI EUA/UE de {fmt_fr(us_eu_caci, 2)}:1, modo Power, para uma razao bruta de computaçao de {fmt_fr(us_eu_raw, 1)}:1, abril de 2026)."
    ),
    sections=[
        ("C.1 Objeto e Problemática", [
            "A IA surgiu desde 2023 como o principal vetor de inovaçao economica e competiçao geopolitica. No snapshot de abril de 2026, os Estados Unidos controlam 78,9 por cento do compute de IA operacional global.",
        ]),
        ("C.2 Quadro Metodológico", [
            "A metodologia baseia-se em tres pilares: diagnostico empirico 2020-2026, indice CACI e matriz de cenarios 2x2.",
            "Um adendo prospectivo formaliza o cenário de 'Grande Desacoplamento 2028' baseado nos Mandatos de Soberania em Nuvem dos EUA.",
        ]),
    ],
    tables=[
        ("Tabela C.1. Síntese das consequências regionais.", "Fonte: Construçao do autor, snapshot abril 2026.", [
            ["Região", "Tier", "Dinâmica Principal", "Ativo Estratégico", "Principal Risco"],
            ["França / Europa", "1", "Dependencia de GPU + nuvem EUA; resposta InvestAI", "Energia nuclear (70% mix), Mistral, ASML, AI Act", "Vendor lock-in geopolitico; F_sov baixo"],
            ["Brasil / Am. do Sul", "2", "Arena de competiçao EUA-China; projetos Scala/TikTok", "Mix 83% renovavel; mercado fintech dinamico", "Tripla fratura (N-S, L-O, intra-regional)"],
        ]),
    ],
    notes=[
        f"Snapshot abril 2026: EUA {fmt_fr(us_share, 1)}% do compute operacional global.",
        "Indice CACI Power Mode validado econometricamente (beta = 0,251).",
    ],
    sources_line="Fontes principais: AIE (2025-2026), Banco Mundial (2025), Epoch AI, McKinsey, Synergy Research.",
    footer="AI for Americans First - Fabrice Pizzi - Anexo C Nota Academica",
    filename="Anexo_C_Nota_Academica_PT-BR.docx"
)

# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------

LANGS = [EN, FR, BR]

def build_all(out_dir: Path) -> None:
    for lp in LANGS:
        log.info("Building Annexe C [%s] -> %s", lp.code, lp.filename)
        doc = init_document()
        add_cover(doc, lang=lp.code, chapter_label=lp.label, chapter_subtitle="Note academique de synthese" if lp.code == "FR" else ("Academic Synthesis Note" if lp.code == "EN" else "Nota Academica de Sintese"))
        add_chapter_header(doc, label=lp.label, title=lp.title, intro=lp.intro)
        
        for title, paragraphs in lp.sections:
            render_section(doc, title, paragraphs)
            
        for caption, source, rows in lp.tables:
            render_table(doc, lang=lp.code, caption=caption, source=source, rows=rows)
            
        # Add sources line
        from annexes_helpers import add_paragraph, GREY
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        add_paragraph(doc, lp.sources_line, align=WD_ALIGN_PARAGRAPH.LEFT, size=9, italic=True, color=GREY, space_after=6)

        render_notes(doc, lang=lp.code, notes=lp.notes)
        render_license(doc, lang=lp.code, page_footer=lp.footer)
        
        lang_sub = lp.code.lower()
        if lang_sub == "pt-br": lang_sub = "br"
        
        target_dir = out_dir / lang_sub
        target_dir.mkdir(parents=True, exist_ok=True)
        
        out = target_dir / lp.filename
        doc.save(out)
        log.info("Saved %s", out)

if __name__ == "__main__":
    build_all(Path(__file__).parent)
