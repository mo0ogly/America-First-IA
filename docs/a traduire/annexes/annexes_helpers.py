"""

import sys
from pathlib import Path
sys.path.append(str(Path(__file__).resolve().parents[1]))
import caci_data_helper
m = caci_data_helper.get_metrics()
us_share = m['us_compute_share']
us_eu_caci = m['us_eu_caci_power_ratio']
us_eu_raw = m['us_eu_raw_ratio']
eu_sov = m['eu_sovereignty_ratio']
caci_scores = {k: v['caci_power_phys'] for k, v in m['country_results'].items()}

def fmt_fr(val, decimals=1):
    return f"{val:.{decimals}f}".replace(".", ",")

def fmt_en(val, decimals=1):
    return f"{val:.{decimals}f}"

Trilingual helpers for the Annexes (EN, FR, PT-BR).

Provides cover, header, section/table/notes/license rendering matching
the visual identity of the thesis.
"""

from __future__ import annotations

import sys
from pathlib import Path
sys.path.append(str(Path(__file__).resolve().parents[1]))
import caci_data_helper
m = caci_data_helper.get_metrics()
us_share = m['us_compute_share']
us_eu_caci = m['us_eu_caci_power_ratio']
us_eu_raw = m['us_eu_raw_ratio']
eu_sov = m['eu_sovereignty_ratio']
caci_scores = {k: v['caci_power_phys'] for k, v in m['country_results'].items()}

def fmt_fr(val, decimals=1):
    return f"{val:.{decimals}f}".replace(".", ",")

def fmt_en(val, decimals=1):
    return f"{val:.{decimals}f}"


import logging
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1A, 0x27, 0x44)
GOLD = RGBColor(0xB8, 0x92, 0x2F)
GREY = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x20, 0x20, 0x20)

LABELS = {
    "EN": {
        "study": "RESEARCH STUDY - FEBRUARY 2026",
        "subtitle": "AI Protectionism, Energy and Semiconductors:\nUS/Europe Divergence Trajectories 2024-2030",
        "analysis": "Integrated Geostrategic and Economic Analysis - {label}",
        "author": "Fabrice Pizzi",
        "university": "Paris-Sorbonne University",
        "master": "Master 2 Economic Intelligence - Intelligence Warfare",
        "stats": "Paris - February 2026  |  11 chapters  |  4 prospective scenarios  |  Global analysis",
        "keywords": "Keywords: artificial intelligence, technological protectionism, semiconductors, export controls, sovereign compute, AI geopolitics, France, United States, China",
        "notes": "Notes",
        "license": "License and Disclaimer. This work, 'AI for Americans First', is made available under the terms of the Creative Commons Attribution - NonCommercial - ShareAlike 4.0 International (CC BY-NC-SA 4.0) license.",
        "license_body": "You are free to share and adapt the material for non-commercial purposes, provided you properly attribute the work to Fabrice Pizzi (Paris-Sorbonne University) and distribute your contributions under the same license. This document is provided for educational and research purposes only.",
        "dashboard": "Public dashboard: https://mo0ogly.github.io/America-First-IA/dashboard/",
        "repo": "Repository: https://github.com/mo0ogly/America-First-IA",
        "chips": [
            f"{fmt_en(us_share, 1)}% of global operational AI compute = USA",
            "1.59x energy cost EU/US (PPP-adjusted)",
            f"{fmt_en(us_eu_caci, 2)}:1 CACI US/EU ratio (Power Mode)",
        ]
    },
    "FR": {
        "study": "ETUDE DE RECHERCHE - FEVRIER 2026",
        "subtitle": "Protectionnisme IA, Energie et Semi-conducteurs :\nTrajectoires de divergence US/Europe 2024-2030",
        "analysis": "Analyse geostrategique et economique integree - {label}",
        "author": "Fabrice Pizzi",
        "university": "Universite Paris-Sorbonne",
        "master": "Master 2 Intelligence Economique - Intelligence Warfare",
        "stats": "Paris - fevrier 2026  |  11 chapitres  |  4 scenarios prospectifs  |  Analyse globale",
        "keywords": "Mots-cles : intelligence artificielle, protectionnisme technologique, semi-conducteurs, controles a l'exportation, compute souverain, geopolitique IA, France, Etats-Unis, Chine",
        "notes": "Notes",
        "license": "Licence et avertissement. Cette oeuvre, 'AI for Americans First', est mise a disposition selon les termes de la licence Creative Commons Attribution - Pas d'Utilisation Commerciale - Partage dans les Memes Conditions 4.0 International (CC BY-NC-SA 4.0).",
        "license_body": "Vous etes libre de partager et adapter le materiel a des fins non commerciales, a condition d'attribuer correctement l'oeuvre a Fabrice Pizzi (Universite Paris-Sorbonne) et de distribuer vos contributions sous la meme licence. Ce document est fourni a des fins educatives et de recherche uniquement.",
        "dashboard": "Tableau de bord public : https://mo0ogly.github.io/America-First-IA/dashboard/",
        "repo": "Depot : https://github.com/mo0ogly/America-First-IA",
        "chips": [
            f"{fmt_fr(us_share, 1)} pct du compute IA operationnel mondial = USA",
            "1,59x cout energie EU/US (ajuste-PPA)",
            f"{fmt_fr(us_eu_caci, 2)}:1 ratio CACI US/EU (Power Mode)",
        ]
    },
    "PT-BR": {
        "study": "ESTUDO DE PESQUISA - FEVEREIRO 2026",
        "subtitle": "Protecionismo de IA, Energia e Semicondutores:\nTrajetorias de Divergencia EUA/Europa 2024-2030",
        "analysis": "Analise Geoestrategica e Economica Integrada - {label}",
        "author": "Fabrice Pizzi",
        "university": "Universidade Paris-Sorbonne",
        "master": "Mestrado 2 Inteligencia Economica - Intelligence Warfare",
        "stats": "Paris - fevereiro 2026  |  11 capitulos  |  4 cenarios prospectivos  |  Analise global",
        "keywords": "Palavras-chave: inteligencia artificial, protecionismo tecnologico, semicondutores, controles de exportacao, computacao soberana, geopolitica da IA, Franca, Estados Unidos, China",
        "notes": "Notas",
        "license": "Licenca e Isencao de Responsabilidade. Este trabalho, 'AI for Americans First', e disponibilizado sob os termos da licenca Creative Commons Atribuicao-NaoComercial-CompartilhaIgual 4.0 Internacional (CC BY-NC-SA 4.0).",
        "license_body": "Voce tem a liberdade de compartilhar e adaptar o material para fins nao comerciais, desde que atribua corretamente o trabalho a Fabrice Pizzi (Universidade Paris-Sorbonne) e distribua suas contribuicoes sob a mesma licenca. Este documento e fornecido apenas para fins educacionais e de pesquisa.",
        "dashboard": "Painel publico: https://mo0ogly.github.io/America-First-IA/dashboard/",
        "repo": "Repositorio: https://github.com/mo0ogly/America-First-IA",
        "chips": [
            f"{fmt_fr(us_share, 1)}% da computacao de IA operacional global = EUA",
            "1,59x custo de energia UE/EUA (ajustado-PPP)",
            f"{fmt_fr(us_eu_caci, 2)}:1 ratio CACI EUA/UE (Power Mode)",
        ]
    }
}

def set_run(run, *, font="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color

def add_paragraph(doc, text, *, align=None, space_after=6, **run_kwargs):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run(run, **run_kwargs)
    return p

def add_heading(doc, text, level):
    sizes = {1: 22, 2: 16, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, font="Calibri", size=sizes.get(level, 11), bold=True, color=NAVY)
    return p

def init_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    return doc

def add_cover(doc, lang: str, chapter_label: str, chapter_subtitle: str):
    L = LABELS.get(lang, LABELS["EN"])
    add_paragraph(doc, "", space_after=0)
    add_paragraph(doc, L["study"], align=WD_ALIGN_PARAGRAPH.CENTER, size=10, bold=True, color=GREY, space_after=4)
    add_paragraph(doc, "AI FOR AMERICANS FIRST", align=WD_ALIGN_PARAGRAPH.CENTER, size=26, bold=True, color=NAVY, space_after=4)
    add_paragraph(doc, L["subtitle"], align=WD_ALIGN_PARAGRAPH.CENTER, size=12, italic=True, color=GREY, space_after=4)
    add_paragraph(doc, L["analysis"].format(label=chapter_label), align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True, color=GREY, space_after=18)

    table = doc.add_table(rows=1, cols=3)
    for i, line in enumerate(L["chips"]):
        cell = table.rows[0].cells[i]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        set_run(run, size=11, bold=True, color=NAVY)

    add_paragraph(doc, "", space_after=8)
    add_paragraph(doc, L["author"], align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, color=NAVY, space_after=2)
    add_paragraph(doc, L["university"], align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=GREY, space_after=2)
    add_paragraph(doc, L["master"], align=WD_ALIGN_PARAGRAPH.CENTER, size=10, italic=True, color=GREY, space_after=14)
    add_paragraph(doc, L["stats"], align=WD_ALIGN_PARAGRAPH.CENTER, size=10, italic=True, color=GREY, space_after=10)
    add_paragraph(doc, L["keywords"], align=WD_ALIGN_PARAGRAPH.CENTER, size=9, italic=True, color=GREY, space_after=24)

def add_chapter_header(doc, label: str, title: str, intro: str):
    add_paragraph(doc, label, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, bold=True, color=GOLD, space_after=2)
    add_paragraph(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True, color=NAVY, space_after=12)
    if intro:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(intro)
        set_run(run, size=11, italic=True, color=GREY)

def render_section(doc, title: str, paragraphs: list[str]):
    first = title.split()[0]
    # Level 2 if 1 dot, level 3 if 2 dots
    level = 2 if first.count(".") == 1 else 3
    # Check if it's a "B.X" or "C.X" style (1 dot)
    if len(first) > 1 and first[1] == "." and first.count(".") == 1:
        level = 2
    
    add_heading(doc, title, level)
    for para in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(para)
        set_run(run, size=11, color=DARK)

def render_table(doc, lang: str, caption: str, source: str, rows: list[list[str]]):
    add_paragraph(doc, caption, size=10, bold=True, color=NAVY, space_after=2)
    add_paragraph(doc, source, size=9, italic=True, color=GREY, space_after=4)
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            set_run(run, size=10, bold=(i == 0), color=NAVY if i == 0 else DARK)
    doc.add_paragraph()

def render_notes(doc, lang: str, notes: list[str]):
    L = LABELS.get(lang, LABELS["EN"])
    add_heading(doc, L["notes"], 2)
    for i, note in enumerate(notes, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run_id = p.add_run(f"{i}. ")
        set_run(run_id, size=9, bold=True, color=GOLD)
        run_txt = p.add_run(note)
        set_run(run_txt, size=9, color=GREY)

def render_license(doc, lang: str, page_footer: str):
    L = LABELS.get(lang, LABELS["EN"])
    doc.add_paragraph()
    license_lines = [
        L["license"],
        L["license_body"],
        L["dashboard"],
        L["repo"],
    ]
    for line in license_lines:
        add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT, size=8, italic=True, color=GREY, space_after=2)
    add_paragraph(doc, page_footer, align=WD_ALIGN_PARAGRAPH.CENTER, size=8, italic=True, color=GREY, space_after=0)
