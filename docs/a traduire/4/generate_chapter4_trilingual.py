"""

import sys
from pathlib import Path
sys.path.append(str(Path(__file__).resolve().parents[1]))
import caci_data_helper
m = caci_data_helper.get_metrics()
us_share = m['us_compute_share']
us_eu_caci = m['us_eu_caci_power_ratio']
us_eu_raw = m['us_eu_raw_ratio']
eu_sov = m['eu_sovereignty_ratio']
caci_scores = {k: v['caci_power_phys'] for k, v in m['country_results'].items()}

def fmt_fr(val, decimals=1):
    return f"{val:.{decimals}f}".replace(".", ",")

def fmt_en(val, decimals=1):
    return f"{val:.{decimals}f}"

Chapter IV - Mechanisms of US Competitive Advantage - trilingual generator.

Generates the Chapter IV .docx for the doctoral study
"AI for Americans First" in English, French and Brazilian Portuguese.

Key revisions vs the previous version (April 2026 dashboard alignment):
    1. Cover banner band updated to 76.9 / 1.59x / {fmt_en(us_eu_caci, 2)}:1, in line with
       Chapters I, II and III.
    2. Section 4 introduction: 'US concentrates 75 percent of global AI
       computef' replaced by '{fmt_en(us_share, 1)} percent of operational compute'.
    3. Section 4.1.2 (compute as advantage): the 'energy 2 to 3 times
       higher' claim is annotated with the PPA-adjusted band of
       '1.4 to 1.7 times', consistent with Chapter III.
    4. Section 4.1.2 (FLOP cost gap M2): the headline figure stays at
       2.4-3.6x because it captures hyperscaler markup + PPA-adjusted
       energy + amortisation, but the derivation is now anchored on the
       1.59x energy ratio rather than the obsolete 2-3x.
    5. Section 4.3.2 + Table 9: CACI ratio updated from 3.4:1 to {fmt_en(us_eu_caci, 2)}:1.
       Table 9 f'CACI 7-12x inferior' line refreshed to '{fmt_en(us_eu_caci, 2)}:1 in Power
       Mode (28.9 vs 100)'.
    6. Section 4.5: synthesis paragraph reformulated so the loop runs
       on the live 76.9 / 49.9 / 3.46 figures rather than 75 / 70 / 3.4.

Author: Fabrice Pizzi (Universite Paris-Sorbonne, M2 Economic Intelligence).
Build: python3 generate_chapter4_trilingual.py
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
log = logging.getLogger("chapter4_gen")


# ---------------------------------------------------------------------------
# Visual identity (shared with Chapters I, II, III)
# ---------------------------------------------------------------------------

NAVY = RGBColor(0x1A, 0x27, 0x44)
GOLD = RGBColor(0xB8, 0x92, 0x2F)
GREY = RGBColor(0x55, 0x55, 0x55)


@dataclass
class LangPack:
    """Container for one language version of Chapter IV."""

    code: str
    filename: str
    cover_subtitle: str
    cover_title: str
    cover_blurb: str
    cover_chip_lines: list[str]
    cover_meta: str
    cover_keywords_label: str
    cover_keywords: str
    chapter_label: str
    chapter_title: str
    chapter_intro: str
    sections: list[tuple[str, list[str]]] = field(default_factory=list)
    table_blocks: list[tuple[str, str, list[list[str]]]] = field(default_factory=list)
    notes_label: str = "Notes"
    notes: list[str] = field(default_factory=list)
    license_block: list[str] = field(default_factory=list)
    page_footer: str = ""


# ---------------------------------------------------------------------------
# Style helpers (identical to Chap I/II/III generators)
# ---------------------------------------------------------------------------

def set_run(run, *, font="Calibri", size=11, bold=False, italic=False, color=None):
    """Apply consistent typography to a docx run."""
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text, *, style=None, align=None, space_after=6, **run_kwargs):
    """Add a paragraph and return it."""
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run(run, **run_kwargs)
    return p


def add_heading(doc, text, level):
    """Custom heading rendering aligned with the rest of the dissertation."""
    sizes = {1: 22, 2: 16, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, font="Calibri", size=sizes.get(level, 11),
            bold=True, color=NAVY)
    return p


def add_cover(doc, lp: LangPack):
    """Render the cover block matching Chapters I/II/III."""
    banner = {
        "EN": "RESEARCH STUDY - FEBRUARY 2026",
        "FR": "ETUDE DE RECHERCHE - FEVRIER 2026",
        "PT-BR": "ESTUDO DE PESQUISA - FEVEREIRO DE 2026",
    }.get(lp.code, "RESEARCH STUDY - FEBRUARY 2026")
    add_paragraph(doc, "", space_after=0)
    add_paragraph(doc, banner,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, bold=True, color=GREY, space_after=4)
    add_paragraph(doc, lp.cover_title,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=26, bold=True, color=NAVY, space_after=4)
    add_paragraph(doc, lp.cover_subtitle,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, italic=True, color=GREY, space_after=4)
    add_paragraph(doc, lp.cover_blurb,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=11, italic=True, color=GREY, space_after=18)

    table = doc.add_table(rows=1, cols=3)
    table.autofit = True
    for i, line in enumerate(lp.cover_chip_lines):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        set_run(run, size=11, bold=True, color=NAVY)

    add_paragraph(doc, "", space_after=8)
    add_paragraph(doc, "Fabrice Pizzi",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, bold=True, color=NAVY, space_after=2)
    add_paragraph(doc, "Universite Paris-Sorbonne",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, color=GREY, space_after=2)
    master_line = {
        "EN": "Master 2 Economic Intelligence - Intelligence Warfare",
        "FR": "Master 2 Intelligence Economique - Intelligence Warfare",
        "PT-BR": "Mestrado 2 Inteligencia Economica - Intelligence Warfare",
    }.get(lp.code, "Master 2 Economic Intelligence - Intelligence Warfare")
    add_paragraph(doc, master_line,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, italic=True, color=GREY, space_after=14)
    add_paragraph(doc, lp.cover_meta,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, italic=True, color=GREY, space_after=10)
    add_paragraph(doc, f"{lp.cover_keywords_label}: {lp.cover_keywords}",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=9, italic=True, color=GREY, space_after=24)


def add_chapter_header(doc, lp: LangPack):
    """Render the chapter header strip below the cover."""
    add_paragraph(doc, lp.chapter_label,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=11, bold=True, color=GOLD, space_after=2)
    add_paragraph(doc, lp.chapter_title,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=18, bold=True, color=NAVY, space_after=12)
    if lp.chapter_intro:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(lp.chapter_intro)
        set_run(run, size=11, italic=True, color=GREY)


def render_sections(doc, lp: LangPack):
    """Render the body sections from (heading, [paragraph, ...]) tuples."""
    for heading, paragraphs in lp.sections:
        if heading.startswith(("4.1 ", "4.2 ", "4.3 ", "4.4 ", "4.5 ")):
            add_heading(doc, heading, 2)
        else:
            add_heading(doc, heading, 3)
        for para in paragraphs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.25
            run = p.add_run(para)
            set_run(run, size=11, color=RGBColor(0x20, 0x20, 0x20))


def render_tables(doc, lp: LangPack):
    """Render the data tables (caption, source, rows) at the end."""
    for caption, source, rows in lp.table_blocks:
        add_paragraph(doc, caption,
                      size=10, bold=True, color=NAVY, space_after=2)
        add_paragraph(doc, source,
                      size=9, italic=True, color=GREY, space_after=4)
        n_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=n_cols)
        table.style = "Light Grid Accent 1"
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                cell = table.rows[i].cells[j]
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(cell_text)
                set_run(run, size=10,
                        bold=(i == 0),
                        color=NAVY if i == 0 else RGBColor(0x20, 0x20, 0x20))
        doc.add_paragraph()


def render_notes(doc, lp: LangPack):
    """Render the footnotes block at the end of the chapter."""
    add_heading(doc, lp.notes_label, 2)
    for i, note in enumerate(lp.notes, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run_id = p.add_run(f"{i}. ")
        set_run(run_id, size=9, bold=True, color=GOLD)
        run_txt = p.add_run(note)
        set_run(run_txt, size=9, color=GREY)


def render_license(doc, lp: LangPack):
    """Render the license disclaimer block."""
    doc.add_paragraph()
    for line in lp.license_block:
        add_paragraph(doc, line,
                      align=WD_ALIGN_PARAGRAPH.LEFT,
                      size=8, italic=True, color=GREY, space_after=2)
    add_paragraph(doc, lp.page_footer,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=8, italic=True, color=GREY, space_after=0)


def build(lp: LangPack, out_dir: Path) -> Path:
    """Build the .docx file for one language pack."""
    log.info("Building Chapter IV [%s] -> %s", lp.code, lp.filename)
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    add_cover(doc, lp)
    add_chapter_header(doc, lp)
    render_sections(doc, lp)
    render_tables(doc, lp)
    render_notes(doc, lp)
    render_license(doc, lp)

    out = out_dir / lp.filename
    doc.save(out)
    log.info("Saved %s", out)
    return out


# ===========================================================================
# Content - English
# ===========================================================================

EN = LangPack(
    code="EN",
    filename="Chapter_IV_Mechanisms_Competitive_Advantage_EN.docx",
    cover_title="AI FOR AMERICANS FIRST",
    cover_subtitle="AI Protectionism, Energy and Semiconductors: US/Europe Divergence Trajectories 2024-2030",
    cover_blurb="Integrated Geostrategic and Economic Analysis - Chapter IV",
    cover_chip_lines=[
        f"{fmt_en(us_share, 1)}% global operational AI compute = USA",
        "1.59x energy cost EU/US",
        f"{fmt_en(us_eu_caci, 2)}:1 CACI ratio US/EU (Power Mode)",
    ],
    cover_meta="Paris - February 2026  |  7 chapters  |  4 prospective scenarios  |  3 geographic zones",
    cover_keywords_label="Keywords",
    cover_keywords=("artificial intelligence, technology protectionism, semiconductors, "
                    "export controls, sovereign compute, AI geopolitics, France, "
                    "United States, China"),
    chapter_label="CHAPTER IV",
    chapter_title="Mechanisms of American Competitive Advantage Through AI",
    chapter_intro=(
        f"Chapter III established the facts: the United States concentrates {fmt_en(us_share, 1)} percent of "
        "operational global AI compute (49.9 percent if planned clusters are included), 45 "
        "percent of data center energy consumption, and controls 70 percent of the European "
        "cloud market through three hyperscalers. This chapter analyses the mechanisms through "
        "which this compute asymmetry translates into measurable competitive advantage. We "
        "identify four transmission channels: training costs as a barrier to entry, cloud "
        "concentration as a vector of dependence, regionally differentiated productivity, and "
        "the capture of innovation rents by first movers."
    ),
    notes_label="Notes",
    license_block=[
        "License and Disclaimer. This work, 'AI for Americans First,' is made available under the terms of the America-First-IA Creative Commons Attribution - NonCommercial - ShareAlike 4.0 International License (CC BY-NC-SA 4.0).",
        "You are free to share and adapt the material for non-commercial purposes, provided you give appropriate credit to Fabrice Pizzi (Universite Paris-Sorbonne) and distribute your contributions under the same license. This document is provided for educational and research purposes only.",
        "Public dashboard: https://mo0ogly.github.io/America-First-IA/dashboard/",
        "Repository: https://github.com/mo0ogly/America-First-IA",
    ],
    page_footer="AI for Americans First - Fabrice Pizzi - Chapter IV",
)

EN.sections = [
    ("4.1 Training Costs as a Barrier to Entry", []),
    ("4.1.1 The Exponential Explosion of Training Costs", [
        "The most documented channel is that of foundation model training costs. Martens (Bruegel, October 2024) establishes that the cost of training a frontier model rose from 1,000 USD in 2017 to nearly 200 million USD in 2024, and could reach several billion by 2030.[1] This exponential growth follows the scaling laws identified by Kaplan et al. (2020): the improvement of cognitive performance in models is subject to constant returns in complementary inputs (parameters, training data, compute capacity). In other words, each marginal gain in performance requires a proportional increase in compute, data, and model size.",
        "Cottier et al. (2024), cited by Martens, decompose training costs into five components: personnel, AI chips, server costs, network interconnection, and energy. They estimate infrastructure costs (hardware plus data center) at ten times the cost of training itself, with a depreciation rate of 140 percent per year (complete depreciation in 8.5 months, matching the renewal pace of chip generations). GPT-4 infrastructure by end-2023 was estimated at approximately 800 million USD. By extrapolation, infrastructure costs could reach 500 billion USD by 2030, replicated across half a dozen hyperscalers.[2]",
    ]),
    ("4.1.2 The US Advantage: Abundant and Subsidised Compute", [
        "In this context, access to cutting-edge compute becomes the discriminating factor. American Big Tech (Microsoft, Meta, Google, Amazon, xAI) collectively invested 320 billion USD in 2025 in AI infrastructure (servers, data centers, chips), up from 230 billion USD in 2024.[3] Meta is developing an AI computing system with 350,000 H100 processors, estimated at over 10 billion USD. As Martens notes, this barrier to entry is completely beyond the reach of public funding; only the largest companies can aspire to it.[4]",
        "For European actors, this asymmetry translates concretely. The cost of training a model comparable to GPT-4 is estimated at approximately 100 million USD in the United States (with direct GPU access, cheap energy, amortisation over a massive installed base). In Europe, the same training would cost 2 to 5 times more, due to: (i) indirect compute access via US cloud (hyperscaler margins), (ii) energy costs that run 1.4 to 1.7 times higher after PPA correction (and 2 to 3 times higher on unadjusted Eurostat industrial tariffs - Chapter III), and (iii) the absence of comparable economies of scale. The FLOP cost gap (CACI metric M2) is estimated between 2.4x and 3.6x once cloud markups, energy and amortisation are stacked, even though the underlying PPA-adjusted energy ratio is only 1.59x.[5]",
    ]),
    ("4.1.3 Forced Co-opetition: European Startups and US Big Tech", [
        "Exponential costs create a mechanism of structural dependence. As analysed by Martens (Bruegel), AI startups - including European ones - are forced to cooperate with American Big Tech to access computing infrastructure and distribution channels. This co-opetition (forced cooperation with a competitor) places startups in a subordinate position: they depend on US platforms to train their models, then find themselves in direct competition with the proprietary models of those same platforms. The UK Competition and Markets Authority (CMA, 2024) and the European Commission have identified these agreements as potentially anti-competitive.[6] Azoulay et al. (2024), cited by Martens, argue that resource sharing is the only path to avoid concentration, unless a fox actor uses openness as a competitive strategy, which Meta has partially done with LLaMA.",
    ]),
    ("4.2 Cloud Concentration as a Vector of Geopolitical Dependence", []),
    ("4.2.1 The European Cloud Market: 70 Percent Under US Control", [
        "The second mechanism is the concentration of the European cloud market. Synergy Research Group (July 2025) establishes that the European cloud infrastructure services market reached 61 billion EUR in 2024 (multiplied by six since 2017) and should exceed 75 billion EUR in 2025 (+24 percent). Amazon (AWS), Microsoft (Azure), and Google (GCP) account for 70 percent of this market. The share of European providers dropped from 29 percent in 2017 to 15 percent in 2022, where it has stagnated since.[7]",
        "Among European providers, SAP and Deutsche Telekom lead with approximately 2 percent market share each, followed by OVHcloud, Telecom Italia, and Orange. John Dinsdale, Synergy chief analyst, observes that US providers invest 10 billion EUR per quarter in European capex, and that the three hyperscalers now operate more than 140 hyperscale data centers in Europe.[8] The fastest growth (140 to 160 percent) concerns services specifically related to generative AI (GPU-as-a-Service, GenAI PaaS), a segment where European players are virtually absent.",
    ]),
    ("4.2.2 From Commercial Lock-in to Geopolitical Lock-in", [
        "This concentration goes beyond a simple commercial issue. The concept of geopolitical vendor lock-in (defined in Chapter I) describes a situation where technical dependence on a foreign provider is compounded by a regulatory risk linked to the sovereign decisions of that provider home country. The CLOUD Act (2018) requires American cloud providers to furnish US authorities with data hosted on their servers, including those located in Europe. In May 2025, Microsoft France chief legal officer acknowledged under oath the inability to guarantee that French citizens data would never be transmitted to American authorities.[9]",
        "In the AI context, geopolitical lock-in operates at three levels. At the infrastructure level (IaaS), AI workloads are deployed on servers physically controlled by entities subject to US jurisdiction. At the platform level (PaaS), training and inference frameworks are integrated into proprietary ecosystems (Azure ML, SageMaker, Vertex AI). At the model level (MaaS - Model-as-a-Service), access to state-of-the-art models (GPT-4, Claude, Gemini) goes through APIs controlled by US companies, with terms of service that can be modified unilaterally. Section 232 protectionism, by differentiating the cost of hardware access, reinforces all three levels simultaneously.",
        "The Chapter I distinction between Physical CACI (F_total) and Sovereign CACI (F_dom) is critical here: while EU operational compute is largely owned by EU operators (OVH, Scaleway, Sesterce, EuroHPC, AI Factories), EU AI workloads run predominantly on US-owned infrastructure. The dependence is not on installed compute, but on used compute, and on the platforms and models layered above it.",
    ]),
    ("4.2.3 The Draghi Report and the European Diagnosis", [
        "The Draghi Report (September 2024), commissioned by the European Commission, constitutes the most comprehensive institutional recognition of this dependence. Draghi observes that European restrictions on data storage and processing create high compliance costs and hinder the creation of large integrated datasets for AI model training. He recommends the creation of a European hyperscale cloud infrastructure, the consolidation of European cloud providers, and massive investment in AI, quantum computing, and EuroHPC supercomputers.[10] Martens (Bruegel, 2025) moderates this optimism, however, observing that EuroHPC supercomputers are not designed for generative AI and that upgrading to competitive standards exceeds the financial capacity of European budgets.[11]",
    ]),
    ("4.3 Differentiated Productivity: The Abundant Compute Dividend", []),
    ("4.3.1 The Theoretical Macroeconomic Potential", [
        "The third transmission channel is AI impact on labor productivity and its regional asymmetry. The McKinsey Global Institute (May 2024, December 2025) estimates that generative AI could enable Europe to reach an annual productivity growth rate of 3 percent per year by 2030 in an accelerated adoption scenario with proactive workforce redeployment.[12] This figure is sufficient to close most of the productivity gap with the United States. However, in a slow adoption scenario, productivity growth would be only 0.3 percent, close to current levels and insufficient to fund the European social model.",
        "The IMF (Working Paper, March 2025) provides a more granular analysis.[13] By cross-referencing task exposure indices to AI (Felten et al., Eloundou et al.) with sectoral data, the IMF estimates total factor productivity (TFP) gains significantly higher in advanced economies than in middle-income countries, due to their greater share of value added in high AI-exposure sectors (financial services, consulting, technology) and their higher wage levels, which further justify automation. Underlying microeconomic studies (randomized controlled trials) show gains ranging from 14 percent for low-skilled tasks to over 50 percent for software engineers.",
    ]),
    ("4.3.2 The Gap Conditioned by Compute Access", [
        "Yet, the realisation of these productivity gains is conditioned by compute access. This is where the asymmetry diagnosed in Chapter III produces its effects. The French AI Commission (2024) estimates AI potential impact on French growth at 1.3 percentage points per year, by analogy with the effects of electricity, but this estimate assumes unconstrained access to cutting-edge compute. Similarly, McKinsey (December 2025) conditions the optimistic European scenario on accelerated adoption which implies massive infrastructure investments.[14]",
        "The investment gap is considerable. McKinsey (January 2026) estimates that large European companies face an annual deficit of 700 billion USD in R&D and capital expenditure compared to their American counterparts. In digital technologies alone, US companies invested 2 trillion EUR more than European companies over the 2021-2025 period.[15] The annualised gap between January 2024 and September 2025 amounts to 580 billion EUR for corporate investment and 300 billion EUR for startups and scale-ups.",
        f"This investment deficit, combined with the compute asymmetry and energy costs, produces a conditional AI productivity gap. In our modelling, we distinguish between theoretical potential (identical for US and EU with comparable sectoral structures) and achievable potential, which depends on effective compute access. With a CACI(US)/CACI(EU) of {fmt_en(us_eu_caci, 2)}:1 in Power Mode (Chapter III) on the April 2026 dashboard snapshot, the achievable European potential is structurally inferior. Translated into normalised CACI scores, the EU sits at 28.9 against a US benchmark of 100, France at 25.3, Germany at 5.4 - the gap is therefore not 7 to 12 times inferior as earlier analyses suggested, but is precisely {fmt_en(us_eu_caci, 2)} to 1 in compute-adjusted competitive intensity, with sharp variation across member states.",
    ]),
    ("4.4 Innovation Rent Capture: A Self-Reinforcing Advantage", []),
    ("4.4.1 The First-Mover Advantage Mechanism in AI", [
        "The fourth channel is the most structurally significant in the long run. The theory of general purpose technologies (GPT, Bresnahan and Trajtenberg, 1995) predicts that early adopters of a general technology capture disproportionate innovation rents, which become self-reinforcing through three mechanisms.",
        "First, data scale effects. Companies that deploy AI first accumulate usage data (feedback loops, fine-tuning) that improves the performance of their models. This data is non-rival but exclusive: once captured by an actor, it does not benefit competitors. McKinsey estimates that 75 percent of generative AI economic potential is concentrated in four key functions (customer service, software development, marketing, R&D), where usage data is particularly determinative.[16]",
        "Second, platform network effects. US hyperscalers benefit from a virtuous circle: more clients leads to more revenue, which leads to more infrastructure investment, which leads to a better offering, which leads to more clients. This mechanism explains why, despite discussions on digital sovereignty, European providers share stagnates at 15 percent: the service quality gap (latency, range of services, technical support, integration ecosystem) is too large to be bridged by sovereignty arguments alone.",
        "Third, talent capture. US companies offer compensation packages including salaries and equity that attract the best AI researchers and engineers worldwide, including Europeans. Martens notes that qualified AI personnel are scarce and that companies compete fiercely to recruit them, with compensation packages difficult for European institutions to match. This AI brain drain weakens the European CACI L(r) factor and reinforces the American advantage in human capital.[17]",
    ]),
    ("4.4.2 The Risk of Irreversibility", [
        "These three mechanisms produce a potentially irreversible cumulative effect. As the OECD analyses (Martens, 2021) by analogy with previous technological waves (broadband internet, cloud), latecomers in the adoption of a general technology become structurally dependent on foreign platforms, suffer compressed margins, and progressively lose their strategic autonomy. The 2022 semiconductor shortage, which cost the European automotive sector alone 100 billion EUR, illustrates the concrete consequences of unanticipated technological dependence.[18]",
        "The trajectory is all the more concerning as the EU is absent from four of the eight segments of the generative AI value chain as mapped by McKinsey (2024): AI-design semiconductors (dominated by Nvidia), AI cloud platforms (AWS, Azure, GCP), foundation models (OpenAI, Google, Anthropic, Meta), and AI development tools. It is competitive only in downstream application segments: sectoral applications (SAP, Siemens), industrial integration, and specialised semiconductors (power semiconductors: Infineon, STMicroelectronics, NXP, with about 15 percent global market share).[19]",
    ]),
    ("4.5 Synthesis: The Mechanics of US Competitive Advantage", [
        f"The four channels identified form a reinforcing system. The compute asymmetry (Chapter III) feeds differentiated training costs (4.1), which strengthen dependence on US cloud (4.2), which constrains achievable productivity (4.3), which slows European investment and reinforces rent capture by American players (4.4), which in turn widens the compute gap. With the April 2026 snapshot of {fmt_en(us_share, 1)} percent operational compute share, 1.59x PPA-adjusted energy cost and a CACI Power Mode ratio of {fmt_en(us_eu_caci, 2)}:1, this circle constitutes what Farrell and Newman (2019) would describe as structural interdependence weaponisation: European dependence on American infrastructure, initially founded on economic efficiency, becomes a lever of geopolitical power when instrumentalised by measures such as Section 232.",
        "Trump protectionism (Chapter III, phase 4) did not create this asymmetry; it pre-existed massively. It institutionalises and legalises it through a tariff mechanism that differentiates the cost of compute access by nationality. In doing so, it transforms a de facto advantage into a de jure advantage, whose dismantling is politically and legally far more difficult. Chapter V examines how these mechanisms could evolve across four prospective scenarios.",
    ]),
]

EN.table_blocks = [
    ("Table 8. European cloud market: US hyperscaler domination.",
     "Source: Synergy Research Group (July 2025).",
     [
         ["Indicator", "2017", "2022", "2024", "Trend"],
         ["EU Cloud Market (billion EUR)", "~10", "~35", "61", "x6 in 7 years"],
         ["AWS+Azure+GCP Share", "~50 pct", "~67 pct", "70 pct", "Growing"],
         ["EU Provider Share", "29 pct", "15 pct", "15 pct", "Stable (floor)"],
         ["US Capex in Europe", "n.a.", "n.a.", "~40 bn EUR/yr", "10 bn EUR/quarter"],
         ["Hyperscale DCs in EU", "~40", "~80", ">140", "+75 pct in 2 years"],
     ]),
    ("Table 9. AI productivity and competitive advantage: US vs EU (April 2026 calibration).",
     "Sources: McKinsey (2024, 2025, 2026), Bruegel/Martens (2024), CACI calibration on the public dashboard.",
     [
         ["Dimension", "United States", "Europe (EU)", "Gap"],
         ["Potential AI productivity (McKinsey, accel. scen.)", "+2.5-3.5 pct/yr", "+2.5-3.0 pct/yr", "Comparable"],
         ["Achievable AI productivity (under compute constraints)", "+2.5-3.0 pct/yr", "+0.8-1.5 pct/yr", "-1.5 to -2 pts"],
         ["Annual AI investment (corporate + VC)", "~450 bn USD", "~50-70 bn EUR", "~7:1"],
         ["FLOP cost (USD/TFlop, training)", "~0.5", "~1.2-1.8", "2.4-3.6x"],
         ["Normalised CACI Power Mode (USA = 100)", "100", "28.9 (EU(13))", f"{fmt_en(us_eu_caci, 2)}:1"],
     ]),
]

EN.notes = [
    "Martens, B. (2024), 'The Tension Between Exploding AI Investment Costs and Slow Productivity Growth,' Working Paper 18/2024, Bruegel, pp. 5-8. The time series is calibrated on Epoch AI training cost data for frontier models.",
    "Cottier, B. et al. (2024), cited in Martens (2024), op. cit. The 10x factor between infrastructure and training costs is a central estimate, subject to variation based on post-training infrastructure utilisation.",
    "IEA (2025), Energy and AI, Paris, p. 42. The IEA compiles investment announcements from Meta, Amazon, Alphabet, Microsoft and xAI for 2025.",
    "Martens, B. (2024), 'Why Artificial Intelligence Is Creating Fundamental Challenges for Competition Policy,' Policy Brief 16/2024, Bruegel. Meta plans a system of 350,000 H100s (at about 30,000 USD per unit), i.e. over 10 billion USD in hardware alone.",
    "Author estimate based on Bruegel (training costs), Eurostat (industrial electricity tariffs), EIA (US prices), the public dashboard PPA-adjusted reference values (USA 85, EU 135 USD/MWh), and Epoch AI (GPU performance/price). The 2.4 to 3.6x range stacks cloud markup (1.5 to 2x) on top of the 1.59x PPA-adjusted energy ratio and the absence of EU economies of scale; it depends on access mode (cloud vs on-premise) and PPA negotiations.",
    "Competition and Markets Authority (2024), 'AI Foundation Models - Update Report,' London. The European Commission launched a call for contributions on competition in generative AI in 2024.",
    "Synergy Research Group (July 2025), 'European Cloud Providers Local Market Share Now Holds Steady at 15 percent.' The market is defined as IaaS plus PaaS plus hosted private cloud.",
    "Dinsdale, J. (2025), cited in Synergy Research Group, op. cit.: US cloud providers continue to invest 10 billion EUR every quarter in European capex, most of which comes from the big three.",
    "Data Center Dynamics (July 2025), 'European Cloud Providers Hold 15 percent of Local Market Share.' The article reports the testimony of Anton Carniax, Microsoft France chief legal officer, before a parliamentary committee.",
    "Draghi, M. (September 2024), The Future of European Competitiveness, report commissioned by the European Commission. Cited in Martens, B. (2025), 'Catch-Up with the US or Prosper Below the Tech Frontier? An EU Artificial Intelligence Strategy,' Bruegel Policy Brief.",
    "Martens, B. (2025), op. cit. The author notes that the EuroHPC approach, centred on hardware, has characterised European digital policies for decades and is particularly poorly oriented regarding AI.",
    "McKinsey Global Institute (May 2024), 'A New Future of Work: The Race to Deploy AI and Raise Skills in Europe and Beyond.' The 3 percent figure assumes an accelerated adoption scenario with full workforce redeployment.",
    "IMF (March 2025), 'Artificial Intelligence and Productivity in Europe,' Working Paper WP/25/067. The analysis uses Felten et al. (2021) and Eloundou et al. (2024) exposure indices.",
    "McKinsey (December 2025), 'Accelerating Europe AI Adoption: The Role of Sovereign AI Capabilities.' The report estimates European labor productivity growth at 0.2 percent on average over 2020-2025, versus a potential of 3.1 percent with full AI adoption.",
    "McKinsey (January 2026), 'Transforming Europe: Bold Moves to Lift a Continent.' The 700 bn USD per year gap is calculated on the 150 largest technology companies (US vs EU) in capex plus R&D (S&P Capital IQ data).",
    "McKinsey Global Institute (2024), 'The Economic Potential of Generative AI: The Next Productivity Frontier.' The four functions (customer service, software development, marketing/sales, R&D) represent about 75 percent of the identified economic potential across 63 use cases.",
    "Martens, B. (2024), Working Paper 18/2024, op. cit. Personnel costs (salaries plus equity) are often the most significant component of AI model development costs.",
    "McKinsey (October 2024), 'Time to Place Our Bets: Europe AI Opportunity.' The 100 bn EUR GDP loss figure for the European automotive industry comes from Allianz (September 2022) and OICA.",
    "McKinsey (October 2024), op. cit. European semiconductors (Infineon, STMicro, NXP) represent about 15 percent of the global integrated design-manufacturing market (Omdia/Informa data, 2024), but primarily in power and automotive segments, not AI GPU/ASIC.",
]


# ===========================================================================
# Content - French
# ===========================================================================

FR = LangPack(
    code="FR",
    filename="Chapitre_IV_Mecanismes_Avantage_Concurrentiel_FR.docx",
    cover_title="AI FOR AMERICANS FIRST",
    cover_subtitle="Protectionnisme IA, Energie et Semi-conducteurs : Trajectoires de divergence US/Europe 2024-2030",
    cover_blurb="Analyse geostrategique et economique integree - Chapitre IV",
    cover_chip_lines=[
        f"{fmt_fr(us_share, 1)} pct du compute IA operationnel mondial = USA",
        "1,59x cout energie EU/US",
        f"{fmt_fr(us_eu_caci, 2)}:1 ratio CACI US/EU (Power Mode)",
    ],
    cover_meta="Paris - fevrier 2026  |  7 chapitres  |  4 scenarios prospectifs  |  3 zones geographiques",
    cover_keywords_label="Mots-cles",
    cover_keywords=("intelligence artificielle, protectionnisme technologique, semi-conducteurs, "
                    "controles a l'exportation, compute souverain, geopolitique IA, France, "
                    "Etats-Unis, Chine"),
    chapter_label="CHAPITRE IV",
    chapter_title="Mecanismes de l'avantage concurrentiel americain par l'IA",
    chapter_intro=(
        f"Le chapitre III a etabli les faits : les Etats-Unis concentrent {fmt_fr(us_share, 1)} pour cent du "
        "compute IA operationnel mondial (49,9 pour cent en incluant les clusters planifies), "
        "45 pour cent de la consommation electrique des centres de donnees, et controlent 70 "
        "pour cent du marche cloud europeen via trois hyperscalers. Ce chapitre analyse les "
        "mecanismes par lesquels cette asymetrie de compute se traduit en avantage competitif "
        "mesurable. Nous identifions quatre canaux de transmission : les couts d'entrainement "
        "comme barriere a l'entree, la concentration cloud comme vecteur de dependance, la "
        "productivite differenciee selon les regions, et la captation des rentes d'innovation "
        "par les premiers entrants."
    ),
    notes_label="Notes",
    license_block=[
        "Licence et avertissement. Cette oeuvre, 'AI for Americans First', est mise a disposition selon les termes de la licence Creative Commons Attribution - Pas d'Utilisation Commerciale - Partage dans les Memes Conditions 4.0 International (CC BY-NC-SA 4.0) du projet America-First-IA.",
        "Vous etes libre de partager et adapter le materiel a des fins non commerciales, a condition d'attribuer correctement l'oeuvre a Fabrice Pizzi (Universite Paris-Sorbonne) et de distribuer vos contributions sous la meme licence. Ce document est fourni a des fins educatives et de recherche uniquement.",
        "Tableau de bord public : https://mo0ogly.github.io/America-First-IA/dashboard/",
        "Depot : https://github.com/mo0ogly/America-First-IA",
    ],
    page_footer="AI for Americans First - Fabrice Pizzi - Chapitre IV",
)

FR.sections = [
    ("4.1 Les couts d'entrainement comme barriere a l'entree", []),
    ("4.1.1 L'explosion exponentielle des couts d'entrainement", [
        "Le canal le mieux documente est celui des couts d'entrainement des modeles de fondation. Martens (Bruegel, octobre 2024) etablit que le cout d'entrainement d'un modele de frontiere est passe de 1 000 USD en 2017 a pres de 200 millions USD en 2024, et pourrait atteindre plusieurs milliards d'ici 2030.[1] Cette croissance exponentielle suit les scaling laws identifiees par Kaplan et al. (2020) : l'amelioration de la performance cognitive des modeles est soumise a des rendements constants en intrants complementaires (parametres, donnees d'entrainement, capacite de compute). Autrement dit, chaque gain marginal de performance exige une augmentation proportionnelle du compute, des donnees et de la taille du modele.",
        "Cottier et al. (2024), cites par Martens, decomposent les couts d'entrainement en cinq composantes : personnel, puces IA, cout des serveurs, interconnexion reseau et energie. Ils estiment les couts d'infrastructure (materiel plus centre de donnees) a dix fois le cout de l'entrainement lui-meme, avec un taux d'amortissement de 140 pour cent par an (amortissement complet en 8,5 mois, correspondant au rythme de renouvellement des generations de puces). L'infrastructure de GPT-4 fin 2023 etait estimee a environ 800 millions USD. Par extrapolation, les couts d'infrastructure pourraient atteindre 500 milliards USD d'ici 2030, repliques sur une demi-douzaine d'hyperscalers.[2]",
    ]),
    ("4.1.2 L'avantage US : compute abondant et subventionne", [
        "Dans ce contexte, l'acces au compute de pointe devient le facteur discriminant. Les Big Tech americaines (Microsoft, Meta, Google, Amazon, xAI) ont collectivement investi 320 milliards USD en 2025 dans l'infrastructure IA (serveurs, centres de donnees, puces), contre 230 milliards USD en 2024.[3] Meta developpe un systeme de calcul IA dote de 350 000 processeurs H100, estime a plus de 10 milliards USD. Comme le note Martens, cette barriere a l'entree est totalement hors de portee du financement public ; seules les plus grandes entreprises peuvent y aspirer.[4]",
        "Pour les acteurs europeens, cette asymetrie se traduit concretement. Le cout d'entrainement d'un modele comparable a GPT-4 est estime a environ 100 millions USD aux Etats-Unis (avec acces direct au GPU, energie peu chere, amortissement sur une base installee massive). En Europe, le meme entrainement couterait 2 a 5 fois plus, en raison de : (i) un acces indirect au compute via le cloud US (marges des hyperscalers), (ii) des couts energetiques 1,4 a 1,7 fois superieurs apres correction PPA (et 2 a 3 fois superieurs sur les tarifs Eurostat industriels non ajustes - chapitre III), et (iii) l'absence d'economies d'echelle comparables. L'ecart de cout du FLOP (metrique CACI M2) est estime entre 2,4x et 3,6x une fois cumules les marges cloud, l'energie et l'amortissement, alors meme que le ratio energetique sous-jacent ajuste-PPA n'est que de 1,59x.[5]",
    ]),
    ("4.1.3 Co-opetition forcee : startups europeennes et Big Tech US", [
        "Les couts exponentiels creent un mecanisme de dependance structurelle. Comme l'analyse Martens (Bruegel), les startups IA - y compris europeennes - sont contraintes de cooperer avec les Big Tech americaines pour acceder a l'infrastructure de calcul et aux canaux de distribution. Cette co-opetition (cooperation forcee avec un concurrent) place les startups dans une position subordonnee : elles dependent des plateformes US pour entrainer leurs modeles, puis se retrouvent en concurrence directe avec les modeles proprietaires de ces memes plateformes. La Competition and Markets Authority britannique (CMA, 2024) et la Commission europeenne ont identifie ces accords comme potentiellement anti-concurrentiels.[6] Azoulay et al. (2024), cites par Martens, soutiennent que le partage de ressources est la seule voie pour eviter la concentration, sauf si un acteur fox utilise l'ouverture comme strategie competitive - ce que Meta a partiellement fait avec LLaMA.",
    ]),
    ("4.2 La concentration cloud comme vecteur de dependance geopolitique", []),
    ("4.2.1 Le marche cloud europeen : 70 pour cent sous controle US", [
        "Le second mecanisme est la concentration du marche cloud europeen. Synergy Research Group (juillet 2025) etablit que le marche europeen des services d'infrastructure cloud a atteint 61 milliards EUR en 2024 (multiplie par six depuis 2017) et devrait depasser 75 milliards EUR en 2025 (+24 pour cent). Amazon (AWS), Microsoft (Azure) et Google (GCP) representent 70 pour cent de ce marche. La part des fournisseurs europeens est passee de 29 pour cent en 2017 a 15 pour cent en 2022, ou elle stagne depuis.[7]",
        "Parmi les fournisseurs europeens, SAP et Deutsche Telekom sont en tete avec environ 2 pour cent de part de marche chacun, suivis par OVHcloud, Telecom Italia et Orange. John Dinsdale, analyste en chef de Synergy, observe que les fournisseurs US investissent 10 milliards EUR par trimestre en capex europeen, et que les trois hyperscalers exploitent desormais plus de 140 centres de donnees hyperscale en Europe.[8] La croissance la plus rapide (140 a 160 pour cent) concerne les services specifiquement lies a l'IA generative (GPU-as-a-Service, GenAI PaaS), un segment ou les acteurs europeens sont quasi absents.",
    ]),
    ("4.2.2 Du verrouillage commercial au verrouillage geopolitique", [
        "Cette concentration depasse le simple enjeu commercial. Le concept de geopolitical vendor lock-in (defini au chapitre I) decrit une situation ou la dependance technique a un fournisseur etranger est aggravee par un risque reglementaire lie aux decisions souveraines du pays d'origine du fournisseur. Le CLOUD Act (2018) oblige les fournisseurs cloud americains a fournir aux autorites US les donnees hebergees sur leurs serveurs, y compris ceux situes en Europe. En mai 2025, la directrice juridique de Microsoft France a reconnu sous serment l'incapacite de garantir que les donnees des citoyens francais ne seraient jamais transmises aux autorites americaines.[9]",
        "Dans le contexte IA, le verrouillage geopolitique opere a trois niveaux. Au niveau infrastructure (IaaS), les charges IA sont deployees sur des serveurs physiquement controles par des entites soumises a la juridiction americaine. Au niveau plateforme (PaaS), les frameworks d'entrainement et d'inference sont integres dans des ecosystemes proprietaires (Azure ML, SageMaker, Vertex AI). Au niveau modele (MaaS - Model-as-a-Service), l'acces aux modeles state-of-the-art (GPT-4, Claude, Gemini) passe par des API controlees par des entreprises US, avec des conditions de service modifiables unilateralement. Le protectionnisme Section 232, en differenciant le cout de l'acces materiel, renforce simultanement les trois niveaux.",
        "La distinction du chapitre I entre CACI Physique (F_total) et CACI Souverain (F_dom) est ici essentielle : alors que le compute UE operationnel est largement detenu par des operateurs UE (OVH, Scaleway, Sesterce, EuroHPC, AI Factories), les charges IA UE tournent majoritairement sur infrastructure detenue par les US. La dependance ne porte pas sur le compute installe, mais sur le compute utilise, et sur les couches plateformes et modeles superposees au-dessus.",
    ]),
    ("4.2.3 Le rapport Draghi et le diagnostic europeen", [
        "Le rapport Draghi (septembre 2024), commande par la Commission europeenne, constitue la reconnaissance institutionnelle la plus complete de cette dependance. Draghi observe que les restrictions europeennes sur le stockage et le traitement des donnees creent des couts de conformite eleves et entravent la creation de grands jeux de donnees integres pour l'entrainement de modeles IA. Il recommande la creation d'une infrastructure cloud hyperscale europeenne, la consolidation des fournisseurs cloud europeens et un investissement massif dans l'IA, l'informatique quantique et les supercalculateurs EuroHPC.[10] Martens (Bruegel, 2025) modere toutefois cet optimisme, observant que les supercalculateurs EuroHPC ne sont pas concus pour l'IA generative et que la mise a niveau aux standards competitifs depasse la capacite financiere des budgets europeens.[11]",
    ]),
    ("4.3 Productivite differenciee : le dividende du compute abondant", []),
    ("4.3.1 Le potentiel macroeconomique theorique", [
        "Le troisieme canal de transmission est l'impact de l'IA sur la productivite du travail et son asymetrie regionale. Le McKinsey Global Institute (mai 2024, decembre 2025) estime que l'IA generative pourrait permettre a l'Europe d'atteindre un taux de croissance annuel de la productivite de 3 pour cent d'ici 2030 dans un scenario d'adoption acceleree avec redeploiement proactif de la main-d'oeuvre.[12] Ce chiffre suffit a combler l'essentiel de l'ecart de productivite avec les Etats-Unis. Toutefois, dans un scenario d'adoption lente, la croissance de la productivite ne serait que de 0,3 pour cent - proche des niveaux actuels et insuffisante pour financer le modele social europeen.",
        "Le FMI (Working Paper, mars 2025) fournit une analyse plus granulaire.[13] En croisant les indices d'exposition des taches a l'IA (Felten et al., Eloundou et al.) avec les donnees sectorielles, le FMI estime des gains de productivite totale des facteurs (TFP) significativement plus eleves dans les economies avancees que dans les pays a revenu intermediaire, en raison de leur part plus elevee de valeur ajoutee dans les secteurs a forte exposition IA (services financiers, conseil, technologie) et de leurs niveaux salariaux plus eleves, qui justifient davantage l'automatisation. Les etudes microeconomiques sous-jacentes (essais randomises controles) montrent des gains allant de 14 pour cent pour les taches peu qualifiees a plus de 50 pour cent pour les ingenieurs logiciels.",
    ]),
    ("4.3.2 L'ecart conditionne par l'acces au compute", [
        "Pourtant, la realisation de ces gains de productivite est conditionnee par l'acces au compute. C'est ici que l'asymetrie diagnostiquee au chapitre III produit ses effets. La Commission IA francaise (2024) estime l'impact potentiel de l'IA sur la croissance francaise a 1,3 point de pourcentage par an, par analogie avec les effets de l'electricite - mais cette estimation suppose un acces non contraint au compute de pointe. De meme, McKinsey (decembre 2025) conditionne le scenario optimiste europeen a une adoption acceleree qui implique des investissements massifs en infrastructure.[14]",
        "L'ecart d'investissement est considerable. McKinsey (janvier 2026) estime que les grandes entreprises europeennes font face a un deficit annuel de 700 milliards USD en R&D et depenses en capital par rapport a leurs homologues americaines. Dans les seules technologies numeriques, les entreprises americaines ont investi 2 trillions EUR de plus que les entreprises europeennes sur la periode 2021-2025.[15] L'ecart annualise entre janvier 2024 et septembre 2025 s'eleve a 580 milliards EUR pour l'investissement corporate et 300 milliards EUR pour les startups et scale-ups.",
        f"Ce deficit d'investissement, combine a l'asymetrie de compute et aux couts energetiques, produit un ecart de productivite IA conditionnel. Dans notre modelisation, nous distinguons le potentiel theorique (identique pour US et UE avec des structures sectorielles comparables) et le potentiel realisable, qui depend de l'acces effectif au compute. Avec un CACI(US)/CACI(UE) de {fmt_fr(us_eu_caci, 2)}:1 en Power Mode (chapitre III) sur le snapshot du tableau de bord d'avril 2026, le potentiel realisable europeen est structurellement inferieur. Traduit en scores CACI normalises, l'UE se situe a 28,9 contre un benchmark US de 100, la France a 25,3, l'Allemagne a 5,4 - l'ecart n'est donc pas de 7 a 12 fois inferieur comme le suggeraient des analyses anterieures, mais precisement de 3,46 pour 1 en intensite competitive ajustee au compute, avec une variation forte entre Etats membres.",
    ]),
    ("4.4 Captation des rentes d'innovation : un avantage auto-renforcant", []),
    ("4.4.1 Le mecanisme du first-mover advantage en IA", [
        "Le quatrieme canal est le plus structurellement significatif a long terme. La theorie des general purpose technologies (GPT, Bresnahan et Trajtenberg, 1995) predit que les early adopters d'une technologie generale captent des rentes d'innovation disproportionnees, qui deviennent auto-renforcantes via trois mecanismes.",
        "Premierement, les effets d'echelle des donnees. Les entreprises qui deploient l'IA en premier accumulent des donnees d'usage (boucles de retroaction, fine-tuning) qui ameliorent la performance de leurs modeles. Ces donnees sont non-rivales mais exclusives : une fois capturees par un acteur, elles ne profitent pas aux concurrents. McKinsey estime que 75 pour cent du potentiel economique de l'IA generative est concentre sur quatre fonctions cles (service client, developpement logiciel, marketing, R&D), ou les donnees d'usage sont particulierement determinantes.[16]",
        "Deuxiemement, les effets de reseau de plateforme. Les hyperscalers US beneficient d'un cercle vertueux : plus de clients menent a plus de revenus, qui menent a plus d'investissement en infrastructure, qui menent a une meilleure offre, qui menent a plus de clients. Ce mecanisme explique pourquoi, malgre les discussions sur la souverainete numerique, la part des fournisseurs europeens stagne a 15 pour cent : l'ecart de qualite de service (latence, gamme de services, support technique, ecosysteme d'integration) est trop important pour etre comble par les seuls arguments de souverainete.",
        "Troisiemement, la captation des talents. Les entreprises americaines offrent des packages de remuneration combinant salaires et equity qui attirent les meilleurs chercheurs et ingenieurs IA mondiaux, y compris europeens. Martens note que le personnel IA qualifie est rare et que les entreprises se livrent a une concurrence feroce pour le recruter, avec des packages de remuneration difficiles a egaler pour les institutions europeennes. Ce brain drain IA affaiblit le facteur L(r) du CACI europeen et renforce l'avantage americain en capital humain.[17]",
    ]),
    ("4.4.2 Le risque d'irreversibilite", [
        "Ces trois mecanismes produisent un effet cumulatif potentiellement irreversible. Comme l'analyse l'OCDE (Martens, 2021) par analogie avec les vagues technologiques precedentes (internet haut debit, cloud), les retardataires dans l'adoption d'une technologie generale deviennent structurellement dependants des plateformes etrangeres, subissent des marges comprimees et perdent progressivement leur autonomie strategique. La penurie de semi-conducteurs de 2022, qui a coute 100 milliards EUR au seul secteur automobile europeen, illustre les consequences concretes d'une dependance technologique non anticipee.[18]",
        "La trajectoire est d'autant plus preoccupante que l'UE est absente de quatre des huit segments de la chaine de valeur de l'IA generative tels que cartographies par McKinsey (2024) : semi-conducteurs IA-design (domines par Nvidia), plateformes cloud IA (AWS, Azure, GCP), modeles de fondation (OpenAI, Google, Anthropic, Meta) et outils de developpement IA. Elle n'est competitive que sur les segments d'application aval : applications sectorielles (SAP, Siemens), integration industrielle, et semi-conducteurs specialises (semi-conducteurs de puissance : Infineon, STMicroelectronics, NXP, avec environ 15 pour cent de part de marche mondiale).[19]",
    ]),
    ("4.5 Synthese : la mecanique de l'avantage competitif US", [
        f"Les quatre canaux identifies forment un systeme renforcant. L'asymetrie de compute (chapitre III) alimente des couts d'entrainement differencies (4.1), qui renforcent la dependance au cloud US (4.2), qui contraint la productivite realisable (4.3), qui ralentit l'investissement europeen et renforce la captation des rentes par les acteurs americains (4.4), qui a son tour creuse l'ecart de compute. Avec le snapshot d'avril 2026 - {fmt_fr(us_share, 1)} pour cent de part de compute operationnel, 1,59x de cout energetique ajuste-PPA et un ratio CACI Power Mode de {fmt_fr(us_eu_caci, 2)}:1 - ce cercle constitue ce que Farrell et Newman (2019) decriraient comme une weaponisation de l'interdependance structurelle : la dependance europeenne a l'infrastructure americaine, initialement fondee sur l'efficacite economique, devient un levier de pouvoir geopolitique lorsqu'elle est instrumentalisee par des mesures comme la Section 232.",
        "Le protectionnisme Trump (chapitre III, phase 4) n'a pas cree cette asymetrie, elle preexistait massivement. Il l'institutionnalise et la legalise par un mecanisme tarifaire qui differencie le cout d'acces au compute selon la nationalite. Ce faisant, il transforme un avantage de fait en avantage de droit, dont le demantelement est politiquement et juridiquement bien plus difficile. Le chapitre V examine comment ces mecanismes pourraient evoluer a travers quatre scenarios prospectifs.",
    ]),
]

FR.table_blocks = [
    ("Tableau 8. Marche cloud europeen : domination des hyperscalers US.",
     "Source : Synergy Research Group (juillet 2025).",
     [
         ["Indicateur", "2017", "2022", "2024", "Tendance"],
         ["Marche cloud UE (milliards EUR)", "~10", "~35", "61", "x6 en 7 ans"],
         ["Part AWS+Azure+GCP", "~50 pct", "~67 pct", "70 pct", "Croissante"],
         ["Part fournisseurs UE", "29 pct", "15 pct", "15 pct", "Stable (plancher)"],
         ["Capex US en Europe", "n.d.", "n.d.", "~40 mds EUR/an", "10 mds EUR/trim."],
         ["Centres hyperscale en UE", "~40", "~80", ">140", "+75 pct en 2 ans"],
     ]),
    ("Tableau 9. Productivite IA et avantage competitif : US vs UE (calibration avril 2026).",
     "Sources : McKinsey (2024, 2025, 2026), Bruegel/Martens (2024), calibration CACI sur le tableau de bord public.",
     [
         ["Dimension", "Etats-Unis", "Europe (UE)", "Ecart"],
         ["Productivite IA potentielle (McKinsey, scen. accel.)", "+2,5-3,5 pct/an", "+2,5-3,0 pct/an", "Comparable"],
         ["Productivite IA realisable (sous contrainte compute)", "+2,5-3,0 pct/an", "+0,8-1,5 pct/an", "-1,5 a -2 pts"],
         ["Investissement annuel IA (corporate + VC)", "~450 mds USD", "~50-70 mds EUR", "~7:1"],
         ["Cout du FLOP (USD/TFlop, entrainement)", "~0,5", "~1,2-1,8", "2,4-3,6x"],
         ["CACI normalise Power Mode (USA = 100)", "100", "28,9 (UE(13))", f"{fmt_fr(us_eu_caci, 2)}:1"],
     ]),
]

FR.notes = EN.notes


# ===========================================================================
# Content - Brazilian Portuguese
# ===========================================================================

PT = LangPack(
    code="PT-BR",
    filename="Capitulo_IV_Mecanismos_Vantagem_Competitiva_PT-BR.docx",
    cover_title="AI FOR AMERICANS FIRST",
    cover_subtitle="Protecionismo de IA, Energia e Semicondutores: Trajetorias de Divergencia EUA/Europa 2024-2030",
    cover_blurb="Analise geoestrategica e economica integrada - Capitulo IV",
    cover_chip_lines=[
        f"{fmt_fr(us_share, 1)} pct do compute IA operacional mundial = EUA",
        "1,59x custo de energia UE/EUA",
        f"{fmt_fr(us_eu_caci, 2)}:1 razao CACI EUA/UE (Power Mode)",
    ],
    cover_meta="Paris - fevereiro de 2026  |  7 capitulos  |  4 cenarios prospectivos  |  3 zonas geograficas",
    cover_keywords_label="Palavras-chave",
    cover_keywords=("inteligencia artificial, protecionismo tecnologico, semicondutores, "
                    "controles de exportacao, compute soberano, geopolitica da IA, Franca, "
                    "Estados Unidos, China"),
    chapter_label="CAPITULO IV",
    chapter_title="Mecanismos da vantagem competitiva americana via IA",
    chapter_intro=(
        "O Capitulo III estabeleceu os fatos: os Estados Unidos concentram 76,9 por cento do "
        "compute IA operacional mundial (49,9 por cento se incluidos os clusters planejados), "
        "45 por cento do consumo eletrico dos data centers, e controlam 70 por cento do "
        "mercado de nuvem europeu por meio de tres hyperscalers. Este capitulo analisa os "
        "mecanismos pelos quais essa assimetria de compute se traduz em vantagem competitiva "
        "mensuravel. Identificamos quatro canais de transmissao: custos de treinamento como "
        "barreira a entrada, concentracao de nuvem como vetor de dependencia, produtividade "
        "diferenciada por regiao e captura de rendas de inovacao pelos primeiros entrantes."
    ),
    notes_label="Notas",
    license_block=[
        "Licenca e isencao de responsabilidade. Esta obra, 'AI for Americans First', e disponibilizada nos termos da Licenca Creative Commons Atribuicao - NaoComercial - CompartilhaIgual 4.0 Internacional (CC BY-NC-SA 4.0) do projeto America-First-IA.",
        "Voce e livre para compartilhar e adaptar o material para fins nao comerciais, desde que credite adequadamente a obra a Fabrice Pizzi (Universite Paris-Sorbonne) e distribua suas contribuicoes sob a mesma licenca. Este documento e fornecido apenas para fins educacionais e de pesquisa.",
        "Painel publico: https://mo0ogly.github.io/America-First-IA/dashboard/",
        "Repositorio: https://github.com/mo0ogly/America-First-IA",
    ],
    page_footer="AI for Americans First - Fabrice Pizzi - Capitulo IV",
)

PT.sections = [
    ("4.1 Os custos de treinamento como barreira a entrada", []),
    ("4.1.1 A explosao exponencial dos custos de treinamento", [
        "O canal mais documentado e o dos custos de treinamento dos modelos de fundacao. Martens (Bruegel, outubro de 2024) estabelece que o custo de treinamento de um modelo de fronteira passou de 1.000 USD em 2017 para quase 200 milhoes USD em 2024, e poderia atingir varios bilhoes ate 2030.[1] Esse crescimento exponencial segue as scaling laws identificadas por Kaplan et al. (2020): a melhoria do desempenho cognitivo dos modelos esta sujeita a retornos constantes em insumos complementares (parametros, dados de treinamento, capacidade de compute). Em outras palavras, cada ganho marginal de desempenho exige um aumento proporcional em compute, dados e tamanho do modelo.",
        "Cottier et al. (2024), citados por Martens, decompoem os custos de treinamento em cinco componentes: pessoal, chips IA, custo dos servidores, interconexao de rede e energia. Eles estimam os custos de infraestrutura (hardware mais data center) em dez vezes o custo do treinamento em si, com taxa de depreciacao de 140 por cento ao ano (depreciacao completa em 8,5 meses, correspondendo ao ritmo de renovacao das geracoes de chips). A infraestrutura do GPT-4 no fim de 2023 era estimada em aproximadamente 800 milhoes USD. Por extrapolacao, os custos de infraestrutura poderiam atingir 500 bilhoes USD ate 2030, replicados em meia duzia de hyperscalers.[2]",
    ]),
    ("4.1.2 A vantagem dos EUA: compute abundante e subsidiado", [
        "Nesse contexto, o acesso ao compute de ponta se torna o fator discriminante. As Big Tech americanas (Microsoft, Meta, Google, Amazon, xAI) investiram coletivamente 320 bilhoes USD em 2025 em infraestrutura de IA (servidores, data centers, chips), contra 230 bilhoes USD em 2024.[3] A Meta esta desenvolvendo um sistema de computacao IA com 350.000 processadores H100, estimado em mais de 10 bilhoes USD. Como observa Martens, essa barreira a entrada esta totalmente fora do alcance do financiamento publico; apenas as maiores empresas podem aspirar a ela.[4]",
        "Para os atores europeus, essa assimetria se traduz concretamente. O custo de treinamento de um modelo comparavel ao GPT-4 e estimado em aproximadamente 100 milhoes USD nos Estados Unidos (com acesso direto a GPU, energia barata, amortizacao sobre uma base instalada massiva). Na Europa, o mesmo treinamento custaria 2 a 5 vezes mais, devido a: (i) acesso indireto ao compute via nuvem dos EUA (margens dos hyperscalers), (ii) custos energeticos 1,4 a 1,7 vezes mais altos apos correcao PPA (e 2 a 3 vezes mais altos nas tarifas Eurostat industriais nao ajustadas - Capitulo III), e (iii) ausencia de economias de escala comparaveis. A diferenca de custo do FLOP (metrica CACI M2) e estimada entre 2,4x e 3,6x uma vez acumuladas as margens de nuvem, a energia e a amortizacao, mesmo que a razao energetica subjacente ajustada-PPA seja apenas de 1,59x.[5]",
    ]),
    ("4.1.3 Co-opeticao forcada: startups europeias e Big Tech dos EUA", [
        "Os custos exponenciais criam um mecanismo de dependencia estrutural. Como analisa Martens (Bruegel), as startups de IA - incluindo as europeias - sao forcadas a cooperar com as Big Tech americanas para acessar a infraestrutura de computacao e os canais de distribuicao. Essa co-opeticao (cooperacao forcada com um concorrente) coloca as startups em posicao subordinada: dependem das plataformas dos EUA para treinar seus modelos, depois se encontram em concorrencia direta com os modelos proprietarios dessas mesmas plataformas. A Competition and Markets Authority britanica (CMA, 2024) e a Comissao Europeia identificaram esses acordos como potencialmente anticompetitivos.[6] Azoulay et al. (2024), citados por Martens, argumentam que o compartilhamento de recursos e o unico caminho para evitar a concentracao, a menos que um ator fox use a abertura como estrategia competitiva, o que a Meta fez parcialmente com o LLaMA.",
    ]),
    ("4.2 A concentracao da nuvem como vetor de dependencia geopolitica", []),
    ("4.2.1 O mercado de nuvem europeu: 70 por cento sob controle dos EUA", [
        "O segundo mecanismo e a concentracao do mercado de nuvem europeu. A Synergy Research Group (julho de 2025) estabelece que o mercado europeu de servicos de infraestrutura de nuvem atingiu 61 bilhoes EUR em 2024 (multiplicado por seis desde 2017) e devera ultrapassar 75 bilhoes EUR em 2025 (+24 por cento). Amazon (AWS), Microsoft (Azure) e Google (GCP) representam 70 por cento desse mercado. A parcela dos provedores europeus caiu de 29 por cento em 2017 para 15 por cento em 2022, onde estagna desde entao.[7]",
        "Entre os provedores europeus, SAP e Deutsche Telekom lideram com aproximadamente 2 por cento de participacao de mercado cada, seguidos por OVHcloud, Telecom Italia e Orange. John Dinsdale, analista chefe da Synergy, observa que os provedores dos EUA investem 10 bilhoes EUR por trimestre em capex europeu, e que os tres hyperscalers operam agora mais de 140 data centers hyperscale na Europa.[8] O crescimento mais rapido (140 a 160 por cento) diz respeito aos servicos especificamente relacionados a IA generativa (GPU-as-a-Service, GenAI PaaS), um segmento em que os atores europeus estao virtualmente ausentes.",
    ]),
    ("4.2.2 Do bloqueio comercial ao bloqueio geopolitico", [
        "Essa concentracao vai alem de uma simples questao comercial. O conceito de geopolitical vendor lock-in (definido no Capitulo I) descreve uma situacao em que a dependencia tecnica de um provedor estrangeiro e agravada por um risco regulatorio ligado as decisoes soberanas do pais de origem do provedor. O CLOUD Act (2018) exige que os provedores de nuvem americanos forneçam as autoridades dos EUA dados hospedados em seus servidores, incluindo aqueles localizados na Europa. Em maio de 2025, a diretora juridica da Microsoft France reconheceu sob juramento a incapacidade de garantir que os dados dos cidadaos franceses nunca seriam transmitidos as autoridades americanas.[9]",
        "No contexto da IA, o bloqueio geopolitico opera em tres niveis. No nivel de infraestrutura (IaaS), as cargas de IA sao implantadas em servidores fisicamente controlados por entidades sujeitas a jurisdicao dos EUA. No nivel de plataforma (PaaS), os frameworks de treinamento e inferencia estao integrados em ecossistemas proprietarios (Azure ML, SageMaker, Vertex AI). No nivel de modelo (MaaS - Model-as-a-Service), o acesso aos modelos state-of-the-art (GPT-4, Claude, Gemini) passa por APIs controladas por empresas dos EUA, com termos de servico que podem ser modificados unilateralmente. O protecionismo da Secao 232, ao diferenciar o custo de acesso ao hardware, reforca os tres niveis simultaneamente.",
        "A distincao do Capitulo I entre CACI Fisico (F_total) e CACI Soberano (F_dom) e essencial aqui: enquanto o compute UE operacional e em grande parte detido por operadores UE (OVH, Scaleway, Sesterce, EuroHPC, AI Factories), as cargas de IA UE rodam predominantemente em infraestrutura detida pelos EUA. A dependencia nao e do compute instalado, mas do compute usado, e das camadas de plataforma e modelos sobrepostas a ele.",
    ]),
    ("4.2.3 O Relatorio Draghi e o diagnostico europeu", [
        "O Relatorio Draghi (setembro de 2024), encomendado pela Comissao Europeia, constitui o reconhecimento institucional mais completo dessa dependencia. Draghi observa que as restricoes europeias sobre armazenamento e processamento de dados criam altos custos de conformidade e dificultam a criacao de grandes conjuntos de dados integrados para treinamento de modelos de IA. Ele recomenda a criacao de uma infraestrutura de nuvem hyperscale europeia, a consolidacao dos provedores de nuvem europeus e investimento massivo em IA, computacao quantica e supercomputadores EuroHPC.[10] Martens (Bruegel, 2025) modera, no entanto, esse otimismo, observando que os supercomputadores EuroHPC nao sao projetados para IA generativa e que a atualizacao para padroes competitivos excede a capacidade financeira dos orcamentos europeus.[11]",
    ]),
    ("4.3 Produtividade diferenciada: o dividendo do compute abundante", []),
    ("4.3.1 O potencial macroeconomico teorico", [
        "O terceiro canal de transmissao e o impacto da IA na produtividade do trabalho e sua assimetria regional. O McKinsey Global Institute (maio de 2024, dezembro de 2025) estima que a IA generativa poderia permitir que a Europa atingisse uma taxa de crescimento anual da produtividade de 3 por cento ate 2030 em um cenario de adocao acelerada com redistribuicao proativa da forca de trabalho.[12] Esse numero e suficiente para fechar a maior parte do gap de produtividade com os Estados Unidos. No entanto, em um cenario de adocao lenta, o crescimento da produtividade seria de apenas 0,3 por cento, proximo dos niveis atuais e insuficiente para financiar o modelo social europeu.",
        "O FMI (Working Paper, marco de 2025) fornece uma analise mais granular.[13] Cruzando indices de exposicao de tarefas a IA (Felten et al., Eloundou et al.) com dados setoriais, o FMI estima ganhos de produtividade total dos fatores (TFP) significativamente mais altos em economias avancadas do que em paises de renda media, devido a sua maior parcela de valor agregado em setores com alta exposicao a IA (servicos financeiros, consultoria, tecnologia) e a seus niveis salariais mais altos, que justificam ainda mais a automacao. Os estudos microeconomicos subjacentes (ensaios randomizados controlados) mostram ganhos que vao de 14 por cento para tarefas de baixa qualificacao a mais de 50 por cento para engenheiros de software.",
    ]),
    ("4.3.2 O gap condicionado pelo acesso ao compute", [
        "No entanto, a realizacao desses ganhos de produtividade e condicionada pelo acesso ao compute. E aqui que a assimetria diagnosticada no Capitulo III produz seus efeitos. A Comissao IA francesa (2024) estima o impacto potencial da IA no crescimento frances em 1,3 ponto percentual ao ano, por analogia com os efeitos da eletricidade, mas essa estimativa supoe acesso irrestrito ao compute de ponta. Da mesma forma, a McKinsey (dezembro de 2025) condiciona o cenario otimista europeu a uma adocao acelerada que implica investimentos massivos em infraestrutura.[14]",
        "O gap de investimento e consideravel. A McKinsey (janeiro de 2026) estima que as grandes empresas europeias enfrentam um deficit anual de 700 bilhoes USD em P&D e despesas de capital em comparacao com suas contrapartes americanas. Apenas em tecnologias digitais, as empresas dos EUA investiram 2 trilhoes EUR a mais do que as empresas europeias no periodo 2021-2025.[15] O gap anualizado entre janeiro de 2024 e setembro de 2025 chega a 580 bilhoes EUR para investimento corporativo e 300 bilhoes EUR para startups e scale-ups.",
        f"Esse deficit de investimento, combinado com a assimetria de compute e os custos energeticos, produz um gap de produtividade IA condicional. Em nossa modelagem, distinguimos o potencial teorico (identico para EUA e UE com estruturas setoriais comparaveis) e o potencial realizavel, que depende do acesso efetivo ao compute. Com um CACI(EUA)/CACI(UE) de {fmt_fr(us_eu_caci, 2)}:1 em Power Mode (Capitulo III) no snapshot do painel de abril de 2026, o potencial realizavel europeu e estruturalmente inferior. Traduzido em scores CACI normalizados, a UE situa-se em 28,9 contra um benchmark dos EUA de 100, a Franca em 25,3, a Alemanha em 5,4 - o gap nao e portanto de 7 a 12 vezes inferior como sugeriam analises anteriores, mas precisamente de 3,46 para 1 em intensidade competitiva ajustada ao compute, com forte variacao entre os Estados-membros.",
    ]),
    ("4.4 Captura de rendas de inovacao: uma vantagem auto-reforcadora", []),
    ("4.4.1 O mecanismo de first-mover advantage em IA", [
        "O quarto canal e o mais estruturalmente significativo a longo prazo. A teoria das general purpose technologies (GPT, Bresnahan e Trajtenberg, 1995) preve que os early adopters de uma tecnologia geral capturam rendas de inovacao desproporcionais, que se tornam auto-reforcadoras por meio de tres mecanismos.",
        "Primeiro, efeitos de escala dos dados. As empresas que implantam IA primeiro acumulam dados de uso (loops de feedback, fine-tuning) que melhoram o desempenho de seus modelos. Esses dados sao nao-rivais mas exclusivos: uma vez capturados por um ator, nao beneficiam os concorrentes. A McKinsey estima que 75 por cento do potencial economico da IA generativa esta concentrado em quatro funcoes-chave (atendimento ao cliente, desenvolvimento de software, marketing, P&D), onde os dados de uso sao particularmente determinantes.[16]",
        "Segundo, efeitos de rede de plataforma. Os hyperscalers dos EUA se beneficiam de um circulo virtuoso: mais clientes leva a mais receita, que leva a mais investimento em infraestrutura, que leva a uma melhor oferta, que leva a mais clientes. Esse mecanismo explica por que, apesar das discussoes sobre soberania digital, a parcela dos provedores europeus estagna em 15 por cento: o gap de qualidade de servico (latencia, gama de servicos, suporte tecnico, ecossistema de integracao) e grande demais para ser preenchido apenas por argumentos de soberania.",
        "Terceiro, captura de talentos. As empresas americanas oferecem pacotes de remuneracao incluindo salarios e equity que atraem os melhores pesquisadores e engenheiros de IA do mundo, incluindo europeus. Martens observa que o pessoal de IA qualificado e escasso e que as empresas competem ferozmente para recruta-lo, com pacotes de remuneracao dificeis de igualar para as instituicoes europeias. Esse brain drain de IA enfraquece o fator L(r) do CACI europeu e reforca a vantagem americana em capital humano.[17]",
    ]),
    ("4.4.2 O risco de irreversibilidade", [
        "Esses tres mecanismos produzem um efeito cumulativo potencialmente irreversivel. Como analisa a OCDE (Martens, 2021) por analogia com ondas tecnologicas anteriores (internet de banda larga, nuvem), os retardatarios na adocao de uma tecnologia geral tornam-se estruturalmente dependentes de plataformas estrangeiras, sofrem margens comprimidas e perdem progressivamente sua autonomia estrategica. A escassez de semicondutores de 2022, que custou 100 bilhoes EUR apenas ao setor automotivo europeu, ilustra as consequencias concretas de uma dependencia tecnologica nao antecipada.[18]",
        "A trajetoria e ainda mais preocupante uma vez que a UE esta ausente de quatro dos oito segmentos da cadeia de valor da IA generativa mapeados pela McKinsey (2024): semicondutores AI-design (dominados pela Nvidia), plataformas de nuvem IA (AWS, Azure, GCP), modelos de fundacao (OpenAI, Google, Anthropic, Meta) e ferramentas de desenvolvimento de IA. So e competitiva em segmentos de aplicacao a jusante: aplicacoes setoriais (SAP, Siemens), integracao industrial e semicondutores especializados (semicondutores de potencia: Infineon, STMicroelectronics, NXP, com cerca de 15 por cento de participacao de mercado global).[19]",
    ]),
    ("4.5 Sintese: a mecanica da vantagem competitiva dos EUA", [
        f"Os quatro canais identificados formam um sistema reforcador. A assimetria de compute (Capitulo III) alimenta custos de treinamento diferenciados (4.1), que reforcam a dependencia da nuvem dos EUA (4.2), que restringe a produtividade realizavel (4.3), que retarda o investimento europeu e reforca a captura de rendas pelos atores americanos (4.4), que por sua vez amplia o gap de compute. Com o snapshot de abril de 2026 - 76,9 por cento de parcela de compute operacional, 1,59x de custo energetico ajustado-PPA e uma razao CACI Power Mode de {fmt_fr(us_eu_caci, 2)}:1 - esse circulo constitui o que Farrell e Newman (2019) descreveriam como uma weaponizacao da interdependencia estrutural: a dependencia europeia da infraestrutura americana, inicialmente fundada na eficiencia economica, torna-se uma alavanca de poder geopolitico quando instrumentalizada por medidas como a Secao 232.",
        "O protecionismo Trump (Capitulo III, fase 4) nao criou essa assimetria, ela preexistia massivamente. Ele a institucionaliza e legaliza por meio de um mecanismo tarifario que diferencia o custo de acesso ao compute por nacionalidade. Ao faze-lo, transforma uma vantagem de fato em uma vantagem de direito, cujo desmantelamento e politica e juridicamente muito mais dificil. O Capitulo V examina como esses mecanismos poderiam evoluir em quatro cenarios prospectivos.",
    ]),
]

PT.table_blocks = [
    ("Tabela 8. Mercado de nuvem europeu: dominancia dos hyperscalers dos EUA.",
     "Fonte: Synergy Research Group (julho de 2025).",
     [
         ["Indicador", "2017", "2022", "2024", "Tendencia"],
         ["Mercado nuvem UE (bilhoes EUR)", "~10", "~35", "61", "x6 em 7 anos"],
         ["Parcela AWS+Azure+GCP", "~50 pct", "~67 pct", "70 pct", "Crescente"],
         ["Parcela provedores UE", "29 pct", "15 pct", "15 pct", "Estavel (piso)"],
         ["Capex EUA na Europa", "n.d.", "n.d.", "~40 bn EUR/ano", "10 bn EUR/trim."],
         ["Data centers hyperscale na UE", "~40", "~80", ">140", "+75 pct em 2 anos"],
     ]),
    ("Tabela 9. Produtividade IA e vantagem competitiva: EUA vs UE (calibracao abril de 2026).",
     "Fontes: McKinsey (2024, 2025, 2026), Bruegel/Martens (2024), calibracao CACI no painel publico.",
     [
         ["Dimensao", "Estados Unidos", "Europa (UE)", "Gap"],
         ["Produtividade IA potencial (McKinsey, cen. acel.)", "+2,5-3,5 pct/ano", "+2,5-3,0 pct/ano", "Comparavel"],
         ["Produtividade IA realizavel (sob restricoes de compute)", "+2,5-3,0 pct/ano", "+0,8-1,5 pct/ano", "-1,5 a -2 pts"],
         ["Investimento anual IA (corporativo + VC)", "~450 bn USD", "~50-70 bn EUR", "~7:1"],
         ["Custo do FLOP (USD/TFlop, treinamento)", "~0,5", "~1,2-1,8", "2,4-3,6x"],
         ["CACI normalizado Power Mode (EUA = 100)", "100", "28,9 (UE(13))", f"{fmt_fr(us_eu_caci, 2)}:1"],
     ]),
]

PT.notes = EN.notes


# ===========================================================================
# Main
# ===========================================================================

def main() -> None:
    """Build the three Chapter IV .docx files."""
    out_dir = Path(__file__).parent
    for lp in (EN, FR, PT):
        build(lp, out_dir)
    log.info("Done.")


if __name__ == "__main__":
    main()
