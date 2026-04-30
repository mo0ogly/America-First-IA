"""
Shared helpers for General Conclusion generator (Trilingual).

Provides cover, header, section/table/notes/license rendering matching
the visual identity used in Chapters I-VII.

Supports EN, FR, and PT-BR localization.
"""

from __future__ import annotations

import logging
from typing import Any

log = logging.getLogger("concl_helpers")

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor, Inches

NAVY = RGBColor(0x1A, 0x27, 0x44)
GOLD = RGBColor(0xB8, 0x92, 0x2F)
GREY = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x20, 0x20, 0x20)

LABELS = {
    "EN": {
        "date": "RESEARCH STUDY - FEBRUARY 2026",
        "subtitle": "AI Protectionism, Energy and Semiconductors: US/Europe Divergence Trajectories 2024-2030",
        "analysis": "Integrated Geostrategy and Economic Analysis",
        "author_info": "Paris - February 2026 | 11 Chapters | 4 Prospective Scenarios | Global Coverage",
        "keywords": "Keywords: artificial intelligence, technological protectionism, semiconductors, export controls, sovereign compute, AI geopolitics, France, United States, China, Brazil, Africa, India",
        "notes": "Notes",
        "license_title": "License and Disclaimer",
        "license_body": "This work, 'AI for Americans First', is made available under the terms of the Creative Commons Attribution-NonCommercial-ShareAlike 4.0 International (CC BY-NC-SA 4.0) license. You are free to share and adapt the material for non-commercial purposes, provided you properly attribute the work to Fabrice Pizzi (Université Paris-Sorbonne) and distribute your contributions under the same license. This document is provided for educational and research purposes only.",
        "dashboard": "Public Dashboard",
        "repo": "Repository",
    },
    "FR": {
        "date": "ETUDE DE RECHERCHE - FEVRIER 2026",
        "subtitle": "Protectionnisme IA, Energie et Semi-conducteurs : Trajectoires de divergence US/Europe 2024-2030",
        "analysis": "Analyse geostrategique et economique integree",
        "author_info": "Paris - fevrier 2026  |  11 chapitres  |  4 scenarios prospectifs  |  Couverture mondiale",
        "keywords": "Mots-cles : intelligence artificielle, protectionnisme technologique, semi-conducteurs, controles a l'exportation, compute souverain, geopolitique IA, France, Etats-Unis, Chine, Bresil, Afrique, Inde",
        "notes": "Notes",
        "license_title": "Licence et avertissement",
        "license_body": "Cette oeuvre, 'AI for Americans First', est mise a disposition selon les termes de la licence Creative Commons Attribution - Pas d'Utilisation Commerciale - Partage dans les Memes Conditions 4.0 International (CC BY-NC-SA 4.0). Vous etes libre de partager et adapter le materiel a des fins non commerciales, a condition d'attribuer correctement l'oeuvre a Fabrice Pizzi (Universite Paris-Sorbonne) et de distribuer vos contributions sous la meme licence. Ce document est fourni a des fins educatives et de recherche uniquement.",
        "dashboard": "Tableau de bord public",
        "repo": "Depot",
    },
    "PT-BR": {
        "date": "ESTUDO DE PESQUISA - FEVEREIRO 2026",
        "subtitle": "Protecionismo de IA, Energia e Semicondutores: Trajetorias de Divergencia EUA/Europa 2024-2030",
        "analysis": "Analise Geoestrategica e Economica Integrada",
        "author_info": "Paris - fevereiro 2026 | 11 Capitulos | 4 Cenarios Prospectivos | Cobertura Global",
        "keywords": "Palavras-chave: inteligencia artificial, protecionismo tecnologico, semicondutores, controles de exportacao, computacao soberana, geopolitica da IA, Franca, Estados Unidos, China, Brasil, Africa, India",
        "notes": "Notas",
        "license_title": "Licenca e Aviso Legal",
        "license_body": "Esta obra, 'AI for Americans First', esta disponivel sob os termos da licenca Creative Commons Atribuicao-NaoComercial-CompartilhaIgual 4.0 Internacional (CC BY-NC-SA 4.0). Voce tem a liberdade de compartilhar e adaptar o material para fins nao comerciais, desde que atribua corretamente a obra a Fabrice Pizzi (Universite Paris-Sorbonne) e distribua suas contribuicoes sob a mesma licenca. Este documento e fornecido apenas para fins educacionais e de pesquisa.",
        "dashboard": "Painel Publico",
        "repo": "Repositorio",
    }
}

def set_run(run, *, font="Calibri", size=11, bold=False, italic=False, color=None):
    """Apply consistent typography to a docx run."""
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color

def add_paragraph(doc, text, *, align=None, space_after=6, **run_kwargs):
    """Add a paragraph with one run."""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run(run, **run_kwargs)
    return p

def add_heading(doc, text, level):
    """Custom heading rendering aligned with the rest of the dissertation."""
    sizes = {1: 22, 2: 16, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, font="Calibri", size=sizes.get(level, 11),
            bold=True, color=NAVY)
    return p

def init_document() -> Document:
    """Create a Document with the standard margins."""
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    return doc

def add_cover(doc, lang: str, *, chapter_label: str, chapter_subtitle: str):
    """Render the General Conclusion cover."""
    L = LABELS.get(lang, LABELS["EN"])
    add_paragraph(doc, "", space_after=0)
    add_paragraph(doc, L["date"],
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, bold=True, color=GREY, space_after=4)
    add_paragraph(doc, "AI FOR AMERICANS FIRST",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=26, bold=True, color=NAVY, space_after=4)
    add_paragraph(doc, L["subtitle"],
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, italic=True, color=GREY, space_after=4)
    add_paragraph(doc,
                  f"{L['analysis']} - {chapter_label}",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=11, italic=True, color=GREY, space_after=18)

    table = doc.add_table(rows=1, cols=3)
    # Stats from April 2026 baseline
    if lang == "FR":
        chips = [
            "76,9 pct du compute IA operationnel mondial = USA",
            "1,59x cout energie EU/US (ajuste-PPA)",
            "3,46:1 ratio CACI US/EU (Power Mode)",
        ]
    elif lang == "PT-BR":
        chips = [
            "76,9 pct do compute IA operacional mundial = EUA",
            "1,59x custo energia UE/EUA (ajustado-PPP)",
            "3,46:1 razao CACI EUA/UE (Power Mode)",
        ]
    else: # EN
        chips = [
            "76.9% of global operational AI compute = USA",
            "1.59x EU/US energy cost (PPP-adjusted)",
            "3.46:1 US/EU CACI ratio (Power Mode)",
        ]
        
    for i, line in enumerate(chips):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        set_run(run, size=11, bold=True, color=NAVY)

    add_paragraph(doc, "", space_after=8)
    add_paragraph(doc, "Fabrice Pizzi",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, bold=True, color=NAVY, space_after=2)
    add_paragraph(doc, "Universite Paris-Sorbonne",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, color=GREY, space_after=2)
    add_paragraph(doc,
                  "Master 2 Intelligence Economique - Intelligence Warfare",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, italic=True, color=GREY, space_after=14)
    add_paragraph(doc, L["author_info"],
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, italic=True, color=GREY, space_after=10)
    add_paragraph(doc, L["keywords"],
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=9, italic=True, color=GREY, space_after=24)

def add_chapter_header(doc, *, label: str, title: str, intro: str):
    """Render the chapter title strip + intro paragraph."""
    add_paragraph(doc, label,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=11, bold=True, color=GOLD, space_after=2)
    add_paragraph(doc, title,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=18, bold=True, color=NAVY, space_after=12)
    if intro:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(intro)
        set_run(run, size=11, italic=True, color=GREY)

def render_section(doc, title: str, paragraphs: list[str]):
    """Render one section: title at the right level, then body paragraphs."""
    # Conclusion often uses 1. 2. 3. for sections
    level = 2
    add_heading(doc, title, level)
    for para in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(para)
        set_run(run, size=11, color=DARK)

def render_table(doc, lang: str, caption: str, source: str, rows: list[list[str]]):
    """Render one captioned table."""
    add_paragraph(doc, caption,
                  size=10, bold=True, color=NAVY, space_after=2)
    add_paragraph(doc, source,
                  size=9, italic=True, color=GREY, space_after=4)
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            set_run(run, size=10,
                    bold=(i == 0),
                    color=NAVY if i == 0 else DARK)
    doc.add_paragraph()

def render_image(doc, image_path: Path, width_inches: float = 6.0):
    """Embed an image into the document."""
    if not image_path.exists():
        log.warning("Image not found: %s", image_path)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    doc.add_paragraph()

def render_notes(doc, lang: str, notes: list[str]):
    """Render the footnotes block at the end of the chapter."""
    L = LABELS.get(lang, LABELS["EN"])
    add_heading(doc, L["notes"], 2)
    for i, note in enumerate(notes, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run_id = p.add_run(f"{i}. ")
        set_run(run_id, size=9, bold=True, color=GOLD)
        run_txt = p.add_run(note)
        set_run(run_txt, size=9, color=GREY)

def render_license(doc, lang: str, page_footer: str):
    """Render the license disclaimer + footer."""
    L = LABELS.get(lang, LABELS["EN"])
    doc.add_paragraph()
    license_lines = [
        f"{L['license_title']}. {L['license_body']}",
        f"{L['dashboard']} : https://mo0ogly.github.io/America-First-IA/dashboard/",
        f"{L['repo']} : https://github.com/mo0ogly/America-First-IA",
    ]
    for line in license_lines:
        add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT,
                      size=8, italic=True, color=GREY, space_after=2)
    add_paragraph(doc, page_footer,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=8, italic=True, color=GREY, space_after=0)
