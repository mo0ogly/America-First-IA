"""
Chapter III - Empirical Diagnosis 2020-2026 - trilingual generator.

Generates the Chapter III .docx file in English, French and Brazilian
Portuguese for the doctoral study "AI for Americans First".

Key revisions vs the previous version (April 2026 dashboard alignment):
    1. Cover banner now reads 76.9 percent operational compute /
       1.59x PPA-adjusted energy ratio / 3.46:1 CACI Power Mode US/EU,
       in line with Chapters I and II.
    2. Section 3.3.1 (regional shares): updated to USA 76.9 / China 12.8
       / EU 4.4 percent (April 2026 snapshot, operational compute only),
       with the US/EU H100-equivalent ratio at 17.6:1 (was 15:1).
    3. Section 3.3.3 (CACI calibration): the PetaFLOP figures are
       refreshed to 2,759,968 PF for the USA and 156,632 PF for the
       EU(13) on operational compute, and the resulting CACI Power Mode
       ratio is set to 3.46:1 (was 3.4:1).
    4. Section 3.1.3 (energy as competitive advantage): the "2 to 3
       times" claim is now annotated with the PPA-adjusted figure of
       1.4 to 1.7x for hyperscaler workloads, consistent with Chapter
       II section 2.4.3.
    5. Section 3.5 (synthesis): item (2) clarifies that the 76.9 percent
       share refers to operational compute, while the share including
       planned capacity is 49.9 percent because of the UAE/Saudi/Korea
       offshoring of US-owned compute. Item (3) aligns the energy
       claim on the 1.4 to 1.7x PPA-adjusted band.
    6. Table 6 (US dominance indicators): all rows refreshed to live
       dashboard values, and the US/EU GPU performance ratio updated
       from "~15:1" to "~17.6:1".

Author: Fabrice Pizzi (Universite Paris-Sorbonne, M2 Economic Intelligence).
Build: python3 generate_chapter3_trilingual.py
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
log = logging.getLogger("chapter3_gen")


# ---------------------------------------------------------------------------
# Visual identity (shared with Chapters I and II)
# ---------------------------------------------------------------------------

NAVY = RGBColor(0x1A, 0x27, 0x44)
GOLD = RGBColor(0xB8, 0x92, 0x2F)
GREY = RGBColor(0x55, 0x55, 0x55)


@dataclass
class LangPack:
    """Container for one language version of Chapter III."""

    code: str
    filename: str
    cover_subtitle: str
    cover_title: str
    cover_blurb: str
    cover_chip_lines: list[str]
    cover_meta: str
    cover_keywords_label: str
    cover_keywords: str
    chapter_label: str
    chapter_title: str
    chapter_intro: str
    sections: list[tuple[str, list[str]]] = field(default_factory=list)
    table_blocks: list[tuple[str, str, list[list[str]]]] = field(default_factory=list)
    notes_label: str = "Notes"
    notes: list[str] = field(default_factory=list)
    license_block: list[str] = field(default_factory=list)
    page_footer: str = ""


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def set_run(run, *, font="Calibri", size=11, bold=False, italic=False, color=None):
    """Apply consistent typography to a docx run."""
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text, *, style=None, align=None, space_after=6, **run_kwargs):
    """Add a paragraph and return it."""
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run(run, **run_kwargs)
    return p


def add_heading(doc, text, level):
    """Custom heading rendering aligned with the rest of the dissertation."""
    sizes = {1: 22, 2: 16, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, font="Calibri", size=sizes.get(level, 11),
            bold=True, color=NAVY)
    return p


def add_cover(doc, lp: LangPack):
    """Render the cover block matching Chapters I and II."""
    banner = {
        "EN": "RESEARCH STUDY - FEBRUARY 2026",
        "FR": "ETUDE DE RECHERCHE - FEVRIER 2026",
        "PT-BR": "ESTUDO DE PESQUISA - FEVEREIRO DE 2026",
    }.get(lp.code, "RESEARCH STUDY - FEBRUARY 2026")
    add_paragraph(doc, "", space_after=0)
    add_paragraph(doc, banner,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, bold=True, color=GREY, space_after=4)
    add_paragraph(doc, lp.cover_title,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=26, bold=True, color=NAVY, space_after=4)
    add_paragraph(doc, lp.cover_subtitle,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, italic=True, color=GREY, space_after=4)
    add_paragraph(doc, lp.cover_blurb,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=11, italic=True, color=GREY, space_after=18)

    table = doc.add_table(rows=1, cols=3)
    table.autofit = True
    for i, line in enumerate(lp.cover_chip_lines):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        set_run(run, size=11, bold=True, color=NAVY)

    add_paragraph(doc, "", space_after=8)
    add_paragraph(doc, "Fabrice Pizzi",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, bold=True, color=NAVY, space_after=2)
    add_paragraph(doc, "Universite Paris-Sorbonne",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, color=GREY, space_after=2)
    master_line = {
        "EN": "Master 2 Economic Intelligence - Intelligence Warfare",
        "FR": "Master 2 Intelligence Economique - Intelligence Warfare",
        "PT-BR": "Mestrado 2 Inteligencia Economica - Intelligence Warfare",
    }.get(lp.code, "Master 2 Economic Intelligence - Intelligence Warfare")
    add_paragraph(doc, master_line,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, italic=True, color=GREY, space_after=14)
    add_paragraph(doc, lp.cover_meta,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=10, italic=True, color=GREY, space_after=10)
    add_paragraph(doc, f"{lp.cover_keywords_label}: {lp.cover_keywords}",
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=9, italic=True, color=GREY, space_after=24)


def add_chapter_header(doc, lp: LangPack):
    """Render the chapter header strip below the cover."""
    add_paragraph(doc, lp.chapter_label,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=11, bold=True, color=GOLD, space_after=2)
    add_paragraph(doc, lp.chapter_title,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=18, bold=True, color=NAVY, space_after=12)
    if lp.chapter_intro:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(lp.chapter_intro)
        set_run(run, size=11, italic=True, color=GREY)


def render_sections(doc, lp: LangPack):
    """Render the body sections from (heading, [paragraph, ...]) tuples."""
    for heading, paragraphs in lp.sections:
        if heading.startswith(("3.1 ", "3.2 ", "3.3 ", "3.4 ", "3.5 ")):
            add_heading(doc, heading, 2)
        else:
            add_heading(doc, heading, 3)
        for para in paragraphs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.25
            run = p.add_run(para)
            set_run(run, size=11, color=RGBColor(0x20, 0x20, 0x20))


def render_tables(doc, lp: LangPack):
    """Render the data tables (caption, source, rows) at the end."""
    for caption, source, rows in lp.table_blocks:
        add_paragraph(doc, caption,
                      size=10, bold=True, color=NAVY, space_after=2)
        add_paragraph(doc, source,
                      size=9, italic=True, color=GREY, space_after=4)
        n_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=n_cols)
        table.style = "Light Grid Accent 1"
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                cell = table.rows[i].cells[j]
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(cell_text)
                set_run(run, size=10,
                        bold=(i == 0),
                        color=NAVY if i == 0 else RGBColor(0x20, 0x20, 0x20))
        doc.add_paragraph()


def render_notes(doc, lp: LangPack):
    """Render the footnotes block at the end of the chapter."""
    add_heading(doc, lp.notes_label, 2)
    for i, note in enumerate(lp.notes, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run_id = p.add_run(f"{i}. ")
        set_run(run_id, size=9, bold=True, color=GOLD)
        run_txt = p.add_run(note)
        set_run(run_txt, size=9, color=GREY)


def render_license(doc, lp: LangPack):
    """Render the license disclaimer block."""
    doc.add_paragraph()
    for line in lp.license_block:
        add_paragraph(doc, line,
                      align=WD_ALIGN_PARAGRAPH.LEFT,
                      size=8, italic=True, color=GREY, space_after=2)
    add_paragraph(doc, lp.page_footer,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=8, italic=True, color=GREY, space_after=0)


def build(lp: LangPack, out_dir: Path) -> Path:
    """Build the .docx file for one language pack."""
    log.info("Building Chapter III [%s] -> %s", lp.code, lp.filename)
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    add_cover(doc, lp)
    add_chapter_header(doc, lp)
    render_sections(doc, lp)
    render_tables(doc, lp)
    render_notes(doc, lp)
    render_license(doc, lp)

    out = out_dir / lp.filename
    doc.save(out)
    log.info("Saved %s", out)
    return out


# ===========================================================================
# Content - English
# ===========================================================================

EN = LangPack(
    code="EN",
    filename="Chapter_III_Empirical_Diagnosis_EN.docx",
    cover_title="AI FOR AMERICANS FIRST",
    cover_subtitle="AI Protectionism, Energy and Semiconductors: US/Europe Divergence Trajectories 2024-2030",
    cover_blurb="Integrated Geostrategic and Economic Analysis - Chapter III",
    cover_chip_lines=[
        "76.9% global operational AI compute = USA",
        "1.59x energy cost EU/US",
        "3.46:1 CACI ratio US/EU (Power Mode)",
    ],
    cover_meta="Paris - February 2026  |  7 chapters  |  4 prospective scenarios  |  3 geographic zones",
    cover_keywords_label="Keywords",
    cover_keywords=("artificial intelligence, technology protectionism, semiconductors, "
                    "export controls, sovereign compute, AI geopolitics, France, "
                    "United States, China"),
    chapter_label="CHAPTER III",
    chapter_title="Empirical Diagnosis 2020-2026",
    chapter_intro=(
        "This chapter establishes the factual foundation of the analysis. It covers four "
        "interdependent dimensions: the energy trajectory of data centers, the evolution of the "
        "semiconductor market, the geographical distribution of installed AI compute, and the "
        "timeline of American regulatory measures. The data presented here constitute the "
        "predetermined elements (in the Schwartz sense) that structure the prospective scenarios "
        "in Chapter V. All time series are sourced, and where data diverge between sources, the "
        "discrepancy is made explicit. All ratios involving compute have been recalibrated on "
        "the live April 2026 snapshot of the public dashboard."
    ),
    notes_label="Notes",
    license_block=[
        "License and Disclaimer. This work, 'AI for Americans First,' is made available under the terms of the America-First-IA Creative Commons Attribution - NonCommercial - ShareAlike 4.0 International License (CC BY-NC-SA 4.0).",
        "You are free to share and adapt the material for non-commercial purposes, provided you give appropriate credit to Fabrice Pizzi (Universite Paris-Sorbonne) and distribute your contributions under the same license. This document is provided for educational and research purposes only.",
        "Public dashboard: https://mo0ogly.github.io/America-First-IA/dashboard/",
        "Repository: https://github.com/mo0ogly/America-First-IA",
    ],
    page_footer="AI for Americans First - Fabrice Pizzi - Chapter III",
)

EN.sections = [
    ("3.1 Data Center Energy Trajectory: Demand Doubling in Six Years", []),
    ("3.1.1 Global Consumption 2020-2024", [
        "The International Energy Agency (IEA, April 2025) estimates global data center electricity consumption at approximately 415 TWh in 2024, representing 1.5 percent of global electricity consumption.[1] This consumption has grown at an average of 12 percent per year since 2017, a rate four times higher than total electricity consumption growth. In 2020, it stood at approximately 270 TWh; the increase thus represents +54 percent over four years.",
        "The geographical distribution is highly concentrated. The United States absorbs 45 percent of global data center consumption (approximately 180 TWh in 2024), followed by China (25 percent, ~102 TWh) and Europe (15 percent, ~70 TWh for the EU according to the European Commission).[2] The four main European countries (Germany, France, United Kingdom, Netherlands) total approximately 41 TWh, or 10 percent of the global total. Total installed global data center capacity reached nearly 100 GW in 2024. Nearly half of US capacity is concentrated in just five regional clusters.",
    ]),
    ("3.1.2 Projections 2024-2030: The IEA Central Scenario", [
        "The IEA central scenario (Base Case) projects global consumption of 945 TWh by 2030, representing a doubling compared to 2024 and equivalent to Japan current electricity consumption. AI is identified as the most important driver of this growth. AI share of data center consumption, estimated at 5 to 15 percent in recent years, could reach 35 to 50 percent by 2030.[3]",
        "Growth is geographically uneven. The United States adds approximately +240 TWh (+130 percent compared to 2024), China +175 TWh (+170 percent), and Europe only +45 TWh (+70 percent). The IEA notes a 20 percent risk of project delays in Europe, linked to grid connection constraints and permitting timelines.[4] In the United States, data centers will account for nearly half of electricity demand growth by 2030; the country will consume more electricity for its data centers than for aluminum, steel, cement, and chemical production combined.",
    ]),
    ("3.1.3 The Energy Factor as Competitive Advantage", [
        "The energy asymmetry constitutes a structural advantage for the United States. The IEA notes that natural gas currently provides 40 percent of US data center electricity, renewables 24 percent, nuclear 20 percent, and coal 15 percent. In Europe, the energy mix is more constrained by climate objectives and higher energy costs. Industrial electricity prices are typically 2 to 3 times higher in Europe than in the United States on the unadjusted Eurostat tariff. Once Power Purchase Agreements negotiated by hyperscalers are factored in, the gap narrows to roughly 1.4 to 1.7 times (1.59x for the EU aggregate at the April 2026 dashboard reference of 135 USD/MWh in the EU vs. 85 USD/MWh in the United States). As documented by the Federal Reserve Board (October 2025), there is a significant negative correlation between energy costs and AI adoption at the European enterprise level, even after PPA correction.[5]",
        "France, however, has a specific asset in this context: its nuclear electricity mix (approximately 65 to 70 percent of production), which offers relatively inexpensive, decarbonized, and baseload-available electricity. RTE estimates an additional need of +10 GW for data centers in France by 2030, which remains compatible with installed capacity subject to investments in the transmission network.",
    ]),
    ("3.1.4 Amplifying Effect: AI Robotics and Industrial Energy Demand", [
        "A factor still poorly quantified in the literature is the impact of AI robotics on energy demand. Autonomous robots require both embedded compute (edge computing) and centralized compute (cloud AI for training and model updates). The widespread deployment of AI automation in industry could add 20 to 30 percent additional energy demand in affected sectors, on top of data center growth. This factor is integrated as a sensitivity variable in our scenarios (Chapter V), as no precise and reliable estimate yet exists. The IEA nonetheless devotes a section of its report to the Jevons rebound effect: even with significant efficiency gains (illustrated by the DeepSeek case), demand growth absorbs the gains.[6]",
    ]),
    ("3.2 Semiconductor Market: AI-Driven Growth", []),
    ("3.2.1 Global Sales Trajectory 2020-2025", [
        "The global semiconductor market has followed a remarkable trajectory since 2020. After a record of 555.9 billion USD in 2022, the sector contracted by 8.2 percent in 2023 (526.8 billion USD), before a spectacular rebound in 2024: global sales reached 630.5 billion USD (revised SIA/WSTS data, March 2025), a 19.1 percent year-on-year increase and a new all-time record.[7]",
        "The year 2025 confirmed the acceleration. The SIA announced in February 2026 global sales of 791.7 billion USD in 2025, a 25.6 percent increase over 2024.[8] The fourth quarter of 2025 (236.6 billion USD) was 37.1 percent higher than the same quarter in 2024. The SIA now projects sales close to 1 trillion USD in 2026, a symbolic threshold that seemed out of reach two years ago.",
        "Two segments are driving this growth. Logic chips (processors, GPUs, ASICs) reached 301.9 billion USD in 2025 (+39.9 percent), becoming the top category by sales. Memory (DRAM, NAND) reached 223.1 billion USD (+34.8 percent). Together, these two segments represent over 66 percent of global sales and directly reflect AI infrastructure demand: servers, data centers, and HPC (High Performance Computing).",
    ]),
    ("3.2.2 The Scope Question: SIA vs McKinsey", [
        "As noted in Chapter II, a notable methodological gap exists between SIA/WSTS data and McKinsey estimates. The SIA records semiconductor sales in the strict sense (components sold by manufacturers). McKinsey (January 2026) adopts an expanded scope including the value of captive designers (Apple, Amazon, Tesla, Google), i.e., the economic value of chips designed internally but manufactured by contract foundries. Under this expanded scope, McKinsey values the market at approximately 775 billion USD in 2024 and projects 1.5 to 1.8 trillion by 2030.[9] The expanded scope is more relevant to our analysis because it captures the full value of the AI chain, including hyperscaler design investments.",
    ]),
    ("3.2.3 Regional Concentration of Manufacturing Capacity", [
        "The geographical distribution of foundry capacity constitutes a key vulnerability factor. Taiwan (TSMC) concentrates most of the leading-edge production (sub-7nm nodes). The United States represents approximately 10 percent of global installed capacity but is growing rapidly (the Chips Act has generated nearly 500 billion USD in announced private investment, with a goal of tripling US capacity by 2032).[10] The EU represents approximately 8 percent of global capacity, a figure the European Chips Act (43 billion EUR) aims to raise to 20 percent by 2030, a target deemed unrealistic by many observers. China, despite restrictions, is progressing from 21 percent toward a 30 percent target of installed capacity by 2030, primarily on mature nodes (28 nm and above).",
    ]),
    ("3.3 Geographical Distribution of Installed AI Compute", []),
    ("3.3.1 American Dominance: Three Quarters of Operational Compute", [
        "The most significant data point for our analysis is the geographical distribution of AI-dedicated computing capacity. The work of Epoch AI (Pilz et al., April 2025) provides the foundational dataset, refreshed monthly on the public dashboard.[11] As of April 2026, the dashboard catalogues 786 clusters, of which 609 are operational and 171 are planned or announced.",
        "The result is unequivocal. On operational compute alone, the United States concentrates 76.9 percent of global H100-equivalent capacity, China holds 12.8 percent, and the EU(13) accounts for 4.4 percent. The remaining countries combine to roughly 6 percent. The operational US/EU H100-equivalent ratio is therefore 17.6:1, a gap of considerable magnitude that has widened slightly since the May 2025 GeoCoded/Sanchez snapshot (74.5 / 14.1 / 4.8 percent).[12]",
        "Two dynamics explain this concentration. First, the private sector share of global AI compute has grown from 40 percent in 2019 to 80 percent in 2025, and the technology companies investing massively in AI are almost exclusively American (Microsoft, Meta, Google, Amazon, xAI). Second, cluster sizes have exploded: systems with more than 10,000 chips were rare in 2019; in 2024, xAI deployed a 200,000-GPU Colossus cluster. The performance of leading AI supercomputers doubles every nine months.[13]",
        "The picture shifts substantially when planned capacity is included. The dashboard F_total view, which counts both operational and announced clusters, shows the US share falling to 49.9 percent. The shift is not driven by foreign actors building independent capacity, but by US hyperscalers offshoring to allied jurisdictions: the announced UAE 5 GW campus and Stargate UAE Phase 2 represent 22.9 million H100-equivalents, Saudi Arabia 7.2 million, and South Korea 5.3 million, all largely owned or operated by US firms. This distinction between Physical CACI (where the cluster sits) and Sovereign CACI (who controls it) is developed in Chapter V. The Chinese share collapses to 0.5 percent in F_total, an artifact of the Epoch AI methodology rather than a real position change: the dataset anonymises Chinese systems and rounds them aggressively, so announced Chinese capacity is structurally under-recorded. The EU(13) F_total share also dips to 3.3 percent because the Fluidstack 1 GW campus in France pulls the total up while the rest of the EU shows few large announcements.",
    ]),
    ("3.3.2 Total Electrical Power of the AI Chip Fleet", [
        "Epoch AI (January 2026) estimates the total electrical capacity of the global AI chip fleet at approximately 30 GW at end-2025, comparable to the peak power consumption of New York State.[14] This estimate is based on quarterly AI chip sales from major manufacturers (Nvidia, AMD, Google TPU), multiplied by their nominal power (TDP) and a 2.5x factor for data center infrastructure. Global AI chip production doubles every seven months, a pace that exceeds all prior forecasts. The five largest US AI server investors (Microsoft, Google, Meta, Amazon, xAI) announced 320 billion USD in investments in 2025, compared to 230 billion USD in 2024.",
    ]),
    ("3.3.3 Implications for the CACI Metric", [
        "These data allow a sharp calibration of the CACI defined in Chapter II. Using the geometric weighted formula CACI = F^0.40 x L^0.20 x R^0.15 / E^0.25 (Power Mode), with the April 2026 dashboard reference values: operational US compute of 2,759,968 PetaFLOP/s (FP16) versus 156,632 PetaFLOP/s for the EU(13), an EU/US PPA-adjusted energy ratio of 1.59x (135 vs. 85 USD/MWh), comparable workforce orders of magnitude (3.5 vs. 9.6 million across the EU(13)), and geopolitical access factors of R(US) = 1.0 versus R(EU) = 0.9, the CACI(US)/CACI(EU) Power Mode ratio settles at 3.46:1.",
        "The gap is substantially attenuated relative to the 17.6:1 raw H100-equivalent advantage because the geometric exponents (F at 0.40, L at 0.20, R at 0.15, E at 0.25) apply diminishing returns to the massive US compute lead and reward the EU partially on labour, geopolitical access and lower energy multipliers. In Intensity Mode (dividing additionally by GDP), the gap narrows further to roughly 1.4:1, which is why the headline ratio used throughout this study is the Power Mode ratio of 3.46:1. Chapter V projects the evolution of this ratio under each scenario.",
    ]),
    ("3.4 Timeline of American Regulatory Measures (2022-2026)", [
        "The sequence of American semiconductor and AI control measures constitutes the guiding thread of our hypothesis. Four phases are identifiable, marking a progressive escalation in the scope and intensity of restrictions.",
    ]),
    ("Phase 1 - The Initial Shock (October 7, 2022)", [
        "The BIS (Bureau of Industry and Security) of the Department of Commerce publishes a final interim rule that radically transforms American export controls on semiconductors. The measures cover three components: (i) controls on advanced computing chips (GPUs above performance thresholds defined by TTP - Total Processing Performance), (ii) controls on semiconductor manufacturing equipment (SME, including EUV and DUV lithography), and (iii) restrictions on US person activities supporting advanced chip manufacturing in China.[15] The target is explicitly China: the stated objective is to prevent Chinese military modernization through AI compute access. The restrictions include three new Foreign Direct Product (FDP) rules that extend American jurisdiction to products manufactured outside the United States if US technologies are used in their production.",
    ]),
    ("Phase 2 - Closing the Gaps (October 2023)", [
        "The BIS publishes two new interim rules that strengthen and broaden the October 2022 controls. Technical thresholds are adjusted to capture Nvidia chips specifically designed to circumvent restrictions (A800, H800). The geographic scope is extended to approximately 40 additional countries (Country Groups D:1, D:4, and D:5), with a differentiated licensing regime by country category.[16] Controls on manufacturing equipment are deepened. It is during this phase that European actors begin to perceive the indirect effects of restrictions, even though the EU is not the primary target.",
    ]),
    ("Phase 3 - The AI Diffusion Rule (January 2025, Biden)", [
        "The Biden administration publishes in January 2025 the AI Diffusion Rule, which represents a qualitative shift. For the first time, restrictions apply not only to physical chips but also to AI model weights and cloud compute access. The rule classifies 120 countries into three categories: (i) trusted allies (broad access), (ii) intermediate countries (quotas), (iii) embargoed countries.[17] European reactions are strong: the European Parliament raises alarm about restrictions threatening the EU ability to train models on its AI Factories. France and Germany are classified as trusted partners, but other member states face caps on importable GPU volumes.",
    ]),
    ("Phase 4 - The Trump Break: Section 232 and Explicit Protectionism (January 2026)", [
        "On January 14, 2026, President Trump signs Proclamation 11002, invoking Section 232 of the Trade Expansion Act of 1962.[18] This action marks a qualitative break with previous phases, for three reasons.",
        "First, the legal instrument changes. Phases 1 to 3 fell under export controls (export regulation, defensive national security logic). Section 232 is a tariff instrument (import duties), whose logic is protectionist: it aims to protect domestic production, not merely restrict an adversary access.",
        "Second, the tariff creates an explicit competitive advantage for American companies. The 25 percent tariff hits advanced GPUs (H200, MI325X) imported unless destined for American domestic uses: US data centers, R&D, startups, public sector, industrial and non-data-center consumer applications. Concretely, an American company using these chips on US soil does not pay the tariff; a foreign company importing the same chips for re-export to China pays 25 percent.[19]",
        "Third, the proclamation announces an imminent expansion. The text provides that the Secretary of Commerce and USTR negotiate within 90 days with semiconductor-producing countries, and that broader tariffs, accompanied by a tariff offset program for companies investing in US production, could be imposed. The Secretary of Commerce must provide by July 2026 a report on the semiconductor market used in American data centers.[20]",
    ]),
    ("3.4.5 Interpretation: From Defensive Control to Offensive Protectionism", [
        "The 2022-2026 sequence reveals a qualitative transformation of American policy. The Biden phases (2022-2025) follow a denial strategy logic: preventing a designated adversary (China) from accessing key technologies, within a multilateral framework (coordination with Japan, Netherlands). The Trump phase (2025-2026) adds a capture strategy logic: generating revenue (25 percent tariff), prioritizing American companies (domestic exemptions), and using compute access as a negotiation lever with third countries.",
        "This is precisely the transformation that validates the central hypothesis of our study: American technology protectionism is no longer limited to denying access to an adversary; it actively constructs a structural competitive advantage for American companies. The Section 232 domestic exemptions mean that, for the first time, the cost of accessing cutting-edge compute is legally differentiated by the user nationality. Even though direct effects on the EU remain limited in January 2026 (the tariff primarily targets re-exportation to China), the proclamation opens the way for an expansion that could directly affect Europe, precisely the territory of scenarios B, C, and D of our analysis.",
    ]),
    ("3.5 Diagnostic Synthesis: The Predetermined Elements of 2026", [
        "The four dimensions of the diagnosis converge toward a coherent picture that structures the scenarios in Chapter V.",
        "(1) AI compute demand is growing exponentially (chip sales doubling in two years, energy consumption doubling projected over six years), and this growth shows no signs of slowing.",
        "(2) This demand is structurally concentrated in the United States. The April 2026 dashboard snapshot puts the US share of operational AI compute at 76.9 percent in H100-equivalents and at 49.9 percent if announced clusters are included; 45 percent of data center electricity consumption; and more than 80 percent of frontier-AI private investment. The concentration on operational compute is increasing over time.",
        "(3) Europe starts from a structurally deficit position: roughly 4.4 percent of operational compute in H100-equivalents, around 15 percent of data center consumption, energy costs that run 1.4 to 1.7 times higher after PPA correction (and 2 to 3 times higher on unadjusted Eurostat industrial tariffs), and 72 to 80 percent dependency on US hyperscalers for AI cloud workloads. The EU investment plans (Chips Act, AI Factories, SNIA) carry budgets at least one order of magnitude smaller than the 2025 Big Tech announcements (320 billion USD for the five US frontier-AI investors alone).",
        "(4) The American regulatory framework crossed a qualitative threshold in January 2026 with the shift from export controls to Section 232 tariffs, creating a legal mechanism for differentiating compute access costs by nationality. The proclamation explicitly signals the possibility of expansion, whose scope and timeline will depend on political variables, precisely the critical uncertainties that our scenarios explore.",
        "The following chapter (Chapter IV) analyzes the mechanisms by which this compute asymmetry translates into measurable competitive advantage for American companies.",
    ]),
]

EN.table_blocks = [
    ("Table 4. Data center electricity consumption by region, 2020-2030.",
     "Source: IEA (2025), Energy and AI.",
     [
         ["Region", "2020", "2024", "2030 (IEA)", "Delta 24-30", "Global Share 2024"],
         ["United States", "~120 TWh", "~180 TWh", "~420 TWh", "+240 TWh", "45 pct"],
         ["China", "~60 TWh", "~102 TWh", "~280 TWh", "+175 TWh", "25 pct"],
         ["Europe (EU)", "~45 TWh", "~70 TWh", "~115 TWh", "+45 TWh", "15 pct"],
         ["World", "~270 TWh", "~415 TWh", "~945 TWh", "+530 TWh", "100 pct"],
     ]),
    ("Table 5. Global semiconductor sales, 2020-2026.",
     "Sources: SIA/WSTS (Feb. 2025, Feb. 2026). The 2026 projection is based on the WSTS autumn 2025 forecast (975.4 billion USD) and the SIA February 2026 statement.",
     [
         ["Year", "2020", "2022", "2023", "2024", "2025", "2026 (proj.)"],
         ["Sales (billion USD, SIA)", "440", "556", "527", "631", "792", "~975-1,000"],
         ["Growth (pct)", "+6.8", "+3.3", "-8.2", "+19.1", "+25.6", "+23-26"],
     ]),
    ("Table 6. Indicators of US dominance in AI compute (April 2026 dashboard snapshot).",
     "Sources: Epoch AI GPU Clusters dataset April 2026; IEA (2025); Sanchez/GeoCoded (2025).",
     [
         ["Indicator", "United States", "China", "EU(13)", "US/EU Ratio"],
         ["Operational H100-eq share", "76.9 pct", "12.8 pct", "4.4 pct", "17.6:1"],
         ["F_total share (op + planned)", "49.9 pct", "0.5 pct", "3.3 pct", "15.1:1"],
         ["Private sector share (2025)", "~65 pct of global", "~12 pct", "~3 pct", "~22:1"],
         ["Data center consumption 2024", "180 TWh", "102 TWh", "70 TWh", "2.6:1"],
         ["AI investment 2025", "320 bn USD (5 Big Tech)", "n.a.", "~20 bn EUR (EU total)", ">15:1"],
         ["CACI Power Mode (April 2026)", "100", "15.7", "28.9", "3.46:1"],
     ]),
    ("Table 7. Timeline of American regulatory measures on semiconductors and AI (2022-2026).",
     "Sources: BIS, White House, Pillsbury Law (2026), Gibson Dunn (2026).",
     [
         ["Date", "Admin.", "Measure", "Scope / Target"],
         ["Oct. 2022", "Biden", "BIS export controls: advanced GPUs, SME, US persons", "China (military)"],
         ["Oct. 2023", "Biden", "Threshold reinforcement + 40-country extension + HBM", "China + 40 countries (D:1/D:4/D:5)"],
         ["Dec. 2024", "Biden", "Wave 3: 24 SME types, HBM, 140 entities, ECAD", "China (full chain)"],
         ["Jan. 2025", "Biden", "AI Diffusion Rule: model weights, cloud, 3-tier countries", "120+ countries (incl. EU effects)"],
         ["Jul. 2025", "Trump", "America AI Action Plan: deregulation, US compute", "US (domestic strategy)"],
         ["Sep. 2025", "Trump", "Announcement: China sales authorized against 25 pct revenue", "China (monetization)"],
         ["Jan. 2026", "Trump", "Section 232: 25 pct tariff advanced GPUs + BIS China license", "Global (US domestic exemption)"],
     ]),
]

EN.notes = [
    "IEA (2025), Energy and AI, Paris, IEA. https://www.iea.org/reports/energy-and-ai. The 415 TWh figure includes all data centers (cloud, enterprise, colocation), not just AI-dedicated data centers. The IEA notes that consumption has grown by 12 percent per year since 2017.",
    "IEA (2025), op. cit., chapter 'Energy demand from AI.' European Commission (November 2025), 'In Focus: Data Centres - An Energy-Hungry Challenge,' energy.ec.europa.eu. The 70 TWh figure for the EU is an IEA estimate in the absence of precise consumption data.",
    "IEA (2025), op. cit.; Carbon Brief (September 2025), 'AI: Five Charts that Put Data-Centre Energy Use - and Emissions - into Context.' The 35-50 percent range comes from a report prepared for the IEA (Kamiya and Coroama, 2025, IEA-4E).",
    "IEA (2025), op. cit. The IEA further notes that in its Headwinds scenario, global consumption would only reach 790 TWh, 40 percent less than the central scenario, illustrating the scale of uncertainty.",
    "Federal Reserve Board (October 2025), State of AI Competition in Advanced Economies. The report documents a significant correlation between industrial energy costs and AI adoption rates, particularly for large European companies. The PPA-adjusted reference values (USA 85, China 92, France 115, Germany 140, EU aggregate 135 USD/MWh) are documented on the public dashboard.",
    "IEA (2025), op. cit., chapter 'Case study: DeepSeek and efficiency gains.' The IEA concludes that algorithmic efficiency gains, while significant, tend to be absorbed by demand growth (Jevons paradox).",
    "SIA (February 2025, revised March 2025), 'Global Semiconductor Sales Increase 19.1 percent in 2024.' The initial figure of 627.6 billion USD was revised to 630.5 billion USD by the WSTS.",
    "SIA (February 2026), 'Global Annual Semiconductor Sales Increase 25.6 percent to 791.7 Billion USD in 2025.' Q4 2025 sales (236.6 billion USD) were 37.1 percent higher than Q4 2024, marking a notable acceleration.",
    "McKinsey (January 2026), 'Hiding in Plain Sight: The Semiconductor Industry Expanding Perimeter.' The gap with the SIA is explained by the inclusion of captive designers and vertically integrated fabless operators.",
    "SIA (July 2025), State of the U.S. Semiconductor Industry Report. The SIA notes that the Chips and Science Act has generated nearly 500 billion USD in announced private investment and is expected to create or support more than 500,000 jobs.",
    "Pilz, K.F., Rahman, R., Sanders, J. and Heim, L. (2025), 'Trends in AI Supercomputers,' arXiv:2504.16026. The dashboard refresh of April 2026 reaches 786 catalogued clusters, of which 609 are operational and 171 are planned, representing 10 to 20 percent of estimated global capacity.",
    "Sanchez, C. (2025), 'GeoCoded Special Report: State of Global AI Compute (2025 Edition),' Sanchez.vc. The May 2025 estimate cross-references Epoch AI and Georgetown University data and produced shares USA 74.5 / China 14.1 / EU 4.8. The April 2026 dashboard snapshot is based on the same data feed but with eleven additional months of cluster commissioning, which explains the small drift in shares.",
    "Pilz et al. (2025), op. cit. Leading AI supercomputer performance doubles every nine months, driven by a 1.6x per year increase in chip count and a 1.6x per year increase in per-chip performance.",
    "Epoch AI (January 2026), 'Global AI Power Capacity Is Now Comparable to Peak Power Usage of New York State.' The 30 GW estimate is based on quarterly AI chip sales x TDP x 2.5x infrastructure factor.",
    "BIS (October 7, 2022), 'Commerce Implements New Export Controls on Advanced Computing and Semiconductor Manufacturing Items to the People Republic of China,' Federal Register. See also GAO (2025), 'Export Controls: Commerce Implemented Advanced Semiconductor Rules,' GAO-25-107386.",
    "Skadden (October 2023), 'BIS Updates October 2022 Semiconductor Export Control Rules.' The new rules notably capture the A800 and H800, Nvidia chips specifically designed to fall just below the October 2022 thresholds.",
    "Carnegie Endowment (May 2025), Winter-Levy, H. and Phillips-Robins, A., 'What Is the AI Diffusion Rule?' The AI Diffusion Rule remained in effect only briefly before being partially replaced by the Trump approach, but it established the precedent of control over models and cloud, not just physical chips.",
    "White House (January 14, 2026), Presidential Proclamation 11002, 'Adjusting Imports of Semiconductors, Semiconductor Manufacturing Equipment, and Their Derivative Products into the United States.' The proclamation invokes Section 232 of the Trade Expansion Act of 1962 (19 U.S.C. 1862).",
    "Pillsbury Law (January 2026), 'Trump Admin Targets Advanced AI Semiconductors, Defers Broader Tariffs.' The legal analysis notes that the combination of Section 232 and the BIS rule of January 15, 2026 operationalizes the September 2025 announcement that the US government would collect 25 percent of AI chip sales to China.",
    "White House (January 14, 2026), Fact Sheet: 'President Donald J. Trump Takes Action on Certain Advanced Computing Chips to Protect America Economic and National Security.' Gibson Dunn (January 2026), 'The Trump Administration New Tariffs on and Export Licensing Requirements for Advanced Semiconductors.'",
]


# ===========================================================================
# Content - French
# ===========================================================================

FR = LangPack(
    code="FR",
    filename="Chapitre_III_Diagnostic_Empirique_FR.docx",
    cover_title="AI FOR AMERICANS FIRST",
    cover_subtitle="Protectionnisme IA, Energie et Semi-conducteurs : Trajectoires de divergence US/Europe 2024-2030",
    cover_blurb="Analyse geostrategique et economique integree - Chapitre III",
    cover_chip_lines=[
        "76,9 pct du compute IA operationnel mondial = USA",
        "1,59x cout energie EU/US",
        "3,46:1 ratio CACI US/EU (Power Mode)",
    ],
    cover_meta="Paris - fevrier 2026  |  7 chapitres  |  4 scenarios prospectifs  |  3 zones geographiques",
    cover_keywords_label="Mots-cles",
    cover_keywords=("intelligence artificielle, protectionnisme technologique, semi-conducteurs, "
                    "controles a l'exportation, compute souverain, geopolitique IA, France, "
                    "Etats-Unis, Chine"),
    chapter_label="CHAPITRE III",
    chapter_title="Diagnostic empirique 2020-2026",
    chapter_intro=(
        "Ce chapitre etablit la base factuelle de l'analyse. Il couvre quatre dimensions "
        "interdependantes : la trajectoire energetique des centres de donnees, l'evolution du "
        "marche des semi-conducteurs, la repartition geographique du compute IA installe et la "
        "chronologie des mesures reglementaires americaines. Les donnees presentees ici "
        "constituent les elements predetermines (au sens de Schwartz) qui structurent les "
        "scenarios prospectifs du chapitre V. Toutes les series temporelles sont sourcees, et "
        "lorsque les sources divergent, l'ecart est explicite. Tous les ratios impliquant le "
        "compute ont ete recalibres sur le snapshot d'avril 2026 du tableau de bord public."
    ),
    notes_label="Notes",
    license_block=[
        "Licence et avertissement. Cette oeuvre, 'AI for Americans First', est mise a disposition selon les termes de la licence Creative Commons Attribution - Pas d'Utilisation Commerciale - Partage dans les Memes Conditions 4.0 International (CC BY-NC-SA 4.0) du projet America-First-IA.",
        "Vous etes libre de partager et adapter le materiel a des fins non commerciales, a condition d'attribuer correctement l'oeuvre a Fabrice Pizzi (Universite Paris-Sorbonne) et de distribuer vos contributions sous la meme licence. Ce document est fourni a des fins educatives et de recherche uniquement.",
        "Tableau de bord public : https://mo0ogly.github.io/America-First-IA/dashboard/",
        "Depot : https://github.com/mo0ogly/America-First-IA",
    ],
    page_footer="AI for Americans First - Fabrice Pizzi - Chapitre III",
)

FR.sections = [
    ("3.1 Trajectoire energetique des centres de donnees : doublement de la demande en six ans", []),
    ("3.1.1 Consommation mondiale 2020-2024", [
        "L'Agence internationale de l'energie (AIE, avril 2025) estime la consommation electrique mondiale des centres de donnees a environ 415 TWh en 2024, soit 1,5 pour cent de la consommation electrique mondiale.[1] Cette consommation a cru en moyenne de 12 pour cent par an depuis 2017, soit un rythme quatre fois superieur a la croissance totale de la consommation d'electricite. En 2020, elle etait d'environ 270 TWh ; l'augmentation represente donc +54 pour cent en quatre ans.",
        "La distribution geographique est tres concentree. Les Etats-Unis absorbent 45 pour cent de la consommation mondiale des centres de donnees (environ 180 TWh en 2024), suivis de la Chine (25 pour cent, ~102 TWh) et de l'Europe (15 pour cent, ~70 TWh pour l'UE selon la Commission europeenne).[2] Les quatre principaux pays europeens (Allemagne, France, Royaume-Uni, Pays-Bas) totalisent environ 41 TWh, soit 10 pour cent du total mondial. La capacite mondiale installee de centres de donnees a atteint pres de 100 GW en 2024. Pres de la moitie de la capacite americaine est concentree dans cinq clusters regionaux.",
    ]),
    ("3.1.2 Projections 2024-2030 : le scenario central de l'AIE", [
        "Le scenario central de l'AIE (Base Case) projette une consommation mondiale de 945 TWh d'ici 2030, soit un doublement par rapport a 2024 et l'equivalent de la consommation electrique actuelle du Japon. L'IA est identifiee comme le principal moteur de cette croissance. La part de l'IA dans la consommation des centres de donnees, estimee a 5 a 15 pour cent ces dernieres annees, pourrait atteindre 35 a 50 pour cent d'ici 2030.[3]",
        "La croissance est geographiquement inegale. Les Etats-Unis ajoutent environ +240 TWh (+130 pour cent par rapport a 2024), la Chine +175 TWh (+170 pour cent), et l'Europe seulement +45 TWh (+70 pour cent). L'AIE note un risque de 20 pour cent de retard des projets en Europe, lie aux contraintes de raccordement au reseau et aux delais d'autorisation.[4] Aux Etats-Unis, les centres de donnees representeront pres de la moitie de la croissance de la demande electrique d'ici 2030 ; le pays consommera plus d'electricite pour ses centres de donnees que pour la production combinee d'aluminium, d'acier, de ciment et de produits chimiques.",
    ]),
    ("3.1.3 Le facteur energetique comme avantage competitif", [
        "L'asymetrie energetique constitue un avantage structurel pour les Etats-Unis. L'AIE note que le gaz naturel fournit actuellement 40 pour cent de l'electricite des centres de donnees americains, les renouvelables 24 pour cent, le nucleaire 20 pour cent et le charbon 15 pour cent. En Europe, le mix energetique est davantage contraint par les objectifs climatiques et des couts energetiques plus eleves. Les prix de l'electricite industrielle sont typiquement 2 a 3 fois plus eleves en Europe qu'aux Etats-Unis sur le tarif Eurostat non ajuste. Une fois pris en compte les Power Purchase Agreements negocies par les hyperscalers, l'ecart se reduit a environ 1,4 a 1,7 fois (1,59x pour l'agregat UE a la reference du tableau de bord d'avril 2026 : 135 USD/MWh dans l'UE contre 85 USD/MWh aux Etats-Unis). Comme le documente la Federal Reserve Board (octobre 2025), il existe une correlation negative significative entre couts energetiques et adoption de l'IA au niveau des entreprises europeennes, meme apres correction PPA.[5]",
        "La France dispose toutefois d'un atout specifique dans ce contexte : son mix electrique nucleaire (environ 65 a 70 pour cent de la production), qui offre une electricite relativement bon marche, decarbonee et disponible en base. RTE estime un besoin supplementaire de +10 GW pour les centres de donnees en France d'ici 2030, qui reste compatible avec la capacite installee sous reserve d'investissements dans le reseau de transport.",
    ]),
    ("3.1.4 Effet amplificateur : robotique IA et demande energetique industrielle", [
        "Un facteur encore peu quantifie dans la litterature est l'impact de la robotique IA sur la demande energetique. Les robots autonomes necessitent a la fois du compute embarque (edge computing) et du compute centralise (cloud IA pour l'entrainement et les mises a jour de modeles). Le deploiement massif de l'automatisation IA dans l'industrie pourrait ajouter 20 a 30 pour cent de demande energetique supplementaire dans les secteurs concernes, en plus de la croissance des centres de donnees. Ce facteur est integre comme variable de sensibilite dans nos scenarios (chapitre V), aucune estimation precise et fiable n'existant encore. L'AIE consacre neanmoins une section de son rapport a l'effet rebond de Jevons : meme avec des gains d'efficacite significatifs (illustres par le cas DeepSeek), la croissance de la demande absorbe les gains.[6]",
    ]),
    ("3.2 Marche des semi-conducteurs : une croissance tiree par l'IA", []),
    ("3.2.1 Trajectoire des ventes mondiales 2020-2025", [
        "Le marche mondial des semi-conducteurs a suivi une trajectoire remarquable depuis 2020. Apres un record de 555,9 milliards USD en 2022, le secteur s'est contracte de 8,2 pour cent en 2023 (526,8 milliards USD), avant un rebond spectaculaire en 2024 : les ventes mondiales ont atteint 630,5 milliards USD (donnees SIA/WSTS revisees, mars 2025), une hausse de 19,1 pour cent en glissement annuel et un nouveau record absolu.[7]",
        "L'annee 2025 a confirme l'acceleration. La SIA a annonce en fevrier 2026 des ventes mondiales de 791,7 milliards USD en 2025, une hausse de 25,6 pour cent par rapport a 2024.[8] Le quatrieme trimestre 2025 (236,6 milliards USD) etait 37,1 pour cent plus eleve que le meme trimestre en 2024. La SIA projette desormais des ventes proches de 1 000 milliards USD en 2026, un seuil symbolique qui semblait hors de portee il y a deux ans.",
        "Deux segments tirent cette croissance. Les puces logiques (processeurs, GPU, ASIC) ont atteint 301,9 milliards USD en 2025 (+39,9 pour cent), devenant la premiere categorie par les ventes. La memoire (DRAM, NAND) a atteint 223,1 milliards USD (+34,8 pour cent). Ensemble, ces deux segments representent plus de 66 pour cent des ventes mondiales et refletent directement la demande d'infrastructure IA : serveurs, centres de donnees et HPC (High Performance Computing).",
    ]),
    ("3.2.2 La question du perimetre : SIA vs McKinsey", [
        "Comme note au chapitre II, un ecart methodologique notable existe entre les donnees SIA/WSTS et les estimations McKinsey. La SIA enregistre les ventes de semi-conducteurs au sens strict (composants vendus par les fabricants). McKinsey (janvier 2026) adopte un perimetre elargi incluant la valeur des designers captifs (Apple, Amazon, Tesla, Google), c'est-a-dire la valeur economique des puces concues en interne mais fabriquees par des fonderies sous contrat. Sous ce perimetre elargi, McKinsey valorise le marche a environ 775 milliards USD en 2024 et projette 1,5 a 1,8 trillions USD d'ici 2030.[9] Le perimetre elargi est plus pertinent pour notre analyse car il capture la pleine valeur de la chaine IA, y compris les investissements de design des hyperscalers.",
    ]),
    ("3.2.3 Concentration regionale de la capacite de fabrication", [
        "La distribution geographique de la capacite de fonderie constitue un facteur cle de vulnerabilite. Taiwan (TSMC) concentre l'essentiel de la production de pointe (noeuds sub-7 nm). Les Etats-Unis representent environ 10 pour cent de la capacite installee mondiale mais croissent rapidement (le Chips Act a genere pres de 500 milliards USD d'investissement prive annonce, avec un objectif de tripler la capacite americaine d'ici 2032).[10] L'UE represente environ 8 pour cent de la capacite mondiale, un chiffre que l'EU Chips Act (43 milliards EUR) vise a porter a 20 pour cent d'ici 2030, un objectif juge irrealiste par de nombreux observateurs. La Chine, malgre les restrictions, progresse de 21 pour cent vers un objectif de 30 pour cent de capacite installee d'ici 2030, principalement sur les noeuds matures (28 nm et au-dessus).",
    ]),
    ("3.3 Repartition geographique du compute IA installe", []),
    ("3.3.1 Domination americaine : trois quarts du compute operationnel", [
        "La donnee la plus significative pour notre analyse est la repartition geographique de la capacite de calcul dediee a l'IA. Les travaux d'Epoch AI (Pilz et al., avril 2025) fournissent le jeu de donnees fondamental, rafraichi mensuellement sur le tableau de bord public.[11] A avril 2026, le tableau de bord catalogue 786 clusters, dont 609 operationnels et 171 planifies ou annonces.",
        "Le resultat est sans ambiguite. Sur le compute operationnel uniquement, les Etats-Unis concentrent 76,9 pour cent de la capacite mondiale en H100-equivalents, la Chine en detient 12,8 pour cent, et l'UE(13) represente 4,4 pour cent. Les autres pays se partagent environ 6 pour cent. Le ratio US/UE en H100-equivalents operationnels est donc de 17,6:1, un ecart de magnitude considerable qui s'est legerement creuse depuis le snapshot GeoCoded/Sanchez de mai 2025 (74,5 / 14,1 / 4,8 pour cent).[12]",
        "Deux dynamiques expliquent cette concentration. Premierement, la part du secteur prive dans le compute IA mondial est passee de 40 pour cent en 2019 a 80 pour cent en 2025, et les entreprises technologiques investissant massivement dans l'IA sont quasi exclusivement americaines (Microsoft, Meta, Google, Amazon, xAI). Deuxiemement, la taille des clusters a explose : les systemes de plus de 10 000 puces etaient rares en 2019 ; en 2024, xAI a deploye un cluster Colossus de 200 000 GPU. La performance des supercalculateurs IA de pointe double tous les neuf mois.[13]",
        "Le tableau change substantiellement lorsque la capacite planifiee est incluse. La vue F_total du tableau de bord, qui compte a la fois les clusters operationnels et annonces, montre la part US tomber a 49,9 pour cent. Le decalage n'est pas du a des acteurs etrangers construisant une capacite independante, mais aux hyperscalers americains delocalisant vers des juridictions alliees : le campus annonce de 5 GW aux EAU et le Stargate UAE Phase 2 representent 22,9 millions de H100-equivalents, l'Arabie saoudite 7,2 millions, et la Coree du Sud 5,3 millions, tous largement detenus ou operes par des entreprises americaines. Cette distinction entre CACI Physique (ou se trouve le cluster) et CACI Souverain (qui le controle) est developpee au chapitre V. La part chinoise s'effondre a 0,5 pour cent en F_total, un artefact de la methodologie Epoch AI plutot qu'un veritable changement de position : le jeu de donnees anonymise les systemes chinois et les arrondit aggressivement, de sorte que la capacite chinoise annoncee est structurellement sous-enregistree. La part UE(13) en F_total descend egalement a 3,3 pour cent car le campus Fluidstack 1 GW en France tire le total vers le haut tandis que le reste de l'UE compte peu d'annonces majeures.",
    ]),
    ("3.3.2 Puissance electrique totale du parc de puces IA", [
        "Epoch AI (janvier 2026) estime la capacite electrique totale du parc mondial de puces IA a environ 30 GW fin 2025, comparable a la consommation de pointe de l'Etat de New York.[14] Cette estimation repose sur les ventes trimestrielles de puces IA des principaux fabricants (Nvidia, AMD, Google TPU), multipliees par leur puissance nominale (TDP) et un facteur 2,5x pour l'infrastructure des centres de donnees. La production mondiale de puces IA double tous les sept mois, un rythme qui depasse toutes les previsions anterieures. Les cinq plus grands investisseurs americains en serveurs IA (Microsoft, Google, Meta, Amazon, xAI) ont annonce 320 milliards USD d'investissements en 2025, contre 230 milliards USD en 2024.",
    ]),
    ("3.3.3 Implications pour la metrique CACI", [
        "Ces donnees permettent une calibration precise du CACI defini au chapitre II. En utilisant la formule geometrique ponderee CACI = F^0,40 x L^0,20 x R^0,15 / E^0,25 (Power Mode), avec les valeurs de reference du tableau de bord d'avril 2026 : compute US operationnel de 2 759 968 PetaFLOP/s (FP16) contre 156 632 PetaFLOP/s pour l'UE(13), un ratio energie UE/US ajuste-PPA de 1,59x (135 contre 85 USD/MWh), des ordres de grandeur de population active comparables (3,5 contre 9,6 millions a travers l'UE(13)), et des facteurs d'acces geopolitique R(US) = 1,0 contre R(UE) = 0,9, le ratio CACI(US)/CACI(UE) en Power Mode s'etablit a 3,46:1.",
        "L'ecart est substantiellement attenue par rapport a l'avantage brut de 17,6:1 en H100-equivalents car les exposants geometriques (F a 0,40 ; L a 0,20 ; R a 0,15 ; E a 0,25) appliquent des rendements decroissants a la massive avance US sur le compute et recompensent partiellement l'UE sur le travail, l'acces geopolitique et des multiplicateurs energetiques plus faibles. En Intensity Mode (en divisant additionnellement par le PIB), l'ecart se resserre encore a environ 1,4:1, raison pour laquelle le ratio principal utilise dans toute cette etude est le ratio Power Mode de 3,46:1. Le chapitre V projette l'evolution de ce ratio dans chaque scenario.",
    ]),
    ("3.4 Chronologie des mesures reglementaires americaines (2022-2026)", [
        "La sequence des mesures americaines de controle des semi-conducteurs et de l'IA constitue le fil directeur de notre hypothese. Quatre phases sont identifiables, marquant une escalade progressive dans l'etendue et l'intensite des restrictions.",
    ]),
    ("Phase 1 - Le choc initial (7 octobre 2022)", [
        "Le BIS (Bureau of Industry and Security) du departement du Commerce publie une regle interimaire definitive qui transforme radicalement les controles a l'exportation americains sur les semi-conducteurs. Les mesures couvrent trois composantes : (i) controles sur les puces de calcul avancees (GPU au-dessus des seuils de performance definis par TTP - Total Processing Performance), (ii) controles sur les equipements de fabrication de semi-conducteurs (SME, y compris les lithographies EUV et DUV), et (iii) restrictions sur les activites de US persons soutenant la fabrication de puces avancees en Chine.[15] La cible est explicitement la Chine : l'objectif annonce est d'empecher la modernisation militaire chinoise via l'acces au compute IA. Les restrictions incluent trois nouvelles regles Foreign Direct Product (FDP) qui etendent la juridiction americaine aux produits fabriques hors des Etats-Unis si des technologies americaines sont utilisees dans leur production.",
    ]),
    ("Phase 2 - Combler les failles (octobre 2023)", [
        "Le BIS publie deux nouvelles regles interimaires qui renforcent et elargissent les controles d'octobre 2022. Les seuils techniques sont ajustes pour capturer les puces Nvidia specifiquement concues pour contourner les restrictions (A800, H800). Le perimetre geographique est etendu a environ 40 pays supplementaires (Country Groups D:1, D:4 et D:5), avec un regime de licences differencie par categorie de pays.[16] Les controles sur les equipements de fabrication sont approfondis. C'est durant cette phase que les acteurs europeens commencent a percevoir les effets indirects des restrictions, meme si l'UE n'est pas la cible principale.",
    ]),
    ("Phase 3 - L'AI Diffusion Rule (janvier 2025, Biden)", [
        "L'administration Biden publie en janvier 2025 l'AI Diffusion Rule, qui represente un changement qualitatif. Pour la premiere fois, les restrictions s'appliquent non seulement aux puces physiques mais aussi aux poids des modeles IA et a l'acces au compute cloud. La regle classifie 120 pays en trois categories : (i) allies de confiance (acces large), (ii) pays intermediaires (quotas), (iii) pays embargoes.[17] Les reactions europeennes sont fortes : le Parlement europeen alerte sur les restrictions menacant la capacite de l'UE a entrainer des modeles sur ses AI Factories. La France et l'Allemagne sont classees comme partenaires de confiance, mais d'autres Etats membres font face a des plafonds sur les volumes de GPU importables.",
    ]),
    ("Phase 4 - La rupture Trump : Section 232 et protectionnisme explicite (janvier 2026)", [
        "Le 14 janvier 2026, le president Trump signe la Proclamation 11002, invoquant la Section 232 du Trade Expansion Act de 1962.[18] Cette action marque une rupture qualitative avec les phases precedentes, pour trois raisons.",
        "Premierement, l'instrument juridique change. Les phases 1 a 3 relevaient des controles a l'exportation (regulation des exportations, logique defensive de securite nationale). La Section 232 est un instrument tarifaire (droits a l'importation), dont la logique est protectionniste : elle vise a proteger la production domestique, et non simplement a restreindre l'acces d'un adversaire.",
        "Deuxiemement, le tarif cree un avantage competitif explicite pour les entreprises americaines. Le tarif de 25 pour cent frappe les GPU avances (H200, MI325X) importes sauf s'ils sont destines a des usages domestiques americains : centres de donnees US, R&D, startups, secteur public, applications industrielles et grand public hors centres de donnees. Concretement, une entreprise americaine utilisant ces puces sur le sol americain ne paie pas le tarif ; une entreprise etrangere important les memes puces pour reexportation vers la Chine paie 25 pour cent.[19]",
        "Troisiemement, la proclamation annonce une expansion imminente. Le texte prevoit que le secretaire au Commerce et l'USTR negocient dans les 90 jours avec les pays producteurs de semi-conducteurs, et que des tarifs plus larges, accompagnes d'un programme de tariff offset pour les entreprises investissant dans la production americaine, pourraient etre imposes. Le secretaire au Commerce doit fournir d'ici juillet 2026 un rapport sur le marche des semi-conducteurs utilises dans les centres de donnees americains.[20]",
    ]),
    ("3.4.5 Interpretation : du controle defensif au protectionnisme offensif", [
        "La sequence 2022-2026 revele une transformation qualitative de la politique americaine. Les phases Biden (2022-2025) suivent une logique de strategie de denial : empecher un adversaire designe (la Chine) d'acceder a des technologies cles, dans un cadre multilateral (coordination avec le Japon, les Pays-Bas). La phase Trump (2025-2026) ajoute une logique de strategie de capture : generer des revenus (tarif de 25 pour cent), prioriser les entreprises americaines (exemptions domestiques) et utiliser l'acces au compute comme levier de negociation avec des pays tiers.",
        "C'est precisement la transformation qui valide l'hypothese centrale de notre etude : le protectionnisme technologique americain ne se limite plus a denier l'acces a un adversaire ; il construit activement un avantage competitif structurel pour les entreprises americaines. Les exemptions domestiques de la Section 232 signifient que, pour la premiere fois, le cout d'acces au compute de pointe est legalement differencie par la nationalite de l'utilisateur. Meme si les effets directs sur l'UE restent limites en janvier 2026 (le tarif cible principalement la reexportation vers la Chine), la proclamation ouvre la voie a une expansion qui pourrait directement affecter l'Europe, precisement le territoire des scenarios B, C et D de notre analyse.",
    ]),
    ("3.5 Synthese diagnostique : les elements predetermines de 2026", [
        "Les quatre dimensions du diagnostic convergent vers un tableau coherent qui structure les scenarios du chapitre V.",
        "(1) La demande de compute IA croit de maniere exponentielle (doublement des ventes de puces en deux ans, doublement projete de la consommation energetique sur six ans), et cette croissance ne montre aucun signe de ralentissement.",
        "(2) Cette demande est structurellement concentree aux Etats-Unis. Le snapshot du tableau de bord d'avril 2026 chiffre la part US a 76,9 pour cent du compute IA operationnel en H100-equivalents et a 49,9 pour cent en incluant les clusters annonces ; 45 pour cent de la consommation electrique des centres de donnees ; et plus de 80 pour cent des investissements prives dans l'IA de frontiere. La concentration sur le compute operationnel s'accroit au fil du temps.",
        "(3) L'Europe part d'une position structurellement deficitaire : environ 4,4 pour cent du compute operationnel en H100-equivalents, environ 15 pour cent de la consommation des centres de donnees, des couts energetiques 1,4 a 1,7 fois superieurs apres correction PPA (et 2 a 3 fois superieurs sur les tarifs Eurostat industriels non ajustes), et 72 a 80 pour cent de dependance aux hyperscalers US pour les charges IA cloud. Les plans d'investissement de l'UE (Chips Act, AI Factories, SNIA) portent des budgets au moins un ordre de grandeur inferieurs aux annonces des Big Tech 2025 (320 milliards USD pour les cinq seuls investisseurs US frontier-AI).",
        "(4) Le cadre reglementaire americain a franchi un seuil qualitatif en janvier 2026 avec le passage des controles a l'exportation aux tarifs Section 232, creant un mecanisme legal de differenciation des couts d'acces au compute par nationalite. La proclamation signale explicitement la possibilite d'une expansion, dont la portee et le calendrier dependront de variables politiques, precisement les incertitudes critiques qu'explorent nos scenarios.",
        "Le chapitre suivant (chapitre IV) analyse les mecanismes par lesquels cette asymetrie de compute se traduit en avantage competitif mesurable pour les entreprises americaines.",
    ]),
]

FR.table_blocks = [
    ("Tableau 4. Consommation electrique des centres de donnees par region, 2020-2030.",
     "Source : AIE (2025), Energy and AI.",
     [
         ["Region", "2020", "2024", "2030 (AIE)", "Delta 24-30", "Part mondiale 2024"],
         ["Etats-Unis", "~120 TWh", "~180 TWh", "~420 TWh", "+240 TWh", "45 pct"],
         ["Chine", "~60 TWh", "~102 TWh", "~280 TWh", "+175 TWh", "25 pct"],
         ["Europe (UE)", "~45 TWh", "~70 TWh", "~115 TWh", "+45 TWh", "15 pct"],
         ["Monde", "~270 TWh", "~415 TWh", "~945 TWh", "+530 TWh", "100 pct"],
     ]),
    ("Tableau 5. Ventes mondiales de semi-conducteurs, 2020-2026.",
     "Sources : SIA/WSTS (fev. 2025, fev. 2026). La projection 2026 est basee sur la prevision automne 2025 du WSTS (975,4 milliards USD) et la declaration SIA de fevrier 2026.",
     [
         ["Annee", "2020", "2022", "2023", "2024", "2025", "2026 (proj.)"],
         ["Ventes (milliards USD, SIA)", "440", "556", "527", "631", "792", "~975-1000"],
         ["Croissance (pct)", "+6,8", "+3,3", "-8,2", "+19,1", "+25,6", "+23-26"],
     ]),
    ("Tableau 6. Indicateurs de domination US dans le compute IA (snapshot tableau de bord avril 2026).",
     "Sources : Epoch AI GPU Clusters dataset avril 2026 ; AIE (2025) ; Sanchez/GeoCoded (2025).",
     [
         ["Indicateur", "Etats-Unis", "Chine", "UE(13)", "Ratio US/UE"],
         ["Part H100-eq operationnel", "76,9 pct", "12,8 pct", "4,4 pct", "17,6:1"],
         ["Part F_total (op + planifie)", "49,9 pct", "0,5 pct", "3,3 pct", "15,1:1"],
         ["Part secteur prive (2025)", "~65 pct mondial", "~12 pct", "~3 pct", "~22:1"],
         ["Conso centres de donnees 2024", "180 TWh", "102 TWh", "70 TWh", "2,6:1"],
         ["Investissement IA 2025", "320 mds USD (5 Big Tech)", "n.d.", "~20 mds EUR (UE total)", ">15:1"],
         ["CACI Power Mode (avril 2026)", "100", "15,7", "28,9", "3,46:1"],
     ]),
    ("Tableau 7. Chronologie des mesures reglementaires americaines sur les semi-conducteurs et l'IA (2022-2026).",
     "Sources : BIS, Maison-Blanche, Pillsbury Law (2026), Gibson Dunn (2026).",
     [
         ["Date", "Admin.", "Mesure", "Perimetre / Cible"],
         ["Oct. 2022", "Biden", "Controles export BIS : GPU avances, SME, US persons", "Chine (militaire)"],
         ["Oct. 2023", "Biden", "Renforcement seuils + extension 40 pays + HBM", "Chine + 40 pays (D:1/D:4/D:5)"],
         ["Dec. 2024", "Biden", "Vague 3 : 24 types SME, HBM, 140 entites, ECAD", "Chine (chaine complete)"],
         ["Jan. 2025", "Biden", "AI Diffusion Rule : poids modeles, cloud, 3 paliers pays", "120+ pays (effets UE inclus)"],
         ["Jul. 2025", "Trump", "America's AI Action Plan : dereglementation, US compute", "USA (strategie domestique)"],
         ["Sep. 2025", "Trump", "Annonce : ventes Chine autorisees contre 25 pct revenus", "Chine (monetisation)"],
         ["Jan. 2026", "Trump", "Section 232 : tarif 25 pct GPU avances + licence BIS Chine", "Mondial (exemption domestique US)"],
     ]),
]

FR.notes = EN.notes


# ===========================================================================
# Content - Brazilian Portuguese
# ===========================================================================

PT = LangPack(
    code="PT-BR",
    filename="Capitulo_III_Diagnostico_Empirico_PT-BR.docx",
    cover_title="AI FOR AMERICANS FIRST",
    cover_subtitle="Protecionismo de IA, Energia e Semicondutores: Trajetorias de Divergencia EUA/Europa 2024-2030",
    cover_blurb="Analise geoestrategica e economica integrada - Capitulo III",
    cover_chip_lines=[
        "76,9 pct do compute IA operacional mundial = EUA",
        "1,59x custo de energia UE/EUA",
        "3,46:1 razao CACI EUA/UE (Power Mode)",
    ],
    cover_meta="Paris - fevereiro de 2026  |  7 capitulos  |  4 cenarios prospectivos  |  3 zonas geograficas",
    cover_keywords_label="Palavras-chave",
    cover_keywords=("inteligencia artificial, protecionismo tecnologico, semicondutores, "
                    "controles de exportacao, compute soberano, geopolitica da IA, Franca, "
                    "Estados Unidos, China"),
    chapter_label="CAPITULO III",
    chapter_title="Diagnostico empirico 2020-2026",
    chapter_intro=(
        "Este capitulo estabelece a base factual da analise. Cobre quatro dimensoes "
        "interdependentes: a trajetoria energetica dos data centers, a evolucao do mercado de "
        "semicondutores, a distribuicao geografica do compute IA instalado e a cronologia das "
        "medidas regulatorias americanas. Os dados aqui apresentados constituem os elementos "
        "predeterminados (no sentido de Schwartz) que estruturam os cenarios prospectivos do "
        "Capitulo V. Todas as series temporais sao referenciadas e, quando as fontes divergem, "
        "a discrepancia e explicitada. Todas as razoes envolvendo compute foram recalibradas "
        "no snapshot de abril de 2026 do painel publico."
    ),
    notes_label="Notas",
    license_block=[
        "Licenca e isencao de responsabilidade. Esta obra, 'AI for Americans First', e disponibilizada nos termos da Licenca Creative Commons Atribuicao - NaoComercial - CompartilhaIgual 4.0 Internacional (CC BY-NC-SA 4.0) do projeto America-First-IA.",
        "Voce e livre para compartilhar e adaptar o material para fins nao comerciais, desde que credite adequadamente a obra a Fabrice Pizzi (Universite Paris-Sorbonne) e distribua suas contribuicoes sob a mesma licenca. Este documento e fornecido apenas para fins educacionais e de pesquisa.",
        "Painel publico: https://mo0ogly.github.io/America-First-IA/dashboard/",
        "Repositorio: https://github.com/mo0ogly/America-First-IA",
    ],
    page_footer="AI for Americans First - Fabrice Pizzi - Capitulo III",
)

PT.sections = [
    ("3.1 Trajetoria energetica dos data centers: dobra da demanda em seis anos", []),
    ("3.1.1 Consumo global 2020-2024", [
        "A Agencia Internacional de Energia (AIE, abril de 2025) estima o consumo eletrico global dos data centers em aproximadamente 415 TWh em 2024, representando 1,5 por cento do consumo eletrico mundial.[1] Esse consumo cresceu em media 12 por cento ao ano desde 2017, ritmo quatro vezes superior ao crescimento total do consumo de eletricidade. Em 2020, era de aproximadamente 270 TWh; o aumento representa portanto +54 por cento em quatro anos.",
        "A distribuicao geografica e altamente concentrada. Os Estados Unidos absorvem 45 por cento do consumo global dos data centers (aproximadamente 180 TWh em 2024), seguidos pela China (25 por cento, ~102 TWh) e pela Europa (15 por cento, ~70 TWh para a UE segundo a Comissao Europeia).[2] Os quatro principais paises europeus (Alemanha, Franca, Reino Unido, Holanda) totalizam aproximadamente 41 TWh, ou 10 por cento do total mundial. A capacidade total instalada global de data centers atingiu cerca de 100 GW em 2024. Quase metade da capacidade dos EUA esta concentrada em apenas cinco clusters regionais.",
    ]),
    ("3.1.2 Projecoes 2024-2030: o cenario central da AIE", [
        "O cenario central da AIE (Base Case) projeta consumo global de 945 TWh ate 2030, representando dobra em relacao a 2024 e equivalente ao atual consumo eletrico do Japao. A IA e identificada como o motor mais importante desse crescimento. A parcela da IA no consumo dos data centers, estimada em 5 a 15 por cento nos ultimos anos, podera atingir 35 a 50 por cento ate 2030.[3]",
        "O crescimento e geograficamente desigual. Os Estados Unidos adicionam aproximadamente +240 TWh (+130 por cento em relacao a 2024), a China +175 TWh (+170 por cento), e a Europa apenas +45 TWh (+70 por cento). A AIE observa um risco de 20 por cento de atrasos de projetos na Europa, ligado a restricoes de conexao a rede e prazos de licenciamento.[4] Nos Estados Unidos, os data centers responderao por quase metade do crescimento da demanda eletrica ate 2030; o pais consumira mais eletricidade para seus data centers do que para a producao combinada de aluminio, aco, cimento e produtos quimicos.",
    ]),
    ("3.1.3 O fator energetico como vantagem competitiva", [
        "A assimetria energetica constitui uma vantagem estrutural para os Estados Unidos. A AIE observa que o gas natural fornece atualmente 40 por cento da eletricidade dos data centers americanos, as renovaveis 24 por cento, o nuclear 20 por cento e o carvao 15 por cento. Na Europa, o mix energetico e mais restringido por objetivos climaticos e custos de energia mais altos. Os precos da eletricidade industrial sao tipicamente 2 a 3 vezes mais altos na Europa do que nos Estados Unidos na tarifa Eurostat nao ajustada. Uma vez considerados os Power Purchase Agreements negociados pelos hyperscalers, a diferenca se reduz para cerca de 1,4 a 1,7 vezes (1,59x para o agregado UE na referencia do painel de abril de 2026: 135 USD/MWh na UE contra 85 USD/MWh nos Estados Unidos). Como documentado pelo Federal Reserve Board (outubro de 2025), existe uma correlacao negativa significativa entre custos energeticos e adocao de IA no nivel das empresas europeias, mesmo apos correcao PPA.[5]",
        "A Franca, no entanto, possui um ativo especifico nesse contexto: seu mix eletrico nuclear (aproximadamente 65 a 70 por cento da producao), que oferece eletricidade relativamente barata, descarbonizada e disponivel em base. A RTE estima uma necessidade adicional de +10 GW para data centers na Franca ate 2030, que permanece compativel com a capacidade instalada sujeita a investimentos na rede de transmissao.",
    ]),
    ("3.1.4 Efeito amplificador: robotica de IA e demanda energetica industrial", [
        "Um fator ainda pouco quantificado na literatura e o impacto da robotica de IA na demanda energetica. Os robos autonomos requerem tanto compute embarcado (edge computing) quanto compute centralizado (nuvem IA para treinamento e atualizacao de modelos). A implantacao em larga escala da automacao IA na industria poderia adicionar 20 a 30 por cento de demanda energetica adicional nos setores afetados, alem do crescimento dos data centers. Esse fator e integrado como variavel de sensibilidade em nossos cenarios (Capitulo V), pois ainda nao existe estimativa precisa e confiavel. A AIE dedica, no entanto, uma secao de seu relatorio ao efeito rebote de Jevons: mesmo com ganhos significativos de eficiencia (ilustrados pelo caso DeepSeek), o crescimento da demanda absorve os ganhos.[6]",
    ]),
    ("3.2 Mercado de semicondutores: crescimento impulsionado pela IA", []),
    ("3.2.1 Trajetoria de vendas globais 2020-2025", [
        "O mercado global de semicondutores seguiu uma trajetoria notavel desde 2020. Apos um recorde de 555,9 bilhoes USD em 2022, o setor se contraiu 8,2 por cento em 2023 (526,8 bilhoes USD), antes de uma recuperacao espetacular em 2024: as vendas globais atingiram 630,5 bilhoes USD (dados SIA/WSTS revisados, marco de 2025), aumento de 19,1 por cento ano a ano e novo recorde absoluto.[7]",
        "O ano de 2025 confirmou a aceleracao. A SIA anunciou em fevereiro de 2026 vendas globais de 791,7 bilhoes USD em 2025, aumento de 25,6 por cento em relacao a 2024.[8] O quarto trimestre de 2025 (236,6 bilhoes USD) foi 37,1 por cento superior ao mesmo trimestre de 2024. A SIA agora projeta vendas proximas de 1 trilhao USD em 2026, um limiar simbolico que parecia inatingivel ha dois anos.",
        "Dois segmentos impulsionam esse crescimento. Os chips logicos (processadores, GPUs, ASICs) atingiram 301,9 bilhoes USD em 2025 (+39,9 por cento), tornando-se a principal categoria por vendas. A memoria (DRAM, NAND) atingiu 223,1 bilhoes USD (+34,8 por cento). Juntos, esses dois segmentos representam mais de 66 por cento das vendas globais e refletem diretamente a demanda por infraestrutura de IA: servidores, data centers e HPC (High Performance Computing).",
    ]),
    ("3.2.2 A questao do escopo: SIA vs McKinsey", [
        "Como observado no Capitulo II, existe uma diferenca metodologica notavel entre os dados SIA/WSTS e as estimativas McKinsey. A SIA registra as vendas de semicondutores em sentido estrito (componentes vendidos pelos fabricantes). A McKinsey (janeiro de 2026) adota um escopo expandido incluindo o valor dos designers cativos (Apple, Amazon, Tesla, Google), ou seja, o valor economico dos chips projetados internamente mas fabricados por fundicoes contratadas. Sob esse escopo expandido, a McKinsey avalia o mercado em aproximadamente 775 bilhoes USD em 2024 e projeta 1,5 a 1,8 trilhoes USD ate 2030.[9] O escopo expandido e mais relevante para a nossa analise pois captura o pleno valor da cadeia de IA, incluindo investimentos de design dos hyperscalers.",
    ]),
    ("3.2.3 Concentracao regional da capacidade de fabricacao", [
        "A distribuicao geografica da capacidade de fundicao constitui um fator chave de vulnerabilidade. Taiwan (TSMC) concentra a maior parte da producao de ponta (nos sub-7 nm). Os Estados Unidos representam aproximadamente 10 por cento da capacidade global instalada mas crescem rapidamente (o Chips Act gerou cerca de 500 bilhoes USD em investimento privado anunciado, com objetivo de triplicar a capacidade dos EUA ate 2032).[10] A UE representa aproximadamente 8 por cento da capacidade global, numero que o Chips Act europeu (43 bilhoes EUR) visa elevar a 20 por cento ate 2030, objetivo considerado irrealista por muitos observadores. A China, apesar das restricoes, progride de 21 por cento em direcao a uma meta de 30 por cento de capacidade instalada ate 2030, principalmente em nos maduros (28 nm e acima).",
    ]),
    ("3.3 Distribuicao geografica do compute IA instalado", []),
    ("3.3.1 Dominancia americana: tres quartos do compute operacional", [
        "O dado mais significativo para nossa analise e a distribuicao geografica da capacidade de computacao dedicada a IA. O trabalho da Epoch AI (Pilz et al., abril de 2025) fornece o conjunto de dados fundamental, atualizado mensalmente no painel publico.[11] Em abril de 2026, o painel cataloga 786 clusters, dos quais 609 sao operacionais e 171 estao planejados ou anunciados.",
        "O resultado e inequivoco. Apenas no compute operacional, os Estados Unidos concentram 76,9 por cento da capacidade global em H100-equivalentes, a China detem 12,8 por cento, e a UE(13) representa 4,4 por cento. Os paises restantes somam aproximadamente 6 por cento. A razao operacional EUA/UE em H100-equivalentes e portanto 17,6:1, uma diferenca de magnitude consideravel que se ampliou ligeiramente desde o snapshot GeoCoded/Sanchez de maio de 2025 (74,5 / 14,1 / 4,8 por cento).[12]",
        "Duas dinamicas explicam essa concentracao. Primeiro, a parcela do setor privado no compute IA global passou de 40 por cento em 2019 para 80 por cento em 2025, e as empresas de tecnologia que investem massivamente em IA sao quase exclusivamente americanas (Microsoft, Meta, Google, Amazon, xAI). Segundo, o tamanho dos clusters explodiu: sistemas com mais de 10.000 chips eram raros em 2019; em 2024, a xAI implantou um cluster Colossus de 200.000 GPUs. O desempenho dos supercomputadores IA de ponta dobra a cada nove meses.[13]",
        "O quadro muda substancialmente quando a capacidade planejada e incluida. A visao F_total do painel, que conta tanto os clusters operacionais quanto os anunciados, mostra a parcela dos EUA caindo para 49,9 por cento. A diferenca nao e impulsionada por atores estrangeiros construindo capacidade independente, mas pelos hyperscalers americanos relocalizando-se para jurisdicoes aliadas: o campus anunciado de 5 GW dos EAU e o Stargate UAE Phase 2 representam 22,9 milhoes de H100-equivalentes, a Arabia Saudita 7,2 milhoes e a Coreia do Sul 5,3 milhoes, todos majoritariamente detidos ou operados por empresas americanas. Essa distincao entre CACI Fisico (onde o cluster esta) e CACI Soberano (quem o controla) e desenvolvida no Capitulo V. A parcela chinesa colapsa para 0,5 por cento em F_total, um artefato da metodologia Epoch AI em vez de uma verdadeira mudanca de posicao: o conjunto de dados anonimiza os sistemas chineses e os arredonda agressivamente, de modo que a capacidade chinesa anunciada e estruturalmente sub-registrada. A parcela da UE(13) em F_total tambem cai para 3,3 por cento porque o campus Fluidstack 1 GW na Franca puxa o total para cima enquanto o resto da UE tem poucos anuncios significativos.",
    ]),
    ("3.3.2 Potencia eletrica total da frota de chips IA", [
        "A Epoch AI (janeiro de 2026) estima a capacidade eletrica total da frota global de chips IA em aproximadamente 30 GW no fim de 2025, comparavel ao consumo de pico do estado de Nova York.[14] Essa estimativa baseia-se nas vendas trimestrais de chips IA dos principais fabricantes (Nvidia, AMD, Google TPU), multiplicadas por sua potencia nominal (TDP) e um fator 2,5x para a infraestrutura dos data centers. A producao global de chips IA dobra a cada sete meses, ritmo que excede todas as previsoes anteriores. Os cinco maiores investidores americanos em servidores IA (Microsoft, Google, Meta, Amazon, xAI) anunciaram 320 bilhoes USD em investimentos em 2025, contra 230 bilhoes USD em 2024.",
    ]),
    ("3.3.3 Implicacoes para a metrica CACI", [
        "Esses dados permitem uma calibracao precisa do CACI definido no Capitulo II. Usando a formula geometrica ponderada CACI = F^0,40 x L^0,20 x R^0,15 / E^0,25 (Power Mode), com os valores de referencia do painel de abril de 2026: compute operacional dos EUA de 2.759.968 PetaFLOP/s (FP16) contra 156.632 PetaFLOP/s para a UE(13), uma razao energia UE/EUA ajustada-PPA de 1,59x (135 contra 85 USD/MWh), ordens de grandeza de forca de trabalho comparaveis (3,5 contra 9,6 milhoes na UE(13)) e fatores de acesso geopolitico R(EUA) = 1,0 contra R(UE) = 0,9, a razao CACI(EUA)/CACI(UE) em Power Mode estabelece-se em 3,46:1.",
        "A diferenca e substancialmente atenuada em relacao a vantagem bruta de 17,6:1 em H100-equivalentes porque os expoentes geometricos (F em 0,40; L em 0,20; R em 0,15; E em 0,25) aplicam rendimentos decrescentes a massiva lideranca dos EUA em compute e recompensam parcialmente a UE em trabalho, acesso geopolitico e multiplicadores energeticos menores. Em Intensity Mode (dividindo adicionalmente pelo PIB), a diferenca se reduz ainda mais para aproximadamente 1,4:1, razao pela qual a razao principal usada em todo este estudo e a razao Power Mode de 3,46:1. O Capitulo V projeta a evolucao dessa razao em cada cenario.",
    ]),
    ("3.4 Cronologia das medidas regulatorias americanas (2022-2026)", [
        "A sequencia das medidas americanas de controle de semicondutores e IA constitui o fio condutor de nossa hipotese. Quatro fases sao identificaveis, marcando uma escalada progressiva no escopo e intensidade das restricoes.",
    ]),
    ("Fase 1 - O choque inicial (7 de outubro de 2022)", [
        "O BIS (Bureau of Industry and Security) do Departamento de Comercio publica uma regra interina final que transforma radicalmente os controles de exportacao americanos sobre semicondutores. As medidas cobrem tres componentes: (i) controles sobre chips de computacao avancada (GPUs acima dos limites de desempenho definidos por TTP - Total Processing Performance), (ii) controles sobre equipamentos de fabricacao de semicondutores (SME, incluindo litografias EUV e DUV), e (iii) restricoes sobre atividades de US persons que apoiem a fabricacao de chips avancados na China.[15] O alvo e explicitamente a China: o objetivo declarado e impedir a modernizacao militar chinesa por meio do acesso ao compute IA. As restricoes incluem tres novas regras Foreign Direct Product (FDP) que estendem a jurisdicao americana a produtos fabricados fora dos Estados Unidos se tecnologias americanas forem utilizadas em sua producao.",
    ]),
    ("Fase 2 - Fechando as brechas (outubro de 2023)", [
        "O BIS publica duas novas regras interinas que reforcam e ampliam os controles de outubro de 2022. Os limites tecnicos sao ajustados para capturar chips Nvidia projetados especificamente para contornar restricoes (A800, H800). O escopo geografico e estendido a aproximadamente 40 paises adicionais (Country Groups D:1, D:4 e D:5), com regime de licenciamento diferenciado por categoria de pais.[16] Os controles sobre equipamentos de fabricacao sao aprofundados. E durante essa fase que os atores europeus comecam a perceber os efeitos indiretos das restricoes, ainda que a UE nao seja o alvo principal.",
    ]),
    ("Fase 3 - A AI Diffusion Rule (janeiro de 2025, Biden)", [
        "A administracao Biden publica em janeiro de 2025 a AI Diffusion Rule, que representa uma mudanca qualitativa. Pela primeira vez, as restricoes se aplicam nao apenas aos chips fisicos mas tambem aos pesos de modelos de IA e ao acesso a compute em nuvem. A regra classifica 120 paises em tres categorias: (i) aliados de confianca (acesso amplo), (ii) paises intermediarios (cotas), (iii) paises embargados.[17] As reacoes europeias sao fortes: o Parlamento Europeu alerta sobre restricoes que ameacam a capacidade da UE de treinar modelos em suas AI Factories. Franca e Alemanha sao classificadas como parceiros confiaveis, mas outros Estados-membros enfrentam tetos sobre os volumes de GPUs importaveis.",
    ]),
    ("Fase 4 - A ruptura Trump: Secao 232 e protecionismo explicito (janeiro de 2026)", [
        "Em 14 de janeiro de 2026, o presidente Trump assina a Proclamacao 11002, invocando a Secao 232 do Trade Expansion Act de 1962.[18] Essa acao marca uma ruptura qualitativa com as fases anteriores, por tres razoes.",
        "Primeiro, o instrumento juridico muda. As fases 1 a 3 estavam sob controles de exportacao (regulacao de exportacao, logica defensiva de seguranca nacional). A Secao 232 e um instrumento tarifario (direitos de importacao), cuja logica e protecionista: visa proteger a producao domestica, e nao apenas restringir o acesso de um adversario.",
        "Segundo, a tarifa cria uma vantagem competitiva explicita para as empresas americanas. A tarifa de 25 por cento atinge GPUs avancados (H200, MI325X) importados, exceto se destinados a usos domesticos americanos: data centers dos EUA, P&D, startups, setor publico, aplicacoes industriais e de consumo fora de data centers. Concretamente, uma empresa americana usando esses chips em solo americano nao paga a tarifa; uma empresa estrangeira importando os mesmos chips para reexportacao para a China paga 25 por cento.[19]",
        "Terceiro, a proclamacao anuncia uma expansao iminente. O texto preve que o Secretario de Comercio e o USTR negociem em 90 dias com paises produtores de semicondutores, e que tarifas mais amplas, acompanhadas por um programa de tariff offset para empresas que invistam na producao americana, possam ser impostas. O Secretario de Comercio deve fornecer ate julho de 2026 um relatorio sobre o mercado de semicondutores usados nos data centers americanos.[20]",
    ]),
    ("3.4.5 Interpretacao: do controle defensivo ao protecionismo ofensivo", [
        "A sequencia 2022-2026 revela uma transformacao qualitativa da politica americana. As fases Biden (2022-2025) seguem uma logica de estrategia de denial: impedir um adversario designado (a China) de acessar tecnologias chave, dentro de um quadro multilateral (coordenacao com Japao, Holanda). A fase Trump (2025-2026) adiciona uma logica de estrategia de capture: gerar receita (tarifa de 25 por cento), priorizar empresas americanas (isencoes domesticas), e usar o acesso ao compute como alavanca de negociacao com paises terceiros.",
        "E precisamente a transformacao que valida a hipotese central de nosso estudo: o protecionismo tecnologico americano nao se limita mais a negar o acesso a um adversario; ele constroi ativamente uma vantagem competitiva estrutural para as empresas americanas. As isencoes domesticas da Secao 232 significam que, pela primeira vez, o custo de acesso ao compute de ponta e legalmente diferenciado pela nacionalidade do usuario. Mesmo que os efeitos diretos sobre a UE permanecam limitados em janeiro de 2026 (a tarifa visa principalmente a reexportacao para a China), a proclamacao abre o caminho para uma expansao que poderia afetar diretamente a Europa, precisamente o territorio dos cenarios B, C e D de nossa analise.",
    ]),
    ("3.5 Sintese diagnostica: os elementos predeterminados de 2026", [
        "As quatro dimensoes do diagnostico convergem para um quadro coerente que estrutura os cenarios do Capitulo V.",
        "(1) A demanda por compute IA cresce exponencialmente (vendas de chips dobrando em dois anos, dobra do consumo energetico projetada em seis anos), e esse crescimento nao mostra sinais de desacelerar.",
        "(2) Essa demanda esta estruturalmente concentrada nos Estados Unidos. O snapshot do painel de abril de 2026 quantifica a parcela dos EUA em 76,9 por cento do compute IA operacional em H100-equivalentes e em 49,9 por cento se incluidos os clusters anunciados; 45 por cento do consumo eletrico dos data centers; e mais de 80 por cento dos investimentos privados em IA de fronteira. A concentracao no compute operacional aumenta com o tempo.",
        "(3) A Europa parte de uma posicao estruturalmente deficitaria: aproximadamente 4,4 por cento do compute operacional em H100-equivalentes, cerca de 15 por cento do consumo dos data centers, custos energeticos 1,4 a 1,7 vezes mais altos apos correcao PPA (e 2 a 3 vezes mais altos nas tarifas Eurostat industriais nao ajustadas), e 72 a 80 por cento de dependencia dos hyperscalers dos EUA para cargas de trabalho de IA em nuvem. Os planos de investimento da UE (Chips Act, AI Factories, SNIA) carregam orcamentos pelo menos uma ordem de magnitude menores do que os anuncios das Big Tech 2025 (320 bilhoes USD apenas para os cinco investidores dos EUA em IA de fronteira).",
        "(4) O quadro regulatorio americano cruzou um limiar qualitativo em janeiro de 2026 com a passagem dos controles de exportacao para as tarifas Secao 232, criando um mecanismo legal de diferenciacao dos custos de acesso ao compute por nacionalidade. A proclamacao sinaliza explicitamente a possibilidade de uma expansao, cujo escopo e calendario dependerao de variaveis politicas, precisamente as incertezas criticas que nossos cenarios exploram.",
        "O capitulo seguinte (Capitulo IV) analisa os mecanismos pelos quais essa assimetria de compute se traduz em vantagem competitiva mensuravel para as empresas americanas.",
    ]),
]

PT.table_blocks = [
    ("Tabela 4. Consumo eletrico dos data centers por regiao, 2020-2030.",
     "Fonte: AIE (2025), Energy and AI.",
     [
         ["Regiao", "2020", "2024", "2030 (AIE)", "Delta 24-30", "Parcela mundial 2024"],
         ["Estados Unidos", "~120 TWh", "~180 TWh", "~420 TWh", "+240 TWh", "45 pct"],
         ["China", "~60 TWh", "~102 TWh", "~280 TWh", "+175 TWh", "25 pct"],
         ["Europa (UE)", "~45 TWh", "~70 TWh", "~115 TWh", "+45 TWh", "15 pct"],
         ["Mundo", "~270 TWh", "~415 TWh", "~945 TWh", "+530 TWh", "100 pct"],
     ]),
    ("Tabela 5. Vendas globais de semicondutores, 2020-2026.",
     "Fontes: SIA/WSTS (fev. 2025, fev. 2026). A projecao 2026 baseia-se na previsao do outono de 2025 do WSTS (975,4 bilhoes USD) e na declaracao SIA de fevereiro de 2026.",
     [
         ["Ano", "2020", "2022", "2023", "2024", "2025", "2026 (proj.)"],
         ["Vendas (bilhoes USD, SIA)", "440", "556", "527", "631", "792", "~975-1000"],
         ["Crescimento (pct)", "+6,8", "+3,3", "-8,2", "+19,1", "+25,6", "+23-26"],
     ]),
    ("Tabela 6. Indicadores de dominancia dos EUA no compute IA (snapshot painel abril de 2026).",
     "Fontes: Epoch AI GPU Clusters dataset abril de 2026; AIE (2025); Sanchez/GeoCoded (2025).",
     [
         ["Indicador", "Estados Unidos", "China", "UE(13)", "Razao EUA/UE"],
         ["Parcela H100-eq operacional", "76,9 pct", "12,8 pct", "4,4 pct", "17,6:1"],
         ["Parcela F_total (op + planejado)", "49,9 pct", "0,5 pct", "3,3 pct", "15,1:1"],
         ["Parcela setor privado (2025)", "~65 pct global", "~12 pct", "~3 pct", "~22:1"],
         ["Consumo data centers 2024", "180 TWh", "102 TWh", "70 TWh", "2,6:1"],
         ["Investimento IA 2025", "320 bn USD (5 Big Tech)", "n.d.", "~20 bn EUR (UE total)", ">15:1"],
         ["CACI Power Mode (abril de 2026)", "100", "15,7", "28,9", "3,46:1"],
     ]),
    ("Tabela 7. Cronologia das medidas regulatorias americanas sobre semicondutores e IA (2022-2026).",
     "Fontes: BIS, Casa Branca, Pillsbury Law (2026), Gibson Dunn (2026).",
     [
         ["Data", "Admin.", "Medida", "Escopo / Alvo"],
         ["Out. 2022", "Biden", "Controles export BIS: GPUs avancados, SME, US persons", "China (militar)"],
         ["Out. 2023", "Biden", "Reforco de limites + extensao a 40 paises + HBM", "China + 40 paises (D:1/D:4/D:5)"],
         ["Dez. 2024", "Biden", "Onda 3: 24 tipos SME, HBM, 140 entidades, ECAD", "China (cadeia completa)"],
         ["Jan. 2025", "Biden", "AI Diffusion Rule: pesos de modelos, nuvem, 3 niveis paises", "120+ paises (efeitos UE incluidos)"],
         ["Jul. 2025", "Trump", "America AI Action Plan: desregulacao, compute dos EUA", "EUA (estrategia domestica)"],
         ["Set. 2025", "Trump", "Anuncio: vendas a China autorizadas contra 25 pct receitas", "China (monetizacao)"],
         ["Jan. 2026", "Trump", "Secao 232: tarifa 25 pct GPUs + licenca BIS China", "Mundial (isencao domestica EUA)"],
     ]),
]

PT.notes = EN.notes


# ===========================================================================
# Main
# ===========================================================================

def main() -> None:
    """Build the three Chapter III .docx files."""
    out_dir = Path(__file__).parent
    for lp in (EN, FR, PT):
        build(lp, out_dir)
    log.info("Done.")


if __name__ == "__main__":
    main()
